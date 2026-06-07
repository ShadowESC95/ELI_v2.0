# === PHASEN6_EXECUTOR_SYNTAX_PATCH ===
from __future__ import annotations
import hashlib
import subprocess
import time
import json
import requests
from datetime import datetime

from eli.cognition.introspection_agent import IntrospectionAgent
import os
import re
import shutil
from pathlib import Path

def _eli_path_get(obj, key, default=None):
    """
    Compatibility helper for ELI path containers.
    Accepts both dict-style path maps and object/namespace-style path maps.
    """
    if obj is None:
        return default
    if isinstance(obj, dict):
        return obj.get(key, default)
    return getattr(obj, key, default)


def _eli_generated_scripts_dir() -> Path:
    """Directory for generated scripts.

    Honors ELI_ARTIFACTS_DIR / ELI_DATA_DIR so writes can be redirected (tests
    isolate to a temp dir; installed deployments use the user data dir) instead of
    always landing in the source tree's artifacts/scripts. Default unchanged:
    <project_root>/artifacts/scripts.
    """
    root = Path(__file__).resolve().parents[2] / "artifacts"
    env = os.environ.get("ELI_ARTIFACTS_DIR") or os.environ.get("ELI_DATA_DIR")
    if env:
        base = Path(env).expanduser()
        root = base if base.is_absolute() else (Path(__file__).resolve().parents[2] / base)
    d = root / "scripts"
    d.mkdir(parents=True, exist_ok=True)
    return d

# === PHASEN_RESPONSE_MODE_HELPER ===
def _eli_attach_response_mode(result, action=None, args=None, meta=None):
    try:
        if isinstance(result, dict):
            _action = action or result.get("action")
            _args = args or result.get("args") or {}
            _meta = dict(meta or result.get("meta") or {})
            result.setdefault("meta", {})
            result["meta"].setdefault("response_mode", classify_response_mode(_action, _args, _meta))
        return result
    except Exception:
        return result

from eli.runtime.response_policy import classify_response_mode
from eli.core.paths import resolve_user_repo_path
from typing import Any, Dict, List, Optional, Tuple

from eli.core.paths import get_paths
import threading

# ── Security gate (lazy singleton) ───────────────────────────────────────────
_security_manager = None

def _get_security_manager():
    """Return the module-level SecurityManager singleton (imported lazily)."""
    global _security_manager
    if _security_manager is None:
        try:
            from eli.runtime.security import SecurityManager
            _security_manager = SecurityManager()
        except Exception:
            _security_manager = None  # degrade gracefully if unavailable
    return _security_manager
# --- Memory bridge: route through canonical Memory class (not memory_db) ---
# memory_db has a different schema (model/dim/vec columns) than Memory (kind/source/confidence/weight).
# Using Memory directly avoids the "no such column: model" crash.

def _get_canonical_memory():
    """Lazy import to avoid circular import."""
    from eli.memory import get_memory
    return get_memory()

def add_memory(text, tags=""):
    """Store via canonical Memory.store_memory()."""
    try:
        mem = _get_canonical_memory()
        tag_list = [t.strip() for t in str(tags).split(",") if t.strip()] if tags else []
        return mem.store_memory(text, tags=tag_list)
    except Exception as e:
        return {"ok": False, "error": str(e)}

def search_memory(query="", k=5, **kwargs):
    """Search via canonical Memory.recall_memory(), with conversation fallback."""
    try:
        mem = _get_canonical_memory()
        results = list(mem.recall_memory(query, limit=k))
        if hasattr(mem, "search_conversations") and len(results) < int(k or 5):
            conv = mem.search_conversations(query, limit=max(1, int(k or 5) - len(results)))
            for row in conv:
                results.append({
                    "timestamp": row.get("timestamp") or 0,
                    "text": row.get("content") or "",
                    "tags": "conversation",
                    "kind": "conversation",
                    "role": row.get("role") or "",
                    "session_id": row.get("session_id") or "",
                })
        results = results[: int(k or 5)]
        return {"ok": True, "results": results, "count": len(results)}
    except Exception as e:
        return {"ok": False, "error": str(e), "results": []}

# --- Compatibility shim: some older code expects search_memory_compat() ---
def search_memory_compat(query: str = "", limit: int = 10, q: str = "", k: int = 0, db_path=None):
    """Back-compat wrapper around Memory.recall_memory()."""
    qq = (q or query or "").strip()
    kk = int(k or limit or 10)
    return search_memory(query=qq, k=kk)


def _browser_user_dir() -> Path:
    try:
        base = Path(get_paths().config_dir)
    except Exception:
        base = Path.home() / ".config" / "eli"
    return (base / "browser").expanduser().resolve()

# === CONVERSATION LOGGING (append-only) ===
def _repo_root():
    try:
        return get_paths().artifacts_dir
    except Exception:
        return None

def _convlog_path():
    try:
        from datetime import datetime, timezone
        art = _repo_root()
        if art is None:
            return None
        d = Path(art) / "conversations"
        d.mkdir(parents=True, exist_ok=True)
        fn = datetime.now(timezone.utc).strftime("%Y%m%d") + ".jsonl"
        return d / fn
    except Exception:
        return None

def _convlog_append(role: str, text: str, meta: dict | None = None) -> None:
    try:
        import json
        from datetime import datetime, timezone

        path = _convlog_path()
        if path is None:
            return

        rec = {
            "ts_iso": datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S"),
            "ts_unix": float(time.time()),
            "role": str(role),
            "text": str(text),
            "meta": meta or {},
        }
        with open(path, "a", encoding="utf-8") as f:
            f.write(json.dumps(rec, ensure_ascii=False) + "\n")
    except Exception:
        return
# === END CONVERSATION LOGGING ===

# === MEMORY ===

def _get_memory_path():
    from eli.core.db_paths import get_user_db_path
    return get_user_db_path()


def _sqlite_init(conn):
    # Canonical Memory manages schema; keep as no-op for compatibility.
    return None


def memory_store(text, tags=None, kind="note", source="chat", confidence=0.7, ts=None):
    try:
        from eli.memory import get_memory
        mem = get_memory(db_path=_get_memory_path())
        if isinstance(tags, str):
            tags = [t.strip() for t in tags.split(",") if t.strip()]
        result = mem.store_memory(
            text=str(text),
            tags=tags or [],
            source=str(source),
            kind=str(kind),
            confidence=float(confidence),
        )
        return {"ok": True, "path": str(_get_memory_path())}
    except Exception as e:
        return {"ok": False, "error": repr(e), "path": str(_get_memory_path())}


def memory_recall(query, limit=20):
    try:
        from eli.memory import get_memory
        mem = get_memory(db_path=_get_memory_path())
        results = mem.recall_memory(str(query), limit=int(limit))
        return {"ok": True, "results": results, "path": str(_get_memory_path())}
    except Exception as e:
        return {"ok": False, "error": repr(e), "results": [], "path": str(_get_memory_path())}


def _ollama_warm_load(model: str, ollama_host: str) -> bool:
    """
    Best-effort warm-load so /api/ps can report digest for lockdown verification.
    Uses /api/chat (non-stream) with a tiny prompt.
    Returns True if the request likely caused the model to load (or is already loadable),
    False otherwise.
    """
    try:
        model = (model or "").strip()
        if not model:
            return False
        host = (ollama_host or "http://localhost:11434").rstrip("/")
        url = host + "/api/chat"

        payload = {
            "model": model,
            "stream": False,
            "messages": [{"role": "user", "content": "ping"}],
            "options": {"num_predict": 1},
        }

        ok, status, data, err = _ollama_http_post_json(url, payload, timeout_s=30)
        if not ok:
            return False

        # If we got a response object at all, the model was at least resolvable.
        # Ollama typically loads the model to answer, so this is a decent warm-load signal.
        if isinstance(data, dict):
            if "message" in data or "done" in data or "created_at" in data:
                return True
        return True
    except Exception:
        return False
# ------------------ END OLLAMA WARM LOAD HELPERS ------------------

# ----------------------------
# Core normalization utilities
# ----------------------------

def _normalize_result(res: object) -> Dict[str, Any]:
    """
    Guarantee GUI/voice always has:
      - res['ok'] (bool)
      - res['content'] (str)
      - res['response'] (str)
    """
    if not isinstance(res, dict):
        txt = "" if res is None else str(res)
        return {"ok": False, "error": "handler_returned_non_dict", "content": txt, "response": txt}

    if "ok" not in res:
        res["ok"] = False

    if "response" in res and "content" not in res:
        res["content"] = res.get("response")
    if "content" in res and "response" not in res:
        res["response"] = res.get("content")

    if "content" not in res or res.get("content") is None:
        fallback = res.get("message") or res.get("error") or ""
        res["content"] = str(fallback)
    if "response" not in res or res.get("response") is None:
        res["response"] = str(res.get("content") or "")

    res["content"] = "" if res["content"] is None else str(res["content"])
    res["response"] = "" if res["response"] is None else str(res["response"])
    return res


def _strip_ollama_artifacts(s: str) -> str:
    s = (s or "")
    s = s.replace("<|\nend\n", "").replace("<|end|>", "").replace("<|end|", "")
    s = re.sub(r"\s+$", "", s)
    return s.strip()


def _sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _shell_command_allowed_fallback(cmd: str) -> bool:
    """Fail-closed allowlist check used when the SecurityManager can't load.

    Mirrors SecurityManager.is_command_allowed so a degraded manager can never
    become fail-OPEN: with neither ELI_FULL_CONTROL nor ELI_ALLOWED_CMDS set,
    nothing runs.
    """
    if os.environ.get("ELI_FULL_CONTROL", "0") == "1":
        return True
    raw = (os.environ.get("ELI_ALLOWED_CMDS") or "").strip()
    if not raw:
        return False  # fail-closed
    allowed = {p for p in raw.replace(",", " ").split() if p}
    if "*" in allowed:
        return True
    return os.path.normpath(cmd) in allowed or cmd in allowed


def _run(argv: List[str], timeout: int = 15) -> Dict[str, Any]:
    # Security gate: check command against allowlist before executing.
    # Fail CLOSED — if the SecurityManager can't load, fall back to an inline
    # allowlist check rather than running unchecked.
    if argv:
        sm = _get_security_manager()
        allowed = sm.is_command_allowed(argv[0]) if sm is not None else _shell_command_allowed_fallback(argv[0])
        if not allowed:
            msg = f"Command blocked by security policy: {argv[0]}"
            return {
                "ok": False,
                "returncode": 1,
                "stdout": "",
                "stderr": msg,
                "cmd": argv,
                "content": msg,
                "response": msg,
            }
    try:
        p = subprocess.run(argv, timeout=timeout, capture_output=True, text=True, check=False)
        ok = (p.returncode == 0)
        msg = "OK" if ok else "Command failed"
        return {
            "ok": ok,
            "returncode": p.returncode,
            "stdout": p.stdout,
            "stderr": p.stderr,
            "cmd": argv,
            "content": msg,
            "response": msg,
        }
    except Exception as e:
        msg = "Command error"
        return {
            "ok": False,
            "returncode": 1,
            "stdout": "",
            "stderr": repr(e),
            "cmd": argv,
            "content": msg,
            "response": msg,
            "error": repr(e),
        }


# ----------------------------
# Grounded runtime audit helpers
# ----------------------------
def _canonical_runtime_file_map() -> Dict[str, Path]:
    root = get_paths().project_root.resolve()
    # Files live under eli/ subdir
    _eli = root / 'eli'
    root = _eli if _eli.exists() else root
    # Canonical runtime file paths under the live `eli/` layout.
    # phaseBW5 fix: replaced legacy `brain/cognition/...` and
    # `tools/automation/...` candidates that no longer correspond
    # to anything on disk and were producing hallucinated "FAIL"
    # entries in RUNTIME_AUDIT output.
    rels = {
        'cognitive_engine': ['kernel/engine.py'],
        'gguf_inference':   ['cognition/gguf_inference.py'],
        'memory':           ['memory/memory.py'],
        'memory_init':      ['memory/__init__.py'],
        'router':           ['execution/router_enhanced.py'],
        'executor':         ['execution/executor_enhanced.py'],
        'gui':              ['gui/eli_pro_audio_gui_MKI.py'],
        'runtime_settings': ['core/runtime_settings.py'],
        'paths':            ['core/paths.py'],
        'proactive_daemon': ['planning/proactive_daemon.py'],
        'agent_bus':        ['cognition/agent_bus.py'],
        'orchestrator':     ['cognition/orchestrator.py'],
        'inference_broker': ['cognition/inference_broker.py'],
        'output_governor':  ['cognition/output_governor.py'],
        'vector_store':     ['memory/vector_store.py'],
    }
    out: Dict[str, Path] = {}
    for key, candidates in rels.items():
        chosen = None
        for rel in candidates:
            cand = (root / rel).resolve()
            if cand.exists():
                chosen = cand
                break
        out[key] = chosen or (root / candidates[0]).resolve()
    return out

def _scan_merge_markers(lines: List[str]) -> List[Dict[str, Any]]:
    issues = []
    marker_re = re.compile(r'^(<{7}(?:\s+.+)?|={7}|>{7}(?:\s+.+)?)$')
    for idx, line in enumerate(lines, 1):
        stripped = line.strip()
        if marker_re.match(stripped):
            issues.append({'type': 'merge_marker', 'line': idx, 'message': stripped[:120]})
    return issues

def _scan_user_specific_paths(lines: List[str]) -> List[Dict[str, Any]]:
    issues = []
    path_re = re.compile(r"(/home/[^/\s]+(?:/[^\s]+)?)|(C:\\Users\\[^\\]+(?:\\[^\s]+)?)")
    for idx, line in enumerate(lines, 1):
        stripped = line.strip()
        if not stripped or stripped.startswith('#'):
            continue
        if ('re.compile(' in line or 're.search(' in line or 're.match(' in line or 'path_re =' in line or '/home/...' in line or '/home/\\S+' in line or 'C:\\Users\\[^\\]+' in line):
            continue
        for match in path_re.finditer(line):
            value = match.group(0)
            if value and ('\\S' not in value) and ('[^' not in value) and ('(?:' not in value):
                issues.append({'type': 'hardcoded_user_path', 'line': idx, 'message': value})
    return issues

def _scan_top_level_symbols(path: Path, text: str) -> List[Dict[str, Any]]:
    import ast
    issues: List[Dict[str, Any]] = []
    try:
        tree = ast.parse(text, filename=str(path))
    except SyntaxError as e:
        issues.append({'type': 'syntax_error', 'line': int(getattr(e, 'lineno', 0) or 0), 'message': str(e)})
        return issues
    seen: Dict[str, List[int]] = {}
    for node in tree.body:
        names: List[str] = []
        if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef, ast.ClassDef)):
            names = [node.name]
        elif isinstance(node, ast.Assign):
            for target in node.targets:
                if isinstance(target, ast.Name):
                    names.append(target.id)
        elif isinstance(node, ast.AnnAssign) and isinstance(node.target, ast.Name):
            names = [node.target.id]
        for name in names:
            seen.setdefault(name, []).append(int(getattr(node, 'lineno', 0) or 0))
    for name, lines in seen.items():
        if len(lines) > 1:
            issues.append({'type': 'duplicate_top_level_symbol', 'line': lines[1], 'message': f'{name} also defined at lines {lines}'})
    return issues

def _audit_python_file(path: Path) -> Dict[str, Any]:
    entry = {'path': str(path), 'status': 'PASS', 'issues': []}
    if not path.exists():
        entry['status'] = 'FAIL'
        entry['issues'].append({'type': 'missing_file', 'line': 0, 'message': 'file not found'})
        return entry
    text = path.read_text(encoding='utf-8', errors='replace')
    lines = text.splitlines()
    issues = []
    issues.extend(_scan_merge_markers(lines))
    issues.extend(_scan_user_specific_paths(lines))
    issues.extend(_scan_top_level_symbols(path, text))
    if not any(i['type'] == 'syntax_error' for i in issues):
        try:
            compile(text, str(path), 'exec')
        except SyntaxError as e:
            issues.append({'type': 'syntax_error', 'line': int(getattr(e, 'lineno', 0) or 0), 'message': str(e)})
    entry['issues'] = sorted(issues, key=lambda x: (int(x.get('line', 0) or 0), x.get('type', '')))
    if entry['issues']:
        entry['status'] = 'FAIL'
    return entry

def _runtime_health_probes() -> List[Dict[str, Any]]:
    """LIVE health checks a static source audit cannot catch: actually EXERCISE
    key subsystems and check DATA INTEGRITY, so real runtime faults surface in an
    audit (a method-level NameError, malformed data, recurring logged failures)
    instead of only appearing once at use time. Each probe: {name, ok, detail}."""
    probes: List[Dict[str, Any]] = []

    def _probe(name, fn):
        try:
            ok, detail = fn()
            probes.append({"name": name, "ok": bool(ok), "detail": str(detail)})
        except Exception as e:
            probes.append({"name": name, "ok": False, "detail": f"{type(e).__name__}: {e}"})

    def _db_path():
        from eli.memory import get_memory
        return getattr(get_memory(), "db_path", None)

    # Plugin manager — instantiate + list (catches the method-level NameError class
    # that import/syntax audits miss: the module imports fine, the fault is at call).
    def _plugins():
        from eli.plugins.manager import get_manager
        return True, f"OK — {len(get_manager().list_installed())} plugin(s) installed"
    _probe("plugin_manager", _plugins)

    # Memory — singleton resolves + db file present.
    def _mem():
        from pathlib import Path as _P
        dbp = _db_path()
        return (bool(dbp and _P(dbp).exists()),
                f"db={_P(dbp).name}" if dbp else "no db path")
    _probe("memory", _mem)

    # Agent bus — registry imports + has agents.
    def _bus():
        from eli.cognition.agent_bus import _ALL_AGENTS
        return bool(_ALL_AGENTS), f"{len(_ALL_AGENTS)} agents registered"
    _probe("agent_bus", _bus)

    # Habit-rule integrity — flag the un-schedulable corruption class.
    def _habits():
        import sqlite3 as _s
        dbp = _db_path()
        if not dbp:
            return True, "no db"
        c = _s.connect(str(dbp))
        try:
            bad = c.execute(
                "SELECT COUNT(*) FROM habit_rules WHERE enabled=1 AND (hour IS NULL "
                "OR minute IS NULL OR TRIM(LOWER(command))=TRIM(LOWER(name)))").fetchone()[0]
            tot = c.execute("SELECT COUNT(*) FROM habit_rules WHERE enabled=1").fetchone()[0]
        finally:
            c.close()
        if bad:
            return False, f"{bad} of {tot} enabled habit rule(s) un-schedulable (NULL time / bare-token command)"
        return True, f"{tot} enabled habit rule(s), all schedulable"
    _probe("habit_integrity", _habits)

    # Recent LIVE failures — surface what the failure log recorded so the audit
    # reflects live problems, not just static source state. Reads the same
    # canonical store the Self-Improve panel uses; resolved entries are filtered
    # out by get_recent_failures, so the audit and the panel agree.
    def _failures():
        from eli.runtime.self_improvement import get_self_improvement
        rows = get_self_improvement().memory.get_recent_failures(limit=5)
        if not rows:
            return True, "no live failures logged"
        summary = "; ".join(
            f"{(r.get('error') or r.get('user_input') or '')[:70]} (×{r.get('occurrence_count', 1)})"
            for r in rows)
        return False, f"{len(rows)} recent live failure(s): {summary}"
    _probe("recent_failures", _failures)

    return probes


def _runtime_audit_report() -> Dict[str, Any]:
    files = _canonical_runtime_file_map()
    entries = []
    for key in ['cognitive_engine', 'gguf_inference', 'memory', 'memory_init', 'router', 'executor', 'gui', 'runtime_settings', 'paths', 'proactive_daemon']:
        entries.append(_audit_python_file(files[key]))
    probes = _runtime_health_probes()
    return {'ok': True, 'entries': entries, 'health_probes': probes}

def _format_runtime_audit(report: Dict[str, Any]) -> str:
    lines = []
    for entry in report.get('entries', []):
        lines.append(f"{entry['status']} {entry['path']}")
        for issue in entry.get('issues', []):
            line_no = int(issue.get('line', 0) or 0)
            prefix = f"  - line {line_no}" if line_no else '  - line ?'
            lines.append(f"{prefix} [{issue.get('type')}] {issue.get('message')}")
    probes = report.get('health_probes') or []
    if probes:
        lines.append("")
        lines.append("Live health probes:")
        for p in probes:
            mark = "✅" if p.get("ok") else "❌"
            lines.append(f"  {mark} {p.get('name')}: {p.get('detail')}")
    return '\n'.join(lines) if lines else 'No runtime files audited.'

def _gpu_status_report() -> Dict[str, Any]:
    query = [
        "nvidia-smi",
        "--query-gpu=name,memory.total,memory.used,memory.free,utilization.gpu,temperature.gpu,power.draw,power.limit,driver_version",
        "--format=csv,noheader,nounits",
    ]
    runtime_snapshot: Dict[str, Any] = {}
    try:
        snap_path = get_paths().artifacts_dir / "runtime_snapshot.json"
        if snap_path.exists():
            runtime_snapshot = json.loads(snap_path.read_text(encoding="utf-8"))
    except Exception:
        runtime_snapshot = {}

    try:
        proc = subprocess.run(
            query,
            timeout=6,
            capture_output=True,
            text=True,
            check=False,
        )
    except Exception as exc:
        msg = f"GPU status unavailable: nvidia-smi could not run ({exc})."
        return {"ok": False, "error": str(exc), "content": msg, "response": msg}

    if proc.returncode != 0:
        msg = (proc.stderr or proc.stdout or "nvidia-smi returned no output").strip()
        msg = f"GPU status unavailable: {msg}"
        return {"ok": False, "returncode": proc.returncode, "content": msg, "response": msg}

    import csv
    rows = list(csv.reader((proc.stdout or "").splitlines()))
    if not rows:
        msg = "GPU status unavailable: nvidia-smi returned no GPU rows."
        return {"ok": False, "content": msg, "response": msg}

    fields = [
        "name", "memory_total_mib", "memory_used_mib", "memory_free_mib",
        "utilization_percent", "temperature_c", "power_draw_w", "power_limit_w",
        "driver_version",
    ]
    parsed = []
    for row in rows:
        values = [cell.strip() for cell in row]
        item = {field: (values[i] if i < len(values) else "") for i, field in enumerate(fields)}
        parsed.append(item)

    first = parsed[0]
    def _num(value: Any) -> Optional[float]:
        try:
            cleaned = re.sub(r"[^0-9.]+", "", str(value or ""))
            return float(cleaned) if cleaned else None
        except Exception:
            return None

    total = _num(first.get("memory_total_mib"))
    used = _num(first.get("memory_used_mib"))
    free = _num(first.get("memory_free_mib"))
    util = _num(first.get("utilization_percent"))
    used_pct = (used / total * 100.0) if total and used is not None else None

    lines = [
        "GPU status:",
        f"- name: {first.get('name') or 'unknown'}",
        f"- driver: {first.get('driver_version') or 'unknown'}",
    ]
    if total is not None and used is not None and free is not None:
        lines.append(
            f"- VRAM: {used:.0f} MiB used / {total:.0f} MiB total "
            f"({free:.0f} MiB free, {used_pct:.1f}% used)"
        )
    else:
        lines.append("- VRAM: unavailable")
    lines.extend([
        f"- GPU utilization: {util:.0f}%" if util is not None else "- GPU utilization: unavailable",
        f"- temperature: {first.get('temperature_c') or 'unknown'} C",
        f"- power: {first.get('power_draw_w') or 'unknown'} W / {first.get('power_limit_w') or 'unknown'} W limit",
    ])

    if runtime_snapshot:
        lines.extend([
            "",
            "ELI selected llama.cpp load parameters:",
            f"- context: {runtime_snapshot.get('n_ctx', 'unknown')}",
            f"- GPU-layer parameter: {runtime_snapshot.get('n_gpu_layers', 'unknown')}",
            f"- batch: {runtime_snapshot.get('n_batch', 'unknown')}",
            f"- CPU threads: {runtime_snapshot.get('n_threads', 'unknown')}",
        ])

    lines.extend([
        "",
        "Performance reading:",
        "- The live runtime is constrained by available VRAM. On a 4 GB GPU, large context and high GPU-layer counts can fail and force fallback.",
        "- If ELI booted with lower selected ctx/GPU-layer parameters than requested, the fallback is expected behavior, not a settings lie.",
        "- Higher free VRAM gives more room for GPU layers or batch size; low free VRAM means slower CPU-heavy inference and longer response times.",
    ])
    msg = "\n".join(lines)
    return {"ok": True, "gpus": parsed, "runtime_snapshot": runtime_snapshot, "content": msg, "response": msg}

def _self_improvement_log_report(limit: int = 5, days: int = 30) -> Dict[str, Any]:
    try:
        from eli.runtime.self_improvement import get_self_improvement
        engine = get_self_improvement()
        rows = engine.analyze_failures(limit=max(1, int(limit)), days=max(1, int(days)), min_cluster_size=1)
    except Exception as exc:
        msg = f"Self-improvement log unavailable: {exc}"
        return {"ok": False, "error": str(exc), "content": msg, "response": msg}

    rows = sorted(rows or [], key=lambda r: float(r.get("timestamp") or 0.0), reverse=True)
    if not rows:
        msg = "No failures are recorded in the self-improvement log for the selected window."
        return {"ok": True, "failures": [], "content": msg, "response": msg}

    last = rows[0]
    ts = float(last.get("timestamp") or 0.0)
    when = datetime.fromtimestamp(ts).isoformat(timespec="seconds") if ts else "unknown"
    err = str(last.get("error") or "").strip() or "(empty error field)"
    user_input = str(last.get("user_input") or "").strip()
    count = int(last.get("occurrence_count") or 1)

    lines = [
        "Last self-improvement failure:",
        f"- when: {when}",
        f"- occurrences: {count}",
        f"- exact_error: {err}",
    ]
    if user_input:
        preview = user_input.replace("\n", "\\n")
        if len(preview) > 500:
            preview = preview[:500] + "..."
        lines.append(f"- triggering_input_preview: {preview}")

    if len(rows) > 1:
        lines.append("")
        lines.append("Recent failures:")
        for idx, row in enumerate(rows[:limit], 1):
            row_ts = float(row.get("timestamp") or 0.0)
            row_when = datetime.fromtimestamp(row_ts).isoformat(timespec="seconds") if row_ts else "unknown"
            row_err = str(row.get("error") or "").strip() or "(empty error field)"
            lines.append(f"{idx}. {row_when} | x{int(row.get('occurrence_count') or 1)} | {row_err[:240]}")

    msg = "\n".join(lines)
    return {"ok": True, "failures": rows[:limit], "content": msg, "response": msg}

def _import_audit_report() -> Dict[str, Any]:
    import importlib
    modules = [
        'eli.kernel.engine',
        'eli.cognition.gguf_inference',
        'eli.memory.memory',
        'eli.memory',
        'eli.execution.router_enhanced',
        'eli.execution.executor_enhanced',
        'eli.gui.eli_pro_audio_gui_MKI',
        'eli.core.runtime_settings',
        'eli.core.paths',
        'eli.planning.proactive_daemon',
    ]
    entries = []
    for mod in modules:
        status = 'PASS'; error = ''; fix = ''
        try:
            importlib.import_module(mod)
        except SystemExit:
            status = "SKIP"; error = "Module calls sys.exit() on import (e.g. missing Qt)"; fix = "Install PySide6 (LGPLv3 — `pip install PySide6`) or suppress sys.exit in GUI module."
        except Exception as e:
            status = 'FAIL'; error = f'{type(e).__name__}: {e}'
            if 'No module named' in error:
                fix = 'Check package path, missing module, or PYTHONPATH wiring.'
            elif 'SyntaxError' in error or 'unterminated string literal' in error:
                fix = 'Fix syntax errors in the target module before import.'
            else:
                fix = 'Inspect traceback and referenced dependency or symbol.'
        entries.append({'module': mod, 'status': status, 'error': error, 'suggested_fix': fix})
    return {'ok': True, 'entries': entries}

def _format_import_audit(report: Dict[str, Any]) -> str:
    lines = []
    for entry in report.get('entries', []):
        lines.append(f"{entry['module']} | {entry['status']} | {entry.get('error', '') or '-'} | {entry.get('suggested_fix', '') or '-'}")
    return '\n'.join(lines) if lines else 'No import audit results.'


def _diagnose_executor_wrappers() -> Dict[str, Any]:
    """Walk the executor wrapper stack and report:
       1. The canonical middleware table (named, ordered)
       2. The legacy ORIG_EXECUTE rebind chain (discovered via _ELI_*_ORIG_EXECUTE globals)
       3. For a small set of representative actions, which layer would handle them
    """
    g = globals()
    report: Dict[str, Any] = {
        "middleware_table": [],
        "rebind_chain": [],
        "action_handlers": {},
    }
    # 1. canonical middleware table — already structured as ((name, fn), ...)
    mw_table = g.get("_ELI_EXECUTOR_MIDDLEWARE_TABLE") or ()
    for entry in mw_table:
        try:
            name, fn = entry
            report["middleware_table"].append({
                "name": str(name),
                "callable": getattr(fn, "__name__", repr(fn)),
                "module": getattr(fn, "__module__", "?"),
            })
        except Exception:
            continue
    # 2. rebind chain — discover all _ELI_*_ORIG_EXECUTE globals
    for key, val in sorted(g.items()):
        if not isinstance(key, str):
            continue
        if not key.startswith("_ELI_") or "ORIG_EXECUTE" not in key:
            continue
        report["rebind_chain"].append({
            "sentinel": key,
            "target_callable": getattr(val, "__name__", repr(val) if val is not None else "<unbound>"),
            "is_callable": bool(callable(val)),
        })
    # 3. current top-level execute identity (which wrapper sits on top)
    top_execute = g.get("execute")
    report["active_top_execute"] = {
        "name": getattr(top_execute, "__name__", repr(top_execute)),
        "module": getattr(top_execute, "__module__", "?"),
    }
    report["active_top_execute_action"] = {
        "name": getattr(g.get("execute_action"), "__name__", repr(g.get("execute_action"))),
        "module": getattr(g.get("execute_action"), "__module__", "?"),
    }
    # 4. canonical base (what the middleware table eventually wraps)
    base = g.get("_ELI_EXECUTOR_CANONICAL_BASE")
    if base is not None:
        report["canonical_base"] = {
            "name": getattr(base, "__name__", repr(base)),
            "module": getattr(base, "__module__", "?"),
        }
    # 5. probe a representative set of actions — which middleware short-circuits which
    probes = [
        "CHAT", "GENERATE_SCRIPT", "FIX_FILE", "SELF_REPORT", "RUNTIME_AUDIT",
        "MEMORY_RECALL", "MEMORY_STATUS", "PERSONAL_MEMORY_SUMMARY",
        "PLAY_MEDIA", "PAUSE_MEDIA", "VOLUME", "OPEN_APP", "CLOSE_APP",
        "REASONING_MODE_STATUS", "GUI_RUNTIME_AUDIT",
    ]
    handlers: Dict[str, str] = {}
    for action in probes:
        # Heuristic: each middleware's filter is in its docstring or name. We don't
        # actually invoke; we just record that the middleware MIGHT handle it.
        handlers[action] = "via canonical middleware table → rebind chain → _execute_impl"
    report["action_handlers"] = handlers
    return report


def _format_executor_wrappers(report: Dict[str, Any]) -> str:
    lines = ["EXECUTOR WRAPPER DIAGNOSTIC", "=" * 60, ""]
    lines.append(f"Top-level execute:        {report.get('active_top_execute', {}).get('name')}")
    lines.append(f"Top-level execute_action: {report.get('active_top_execute_action', {}).get('name')}")
    base = report.get("canonical_base")
    if base:
        lines.append(f"Canonical base:           {base.get('name')}")
    lines.append("")
    mw = report.get("middleware_table") or []
    lines.append(f"Canonical middleware table ({len(mw)} entries, top-to-bottom):")
    for i, entry in enumerate(mw, 1):
        lines.append(f"  {i:>2}. {entry['name']:<32} → {entry['callable']}")
    lines.append("")
    rebinds = report.get("rebind_chain") or []
    lines.append(f"Legacy ORIG_EXECUTE sentinels ({len(rebinds)} found):")
    for entry in rebinds:
        status = "callable" if entry["is_callable"] else "unbound"
        lines.append(f"  - {entry['sentinel']:<48} [{status}] → {entry['target_callable']}")
    lines.append("")
    lines.append("Resolution order for any action:")
    lines.append("  1. Top-level execute() — outermost wrapper")
    lines.append("  2. Each middleware in the table (in order)")
    lines.append("  3. _ELI_EXECUTOR_CANONICAL_BASE — the legacy rebind chain")
    lines.append("  4. _execute_impl — the original 160+ action dispatcher")
    return "\n".join(lines)


def _resolve_runtime_paths_report() -> Dict[str, Any]:
    p = get_paths()
    project_root = Path(p.project_root).expanduser().resolve()
    entries = {
        'project_root': str(project_root),
        'artifacts': str(Path(p.artifacts_dir).expanduser().resolve()),
        'config': str(Path(p.config_dir).expanduser().resolve()),
        'notes': str(Path(getattr(p, 'notes_dir', '')).expanduser().resolve()) if getattr(p, 'notes_dir', '') else '',
        'models': str(Path(p.models_dir).expanduser().resolve()),
        'user_db': str(Path(p.user_db).expanduser().resolve()),
        'agent_db': str(Path(p.agent_db).expanduser().resolve()),
    }
    flags = {}
    for key, value in entries.items():
        resolved = Path(value).expanduser().resolve() if value else project_root
        try:
            resolved.relative_to(project_root)
            inside_project = True
        except Exception:
            inside_project = False
        machine_absolute = resolved.is_absolute()
        home_bound = str(resolved).startswith(str(Path.home().expanduser().resolve()))
        user_specific = bool(home_bound and not inside_project)
        non_redistributable = bool(user_specific)
        flags[key] = {
            'path': str(resolved),
            'inside_project': inside_project,
            'machine_absolute': machine_absolute,
            'user_specific': user_specific,
            'non_redistributable': non_redistributable,
        }
    critical_files = {
        key: str(path.expanduser().resolve())
        for key, path in _canonical_runtime_file_map().items()
    }
    import_audit = _import_audit_report()
    return {
        'ok': True,
        'entries': flags,
        'critical_files': critical_files,
        'import_audit': import_audit,
    }

def _format_runtime_paths(report: Dict[str, Any]) -> str:
    lines = []
    critical_files = report.get('critical_files') or {}
    if critical_files:
        lines.append("Critical runtime files:")
        for key, path in critical_files.items():
            lines.append(f"- {key}: {path}")

    import_audit = report.get('import_audit') or {}
    if import_audit.get('entries'):
        if lines:
            lines.append("")
        lines.append("Current import status:")
        for entry in import_audit.get('entries', []):
            error = entry.get('error') or '-'
            lines.append(f"- {entry.get('module')}: {entry.get('status')} ({error})")

    if lines:
        lines.append("")
    lines.append("Runtime directories:")
    for key, meta in report.get('entries', {}).items():
        lines.append(
            f"- {key}: {meta['path']} | "
            f"inside_project={meta.get('inside_project')} "
            f"user_specific={meta.get('user_specific')} "
            f"non_redistributable={meta.get('non_redistributable')}"
        )
    return "\n".join(lines)

def _gui_runtime_audit_report() -> Dict[str, Any]:
    gui_path = _canonical_runtime_file_map()['gui']
    entry = {'path': str(gui_path), 'status': 'PASS', 'issues': [], 'evidence': {}, 'snippets': {}}
    if not gui_path.exists():
        entry['status'] = 'FAIL'; entry['issues'].append({'type': 'missing_file', 'line': 0, 'message': 'file not found'})
        return {'ok': True, 'entry': entry}
    text = gui_path.read_text(encoding='utf-8', errors='replace')
    lines = text.splitlines()
    patterns = {
        'router': 'router_enhanced',
        'executor': 'executor_enhanced',
        'cognition': 'cognitive_engine',
        'proactive_import': 'start_daemon',
        'send_worker': 'threading.Thread(target=generate_worker',
    }
    for key, needle in patterns.items():
        hits = [i for i, line in enumerate(lines, 1) if needle in line]
        entry['evidence'][key] = hits
        # Store actual code at the first 3 matching lines so the model reports
        # what is really there rather than fabricating plausible-looking code.
        entry['snippets'][key] = [
            f"  line {lineno}: {lines[lineno-1].strip()}"
            for lineno in hits[:3]
        ]
        if not hits:
            entry['status'] = 'FAIL'
            entry['issues'].append({'type': 'missing_hook', 'line': 0, 'message': f'{key} hook not found: {needle}'})
    entry['issues'].extend(_scan_user_specific_paths(lines))
    if entry['issues']:
        entry['status'] = 'FAIL'
    return {'ok': True, 'entry': entry}

def _format_gui_runtime_audit(report: Dict[str, Any]) -> str:
    entry = report.get('entry', {})
    out = [f"{entry.get('status', 'FAIL')} {entry.get('path', '')}"]
    evidence = entry.get('evidence') or {}
    snippets = entry.get('snippets') or {}
    for key, hits in evidence.items():
        out.append(f"  - {key}: lines {hits if hits else 'MISSING'}")
        for snip in snippets.get(key, []):
            out.append(f"    {snip}")
    for issue in entry.get('issues', []):
        line_no = int(issue.get('line', 0) or 0)
        prefix = f"  - line {line_no}" if line_no else '  - line ?'
        out.append(f"{prefix} [{issue.get('type')}] {issue.get('message')}")
    return '\n'.join(out)

def _explain_memory_runtime_report() -> Dict[str, Any]:
    import sqlite3
    from eli.memory import (
        get_memory as _executor_get_memory,
        get_memory_status as _executor_get_memory_status,
        resolve_db_paths as _executor_resolve_db_paths,
    )

    paths = _executor_resolve_db_paths()
    mem = _executor_get_memory()
    status_candidates: List[Dict[str, Any]] = []
    for candidate in (
        getattr(mem, 'db_path', None),
        _eli_path_get(paths, "memory_db"),
        _eli_path_get(paths, "user_db"),
    ):
        if candidate in (None, ""):
            continue
        try:
            status_candidates.append(_executor_get_memory_status(candidate))
        except Exception:
            pass
    if status_candidates:
        status = max(
            status_candidates,
            key=lambda item: int(item.get("memory_entries", 0) or 0)
            + int(item.get("conversation_turns", 0) or 0),
        )
    else:
        status = _executor_get_memory_status(getattr(mem, 'db_path', None))

    name_guess = None
    try:
        from eli.kernel.state import get_user_name as _gun_exec
        name_guess = _gun_exec().strip() or None
    except Exception:
        pass

    try:
        recent = mem.get_recent_conversation(limit=6)
    except Exception:
        recent = []

    def _table_counts(path: Any) -> Dict[str, int]:
        out: Dict[str, int] = {}
        try:
            p = Path(path)
            if not p.exists():
                return out
            conn = sqlite3.connect(str(p))
            rows = conn.execute(
                "SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%' ORDER BY name"
            ).fetchall()
            for (table_name,) in rows:
                try:
                    count_row = conn.execute(f'SELECT COUNT(*) FROM "{table_name}"').fetchone()
                    out[str(table_name)] = int((count_row or [0])[0] or 0)
                except Exception:
                    out[str(table_name)] = -1
            conn.close()
        except Exception:
            return out
        return out

    def _identity_hits_from_sqlite(*db_paths: Any, limit: int = 5) -> List[str]:
        hits: List[str] = []
        wanted_tables = ("memories", "memory", "semantic")
        wanted_cols = ("text", "content", "fact", "value", "summary")
        for raw_path in db_paths:
            try:
                p = Path(raw_path)
                if not p.exists():
                    continue
                conn = sqlite3.connect(str(p))
                tables = {
                    row[0] for row in conn.execute(
                        "SELECT name FROM sqlite_master WHERE type='table'"
                    ).fetchall()
                }
                for table in wanted_tables:
                    if table not in tables:
                        continue
                    cols = [row[1] for row in conn.execute(f'PRAGMA table_info("{table}")').fetchall()]
                    text_col = next((c for c in wanted_cols if c in cols), None)
                    if not text_col:
                        continue
                    query = (
                        f'SELECT "{text_col}" FROM "{table}" '
                        f'WHERE lower("{text_col}") LIKE "%name%" '
                        f'OR lower("{text_col}") LIKE "%identity%" '
                        f'OR lower("{text_col}") LIKE "%preference%" '
                        f'ORDER BY rowid DESC LIMIT ?'
                    )
                    for (value,) in conn.execute(query, (limit,)).fetchall():
                        cleaned = " ".join(str(value or "").split())
                        if cleaned:
                            hits.append(cleaned[:160])
                            if len(hits) >= limit:
                                conn.close()
                                return hits
                conn.close()
            except Exception:
                continue
        return hits

    vector_status: Dict[str, Any] = {}
    try:
        from eli.memory import vector_store as _vs
        idx_path, meta_path = _vs._get_index_paths()
        vector_status = {
            "faiss_available": bool(getattr(_vs, "FAISS_AVAILABLE", False)),
            "index_path": str(idx_path),
            "meta_path": str(meta_path),
            "index_exists": Path(idx_path).exists(),
            "meta_exists": Path(meta_path).exists(),
            "embedding_model": str((get_paths().project_root / "models" / "embeddings" / "nomic-embed-text-v1.5.Q4_K_M.gguf").resolve()),
        }
        if vector_status["faiss_available"] and vector_status["index_exists"]:
            import faiss
            vector_status["ntotal"] = int(faiss.read_index(str(idx_path)).ntotal)
    except Exception as e:
        vector_status = {"ok": False, "error": repr(e)}

    schema = {
        "active_db": _table_counts(status.get("db_path")),
        "user_db": _table_counts(paths.user_db),
        "agent_db": _table_counts(paths.agent_db),
        "memory_db": _table_counts(_eli_path_get(paths, "memory_db")),
    }
    identity_hits = _identity_hits_from_sqlite(_eli_path_get(paths, "memory_db"), paths.user_db)

    # Enumerate ALL physical sqlite3 files actually on disk in the db directory —
    # not just the logical roles. The role resolver aliases active_db/user_db/
    # memory_db to the SAME user.sqlite3 and never sees siblings like
    # coding_memory.sqlite3 or system_index.sqlite3, so the report claimed "2
    # databases" when there are 4 on disk (user-reported, 2026-06-06: "you have at least
    # 4 sqlite databases"). Ground the count in the filesystem.
    db_files: List[Dict[str, Any]] = []
    try:
        _db_dir = Path(paths.user_db).parent
        for _f in sorted(_db_dir.glob("*.sqlite3")):
            try:
                _tc = _table_counts(_f)
                db_files.append({
                    "name": _f.name,
                    "path": str(_f),
                    "size_mb": round(_f.stat().st_size / (1024 * 1024), 3),
                    "tables": len(_tc),
                    "table_names": sorted(_tc.keys()),
                })
            except Exception:
                continue
    except Exception:
        pass

    # ── Live mechanism probes (#5/Option 4) ───────────────────────────────
    # Re-derive every claim about ELI's own internals each call rather than
    # hardcoding prose: ELI can add/remove its own modules and tables, so a
    # static description would lie. Probe the filesystem for each mechanism
    # module and the live schema for FTS5 mirrors + knowledge-graph counts.
    def _all_live_tables() -> set:
        names: set = set()
        for _f in db_files:
            for _n in (_f.get("table_names") or []):
                names.add(str(_n))
        for _tbls in (schema or {}).values():
            for _n in (_tbls or {}):
                names.add(str(_n))
        return names

    _live_tables = _all_live_tables()
    # FTS5 virtual tables are named "<x>_fts"; sqlite also creates shadow tables
    # (_fts_config/_data/_docsize/_idx) — exclude those, keep the real mirrors.
    fts_tables = sorted(
        t for t in _live_tables
        if t.endswith("_fts") and not re.search(r"_fts_(config|data|docsize|idx)$", t)
    )

    def _live_count(table: str) -> int:
        best = -1
        for src in (schema or {}).values():
            if src and table in src:
                best = max(best, int(src.get(table) or 0))
        return best

    kg = {
        "entities": _live_count("kg_entities"),
        "relations": _live_count("kg_relations"),
        "present": ("kg_entities" in _live_tables or "kg_relations" in _live_tables),
    }

    _proj_root = get_paths().project_root
    def _probe_module(relpath: str) -> Dict[str, Any]:
        try:
            return {"path": relpath, "present": (Path(_proj_root) / relpath).exists()}
        except Exception:
            return {"path": relpath, "present": False}

    mechanisms = {
        "hyde": _probe_module("eli/cognition/hyde.py"),
        "orchestrator": _probe_module("eli/cognition/orchestrator.py"),
        "agent_bus": _probe_module("eli/cognition/agent_bus.py"),
        "vector_store_faiss": _probe_module("eli/memory/vector_store.py"),
        "knowledge_graph": _probe_module("eli/memory/knowledge_graph.py"),
        "plan_graph_dag": _probe_module("eli/coding/plan_graph.py"),
    }

    return {
        'ok': True,
        'paths': {
            'active_db': str(status.get("db_path", getattr(mem, "db_path", "unknown"))),
            'user_db': str(paths.user_db),
            'agent_db': str(paths.agent_db),
            'memory_db': str(_eli_path_get(paths, "memory_db")),
        },
        'db_files': db_files,
        'status': status,
        'schema': schema,
        'vector_status': vector_status,
        'fts_tables': fts_tables,
        'kg': kg,
        'mechanisms': mechanisms,
        'name_guess': name_guess,
        'identity_hits': [
            str((h.get('text') or h.get('content') or '') if isinstance(h, dict) else h).strip()[:160]
            for h in identity_hits[:5]
        ],
        'recent_turn_count': min(6, int(status.get("conversation_turns", 0) or len(recent or []))),
    }

def _format_memory_runtime(report: Dict[str, Any]) -> str:
    if not report.get("ok"):
        return f"Memory runtime inspection failed: {report.get('error', 'unknown error')}"

    paths = report.get("paths", {}) or {}
    status = report.get("status", {}) or {}
    name_guess = str(report.get("name_guess") or "").strip()
    identity_hits = report.get("identity_hits") or []
    recent_turn_count = int(report.get("recent_turn_count") or 0)

    def _clean_hit(value: Any) -> str:
        return " ".join(str(value or "").split())

    lines: List[str] = []

    if name_guess:
        lines.append(
            f"Yes — I do have grounded memory evidence for you. "
            f"The strongest current name signal in runtime memory is: {name_guess}."
        )
        lines.append("")
    else:
        lines.append(
            "I do have grounded runtime memory state, but I do not currently have a strong "
            "enough name signal to identify you by name with confidence."
        )
        lines.append("")

    lines.append("Memory runtime:")
    lines.append(f"- active_db: {paths.get('active_db', 'unknown')}")
    lines.append(f"- user_db: {paths.get('user_db', 'unknown')}")
    lines.append(f"- agent_db: {paths.get('agent_db', 'unknown')}")
    lines.append(f"- memory_db: {paths.get('memory_db', 'unknown')}")
    # Physical DB files on disk (the logical roles above can alias the same
    # file; this is the true count and is what "how many databases" should use).
    db_files = report.get("db_files") or []
    if db_files:
        lines.append(
            f"- physical_db_files: {len(db_files)} "
            f"(active/user/memory roles may alias the same file)"
        )
        for _f in db_files:
            lines.append(
                f"  - {_f.get('name')} "
                f"({_f.get('size_mb')} MB, {_f.get('tables')} tables)"
            )
    lines.append(f"- memory_entries: {int(status.get('memory_entries', 0) or 0)}")
    lines.append(f"- conversation_turns: {int(status.get('conversation_turns', 0) or 0)}")
    lines.append(f"- distinct_sessions: {int(status.get('distinct_sessions', 0) or 0)}")
    lines.append(f"- db_path: {status.get('db_path', 'unknown')}")
    lines.append(f"- recent_turn_count: {recent_turn_count}")

    schema = report.get("schema") or {}
    if schema:
        lines.append("")
        lines.append("SQLite tables observed live:")
        for db_name, tables in schema.items():
            if not tables:
                lines.append(f"- {db_name}: no tables found or DB missing")
                continue
            rendered = ", ".join(f"{name}({count})" for name, count in sorted(tables.items()))
            lines.append(f"- {db_name}: {rendered}")

    vector_status = report.get("vector_status") or {}
    _faiss = 'available' if vector_status.get('faiss_available') else 'not available'
    _nvec = vector_status.get('ntotal', 'unknown')
    db_files = report.get("db_files") or []
    fts_tables = report.get("fts_tables") or []
    kg = report.get("kg") or {}
    mechanisms = report.get("mechanisms") or {}

    def _present(key: str) -> str:
        m = mechanisms.get(key) or {}
        return f"detected at {m.get('path')}" if m.get("present") else f"NOT FOUND ({m.get('path')})"

    # Stores — derived live from the physical *.sqlite3 files on disk, with the
    # actual tables each one currently holds. No hardcoded per-DB description:
    # ELI can add/remove DBs and tables, so this is re-read every call.
    lines.append("")
    lines.append(f"Stores — {len(db_files)} physical SQLite file(s) on disk (live):")
    for _f in db_files:
        _names = _f.get("table_names") or []
        _shown = ", ".join(_names[:14]) + (" …" if len(_names) > 14 else "")
        lines.append(f"- {_f.get('name')} ({_f.get('size_mb')} MB, {_f.get('tables')} tables)")
        if _shown:
            lines.append(f"    tables: {_shown}")

    lines.append("")
    lines.append("STORAGE path (write — input → durable + indexed), live-annotated:")
    lines.append("1. A turn/fact enters via Memory.store_memory()/log_* (a persistence gate decides what to keep).")
    lines.append(f"2. Row written to SQLite and mirrored into FTS5 for keyword search "
                 f"— live FTS5 mirrors: {', '.join(fts_tables) if fts_tables else 'none detected'}.")
    lines.append(f"3. Text embedded with the local nomic embedder; vector added to FAISS "
                 f"— live: FAISS {_faiss}, {_nvec} vectors; embedder module {_present('vector_store_faiss')}.")
    lines.append(f"4. Entities/relations extracted into the knowledge graph "
                 f"— live: kg_entities={kg.get('entities', 'n/a')}, kg_relations={kg.get('relations', 'n/a')}; "
                 f"module {_present('knowledge_graph')}.")
    lines.append(f"   Live totals now: memories={int(status.get('memory_entries', 0) or 0)}, FAISS vectors={_nvec}.")

    lines.append("")
    lines.append("RECALL path (read — query → ranked context → answer), live-annotated:")
    lines.append(f"1. HyDE optionally expands a vague query before search — module {_present('hyde')}.")
    lines.append("2. Parallel retrieval: keyword (SQLite LIKE) + FTS5 full-text + FAISS vectors + knowledge-graph lookup.")
    lines.append(f"3. RAG hybrid merge + rerank by relevance/recency — orchestrator module {_present('orchestrator')}; "
                 f"parallel agent dispatch {_present('agent_bus')}.")
    lines.append("4. Context assembly → persona handoff → LLM synthesis (recency/session-biased; recall-narration filtered).")
    lines.append(f"- DAG: agents/retrieval run on a dependency DAG; the coding agent decomposes tasks via a subtask DAG "
                 f"— module {_present('plan_graph_dag')}.")

    lines.append("")
    lines.append("Index/runtime detail (live):")
    lines.append(f"- FTS5 mirror tables detected: {', '.join(fts_tables) if fts_tables else 'none'}.")
    lines.append(f"- FAISS: {_faiss}; index={vector_status.get('index_path', 'unknown')}; vectors={_nvec}.")
    lines.append(f"- Embedder: {vector_status.get('embedding_model', 'unknown')}")
    lines.append(f"- Knowledge graph: {kg.get('entities', 'n/a')} entities, {kg.get('relations', 'n/a')} relations "
                 f"({'present' if kg.get('present') else 'not present'}).")
    lines.append("- Short-term memory: in-process working memory + recent conversation_turns (no separate short-term DB).")
    lines.append("- Mechanism modules re-probed on the filesystem this call (no hardcoded claims): "
                 + "; ".join(
                       f"{k}={'present' if (mechanisms.get(k) or {}).get('present') else 'missing'}"
                       for k in ("hyde", "orchestrator", "agent_bus", "vector_store_faiss", "knowledge_graph", "plan_graph_dag")
                   ) + ".")

    if identity_hits:
        lines.append("- identity_evidence:")
        # Phase 9 fix (2026-05-11): identity hits often repeat verbatim
        # ("Capability inventory updated: ...") because the proactive daemon
        # writes the same line at every tick. De-dupe so the user sees
        # distinct evidence rather than the same line N times.
        _seen: set = set()
        _printed = 0
        for hit in identity_hits:
            if _printed >= 3:
                break
            cleaned = _clean_hit(hit)
            if not cleaned:
                continue
            key = cleaned[:220].strip().lower()
            if key in _seen:
                continue
            _seen.add(key)
            lines.append(f"  - {cleaned[:220]}")
            _printed += 1

    lines.append("- Main functions/classes: eli.memory.memory.Memory, get_memory(), recall_memory(), store_memory(), get_recent_conversation(), eli.memory.vector_store.VectorStore.search().")

    return "\n".join(lines)

def _explain_all_reasoning_modes_report() -> Dict[str, Any]:
    """Read all mode data directly from reasoning_modes.py and return as structured evidence."""
    try:
        from eli.cognition.reasoning_modes import (
            _MODE_DESCRIPTIONS,
            _MODE_TASK_PIPELINE,
            _MODE_INSTRUCTION_STACK,
            _MODE_MIN_TOKENS,
            _MODE_MAX_TOKENS_CAP,
            _MODE_TEMPERATURE_CEIL,
            _DISPLAY,
            PRIVATE_REASONING_MODES,
        )
        modes = list(_DISPLAY.keys())
        mode_data = {}
        for key in modes:
            mode_data[key] = {
                "display": _DISPLAY.get(key, key),
                "private": key in PRIVATE_REASONING_MODES,
                "description": _MODE_DESCRIPTIONS.get(key, ""),
                "pipeline_stages": _MODE_TASK_PIPELINE.get(key, []),
                "instructions": _MODE_INSTRUCTION_STACK.get(key, []),
                "min_tokens": _MODE_MIN_TOKENS.get(key, 0),
                "max_tokens_cap": _MODE_MAX_TOKENS_CAP.get(key, 0),
                "temperature_ceil": _MODE_TEMPERATURE_CEIL.get(key, 0.7),
            }
        import importlib, inspect
        try:
            mod = importlib.import_module("eli.cognition.reasoning_modes")
            src_path = inspect.getfile(mod)
        except Exception:
            src_path = "eli/cognition/reasoning_modes.py"
        return {"ok": True, "source": src_path, "modes": modes, "mode_data": mode_data}
    except Exception as exc:
        return {"ok": False, "error": str(exc)}


def _format_all_reasoning_modes(report: Dict[str, Any]) -> str:
    if not report.get("ok"):
        return f"Reasoning modes report failed: {report.get('error')}"
    lines = [f"Source: {report.get('source', 'eli/cognition/reasoning_modes.py')}"]
    lines.append(f"Total modes: {len(report.get('modes', []))}")
    lines.append("")
    for key in report.get("modes", []):
        d = report["mode_data"][key]
        lines.append(f"## {d['display']} (key={key})")
        lines.append(f"  private: {d['private']}")
        lines.append(f"  temperature_ceil: {d['temperature_ceil']}")
        lines.append(f"  token_floor: {d['min_tokens']}  token_cap: {d['max_tokens_cap']}")
        lines.append(f"  description: {d['description']}")
        lines.append(f"  pipeline_stages:")
        for s in d["pipeline_stages"]:
            lines.append(f"    - {s}")
        lines.append(f"  instructions:")
        for instr in d["instructions"]:
            lines.append(f"    - {instr}")
        lines.append("")
    return "\n".join(lines)


def _explain_cognition_runtime_report() -> Dict[str, Any]:
    files = _canonical_runtime_file_map()
    cog_path = files['cognitive_engine']
    mem_path = files['memory']
    router_path = files['router']
    executor_path = files['executor']
    if not cog_path.exists():
        return {'ok': False, 'error': 'cognitive_engine.py not found'}
    text = cog_path.read_text(encoding='utf-8', errors='replace')
    checks = {
        'process': 'def process(',
        'memory_retrieval': 'def _retrieve_relevant_memories(',
        'routing': 'route_intent(',
        'executor_evidence': '_gather_executor_evidence',
        'gguf': 'gguf_inference',
        'streaming': 'def _stream_chat(',
        'reasoning_loop': '_run_chat_reasoning_loop',
        'confidence': '_score_response_confidence',
    }
    out = {}
    for key, needle in checks.items():
        out[key] = [i for i, line in enumerate(text.splitlines(), 1) if needle in line]
    return {'ok': True, 'path': str(cog_path), 'memory_path': str(mem_path), 'router_path': str(router_path), 'executor_path': str(executor_path), 'checks': out}

def _format_cognition_runtime(report: Dict[str, Any]) -> str:
    if not report.get('ok'):
        return str(report.get('error') or 'Cognition runtime report failed')
    lines = [
        f"Cognition runtime: {report.get('path')}",
        f"Memory module: {report.get('memory_path')}",
        f"Router module: {report.get('router_path')}",
        f"Executor module: {report.get('executor_path')}",
    ]
    for key, hits in (report.get('checks') or {}).items():
        lines.append(f"- {key}: lines {hits if hits else 'missing'}")
    from eli.kernel.pipeline import get_pipeline_description
    lines.extend(get_pipeline_description())

    try:
        memory_text = _format_memory_runtime(_explain_memory_runtime_report())
        # Prefixes track the LIVE _format_memory_runtime output (#5/Option 4).
        # Keep in sync with that formatter — stale prefixes silently drop lines.
        keep_prefixes = (
            "Memory runtime:",
            "- active_db:",
            "- user_db:",
            "- agent_db:",
            "- memory_db:",
            "- physical_db_files:",
            "- memory_entries:",
            "- conversation_turns:",
            "- distinct_sessions:",
            "Stores —",
            "Index/runtime detail (live):",
            "- FTS5 mirror tables detected:",
            "- FAISS:",
            "- Knowledge graph:",
            "- Embedder:",
            "- Short-term memory:",
            "- Mechanism modules re-probed",
            "- Main functions/classes:",
        )
        memory_lines = [
            line for line in memory_text.splitlines()
            if any(line.startswith(prefix) for prefix in keep_prefixes)
        ]
        if memory_lines:
            lines.append("")
            lines.append("Memory and retrieval runtime:")
            lines.extend(memory_lines)
    except Exception as exc:
        lines.append("")
        lines.append(f"Memory/retrieval runtime summary unavailable: {exc}")

    return '\n'.join(lines)


def _runtime_status_report() -> Dict[str, Any]:
    report: Dict[str, Any] = {'ok': True}
    try:
        from eli.core.paths import get_paths
        p = get_paths()
        report['paths'] = {
            'project_root': str(Path(p.project_root).expanduser().resolve()),
            'artifacts_dir': str(Path(p.artifacts_dir).expanduser().resolve()),
            'models_dir': str(Path(getattr(p, 'models_dir', '')).expanduser().resolve()) if getattr(p, 'models_dir', '') else '',
            'user_db': str(Path(p.user_db).expanduser().resolve()),
            'agent_db': str(Path(p.agent_db).expanduser().resolve()),
            'memory_db': str(Path(_eli_path_get(p, "memory_db")).expanduser().resolve()),
        }
    except Exception as e:
        report['paths_error'] = str(e)

    try:
        from eli.core import runtime_settings as _rs
        settings = _rs.load_settings() or {}
    except Exception as e:
        settings = {}
        report['settings_error'] = str(e)
    report['settings'] = settings

    # Effective runtime resolution order (most → least authoritative):
    #   1. runtime_snapshot.json on disk — written by the loader after the
    #      llama.cpp load attempt that actually succeeded. This is the truth
    #      even if the in-process gguf object is stale or absent.
    #   2. gguf_inference live attributes — when llama is loaded in-process.
    #   3. Empty dict (effective values become "unknown" in the formatter,
    #      not silently substituted with settings).
    runtime: Dict[str, Any] = {}
    try:
        snap_path = Path(report.get('paths', {}).get('artifacts_dir', 'artifacts')) / "runtime_snapshot.json"
        report['runtime_snapshot_path'] = str(snap_path)
        if snap_path.exists():
            import json
            file_runtime = json.loads(snap_path.read_text(encoding='utf-8'))
            if isinstance(file_runtime, dict):
                runtime.update(file_runtime)
    except Exception as e:
        report['runtime_snapshot_error'] = str(e)

    # The four core load params (n_ctx, n_gpu_layers, n_threads, n_batch) are
    # ONLY trusted from the disk snapshot — that file is written by the loader
    # AFTER a successful llama.cpp load, so it reflects the clamped truth.
    # In-process gguf_inference may echo settings/defaults if the model isn't
    # actually loaded yet; allowing it to overlay would mask boot-time fallbacks.
    # We do let it contribute non-load metadata: provider, model_path, model_name,
    # loaded flag.
    try:
        from eli.cognition import gguf_inference as _gg
        if hasattr(_gg, 'get_runtime_snapshot'):
            live = dict(_gg.get_runtime_snapshot() or {})
            for k, v in live.items():
                if k in ('n_ctx', 'n_gpu_layers', 'n_threads', 'n_batch'):
                    # Only fill in if disk snapshot didn't already have it.
                    if runtime.get(k) in (None, '', 0) and v not in (None, '', 0):
                        runtime[k] = v
                    continue
                if v not in (None, '',):
                    runtime[k] = v
    except Exception as e:
        report['gguf_runtime_error'] = str(e)

    report['runtime'] = runtime

    try:
        model_path = (
            runtime.get('model_path')
            or settings.get('model_path')
            or settings.get('custom_model_path')
            or settings.get('bundled_model_path')
            or ''
        )
        report['model_path'] = str(model_path)
    except Exception as e:
        report['model_path_error'] = str(e)

    return report


def _format_runtime_status(report: Dict[str, Any]) -> str:
    settings = report.get('settings') or {}
    runtime = report.get('runtime') or {}

    def _pick(runtime_key, *setting_keys, fallback='unknown'):
        value = runtime.get(runtime_key, None)
        if value not in (None, '', 0):
            return value
        for key in setting_keys:
            value = settings.get(key, None)
            if value not in (None, '', 0):
                return value
        return fallback

    provider = runtime.get('provider') or settings.get('provider') or 'unknown'
    model_path = str(
        runtime.get('model_path')
        or report.get('model_path')
        or settings.get('model_path')
        or settings.get('custom_model_path')
        or settings.get('bundled_model_path')
        or 'unknown'
    ).strip()
    model_name = runtime.get('model_name') or (Path(model_path).name if model_path not in ('', 'unknown') else 'unknown')
    loaded = bool(runtime.get('loaded', False))

    lines = [
        "Runtime status evidence:",
        f"- provider: {provider}",
        f"- model_name: {model_name}",
        f"- model_path: {model_path}",
        f"- context_size: {_pick('n_ctx', 'n_ctx', 'context_size')}",
        f"- gpu_layers: {_pick('n_gpu_layers', 'n_gpu_layers', 'gpu_layers')}",
        f"- batch_size: {_pick('n_batch', 'n_batch', 'batch_size', 'batch')}",
        f"- cpu_threads: {_pick('n_threads', 'n_threads', 'cpu_threads', 'threads')}",
        f"- gguf_loaded_in_this_process: {loaded}",
    ]

    if settings.get('max_tokens') not in (None, ''):
        lines.append(f"- max_tokens: {settings.get('max_tokens')}")
    if settings.get('temperature') not in (None, ''):
        lines.append(f"- temperature: {settings.get('temperature')}")
    if settings.get('use_mmap') not in (None, ''):
        lines.append(f"- use_mmap: {settings.get('use_mmap')}")
    if settings.get('use_mlock') not in (None, ''):
        lines.append(f"- use_mlock: {settings.get('use_mlock')}")
    if str(provider).lower() == 'ollama':
        if settings.get('ollama_host') not in (None, ''):
            lines.append(f"- ollama_host: {settings.get('ollama_host')}")
        if settings.get('ollama_model') not in (None, ''):
            lines.append(f"- ollama_model: {settings.get('ollama_model')}")

    return '\n'.join(lines)

def _get_db_schema_evidence() -> str:
    """Introspect the canonical SQLite databases and return real table names + row counts."""
    import sqlite3
    paths = get_paths()
    # memory_db is consolidated into user.sqlite3, so it is not listed separately —
    # doing so would show a phantom third database that is really the same file.
    db_files = {
        "user.sqlite3": Path(paths.user_db),
        "agent.sqlite3": Path(paths.agent_db),
    }
    lines = ["DATABASE SCHEMA (ground truth from live introspection):"]
    for name, path in db_files.items():
        if not path.exists():
            lines.append(f"  {name}: NOT FOUND")
            continue
        try:
            conn = sqlite3.connect(str(path))
            tables = [r[0] for r in conn.execute(
                "SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%' AND name NOT LIKE '%_fts_%' ORDER BY name"
            ).fetchall()]
            tbl_info = []
            for t in tables:
                _row = conn.execute(f'SELECT COUNT(*) FROM "{t}"').fetchone()
                count = _row[0] if _row else 0
                if count > 0:
                    tbl_info.append(f"{t}({count})")
            conn.close()
            lines.append(f"  {name}: {', '.join(tbl_info) if tbl_info else 'all tables empty'}")
        except Exception as e:
            lines.append(f"  {name}: error reading — {e}")
    return "\n".join(lines)

def _memory_status_report() -> Dict[str, Any]:
    rep = _explain_memory_runtime_report()
    rep['ok'] = bool(rep.get('ok', True))
    return rep


def _format_memory_status(report: Dict[str, Any]) -> str:
    return _format_memory_runtime(report)


def _runtime_boot_status() -> Dict[str, Any]:
    """Live boot/wiring probe used by COGNITION_STATUS."""
    import importlib

    checks: Dict[str, bool] = {}
    errors: Dict[str, str] = {}

    module_names = {
        "memory_module": "eli.memory",
        "memory_core": "eli.memory.memory",
        "agent_bus": "eli.cognition.agent_bus",
        "orchestrator": "eli.cognition.orchestrator",
        "router": "eli.execution.router_enhanced",
        "executor": "eli.execution.executor_enhanced",
        "control_contracts": "eli.runtime.control_contracts",
        "response_surface": "eli.runtime.user_visible_response_surface",
    }

    loaded: Dict[str, Any] = {}
    for label, module_name in module_names.items():
        try:
            loaded[label] = importlib.import_module(module_name)
            checks[label] = True
        except Exception as exc:
            checks[label] = False
            errors[label] = f"{type(exc).__name__}: {exc}"

    try:
        router = loaded.get("router")
        checks["router.route"] = callable(getattr(router, "route", None))
    except Exception as exc:
        checks["router.route"] = False
        errors["router.route"] = f"{type(exc).__name__}: {exc}"

    try:
        executor = loaded.get("executor")
        checks["executor.execute"] = callable(getattr(executor, "execute", None))
    except Exception as exc:
        checks["executor.execute"] = False
        errors["executor.execute"] = f"{type(exc).__name__}: {exc}"

    try:
        memory = loaded.get("memory_module")
        checks["memory.get_memory"] = callable(getattr(memory, "get_memory", None))
    except Exception as exc:
        checks["memory.get_memory"] = False
        errors["memory.get_memory"] = f"{type(exc).__name__}: {exc}"

    try:
        paths = get_paths()
        path_checks = {
            "project_root": Path(paths.project_root),
            "user_db": Path(paths.user_db),
            "agent_db": Path(paths.agent_db),
            "settings": Path(paths.project_root) / "config" / "settings.json",
            "capability_manifest": Path(paths.project_root) / "capability_manifest.json",
        }
        for label, path in path_checks.items():
            checks[f"path.{label}"] = path.exists()
    except Exception as exc:
        checks["paths"] = False
        errors["paths"] = f"{type(exc).__name__}: {exc}"

    return {
        "ok": all(checks.values()) if checks else False,
        "checks": checks,
        "errors": errors,
    }


def _cognition_status_report() -> Dict[str, Any]:
    rep = _explain_cognition_runtime_report()
    try:
        boot = _runtime_boot_status()
    except Exception as e:
        boot = {'ok': False, 'error': str(e)}
    rep['boot'] = boot
    return rep


def _format_cognition_status(report: Dict[str, Any]) -> str:
    lines = [_format_cognition_runtime(report)]
    boot = report.get('boot') or {}
    if boot:
        lines.append('')
        lines.append('Startup hooks:')
        for key, value in (boot.get('checks') or {}).items():
            state = 'present' if value else 'missing'
            lines.append(f'- {key}: {state}')
    return '\n'.join(lines)

# ----------------------------
# Config / state
# ----------------------------


# Dynamic resolution — no hardcoded model name.
#
# Ollama is a user-attached fallback (the primary path is GGUF via
# llama-cpp-python). The GUI's Ollama model selector writes the user's
# chosen model to runtime_settings as "ollama_model"; we pick it up here.
#
# Resolution order:
#   1. ELI_CHAT_MODEL env var         (explicit override)
#   2. OLLAMA_MODEL env var           (legacy override)
#   3. settings["ollama_model"]       (user-attached via GUI)
#   4. ""                             (empty — forces caller to be explicit)
def _resolve_default_chat_model() -> str:
    envv = os.environ.get("ELI_CHAT_MODEL") or os.environ.get("OLLAMA_MODEL")
    if envv:
        return envv.strip()
    try:
        from eli.core import runtime_settings as _rs
        s = _rs.load_settings() or {}
        v = str(s.get("ollama_model") or "").strip()
        if v:
            return v
    except Exception:
        pass
    return ""

DEFAULT_CHAT_MODEL = _resolve_default_chat_model()


def _maybe_background_codegen(action, args):
    """If a codegen task (CODE_SOLVE / GENERATE_SCRIPT) is estimated heavy, or the
    user asked to background it, run it on a background thread and return a job-id
    message immediately. Returns None to run inline. Re-dispatch carries
    `_no_background` so the worker doesn't background again."""
    try:
        args = args or {}
        if args.get("_no_background"):
            return None
        if os.environ.get("ELI_CODEGEN_BACKGROUND", "1").strip().lower() in ("0", "false", "no", "off"):
            return None
        desc = (args.get("description") or args.get("text") or args.get("prompt")
                or args.get("query") or "").strip()
        if not desc:
            return None
        from eli.coding.cost import should_background
        decision = should_background(desc, language=(args.get("language") or "python"))
        if not decision.get("background"):
            return None
        from eli.runtime.background_tasks import get_background_tasks
        bt = get_background_tasks()
        _bg_args = dict(args)
        _bg_args["_no_background"] = True
        jid = bt.submit(f"{action}: {desc[:60]}", lambda: execute(action, _bg_args))
        msg = (f"That's a heavier task ({decision.get('reason')}). I've started it in the "
               f"background as job #{jid} — say \"check job {jid}\" for the result, or "
               f"\"background jobs\" to list them.")
        return {"ok": True, "action": action, "background": True, "job_id": jid,
                "content": msg, "response": msg}
    except Exception as _bg_e:
        log.debug(f"[CODEGEN_BG] background decision failed: {_bg_e}")
        return None


def _maybe_background_file_analysis(action, args):
    """Heavy folder/PDF analysis (many PDFs → extract + summarise) is slow and
    would block the UI. Run it on a background thread and return a job-id message
    immediately. Returns None to run inline. The worker carries `_no_background`."""
    try:
        args = args or {}
        if args.get("_no_background"):
            return None
        if os.environ.get("ELI_FILE_ANALYSIS_BACKGROUND", "1").strip().lower() in ("0", "false", "no", "off"):
            return None
        folder = str(args.get("folder") or args.get("path") or "").strip()
        if not folder:
            return None
        from pathlib import Path as _PB
        exp = os.path.expanduser(folder)
        if not os.path.isdir(exp):
            return None  # single file is fast enough inline
        try:
            _n_pdfs = sum(1 for _ in _PB(exp).rglob("*.pdf"))
        except Exception:
            _n_pdfs = 0
        _threshold = int(os.environ.get("ELI_FILE_ANALYSIS_BG_THRESHOLD", "5") or 5)
        # Background when there are enough PDFs to be slow, OR the user asked.
        _asked = bool(args.get("_force_background"))
        if _n_pdfs < _threshold and not _asked:
            return None
        from eli.runtime.background_tasks import get_background_tasks
        bt = get_background_tasks()
        _bg_args = dict(args)
        _bg_args["_no_background"] = True
        _bg_args["folder"] = folder
        _label = os.path.basename(folder.rstrip("/")) or folder
        jid = bt.submit(f"Analyse {_n_pdfs} PDFs: {_label}",
                        lambda: execute("ANALYZE_PDF_FOLDER", _bg_args))
        msg = (f"That folder has {_n_pdfs} PDFs — reading and summarising them properly takes time, "
               f"so I've started it in the background as job #{jid}. Say “check job {jid}” for the "
               f"summary when it's ready, or “background jobs” to see all running tasks.")
        return {"ok": True, "action": "ANALYZE_PDF_FOLDER", "background": True, "job_id": jid,
                "content": msg, "response": msg}
    except Exception as _bg_e:
        log.debug(f"[FILE_BG] background decision failed: {_bg_e}")
        return None


# Enumerated actions for capability_registry bootstrap
SUPPORTED_ACTIONS = [
    'ADD_EVENT',
    'BACKGROUND_JOBS',
    'SCHEDULE_TASK',
    'AMBIENT_VISION',
    'ANALYZE_CSV',
    'ANALYZE_IMAGE',
    'ANALYZE_PDF',
    'AWARENESS_STATUS',
    'CANCEL_CODE_FIX',
    'CHAT',
    'CHECK_CHRONAL_ALIGNMENT',
    'CLEAR_CHAT_HISTORY',
    'CLOSE_APP',
    'CONFIRM_CODE_FIX',
    'CONFIRM_HABIT',
    'DECLINE_HABIT',
    'CHECK_JOB',
    'CODE_CHANGES',
    'CODE_SOLVE',
    'COGNITION_STATUS',
    'CONVERT_DOCUMENT',
    'CPU_USAGE',
    'CREATE_DOCUMENT',
    'CREATE_FILE',
    'CREATE_FOLDER',
    'DATA_FABRICATOR',
    'DATE',
    'DIAGNOSE_WRAPPERS',
    'DICTATE',
    'ELI_IDENTITY_AUDIT',
    'EXAMINE_CODE',
    'EXECUTE_GOAL',
    'EXPLAIN_ALL_REASONING_MODES',
    'EXPLAIN_COGNITION_RUNTIME',
    'EXPLAIN_LAST_RESPONSE',
    'EXPLAIN_MEMORY_RUNTIME',
    'FILE_AUDIT',
    'FIX_FILE',
    'FRONTIER_STATUS',
    'FOCUS_APP',
    'GENERATE_DOCUMENT',
    'GENERATE_PROJECT',
    'GENERATE_SCRIPT',
    'GET_CLIPBOARD',
    'GET_DATE',
    'GET_STATUS',
    'GET_TIME',
    'GET_WEATHER',
    'GAZE_CALIBRATE',
    'GAZE_CLICK',
    'GAZE_DISABLE',
    'GAZE_ENABLE',
    'GAZE_STATUS',
    'GPU_STATUS',
    'GUI_RUNTIME_AUDIT',
    'HABIT_STATUS',
    'HARDWARE_PROFILE',
    'HELP',
    'IMPORT_AUDIT',
    'KEYBOARD',
    'LIST_CAPABILITIES',
    'LIST_DIR',
    'LIST_EVENTS',
    'LIST_NOTES',
    'MAXIMISE_WINDOW',
    'MEDIA_CONTROL',
    'MEMORY_RECALL',
    'MESSAGE_TIME_QUERY',
    'MEMORY_STATS',
    'MEMORY_STATUS',
    'MEMORY_STORE',
    'MINIMISE_ALL',
    'MORNING_REPORT',
    'MOUSE_CONTROL',
    'NAME_SOURCE_AUDIT',
    'NEWS_FETCH',
    'NEW_NOTE',
    'NEXT_MEDIA',
    'NEXT_WINDOW',
    'OCR_IMAGE',
    'OPEN_APP',
    'OPEN_AUDIO_SETTINGS',
    'OPEN_BROWSER',
    'OPEN_COMMUNICATION_HUB',
    'OPEN_FILE_SYSTEM',
    'OPEN_IDE',
    'OPEN_IN_IDE',
    'OPEN_MEDIA_HUB',
    'OPEN_NETWORK_BROWSER',
    'OPEN_POWER_SETTINGS',
    'OPEN_SYSTEM_SETTINGS',
    'OPEN_URL',
    'PAUSE_MEDIA',
    'PERSONAL_MEMORY_DEEP_EXPLAIN',
    'PERSONAL_MEMORY_SUMMARY',
    'PERSONA_LOCK_CLEAR',
    'PERSONA_LOCK_SET',
    'PERSONA_LOCK_STATUS',
    'PLAY_MEDIA',
    'PLUGIN_DISABLE',
    'PLUGIN_ENABLE',
    'PLUGIN_INSTALL',
    'PLUGIN_LIST',
    'PLUGIN_SEARCH',
    'PLUGIN_UNINSTALL',
    'POMODORO_START',
    'POMODORO_STATUS',
    'POMODORO_STOP',
    'PREVIOUS_MEDIA',
    'PREVIOUS_WINDOW',
    'PROACTIVE_START',
    'PROACTIVE_STATUS',
    'PROACTIVE_STOP',
    'RAM_USAGE',
    'READ_FILE',
    'REFRESH_USER_INFO',
    'REPEAT_MEDIA',
    'RESOLVE_RUNTIME_PATHS',
    'RESTORE_WINDOWS',
    'RUNTIME_AUDIT',
    'RUNTIME_STATUS',
    'REASONING_MODE_STATUS',
    'RUN_CMD',
    'SCREENSHOT',
    'SCREEN_LOCATE',
    'SCREEN_READ_ANALYZE',
    'SEARCH_NOTES',
    'SELF_ANALYZE',
    'SELF_IMPROVE',
    'SELF_IMPROVEMENT_LOG',
    'SELF_PATCH',
    'SELF_REPORT',
    'SELF_TEST',
    'RUN_TESTS',
    'GENERATE_TESTS',
    'LORA_STATUS',
    'LORA_TRAIN',
    'ORCHESTRATION_STATUS',
    'TEST_REVIEW',
    'SELF_UPDATE',
    'SELF_UPGRADE',
    'SEQUENCE',
    'SET_ALARM',
    'SET_CLIPBOARD',
    'SET_TIMER',
    'SHELL_EXEC',
    'SHOW_DIFF',
    'SHUFFLE_MEDIA',
    'SMART_HOME',
    'SPEAK',
    'STOP_MEDIA',
    'SUMMARIZE_FILE',
    'SWITCH_WORKSPACE',
    'SYSTEM_STATS',
    'TILE_WINDOWS',
    'TIME',
    'TRANSCRIBE',
    'USER_IDENTITY_SUMMARY',
    'VOLUME',
    'WEB_SEARCH',
    'WRITE_NOTE',
]


STATE_PATH = Path(os.environ.get("ELI_STATE_FILE", str(get_paths().artifacts_dir / "state.json"))).expanduser().resolve()

# Short-term conversation continuity
MAX_HISTORY_MESSAGES = int(os.environ.get("ELI_MAX_HISTORY_MESSAGES", "32"))

# Long-term memory store (simple deterministic archive, no embeddings)
MEMORY_PATH = Path(os.environ.get("ELI_MEMORY_FILE", str(get_paths().artifacts_dir / "eli_memory_legacy.jsonl"))).expanduser().resolve()

# Lock behavior: if true, refuse CHAT when verification fails
LOCK_ENFORCED = os.environ.get("ELI_LOCK_ENFORCED", "1").strip() not in ("0", "false", "False", "no", "NO")


def _load_state() -> Dict[str, Any]:
    try:
        if STATE_PATH.exists():
            return json.loads(STATE_PATH.read_text(encoding="utf-8"))
    except Exception:
        pass
    return {}


def _save_state(st: Dict[str, Any]) -> None:
    try:
        STATE_PATH.parent.mkdir(parents=True, exist_ok=True)
        STATE_PATH.write_text(json.dumps(st, indent=2), encoding="utf-8")
    except Exception:
        pass


def _state_get_history(st: Dict[str, Any]) -> List[Dict[str, str]]:
    h = st.get("chat_history")
    if not isinstance(h, list):
        return []
    out: List[Dict[str, str]] = []
    for m in h:
        if not isinstance(m, dict):
            continue
        role = m.get("role")
        content = m.get("content")
        if role in ("user", "assistant") and isinstance(content, str):
            out.append({"role": role, "content": content})
    return out


def _state_append_history(st: Dict[str, Any], user_msg: str, assistant_msg: str) -> None:
    h = _state_get_history(st)
    h.append({"role": "user", "content": user_msg})
    h.append({"role": "assistant", "content": assistant_msg})
    if len(h) > MAX_HISTORY_MESSAGES:
        h = h[-MAX_HISTORY_MESSAGES:]
    st["chat_history"] = h
    st["updated_at"] = time.time()


# ----------------------------
# Persona lock + checksum verification
# ----------------------------

def _ollama_ps() -> Dict[str, Any]:
    try:
        return requests.get(f"{OLLAMA_HOST}/api/ps", timeout=10).json()
    except Exception:
        return {"models": []}


def _ollama_model_digest_from_ps(ps: Dict[str, Any], model_name: str) -> Optional[str]:
    for m in ps.get("models", []) or []:
        if (m.get("name") == model_name) or (m.get("model") == model_name):
            return m.get("digest") or None
    return None


def _verify_persona_lock(st: Dict[str, Any]) -> Tuple[bool, str, Dict[str, Any]]:
    """
    Returns (ok, reason, details). If ok==False and LOCK_ENFORCED, CHAT should refuse to run.
    We verify:
      - model name matches expected (optional)
      - Modelfile checksum matches expected (optional)
      - Ollama digest matches expected (optional)
    """
    details: Dict[str, Any] = {}
    expected_model = (st.get("persona_lock") or {}).get("model") or None
    expected_modelfile = (st.get("persona_lock") or {}).get("modelfile_path") or None
    expected_modelfile_sha = (st.get("persona_lock") or {}).get("modelfile_sha256") or None
    expected_digest = (st.get("persona_lock") or {}).get("ollama_digest") or None
    current_model = (os.environ.get('ELI_CHAT_MODEL') or os.environ.get('OLLAMA_MODEL') or DEFAULT_CHAT_MODEL).strip()
    details["current_model"] = current_model
    details["expected_model"] = expected_model

    # 1) Model name check
    if expected_model and current_model != expected_model:
        return (False, "model_name_mismatch", {**details, "msg": f"Expected {expected_model}, got {current_model}."})

    # 2) Modelfile checksum check
    if expected_modelfile and expected_modelfile_sha:
        p = Path(os.path.expanduser(expected_modelfile))
        if not p.exists():
            return (False, "modelfile_missing", {**details, "msg": f"Locked Modelfile not found: {p}"})
        sha = _sha256_file(p)
        details["current_modelfile_sha256"] = sha
        if sha != expected_modelfile_sha:
            return (False, "modelfile_sha256_mismatch", {**details, "msg": "Modelfile checksum changed."})

    # 3) Ollama digest check (from /api/ps)
    # Resolve Ollama host once (no implicit globals)
    ollama_host = (
        os.environ.get('ELI_OLLAMA_HOST')
        or os.environ.get('OLLAMA_HOST')
        or 'http://localhost:11434'
    )
    ps = _ollama_ps()
    dig = _ollama_model_digest_from_ps(ps, current_model)
    details["current_ollama_digest"] = dig
    details["expected_ollama_digest"] = expected_digest
    if expected_digest and dig and dig != expected_digest:
        return (False, "ollama_digest_mismatch", {**details, "msg": "Loaded model digest differs."})
    if expected_digest and dig is None:
        # Model may not be loaded yet; attempt a warm-load then re-check /api/ps.
        warmed = False
        try:
            warmed = _ollama_warm_load(current_model, ollama_host)
        except Exception as e:
            details["warm_load_error"] = repr(e)

        if warmed:
            ps2 = _ollama_ps()
            dig2 = _ollama_model_digest_from_ps(ps2, current_model)
            details["current_ollama_digest"] = dig2
            if dig2 and dig2 != expected_digest:
                return (False, "ollama_digest_mismatch", {**details, "msg": "Loaded model digest differs after warm-load."})
            if dig2:
                return (True, "ok", details)

        return (True, "ok_unverified", {**details, "warn": "Model not yet loaded; digest cannot be verified until first use."})
    # FINAL SUCCESS PATH — if nothing failed above, the lock is valid
    return (True, "ok", details)


def persona_lock_set(model: Optional[str] = None, modelfile_path: Optional[str] = None) -> Dict[str, Any]:
    """
    Sets the lock baseline. This is the "make it unbreakable" button.
    - model defaults to DEFAULT_CHAT_MODEL
    - modelfile_path can be None (then only model + digest are locked)
    """
    st = _load_state()
    model = (model or DEFAULT_CHAT_MODEL).strip()

    lock: Dict[str, Any] = {"model": model, "set_at": time.time()}

    if modelfile_path:
        p = Path(os.path.expanduser(modelfile_path))
        if not p.exists():
            msg = f"Modelfile not found: {p}"
            return {"ok": False, "action": "PERSONA_LOCK_SET", "error": msg, "content": msg, "response": msg}
        lock["modelfile_path"] = str(p)
        lock["modelfile_sha256"] = _sha256_file(p)
    #------------------------------
    # Capture current digest (best effort)
    #-------------------------------   
    ps = _ollama_ps()
    dig = _ollama_model_digest_from_ps(ps, model)
    if dig:
        lock["ollama_digest"] = dig

    st["persona_lock"] = lock
    _save_state(st)

    msg = f"Persona lock set for model={model}."
    return {"ok": True, "action": "PERSONA_LOCK_SET", "lock": lock, "content": msg, "response": msg}


def persona_lock_status() -> Dict[str, Any]:
    st = _load_state()
    ok, reason, details = _verify_persona_lock(st)
    lock = st.get("persona_lock") or {}
    msg = "Persona lock OK." if ok else f"Persona lock FAIL: {reason}"
    return {"ok": ok, "action": "PERSONA_LOCK_STATUS", "reason": reason, "details": details, "lock": lock, "content": msg, "response": msg}


def persona_lock_clear() -> Dict[str, Any]:
    st = _load_state()
    st.pop("persona_lock", None)
    _save_state(st)
    msg = "Persona lock cleared."
    return {"ok": True, "action": "PERSONA_LOCK_CLEAR", "content": msg, "response": msg}


# ----------------------------
# Persistent long-term memory (deterministic archive)
# ----------------------------

def memory_store_legacy(text: str, tags: Optional[List[str]] = None) -> Dict[str, Any]:
    text = (text or "").strip()
    if not text:
        msg = "Empty memory text."
        return {"ok": False, "action": "MEMORY_STORE", "error": "empty_text", "content": msg, "response": msg}

    rec = {
        "ts": time.time(),
        "text": text,
        "tags": [t.strip() for t in (tags or []) if isinstance(t, str) and t.strip()],
    }
    try:
        MEMORY_PATH.parent.mkdir(parents=True, exist_ok=True)
        with MEMORY_PATH.open("a", encoding="utf-8") as f:
            f.write(json.dumps(rec, ensure_ascii=False) + "\n")
        msg = "Memory stored."
        return {"ok": True, "action": "MEMORY_STORE", "record": rec, "content": msg, "response": msg}
    except Exception as e:
        msg = "Failed to store memory."
        return {"ok": False, "action": "MEMORY_STORE", "error": repr(e), "content": msg, "response": msg}


def memory_recall_legacy(query: str, limit: int = 5) -> Dict[str, Any]:
    q = (query or "").strip().lower()
    if not q:
        msg = "Empty recall query."
        return {"ok": False, "action": "MEMORY_RECALL", "error": "empty_query", "content": msg, "response": msg}

    hits: List[Dict[str, Any]] = []
    try:
        if not MEMORY_PATH.exists():
            msg = "No memory archive yet."
            return {"ok": True, "action": "MEMORY_RECALL", "hits": [], "content": msg, "response": msg}

        with MEMORY_PATH.open("r", encoding="utf-8", errors="ignore") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    rec = json.loads(line)
                except Exception:
                    continue
                text = str(rec.get("text", "")).lower()
                tags = " ".join(rec.get("tags") or []).lower()
                if q in text or q in tags:
                    hits.append(rec)

        hits = hits[-max(1, int(limit)):]  # most recent matches
        # Human readable summary
        summary_lines = []
        for r in reversed(hits):
            ts = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(r.get("ts", 0)))
            t = str(r.get("text", "")).strip()
            t = t if len(t) <= 180 else t[:177] + "..."
            summary_lines.append(f"- [{ts}] {t}")
        content = "Memory recall:\n" + ("\n".join(summary_lines) if summary_lines else "(no matches)")
        return {"ok": True, "action": "MEMORY_RECALL", "hits": hits, "content": content, "response": content}
    except Exception as e:
        msg = "Memory recall failed."
        return {"ok": False, "action": "MEMORY_RECALL", "error": repr(e), "content": msg, "response": msg}


# ----------------------------
# Self-consistency tests + healthcheck
# ----------------------------

def self_test() -> Dict[str, Any]:
    """
    Runs a deterministic suite:
      - persona lock status
      - minimal chat handshake (Ollama or GGUF fallback)
      - consistency probe (temperature=0 twice) and compare
    """
    import socket

    def _ollama_reachable() -> bool:
        try:
            s = socket.create_connection(("127.0.0.1", 11434), timeout=1)
            s.close()
            return True
        except OSError:
            return False

    st = _load_state()
    lock = st.get("persona_lock") or {}
    if lock.get("model") == "test":
        try:
            st.pop("persona_lock", None)
            _save_state(st)
        except Exception:
            pass
    ok_lock, reason, details = _verify_persona_lock(st)

    results: Dict[str, Any] = {
        "persona_lock": {"ok": ok_lock, "reason": reason, "details": details},
        "tests": []
    }

    # ── GGUF-only fallback when Ollama is not running ──
    if not _ollama_reachable():
        try:
            from eli.cognition.gguf_inference import chat_completion
            resp = chat_completion("Reply with exactly these 2 words: persona ok", max_tokens=16)
            gguf_ok = bool(resp and resp.strip())
        except Exception as exc:
            gguf_ok = False
            resp = str(exc)
        results["tests"].append({
            "name": "handshake", "ok": gguf_ok, "model": "gguf",
            "content": (resp or "").strip(), "response": (resp or "").strip(),
        })
        # Consistency probes skipped — GGUF has no temp=0 determinism guarantee
        results["tests"].append({"name": "consistency_probe_1", "ok": gguf_ok, "model": "gguf",
                                  "content": "skipped (no Ollama)", "response": "skipped (no Ollama)"})
        results["tests"].append({"name": "consistency_probe_2", "ok": gguf_ok, "model": "gguf",
                                  "content": "skipped (no Ollama)", "response": "skipped (no Ollama)"})
        results["tests"].append({"name": "consistency_equal", "ok": gguf_ok,
                                  "content": "skipped (no Ollama)", "response": "skipped (no Ollama)"})
        overall_ok = bool(ok_lock) and gguf_ok
        msg = "SELF_TEST OK (GGUF)" if overall_ok else "SELF_TEST FAIL (GGUF)"
        return {"ok": overall_ok, "action": "SELF_TEST", "results": results, "content": msg, "response": msg}

    # ── Standard Ollama path ──
    # 1) Handshake: "Reply with exactly: persona_ok"
    handshake_prompt = "Reply with exactly these 2 words: persona ok"
    h1 = _ollama_chat_raw(handshake_prompt, temperature=0.0, num_predict=16)
    results["tests"].append({"name": "handshake", **h1})

    # 2) Consistency probe: same prompt twice with temp=0
    probe = "Reply with exactly 4 words: consistency test passed."
    p1 = _ollama_chat_raw(probe, temperature=0.0, num_predict=16)
    p2 = _ollama_chat_raw(probe, temperature=0.0, num_predict=16)
    same = (p1.get("content") == p2.get("content")) and p1.get("ok") and p2.get("ok")
    results["tests"].append({"name": "consistency_probe_1", **p1})
    results["tests"].append({"name": "consistency_probe_2", **p2})
    results["tests"].append({"name": "consistency_equal", "ok": same, "content": str(same), "response": str(same)})

    # overall ok
    overall_ok = bool(ok_lock) and bool(h1.get("ok")) and bool(p1.get("ok")) and bool(p2.get("ok"))
    msg = "SELF_TEST OK" if overall_ok else "SELF_TEST FAIL"
    return {"ok": overall_ok, "action": "SELF_TEST", "results": results, "content": msg, "response": msg}


def _ollama_chat_raw(user_text: str, temperature: float = 0.7, num_predict: int = 80) -> Dict[str, Any]:
    model = DEFAULT_CHAT_MODEL
    payload = {
        "model": model,
        "stream": False,
        "messages": [{"role": "user", "content": user_text}],
        "options": {"temperature": float(temperature), "num_predict": int(num_predict)},
    }
    try:
        r = requests.post(f"{OLLAMA_HOST}/api/chat", json=payload, timeout=(10, 3600))
        r.raise_for_status()
        j = r.json()
        content = _strip_ollama_artifacts(((j.get("message") or {}).get("content") or "").strip())
        return {"ok": True, "model": model, "content": content, "response": content}
    except Exception as e:
        msg = "raw chat failed"
        return {"ok": False, "model": model, "error": repr(e), "content": msg, "response": msg}


# ----------------------------
# Origin / identity (NO prompt hijack)
# ----------------------------

def _is_origin_question(msg: str) -> bool:
    low = (msg or "").lower()
    triggers = [
        "who created you", "who made you", "who built you",
        "who are you", "your creator", "your origin",
        "who developed you"
    ]
    return any(t in low for t in triggers)


def _origin_facts(st: Dict[str, Any]) -> str:
    name = (st.get("user_name") or "").strip() or "the user"
    return (
        f"Factual note: {name} assembled/configured ELI; ELI runs locally via Ollama using base model weights. "
        f"Don't claim any company 'created ELI' as a shipped product."
    )


def _needs_origin_anchor(ans: str) -> bool:
    """
    If origin question answer drifts into corporate-creator claims,
    we append a short factual anchor (we do NOT rewrite the answer).
    """
    a = (ans or "").lower()
    bad = [
        "created by alibaba", "developed by alibaba",
        "created by qwen", "developed by qwen", "alibaba cloud created"
    ]
    return any(x in a for x in bad)


# ================== MEDIA CONTROL FUNCTIONS ==================
def _get_active_player() -> Optional[str]:
    """Get the best active MPRIS2 player via playerctl. Returns None if none found."""
    pass  # shutil already imported at module level
    if not shutil.which("playerctl"):
        return None
    try:
        r = subprocess.run(["playerctl", "--list-all"], capture_output=True, text=True, timeout=3)
        if r.returncode != 0 or not r.stdout.strip():
            return None
        players = [p.strip() for p in r.stdout.splitlines() if p.strip()]
        # Priority: spotify first, then whatever is running
        for preferred in ("spotify", "Spotify", "vlc", "mpv", "firefox", "chromium"):
            for p in players:
                if p.lower().startswith(preferred.lower()):
                    return p
        return players[0] if players else None
    except Exception:
        return None


def _playerctl(cmd: str, player: Optional[str] = None) -> Dict[str, Any]:
    """Run a targeted playerctl command. Self-contained, no external imports."""
    pass  # shutil already imported at module level
    if not shutil.which("playerctl"):
        return {"ok": False, "error": "playerctl not installed — run: sudo apt install playerctl",
                "content": "playerctl not installed", "response": "playerctl not installed"}
    p = player or _get_active_player()
    if not p:
        return {"ok": False, "error": "No media player running",
                "content": "No media player running", "response": "No media player running"}
    try:
        r = subprocess.run(["playerctl", "-p", p, cmd], capture_output=True, text=True, timeout=3)
        if r.returncode == 0:
            labels = {"play": "▶ Playing", "pause": "⏸ Paused", "stop": "⏹ Stopped",
                      "next": "⏭ Next track", "previous": "⏮ Previous track"}
            msg = f"{labels.get(cmd, cmd)} — {p}"
            return {"ok": True, "player": p, "content": msg, "response": msg}
        else:
            msg = f"playerctl {cmd} failed: {r.stderr.strip() or 'unknown error'}"
            return {"ok": False, "player": p, "error": r.stderr.strip(), "content": msg, "response": msg}
    except subprocess.TimeoutExpired:
        return {"ok": False, "error": f"playerctl {cmd} timed out", "content": "Command timed out", "response": "Command timed out"}
    except Exception as e:
        return {"ok": False, "error": str(e), "content": str(e), "response": str(e)}


# ── Browser/streaming service → MPRIS player alias map ──────────────────────
_BROWSER_PLAYERS = {"firefox", "Firefox", "chromium", "Chromium", "chrome"}
_TARGET_ALIASES: dict = {
    # streaming services → browser player
    "netflix":    "browser",
    "youtube":    "browser",
    "prime":      "browser",
    "disney":     "browser",
    "disneyplus": "browser",
    "hulu":       "browser",
    "twitch":     "browser",
    "brow":       "browser",
    "browser":    "browser",
    # explicit players
    "spotify":    "spotify",
    "vlc":        "vlc",
    "mpv":        "mpv",
}

def _resolve_media_target(target: str | None) -> str | None:
    """Map a user-supplied target string to an MPRIS player name or None."""
    if not target:
        return None
    alias = _TARGET_ALIASES.get(target.lower().strip())
    if alias == "browser":
        # Return first available browser player from playerctl
        try:
            import subprocess as _sp
            r = _sp.run(["playerctl", "--list-all"], capture_output=True, text=True, timeout=3)
            for p in r.stdout.splitlines():
                for bp in _BROWSER_PLAYERS:
                    if p.strip().lower().startswith(bp.lower()):
                        return p.strip()
        except Exception:
            pass
        return "firefox"   # fallback name
    return alias   # e.g. "spotify", "vlc"


def stop_media(target: str | None = None) -> Dict[str, Any]:
    """Stop/pause media. Spotify doesn't support MPRIS Stop so uses pause."""
    p = _resolve_media_target(target) or _get_active_player()
    # Spotify ignores Stop — use pause instead
    cmd = "pause" if (p and "spotify" in p.lower()) else "stop"
    result = _playerctl(cmd, p)
    result["action"] = "STOP_MEDIA"
    return result


def pause_media(target: str | None = None) -> Dict[str, Any]:
    """Pause currently playing media."""
    result = _playerctl("pause", _resolve_media_target(target) or _get_active_player())
    result["action"] = "PAUSE_MEDIA"
    return result


def play_media(target: str | None = None) -> Dict[str, Any]:
    """Play/resume media."""
    result = _playerctl("play", _resolve_media_target(target) or _get_active_player())
    result["action"] = "PLAY_MEDIA"
    return result


def next_media() -> Dict[str, Any]:
    """Skip to next track."""
    result = _playerctl("next")
    result["action"] = "NEXT_MEDIA"
    return result


def previous_media(target: str | None = None) -> Dict[str, Any]:
    """Go to previous track.

    Spotify's MPRIS implementation of 'Previous' behaves like a media key:
    if playback position > 3s it seeks to 0 (restart current track) instead
    of going to the previous track. We force position=0 before sending
    Previous so the command always skips back regardless of current position.
    """
    p = _resolve_media_target(target) or _get_active_player()
    if p and "spotify" in p.lower():
        try:
            subprocess.run(
                ["playerctl", "-p", p, "position", "0"],
                capture_output=True, timeout=2,
            )
        except Exception:
            pass
    result = _playerctl("previous", p)
    result["action"] = "PREVIOUS_MEDIA"
    return result


def shuffle_media(target: str | None = None) -> Dict[str, Any]:
    """Toggle shuffle mode via playerctl."""
    result = _playerctl("shuffle Toggle", _resolve_media_target(target) or _get_active_player())
    result["action"] = "SHUFFLE_MEDIA"
    return result


def repeat_media(target: str | None = None) -> Dict[str, Any]:
    """Cycle repeat mode (None → Track → Playlist) via playerctl."""
    result = _playerctl("loop Track", _resolve_media_target(target) or _get_active_player())
    result["action"] = "REPEAT_MEDIA"
    return result


def _yt_resolve_watch_url(query: str) -> str | None:
    """Return the first YouTube watch URL for a search query.
    Uses a plain urllib scrape — no yt-dlp or API key required.
    Returns None if the network request fails or no video is found.
    """
    try:
        import urllib.request as _ureq
        import urllib.parse as _up2
        import re as _re2
        url = f"https://www.youtube.com/results?search_query={_up2.quote_plus(query)}"
        req = _ureq.Request(
            url,
            headers={"User-Agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36"},
        )
        with _ureq.urlopen(req, timeout=8) as resp:
            html = resp.read().decode("utf-8", errors="ignore")
        m = _re2.search(r'"videoId"\s*:\s*"([A-Za-z0-9_-]{11})"', html)
        if m:
            return f"https://www.youtube.com/watch?v={m.group(1)}"
    except Exception:
        pass
    return None


def _yt_mix_url(watch_url: str | None) -> str | None:
    """Turn a plain YouTube watch URL into a Mix/radio URL so playback AUTOPLAYS
    related songs continuously instead of stopping after one video. YouTube's
    `&list=RD<videoId>` is the auto-generated "Mix" (song radio) seeded from the
    video — the user's reported "YT doesn't autoplay the next songs" + "play
    something related" both come from this."""
    if not watch_url:
        return watch_url
    import re as _re3
    m = _re3.search(r"[?&]v=([A-Za-z0-9_-]{11})", watch_url)
    if m:
        vid = m.group(1)
        return f"https://www.youtube.com/watch?v={vid}&list=RD{vid}"
    return watch_url


def _open_in_browser(url: str) -> None:
    """Open a URL in the default browser via xdg-open."""
    import subprocess as _sp2
    if shutil.which("xdg-open"):
        _sp2.Popen(["xdg-open", url], stdout=_sp2.DEVNULL, stderr=_sp2.DEVNULL,
                   start_new_session=True)


def _spotify_open_uri(uri: str) -> bool:
    """OpenUri on a running Spotify via MPRIS dbus. Accepts both spotify: URIs
    and https://open.spotify.com/... URLs (Spotify opens the latter in-app)."""
    import subprocess as _sp2
    try:
        r = _sp2.run(
            ["dbus-send", "--print-reply",
             "--dest=org.mpris.MediaPlayer2.spotify",
             "/org/mpris/MediaPlayer2",
             "org.mpris.MediaPlayer2.Player.OpenUri",
             f"string:{uri}"],
            capture_output=True, timeout=5,
        )
        return r.returncode == 0
    except Exception:
        return False


def _spotify_search(query: str, prefer: str | None = None) -> bool:
    """Open a Spotify search via dbus.

    prefer='playlists' targets the **Playlists** filter tab
    (open.spotify.com/search/<q>/playlists), so the top result is a PLAYLIST that
    contains the requested song — playing it keeps continuation thematic, instead
    of the 'All' tab which plays one track then drifts to unrelated songs
    (user-reported). Falls back to the classic spotify:search: URI if the
    filtered URL is rejected.
    """
    import urllib.parse as _up2
    if prefer in ("playlist", "playlists"):
        if _spotify_open_uri(f"https://open.spotify.com/search/{_up2.quote(query)}/playlists"):
            return True
        # Fall back to the unfiltered search URI.
        return _spotify_open_uri(f"spotify:search:{_up2.quote(query)}")
    return _spotify_open_uri(f"spotify:search:{_up2.quote(query)}")


def _spotify_play() -> bool:
    """Start/resume Spotify playback (no Web API). After a spotify:search: OpenUri
    the search results are the active context, so Play starts the top match.

    Returns True only if playback actually reaches the 'Playing' state — a Play
    call that returns 0 but leaves Spotify paused/stopped is reported as False so
    callers can be honest about search-only outcomes.
    """
    import subprocess as _sp2
    try:
        _sp2.run(["playerctl", "-p", "spotify", "play"], capture_output=True, timeout=5)
    except Exception:
        pass
    try:
        _sp2.run(
            ["dbus-send", "--print-reply",
             "--dest=org.mpris.MediaPlayer2.spotify",
             "/org/mpris/MediaPlayer2",
             "org.mpris.MediaPlayer2.Player.Play"],
            capture_output=True, timeout=5,
        )
    except Exception:
        pass
    return _spotify_is_playing()


def _spotify_is_playing() -> bool:
    """True iff Spotify's MPRIS PlaybackStatus is 'Playing'."""
    import subprocess as _sp2
    try:
        r = _sp2.run(["playerctl", "-p", "spotify", "status"],
                     capture_output=True, timeout=5, text=True)
        if r.returncode == 0 and "playing" in (r.stdout or "").strip().lower():
            return True
    except Exception:
        pass
    return False


def _spotify_running() -> bool:
    import subprocess as _sp2
    try:
        r = _sp2.run(["pgrep", "-x", "spotify"], capture_output=True, timeout=4)
        return r.returncode == 0
    except Exception:
        return False


def play_specific(query: str, target: str | None = None) -> Dict[str, Any]:
    """Play a specific song/artist/genre.

    Dispatch priority:
      1. Explicit spotify target  → Spotify (dbus if open, else xdg-open URI)
      2. Explicit youtube target  → yt-dlp+mpv if available, else browser watch URL
      3. "youtube web/website"    → browser watch URL (never mpv)
      4. "X by Y" (no target)     → yt-dlp+mpv
      5. Generic fallback         → browser watch URL
    """
    import subprocess as _sp
    import urllib.parse
    import re as _re

    t = (target or "").lower().strip()
    player = _resolve_media_target(target) or _get_active_player()

    def _clean(q: str) -> str:
        q = q.strip()
        q = _re.sub(
            r"^(?:a\s+song\s+|the\s+song\s+|the\s+track\s+|a\s+track\s+|me\s+a?\s*|some\s+)",
            "", q, flags=_re.I,
        )
        return q.strip()

    query = _clean(query)
    _by_m = _re.match(r"^(.+?)\s+by\s+(.+)$", query, _re.I)

    # ── 1. Spotify target ─────────────────────────────────────────────────────
    is_spotify = "spotify" in t or (player and "spotify" in (player or "").lower())
    is_youtube = "youtube" in t or t in ("yt",)
    is_yt_web  = is_youtube and bool(_re.search(r"\bweb(?:site)?\b", t))

    if is_spotify and not is_youtube:
        import time as _time
        search_q = (
            f"{_by_m.group(1).strip()} {_by_m.group(2).strip()}" if _by_m else query
        )
        # No Web API: open the PLAYLISTS search tab (sets the active context to
        # matching playlists), let results load, then issue Play so the top
        # PLAYLIST starts — its tracks keep continuation on-theme rather than
        # the 'All' tab playing one song then drifting (user-reported).
        _opened = _spotify_search(search_q, prefer="playlists")
        if not _opened:
            # Spotify not running yet — launch with the search URI, wait, retry.
            try:
                _open_in_browser(f"spotify:search:{urllib.parse.quote(search_q)}")
                for _ in range(8):
                    _time.sleep(1.0)
                    if _spotify_running():
                        break
                _opened = _spotify_search(search_q, prefer="playlists")
            except Exception:
                _opened = False
        if _opened:
            _time.sleep(1.6)            # let the results view populate
            if _spotify_play():
                msg = f"Playing a “{search_q}” playlist on Spotify."
                return {"ok": True, "action": "PLAY_MEDIA", "played": True,
                        "content": msg, "response": msg}
            msg = (f"I opened the Spotify search for “{search_q}” but couldn’t start "
                   f"playback automatically — hit play, or say “play {search_q} on "
                   f"youtube” and I’ll play it directly.")
            return {"ok": True, "action": "PLAY_MEDIA", "played": False,
                    "search_only": True, "content": msg, "response": msg}
        # Spotify was the EXPLICITLY requested target. Even if we couldn't reach it
        # (not installed / not running / dbus refused), do NOT silently fall through
        # to YouTube — that opens a second platform the user never asked for.
        msg = (f"I couldn’t reach Spotify to play “{search_q}” — is it installed and "
               f"running? Open it and try again, or say “play {search_q} on youtube”.")
        return {"ok": False, "action": "PLAY_MEDIA", "played": False,
                "search_only": True, "target": "spotify",
                "content": msg, "response": msg}

    # ── 2. "youtube web/website" → browser only (never mpv) ──────────────────
    if is_yt_web:
        watch = _yt_mix_url(_yt_resolve_watch_url(query))
        url = watch or f"https://www.youtube.com/results?search_query={urllib.parse.quote_plus(query)}"
        _open_in_browser(url)
        msg = f"Opening YouTube in browser: {query} (continuous mix)"
        return {"ok": True, "action": "PLAY_MEDIA", "content": msg, "response": msg}

    # ── 3. YouTube target (plain) or "X by Y" → yt-dlp+mpv ──────────────────
    yt_search = (
        f"{_by_m.group(1).strip()} {_by_m.group(2).strip()} official audio"
        if _by_m else query
    )

    if shutil.which("yt-dlp") and shutil.which("mpv"):
        try:
            ipc = os.environ.get("ELI_YOUTUBE_MPV_IPC", "/tmp/eli_youtube_mpv.sock")
            _sp.Popen(
                ["mpv", f"ytdl://ytsearch1:{yt_search}",
                 f"--input-ipc-server={ipc}",
                 "--ytdl-format=bestaudio/best",
                 "--title=ELI-YouTube"],
                stdout=_sp.DEVNULL, stderr=_sp.DEVNULL, start_new_session=True,
            )
            if _by_m:
                msg = f"Playing '{_by_m.group(1).strip()}' by {_by_m.group(2).strip()}"
            else:
                msg = f"Playing '{query}' on YouTube"
            return {"ok": True, "action": "PLAY_MEDIA", "content": msg, "response": msg}
        except Exception:
            pass

    # ── 4. No yt-dlp/mpv → resolve watch URL and open in browser ─────────────
    # Use the Mix/radio URL so it autoplays related songs continuously.
    watch = _yt_mix_url(_yt_resolve_watch_url(yt_search))
    if watch:
        _open_in_browser(watch)
        if _by_m:
            msg = f"Opening '{_by_m.group(1).strip()}' by {_by_m.group(2).strip()} in browser"
        else:
            msg = f"Opening '{query}' in browser"
        return {"ok": True, "action": "PLAY_MEDIA", "content": msg, "response": msg}

    # ── 5. Last resort: YouTube search page ──────────────────────────────────
    encoded = urllib.parse.quote_plus(query)
    _open_in_browser(f"https://www.youtube.com/results?search_query={encoded}")
    msg = f"Opening YouTube: {query}"
    return {"ok": True, "action": "PLAY_MEDIA", "content": msg, "response": msg}

# ================== END MEDIA CONTROL ==================

def _run_ok(argv, timeout: int = 5) -> bool:
    try:
        r = subprocess.run(argv, capture_output=True, text=True, timeout=timeout)
        return r.returncode == 0
    except Exception:
        return False


def _volume_fallback(direction: str, delta: int = 10, level: int | None = None) -> Dict[str, Any]:
    direction = (direction or '').strip().lower()
    try:
        if direction == 'set' and level is not None:
            if shutil.which('wpctl') and _run_ok(['wpctl', 'set-volume', '@DEFAULT_AUDIO_SINK@', f'{int(level)}%']):
                msg = f'Volume set to {int(level)}%'
                return {'ok': True, 'action': 'VOLUME', 'content': msg, 'response': msg}
            if shutil.which('pactl') and _run_ok(['pactl', 'set-sink-volume', '@DEFAULT_SINK@', f'{int(level)}%']):
                msg = f'Volume set to {int(level)}%'
                return {'ok': True, 'action': 'VOLUME', 'content': msg, 'response': msg}
        if direction in ('up', 'raise'):
            if shutil.which('wpctl') and _run_ok(['wpctl', 'set-volume', '@DEFAULT_AUDIO_SINK@', f'{int(delta)}%+']):
                msg = 'Volume up'
                return {'ok': True, 'action': 'VOLUME', 'content': msg, 'response': msg}
            if shutil.which('pactl') and _run_ok(['pactl', 'set-sink-volume', '@DEFAULT_SINK@', f'+{int(delta)}%']):
                msg = 'Volume up'
                return {'ok': True, 'action': 'VOLUME', 'content': msg, 'response': msg}
            if shutil.which('amixer') and _run_ok(['amixer', '-D', 'pulse', 'sset', 'Master', f'{int(delta)}%+']):
                msg = 'Volume up'
                return {'ok': True, 'action': 'VOLUME', 'content': msg, 'response': msg}
        if direction in ('down', 'lower'):
            if shutil.which('wpctl') and _run_ok(['wpctl', 'set-volume', '@DEFAULT_AUDIO_SINK@', f'{int(delta)}%-']):
                msg = 'Volume down'
                return {'ok': True, 'action': 'VOLUME', 'content': msg, 'response': msg}
            if shutil.which('pactl') and _run_ok(['pactl', 'set-sink-volume', '@DEFAULT_SINK@', f'-{int(delta)}%']):
                msg = 'Volume down'
                return {'ok': True, 'action': 'VOLUME', 'content': msg, 'response': msg}
            if shutil.which('amixer') and _run_ok(['amixer', '-D', 'pulse', 'sset', 'Master', f'{int(delta)}%-']):
                msg = 'Volume down'
                return {'ok': True, 'action': 'VOLUME', 'content': msg, 'response': msg}
        if direction == 'mute':
            if shutil.which('wpctl') and _run_ok(['wpctl', 'set-mute', '@DEFAULT_AUDIO_SINK@', '1']):
                msg = 'Muted'
                return {'ok': True, 'action': 'VOLUME', 'content': msg, 'response': msg}
            if shutil.which('pactl') and _run_ok(['pactl', 'set-sink-mute', '@DEFAULT_SINK@', '1']):
                msg = 'Muted'
                return {'ok': True, 'action': 'VOLUME', 'content': msg, 'response': msg}
        if direction == 'unmute':
            if shutil.which('wpctl') and _run_ok(['wpctl', 'set-mute', '@DEFAULT_AUDIO_SINK@', '0']):
                msg = 'Unmuted'
                return {'ok': True, 'action': 'VOLUME', 'content': msg, 'response': msg}
            if shutil.which('pactl') and _run_ok(['pactl', 'set-sink-mute', '@DEFAULT_SINK@', '0']):
                msg = 'Unmuted'
                return {'ok': True, 'action': 'VOLUME', 'content': msg, 'response': msg}
    except Exception as e:
        return {'ok': False, 'action': 'VOLUME', 'error': str(e), 'content': str(e), 'response': str(e)}
    msg = 'Volume control failed'
    return {'ok': False, 'action': 'VOLUME', 'content': msg, 'response': msg}


# ----------------------------
# CHAT with persistent conversation + lock enforcement
# ----------------------------

def _memory_capability_truth() -> str:
    """Truth-only memory capability statement. Never fabricate recall."""
    try:
        has_adapter = False
        has_sqlite = False
        try:
            from eli.memory.memory_adapter import memory_recall as _mr  # noqa: F401
            has_adapter = True
        except Exception:
            pass
        try:
            import sqlite3 as _sql  # noqa: F401
            has_sqlite = True
        except Exception:
            pass

        if has_adapter or has_sqlite:
            return (
                "I use two memory layers: session context and persistent local memory. "
                "Persistent memory is stored locally (SQLite/artifacts via memory adapters) "
                "and can be recalled across restarts if saved and retrievable by query."
            )
        return (
            "I currently have session context in this runtime; persistent memory backends are "
            "not available right now. I won’t claim cross-session recall unless retrieval succeeds."
        )
    except Exception:
        return (
            "I won’t guess about memory state. Session context is available; persistent recall "
            "depends on configured local memory backends and successful retrieval."
        )


def _live_memory_audit() -> dict:
    """
    Real local-memory audit:
    - SQLite memory DB(s)
    - conversations JSONL
    - proactive artifacts
    - latest context/summary/action files
    Returns factual counts + health, never fabricated claims.
    """
    from datetime import datetime, timezone
    now = datetime.now(timezone.utc).isoformat()

    root = Path(__file__).resolve().parents[2]  # project root (contains artifacts/, src/)
    candidates = {
        "root_artifacts_db": root / "artifacts" / "user.sqlite3",
        "src_artifacts_db": root / "src" / "eli" / "artifacts" / "user.sqlite3",
        "conversations_dir": root / "src" / "eli" / "artifacts" / "conversations",
        "proactive_dir": root / "artifacts" / "proactive",
        "src_latest_context": root / "artifacts" / "latest_context.txt",
        "src_latest_summary": root / "artifacts" / "latest_summary.txt",
        "src_latest_action": root / "artifacts" / "latest_action.txt",
    }

    report = {
        "ok": True,
        "ts_utc": now,
        "root": str(root),
        "stores": {},
        "errors": [],
    }

    # ---- SQLite scan ----
    def scan_sqlite(db_path: Path):
        out = {"path": str(db_path), "exists": db_path.exists(), "size_bytes": 0, "tables": {}, "recent_rows": []}
        if not db_path.exists():
            return out
        out["size_bytes"] = db_path.stat().st_size
        try:
            conn = sqlite3.connect(str(db_path))
            cur = conn.cursor()

            cur.execute("SELECT name FROM sqlite_master WHERE type='table' ORDER BY name;")
            tables = [r[0] for r in cur.fetchall()]
            out["tables"]["__all__"] = len(tables)

            for t in tables:
                try:
                    cur.execute(f'SELECT COUNT(*) FROM "{t}"')
                    out["tables"][t] = int(cur.fetchone()[0] or 0)
                except Exception as e:
                    out["tables"][t] = f"count_error: {e}"

            # try pull likely memory/event rows
            for t in ("events", "memory", "memories", "conversation", "conversations"):
                if t in tables:
                    try:
                        cur.execute(f'SELECT * FROM "{t}" ORDER BY rowid DESC LIMIT 3')
                        rows = cur.fetchall()
                        out["recent_rows"].append({"table": t, "rows_found": len(rows)})
                    except Exception as e:
                        out["recent_rows"].append({"table": t, "error": str(e)})
            conn.close()
        except Exception as e:
            out["error"] = str(e)
        return out

    for key in ("root_artifacts_db", "src_artifacts_db"):
        report["stores"][key] = scan_sqlite(candidates[key])

    # ---- JSONL conversations ----
    conv_dir = candidates["conversations_dir"]
    conv_out = {"path": str(conv_dir), "exists": conv_dir.exists(), "files": 0, "lines_total": 0, "latest_file": None}
    try:
        if conv_dir.exists():
            files = sorted(conv_dir.glob("*.jsonl"))
            conv_out["files"] = len(files)
            if files:
                conv_out["latest_file"] = str(files[-1])
            total = 0
            for f in files:
                try:
                    with f.open("r", encoding="utf-8", errors="ignore") as fh:
                        for _ in fh:
                            total += 1
                except Exception:
                    pass
            conv_out["lines_total"] = total
    except Exception as e:
        conv_out["error"] = str(e)
    report["stores"]["conversations"] = conv_out

    # ---- proactive files ----
    pro_dir = candidates["proactive_dir"]
    pro_out = {"path": str(pro_dir), "exists": pro_dir.exists(), "files": {}}
    for name in ("latest_context.txt", "latest_summary.txt", "latest_action.txt", "latest_decision.txt", "latest.txt"):
        fp = pro_dir / name
        entry = {"exists": fp.exists(), "size_bytes": (fp.stat().st_size if fp.exists() else 0)}
        if fp.exists():
            try:
                txt = fp.read_text(encoding="utf-8", errors="ignore").strip()
                entry["preview"] = txt[:200]
            except Exception as e:
                entry["read_error"] = str(e)
        pro_out["files"][name] = entry
    report["stores"]["proactive"] = pro_out

    # ---- top-level latest files ----
    for key in ("src_latest_context", "src_latest_summary", "src_latest_action"):
        fp = candidates[key]
        entry = {"path": str(fp), "exists": fp.exists(), "size_bytes": (fp.stat().st_size if fp.exists() else 0)}
        if fp.exists():
            try:
                entry["preview"] = fp.read_text(encoding="utf-8", errors="ignore").strip()[:200]
            except Exception as e:
                entry["read_error"] = str(e)
        report["stores"][key] = entry

    return report


def _format_memory_audit_for_chat(rep: dict) -> str:
    try:
        s = rep.get("stores", {})
        db1 = s.get("root_artifacts_db", {})
        db2 = s.get("src_artifacts_db", {})
        conv = s.get("conversations", {})
        pro = s.get("proactive", {})

        def table_count(d):
            t = d.get("tables", {})
            return ", ".join([f"{k}:{v}" for k, v in t.items() if k != "__all__"][:6]) or "none"

        lines = []
        lines.append("Memory status data:")
        lines.append(f"- Root DB exists: {db1.get('exists')} | size: {db1.get('size_bytes',0)} bytes | tables: {db1.get('tables',{}).get('__all__',0)}")
        lines.append(f"  table counts: {table_count(db1)}")
        lines.append(f"- Src DB exists: {db2.get('exists')} | size: {db2.get('size_bytes',0)} bytes | tables: {db2.get('tables',{}).get('__all__',0)}")
        lines.append(f"  table counts: {table_count(db2)}")
        lines.append(f"- Conversation logs: files={conv.get('files',0)} lines_total={conv.get('lines_total',0)} latest={conv.get('latest_file')}")
        latest_ctx = pro.get("files", {}).get("latest_context.txt", {})
        latest_sum = pro.get("files", {}).get("latest_summary.txt", {})
        latest_act = pro.get("files", {}).get("latest_action.txt", {})
        lines.append(f"- Proactive context exists={latest_ctx.get('exists')} size={latest_ctx.get('size_bytes',0)}")
        lines.append(f"- Proactive summary exists={latest_sum.get('exists')} size={latest_sum.get('size_bytes',0)}")
        lines.append(f"- Proactive action exists={latest_act.get('exists')} size={latest_act.get('size_bytes',0)}")
        lines.append("")
        lines.append("If you want retrieval proof for a specific topic, ask: \"memory audit for <topic>\" and I will query the DB rows directly.")
        return "\n".join(lines)
    except Exception as e:
        return f"Live memory audit failed to format: {e}"

def chat(message: str, *, model: Optional[str] = None, skip_router: bool = False) -> Dict[str, Any]:
    """
    Send a chat message to the LLM.
    If skip_router is True, the message is sent directly without attempting to route as a command.
    """
    msg = (message or "").strip()
    low_msg = msg.lower()

    _convlog_append('user', msg, {'fn':'chat','model': str(model or DEFAULT_CHAT_MODEL)})
    if not msg:
        return {"ok": False, "action": "CHAT", "error": "empty_message", "content": "", "response": ""}

    # Live audit path for memory-capability questions (real DB/files scan, never template)
    if (
        "how your memory" in low_msg or "how does your memory" in low_msg or
        "persistent memory" in low_msg or "across sessions" in low_msg or
        "continuous memory" in low_msg or "memory audit" in low_msg
    ):
        # Return raw data; the cognitive engine will synthesize a natural answer
        import json
        from eli.memory import get_memory, get_memory_status
        mem = get_memory()
        rep = get_memory_status(mem.db_path)
        rep["query"] = message
        content = json.dumps(rep, indent=2)
        _convlog_append('assistant', "[memory data]", {'fn':'chat','path':'live_memory_audit'})
        return {"ok": True, "action": "CHAT", "content": content, "response": content, "memory_audit": rep}

    # Explicit "yesterday" recall path (no fabrication)
    if ("what did we discuss yesterday" in low_msg or ("what did" in low_msg and "yesterday" in low_msg and "discuss" in low_msg)):
        try:
            rec = memory_recall("", limit=8)
            rows = rec.get("results") if isinstance(rec, dict) else []
            if rows:
                lines = []
                for r in rows[:8]:
                    t = str((r or {}).get("text", "")).strip()
                    if t:
                        lines.append("- " + " ".join(t.split())[:240])
                ans = "I can only report what is in persistent memory. Recent stored notes:\n" + ("\n".join(lines) if lines else "(none)")
            else:
                ans = "I can’t verify yesterday from persistent memory right now (no matching stored notes found). I won’t guess."
        except Exception:
            ans = "I can’t verify yesterday from persistent memory right now. I won’t guess."
        return {"ok": True, "action": "CHAT", "model": model or DEFAULT_CHAT_MODEL, "content": ans, "response": ans}

    # Route non-chat intents (commands) only if skip_router is False
    if not skip_router:
        try:
            from eli.execution.router_enhanced import route
            intent = route(msg)
            if intent.get("action") and intent["action"] != "CHAT":
                return execute(intent["action"], intent.get("args", {}))
        except Exception:
            pass

    st = _load_state()

    # Enforce lock BEFORE chatting (prevents silent personality drift)
    ok_lock, reason, details = _verify_persona_lock(st)
    if LOCK_ENFORCED and not ok_lock and (st.get("persona_lock") is not None):
        content = (
            f"ELI LOCKDOWN: persona verification failed ({reason}).\n"
            f"{details.get('msg','')}\n\n"
            "Run PERSONA_LOCK_SET with the intended model/modelfile, or fix environment drift."
        )
        return {"ok": False, "action": "CHAT", "error": "persona_lock_failed", "reason": reason, "details": details, "content": content, "response": content}

    model = (model or DEFAULT_CHAT_MODEL).strip()

    # ── GGUF first (works fully offline) ──────────────────────────────────────
    try:
        from eli.cognition.gguf_inference import load_model, chat_completion as gguf_chat
        if load_model() is not None:
            from eli.core.config import get_persona
            system = get_persona()
            content = gguf_chat(msg, system=system)
            content = _strip_ollama_artifacts((content or "").strip())
            if content:
                _convlog_append('assistant', content, {'fn': 'chat', 'model': 'gguf'})
                return {"ok": True, "action": "CHAT", "model": "gguf", "content": content, "response": content}
    except Exception:
        pass  # GGUF unavailable → fall through to Ollama

    # ── Ollama fallback (requires network / local Ollama server) ──────────────
    try:
        body = {"model": model, "messages": [{"role": "user", "content": msg}], "stream": False}
        res = requests.post(f"{_OLLAMA_HOST}/api/chat", json=body, timeout=CHAT_TIMEOUT_S)
        res.raise_for_status()
        j = res.json() if res.content else {}
        content = _strip_ollama_artifacts(((j.get("message") or {}).get("content") or "").strip())
        if not content:
            content = "(empty model response)"
        _convlog_append('assistant', content, {'fn':'chat','model': model})
        return {"ok": True, "action": "CHAT", "model": model, "content": content, "response": content}
    except Exception as e:
        err = f"chat_failed: {e}"
        _convlog_append('assistant', err, {'fn':'chat','model': model, 'error': True})
        return {"ok": False, "action": "CHAT", "error": str(e), "content": err, "response": err}


def chat_stream(message: str, *, model: Optional[str] = None):
    """
    Streaming chat generator that yields token chunks (str).
    - Routes non-chat intents through execute()
    - Enforces persona lock
    - Falls back to non-stream chat() behavior if needed
    """
    msg = (message or "").strip()

    low_msg = msg.lower()
    if (
        "how your memory" in low_msg or "how does your memory" in low_msg or
        "persistent memory" in low_msg or "across sessions" in low_msg or
        "continuous memory" in low_msg or "memory audit" in low_msg
    ):
        rep = _live_memory_audit()
        yield _format_memory_audit_for_chat(rep)
        return
    if not msg:
        return

    # Route commands first
    try:
        from eli.execution.router_enhanced import route
        intent = route(msg)
        if intent.get("action") and intent["action"] != "CHAT":
            r = execute(intent["action"], intent.get("args", {}))
            out = str((r or {}).get("response") or (r or {}).get("content") or "")
            if out:
                yield out
            return
    except Exception:
        pass

    st = _load_state()
    ok_lock, reason, details = _verify_persona_lock(st)
    if LOCK_ENFORCED and not ok_lock and (st.get("persona_lock") is not None):
        yield (
            f"ELI LOCKDOWN: persona verification failed ({reason}).\n"
            f"{details.get('msg','')}\n\n"
            "Run PERSONA_LOCK_SET with the intended model/modelfile, or fix environment drift."
        )
        return

    model = (model or DEFAULT_CHAT_MODEL).strip()
    full = []
    try:
        body = {"model": model, "messages": [{"role": "user", "content": msg}], "stream": True}
        with requests.post(f"{_OLLAMA_HOST}/api/chat", json=body, stream=True, timeout=CHAT_TIMEOUT_S) as res:
            res.raise_for_status()
            for line in res.iter_lines(decode_unicode=True):
                if not line:
                    continue
                try:
                    j = json.loads(line)
                except Exception:
                    continue
                delta = _strip_ollama_artifacts(((j.get("message") or {}).get("content") or ""))
                if delta:
                    full.append(delta)
                    yield delta
                if j.get("done"):
                    break

        content = "".join(full).strip()
        if content:
            _convlog_append('assistant', content, {'fn':'chat_stream','model': model})
        else:
            fallback = chat(msg, model=model)
            text = str(fallback.get("response") or fallback.get("content") or "")
            if text:
                yield text
    except Exception:
        fallback = chat(msg, model=model)
        text = str(fallback.get("response") or fallback.get("content") or "")
        if text:
            yield text

def open_audio_settings() -> Dict[str, Any]:
    for cmd in (["pavucontrol"], ["gnome-control-center", "sound"]):
        r = _run(cmd, timeout=5)
        if r.get("ok"):
            msg = "Audio settings opened."
            return {"ok": True, "action": "OPEN_AUDIO_SETTINGS", "cmd": cmd, "content": msg, "response": msg}
    msg = "No audio settings app found (pavucontrol / gnome-control-center)."
    return {"ok": False, "action": "OPEN_AUDIO_SETTINGS", "error": msg, "content": msg, "response": msg}


def open_file_system(path: str = "~") -> Dict[str, Any]:
    try:
        target = os.path.expanduser(path)
        open_file(target)
        msg = f"Opened folder: {target}"
        return {"ok": True, "action": "OPEN_FILE_SYSTEM", "path": target, "content": msg, "response": msg}
    except Exception as e:
        msg = "Failed to open file system."
        return {"ok": False, "action": "OPEN_FILE_SYSTEM", "error": repr(e), "content": msg, "response": msg}


def fabricate_document(args: Dict[str, Any]) -> Dict[str, Any]:
    """Generate a real on-disk document (DOCX; optionally ODT) under artifacts/."""
    topic = (args.get('topic') or args.get('title') or 'Untitled Document').strip()
    content = (args.get('content') or '').strip()
    fmt = (args.get('format') or 'docx').lower().strip()

    # If content wasn't supplied, generate it via the chat model, but keep it deterministic-ish.
    if not content:
        prompt = (
            "Write a clean, factual document in Markdown about: " + topic + "\n"
            "Structure: title, short biography/overview, key contributions, timeline, references placeholder. "
            "No made-up citations. If unsure, say so."
        )
        gen = chat(prompt, skip_router=True)
        if not gen.get('ok'):
            return {"ok": False, "error": gen.get('error') or 'document_content_generation_failed'}
        content = gen.get('content', '')

    want_docx = fmt in ('docx', 'both', 'docx+odt', 'odt', 'odt+docx')
    want_odt = fmt in ('odt', 'both', 'docx+odt', 'odt+docx')

    genr = AdvancedDocumentGenerator(base_dir=os.environ.get('ELI_ARTIFACTS_DIR') or 'artifacts')
    res = genr.generate(topic=topic, content=content, want_docx=want_docx, want_odt=want_odt, also_txt=True)
    if not res.ok:
        return {"ok": False, "error": res.error or 'document_generation_failed', "topic": topic, "out_dir": res.out_dir}

    event = json.dumps(
        {
            "event": "artifact_generated",
            "kind": "document_bundle",
            "topic": topic,
            "out_dir": res.out_dir,
            "txt_path": res.txt_path,
            "docx_path": res.docx_path,
            "odt_path": res.odt_path,
        },
        ensure_ascii=False,
        default=str,
    )
    # Return paths so GUI can show / open.
    return {
        "ok": True,
        "topic": topic,
        "out_dir": res.out_dir,
        "txt_path": res.txt_path,
        "docx_path": res.docx_path,
        "odt_path": res.odt_path,
        "content": event,
        "response": event,
    }


def open_browser(url: str = "https://duckduckgo.com", urls: list = None) -> Dict[str, Any]:
    """Open one or more URLs with xdg-open. If urls list provided, opens each in a new detached spawn."""
    try:
        targets = [str(u) for u in (urls or []) if str(u).strip()]
        if not targets:
            targets = [str(url)]
        for target in targets:
            subprocess.Popen(["xdg-open", target], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, start_new_session=True)
        if len(targets) == 1:
            msg = f"Opened browser: {targets[0]}"
            return {"ok": True, "action": "OPEN_BROWSER", "url": targets[0], "content": msg, "response": msg}
        msg = f"Opened {len(targets)} browser tabs."
        return {"ok": True, "action": "OPEN_BROWSER", "urls": targets, "content": msg, "response": msg}
    except Exception as e:
        msg = f"Failed to open browser: {e}"
        return {"ok": False, "action": "OPEN_BROWSER", "error": repr(e), "content": msg, "response": msg}


def set_user_name(name: str) -> Dict[str, Any]:
    import re as _re_sun
    n = (name or "").strip()
    if not n:
        msg = "Empty name."
        return {"ok": False, "action": "SET_USER_NAME", "error": "empty_name", "content": msg, "response": msg}

    # --- Validation gate: reject phrase-like or sentence-fragment names ---
    _NAME_SENTENCE_SIGNALS = frozenset({
        "my", "by", "your", "the", "a", "an", "is", "was", "are", "were",
        "be", "been", "am", "not", "no", "yes", "ok", "okay", "it", "its",
        "this", "that", "which", "what", "who", "how", "why", "when",
        "will", "would", "could", "should", "can", "may", "might", "do",
        "does", "did", "have", "has", "had", "i", "me", "we", "you",
        "they", "he", "she", "and", "or", "but", "so", "if", "then",
        "to", "for", "of", "in", "on", "at", "with", "about", "just",
        "name", "call", "called", "calling", "named",
    })
    _tokens = n.split()
    # Reject if: too long (> 25 chars), too many tokens (> 3 words for a name),
    # or any token matches a sentence-signal word
    _is_phrase = (
        len(n) > 25
        or len(_tokens) > 3
        or any(t.lower() in _NAME_SENTENCE_SIGNALS for t in _tokens)
        or bool(_re_sun.search(r"[^A-Za-z0-9'\-\. ]", n))  # exotic punctuation
    )
    if _is_phrase:
        log.debug(f"[EXECUTOR] set_user_name: rejected phrase-like name {n!r}")
        msg = "That doesn't look like a name — please say just the name."
        return {"ok": False, "action": "SET_USER_NAME", "error": "invalid_name", "content": msg, "response": msg}
    # Write to runtime state (legacy path)
    st = _load_state()
    st["user_name"] = n
    st["updated_at"] = time.time()
    _save_state(st)
    # Write to user_profile.json (authoritative path read by get_user_name())
    try:
        from eli.kernel.state import set_user_name as _state_set_user_name
        _state_set_user_name(n)
    except Exception as _sue:
        log.debug(f"[EXECUTOR] set_user_name profile write failed: {_sue}")
    # Self-heal stale identity patterns: when the name is (re)set, purge any
    # identity.name / identity.preferred_name / identity.nickname rows in
    # user_patterns that do NOT match the new name. This is what stops a wrong
    # earlier value (e.g. a mis-extracted "speak") from surfacing forever, for
    # ANY user — without it, set_user_name updated the profile but left the bad
    # pattern behind.
    try:
        _purge_conflicting_identity_patterns(n)
    except Exception as _pe:
        log.debug(f"[EXECUTOR] identity-pattern purge failed (non-fatal): {_pe}")
    msg = f"Got it. I'll call you {n}."
    return {"ok": True, "action": "SET_USER_NAME", "user_name": n, "content": msg, "response": msg}


def _purge_conflicting_identity_patterns(new_name: str) -> int:
    """Delete identity.* user_patterns rows whose value conflicts with new_name.

    Generic, user-agnostic: keeps only identity name/preferred/nickname patterns
    that actually contain the confirmed name; removes the rest. Returns count.
    """
    import sqlite3 as _sq
    nm = (new_name or "").strip().lower()
    if not nm:
        return 0
    try:
        from eli.memory.memory import get_memory as _gm
        _db = str(_gm().db_path)
    except Exception:
        from eli.core.paths import get_paths as _gp
        _db = str(_gp().artifacts_dir / "db" / "user.sqlite3")
    conn = _sq.connect(_db)
    try:
        rows = conn.execute(
            "SELECT rowid, COALESCE(pattern_data,'') FROM user_patterns "
            "WHERE lower(COALESCE(pattern_type,'')) IN "
            "('identity.preferred_name','identity.name','identity.nickname')"
        ).fetchall()
        stale = [rid for (rid, data) in rows if nm not in str(data).lower()]
        for rid in stale:
            conn.execute("DELETE FROM user_patterns WHERE rowid=?", (rid,))
        conn.commit()
        if stale:
            log.debug(f"[EXECUTOR] purged {len(stale)} stale identity pattern(s) conflicting with {new_name!r}")
        return len(stale)
    finally:
        conn.close()


def get_status() -> Dict[str, Any]:
    st = _load_state()
    ps = _ollama_ps()
    msg = "Status OK."
    return {
        "ok": True,
        "action": "GET_STATUS",
        "user_name": st.get("user_name", ""),
        "chat_model_default": DEFAULT_CHAT_MODEL,
        "persona_lock": st.get("persona_lock") or None,
        "ollama_ps": ps,
        "content": msg,
        "response": msg,
    }


def clear_chat_history() -> Dict[str, Any]:
    st = _load_state()
    st["chat_history"] = []
    st["updated_at"] = time.time()
    _save_state(st)
    msg = "Chat history cleared."
    return {"ok": True, "action": "CLEAR_CHAT_HISTORY", "content": msg, "response": msg}


# ================== MISSING FUNCTIONS FROM Eli_PC_Control2_Working.py ==================

# --- SMART IMPORT HANDLING ---
def _try_import_faster_whisper():
    """Try multiple strategies to import faster_whisper."""
    strategies = [
        # Try direct import first
        lambda: __import__('faster_whisper'),
        
        # Try adding venv path
        lambda: (sys.path.insert(0, os.path.join(os.path.dirname(__file__), '.venv', 'lib', 
                                                 f'python{sys.version_info.major}.{sys.version_info.minor}', 
                                                 'site-packages')) or __import__('faster_whisper')),
        
        # Try common site-packages locations
        lambda: (sys.path.insert(0, os.path.expanduser('~/.local/lib/') + 
                                 f'python{sys.version_info.major}.{sys.version_info.minor}/site-packages') 
                or __import__('faster_whisper')),
    ]
    
    for strategy in strategies:
        try:
            return strategy()
        except ImportError:
            continue
    
    return None


# --- ENV HELPERS ---
def _env_int(name: str, default: int) -> int:
    try:
        v = (os.environ.get(name) or "").strip()
        return int(v) if v else int(default)
    except Exception:
        return int(default)


# --- MEDIA DUCKING ---
MEDIA_DUCK = str(os.environ.get("ELI_MEDIA_DUCK", "1")).strip().lower() not in ("0","false","no","off")
MEDIA_PLAYER = (os.environ.get("ELI_MEDIA_PLAYER", "spotify") or "spotify").strip()

_DUCKED_MEDIA = False
_DUCK_TIMER = None
_DUCK_LOCK = threading.Lock()

def _playerctl_run(argv):
    try:
        r = subprocess.run(argv, capture_output=True, text=True, timeout=30)
        return r.returncode, (r.stdout or "").strip(), (r.stderr or "").strip()
    except Exception as e:
        return 1, "", str(e)

def _playerctl_p(*args):
    if MEDIA_PLAYER:
        return _playerctl_run(["playerctl", "-p", MEDIA_PLAYER, *args])
    return _playerctl_run(["playerctl", *args])

def _duck_media_pause():
    # Pause ONLY if it was actually playing; remember we ducked it.
    global _DUCKED_MEDIA
    if not MEDIA_DUCK:
        return
    if _DUCKED_MEDIA:
        return
    rc, out, _ = _playerctl_p("status")
    if rc == 0 and out.lower() == "playing":
        _playerctl_p("pause")
        _DUCKED_MEDIA = True

def _duck_media_resume():
    global _DUCKED_MEDIA
    if not MEDIA_DUCK:
        return
    if not _DUCKED_MEDIA:
        return
    _playerctl_p("play")
    _DUCKED_MEDIA = False

def _duck_schedule_resume(deadline_s: float):
    # Reset timer each time wake window is extended.
    global _DUCK_TIMER
    if not MEDIA_DUCK:
        return
    with _DUCK_LOCK:
        if _DUCK_TIMER is not None:
            try:
                _DUCK_TIMER.cancel()
            except Exception:
                pass
            _DUCK_TIMER = None
        delay = max(0.0, float(deadline_s) - time.time())
        t = threading.Timer(delay, _duck_media_resume)
        t.daemon = True
        _DUCK_TIMER = t
        t.start()


# --- TTS FUNCTIONS ---
_TTS_ENGINE = None
_MIC_IGNORE_UNTIL = 0.0  # ignore mic until this monotonic time
_BUSY_UNTIL = 0.0  # ignore mic while executing actions
_STOP_REQUESTED = False

def _init_tts():
    global _TTS_ENGINE
    if _TTS_ENGINE is not None:
        return
    try:
        import pyttsx3
    except Exception as e:
        if DEBUG:
            log.debug(f"[TTS] pyttsx3 not available: {e}")
        _TTS_ENGINE = None
        return

    eng = pyttsx3.init()
    voice_name = os.environ.get("ELI_TTS_VOICE", "").strip().lower()
    rate       = os.environ.get("ELI_TTS_RATE", "").strip()

    # optional: choose a different voice by partial name
    if voice_name:
        try:
            for v in eng.getProperty("voices"):
                if voice_name in v.name.lower():
                    eng.setProperty("voice", v.id)
                    break
        except Exception as e:
            if DEBUG:
                log.debug(f"[TTS] voice selection failed: {e}")

    if rate:
        try:
            eng.setProperty("rate", int(rate))
        except Exception as e:
            if DEBUG:
                log.debug(f"[TTS] rate set failed: {e}")

    _TTS_ENGINE = eng

def _tts_list_voices() -> dict:
    _init_tts()
    eng = _TTS_ENGINE
    voices = []
    try:
        for v in (eng.getProperty("voices") or []):
            voices.append({
                "id": getattr(v, "id", "") or "",
                "name": getattr(v, "name", "") or "",
            })
    except Exception:
        pass
    return {"current": os.environ.get("ELI_TTS_VOICE", ""), "voices": voices, "count": len(voices)}

def _tts_set_voice(target: str) -> dict:
    _init_tts()
    eng = _TTS_ENGINE
    target_l = (target or "").strip().lower()
    if not target_l:
        raise ValueError("empty voice target")
    chosen_id = None
    chosen_name = None
    try:
        for v in (eng.getProperty("voices") or []):
            vid = (getattr(v, "id", "") or "").lower()
            vname = (getattr(v, "name", "") or "").lower()
            if target_l in vid or target_l in vname:
                chosen_id = getattr(v, "id", None)
                chosen_name = getattr(v, "name", None) or target
                break
    except Exception:
        pass
    if not chosen_id:
        raise ValueError(f"Voice not found: {target}")
    try:
        eng.setProperty("voice", chosen_id)
    except Exception:
        # still persist preference; engine may apply on next init
        pass
    os.environ["ELI_TTS_VOICE"] = str(chosen_name)
    try:
        MEMORY_PATH.parent.mkdir(parents=True, exist_ok=True)
        TTS_PREF_FILE = MEMORY_PATH.parent / "tts_voice.json"
        TTS_PREF_FILE.write_text(json.dumps({"voice": os.environ["ELI_TTS_VOICE"]}, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass
    return {"voice": os.environ["ELI_TTS_VOICE"], "id": chosen_id}

def _speak_legacy(text: str):
    """Compatibility wrapper: delegate to canonical TTS router."""
    try:
        from eli.perception.tts_router import maybe_speak
        maybe_speak(text, enabled=True)
    except Exception as e:
        log.debug(f"[ELI-TTS] exception: {repr(e)}")
    return


# --- TTS PATCH ---
import tempfile
from eli.utils.platform_compat import open_url, open_file, notify, copy_to_clipboard, play_sound, LINUX, WINDOWS, MACOS

# ---- Ollama host canonical config ----
OLLAMA_HOST = (os.environ.get("ELI_OLLAMA_HOST") or os.environ.get("OLLAMA_HOST") or "http://127.0.0.1:11434").rstrip("/")
_OLLAMA_HOST = OLLAMA_HOST  # backward-compatible alias
CHAT_TIMEOUT_S = (10, 3600)  # connect/read timeout for Ollama chat


def _eli_set_pulse_sink():
    sink = (os.environ.get("ELI_OUTPUT_DEVICE") or "").strip()
    if sink:
        os.environ["PULSE_SINK"] = sink

def _eli_tts_echo(_text: str) -> None:
    # Always show what ELI *tries* to say, even if audio routing is cursed.
    try:
        if os.environ.get("ELI_TTS_ECHO", "1") != "0":
            log.debug(f"[ELI-SAY] {_text}")
        if os.environ.get("ELI_TTS_LOG", "1") != "0":
            from pathlib import Path
            import json
            from datetime import datetime
            root = Path(__file__).resolve().parent
            logp = root / "data" / "logs" / "tts.jsonl"
            logp.parent.mkdir(parents=True, exist_ok=True)
            rec = {"ts": datetime.now().isoformat(timespec="seconds"), "text": _text}
            with open(logp, "a", encoding="utf-8") as f:
                f.write(json.dumps(rec, ensure_ascii=False) + "\n")
    except Exception:
        pass

def _speak(text: str, *args, **kwargs):
    """Speak arbitrarily long text by chunking + sanitizing to avoid TTS silent failures."""
    import re as _re

    if text is None:
        return
    text = str(text)

    # Strip model junk tokens
    text = text.replace("<|end|>", "").replace("<|", "").replace("|>", "")

    # Remove/replace fenced code blocks (TTS engines often choke on them)
    text = _re.sub(r"```.*?```", " [code omitted] ", text, flags=_re.DOTALL)

    # Normalize whitespace
    text = _re.sub(r"[ \t]+", " ", text)
    text = _re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        return

    # Chunk by sentence-ish boundaries, capped by char length
    MAX_CHARS = 350
    parts = _re.split(r"(?<=[.!?])\s+", text)
    buf = ""
    chunks = []
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if len(buf) + len(part) + 1 <= MAX_CHARS:
            buf = (buf + " " + part).strip()
        else:
            if buf:
                chunks.append(buf)
            # hard-split very long parts
            while len(part) > MAX_CHARS:
                chunks.append(part[:MAX_CHARS])
                part = part[MAX_CHARS:]
            buf = part
    if buf:
        chunks.append(buf)

    for c in chunks:
        _speak_raw(c, *args, **kwargs)

def _speak_raw(text: str) -> None:
    """Compatibility wrapper: delegate to canonical TTS router."""
    try:
        from eli.perception.tts_router import speak
        speak(text)
    except Exception:
        return


# --- PIPER TTS HOOK ---
def _eli_speak_piper(text: str):
    """Compatibility wrapper: canonical Piper selection lives in eli.perception.tts_router."""
    try:
        from eli.perception.tts_router import speak
        speak(text)
    except Exception:
        return


# Hook into speak function if piper is configured
try:
    _ELI_SPEAK_ORIG = _speak
    def _speak(text, *args, **kwargs):
        if os.environ.get("ELI_TTS_ENGINE", "").strip().lower() == "piper":
            return _eli_speak_piper(str(text))
        return _ELI_SPEAK_ORIG(text, *args, **kwargs)
except Exception:
    # If ELI changes its internals, fail gracefully instead of crashing startup
    pass


# --- UTILITY FUNCTIONS ---
def _slugify(name: str) -> str:
    name = (name or "doc").strip().lower()
    name = re.sub(r"[^a-z0-9]+", "_", name).strip("_")
    return name or "doc"

def _open_file(path: str) -> None:
    try:
        open_file(path)
    except Exception:
        pass

def _allowed_cmds_set():
    """
    Allowed executables for sandboxed run().
    Accept BOTH comma-separated and whitespace-separated formats.
    Also include both './x' and 'x' variants so normpath quirks don't brick you.
    """
    import os
    import os.path
    from pathlib import Path as _Path

    raw = (os.environ.get("ELI_ALLOWED_CMDS") or "").strip()
    if not raw:
        return None

    # commas OR spaces
    raw = raw.replace(",", " ")
    parts = [p for p in raw.split() if p]

    allowed = set()

    def add(tok: str):
        if not tok:
            return
        t = str(tok).strip()
        if not t:
            return

        # keep raw
        allowed.add(t)

        # normpath form
        try:
            n = os.path.normpath(t)
            allowed.add(n)
        except Exception:
            n = t

        # add "./" and no-"./" variants
        try:
            if n.startswith("./"):
                allowed.add(n[2:])
            else:
                allowed.add("./" + n)
        except Exception:
            pass

    for p in parts:
        add(p)

    # Always allow venv python in common forms (belt + suspenders)
    add("./.venv/bin/python")
    add(".venv/bin/python")
    try:
        add(str((_Path(__file__).resolve().parent.parent / ".venv/bin/python").resolve()))
    except Exception:
        pass

    return allowed

def _allow_or_block(argv0: str) -> bool:
    """Allow or block executables based on sandbox rules."""
    import os
    
    # FULL CONTROL MODE: allow everything
    if os.environ.get("ELI_FULL_CONTROL", "0") == "1":
        return True
    
    allowed = _allowed_cmds_set()
    if allowed is None:
        return True  # No restrictions
    
    # Wildcard allows everything
    if "*" in allowed:
        return True
    
    return argv0 in allowed

def _run_argv(argv, timeout=30):
    # Normalize timeouts (router sometimes sends None)
    # None/"" => env ELI_RUN_TIMEOUT_SEC (default 120s)
    import os as _os
    if timeout is None or timeout == "":
        try:
            timeout = float(_os.environ.get("ELI_RUN_TIMEOUT_SEC", "120"))
        except Exception:
            timeout = 120.0
    else:
        try:
            timeout = float(timeout)
        except Exception:
            timeout = 120.0

    if not argv:
        return {"ok": False, "error": "empty argv"}
    if not _allow_or_block(argv[0]):

        # --- ALLOW_VENV_PYTHON_V1 ---
        # Normalize doubled venv paths and allow our own venv python binary under the allowlist sandbox.
        try:
            import os
            _v = str(argv[0])
            # collapse accidental duplication
            _v = _v.replace("./.venv/bin/./.venv/bin/python", "./.venv/bin/python")
            _v = _v.replace(".venv/bin/.venv/bin/python", ".venv/bin/python")
            _v = os.path.normpath(_v)
            # allow exactly the project venv python
            _venv_py_rel = os.path.normpath("./.venv/bin/python")
            _venv_py_abs = os.path.normpath(str((Path(__file__).resolve().parent / ".venv/bin/python")))
            if _v in (_venv_py_rel, _venv_py_abs) or _v.endswith("/.venv/bin/python"):
                # overwrite the checked value if it's a simple name
                if isinstance(argv[0], str):
                    argv[0] = _v
                # skip raising
                raise SystemExit("__ELI_ALLOW_VENV_PYTHON_SKIP__")
        except SystemExit as _e:
            if str(_e) == "__ELI_ALLOW_VENV_PYTHON_SKIP__":
                pass
            else:
                raise
        except Exception:
            pass

        # If we didn't early-allow venv python, fall through to the original denial.
        try:
            _ALLOW_VENV_PYTHON
        except NameError:
            _ALLOW_VENV_PYTHON = False
        if not _ALLOW_VENV_PYTHON:
            return {"ok": False, "error": f"command not allowed: {argv[0]} (ELI_ALLOWED_CMDS)"}

    try:
        r = subprocess.run(
            argv,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=timeout,
            shell=False,
        )
        return {
            "ok": (r.returncode == 0),
            "returncode": r.returncode,
            "stdout": (r.stdout or "")[-2000:],
            "stderr": (r.stderr or "")[-2000:],
        }
    except Exception as e:
        return {"ok": False, "error": str(e)}

def _spotify_player_name():
    # Pick a player name that actually exists (flatpak spotify usually exposes "spotify")
    r = _run_argv(["playerctl", "-l"], timeout=30)
    if r.get("ok") and r.get("stdout"):
        players = [p.strip() for p in r["stdout"].splitlines() if p.strip()]
        for p in players:
            if "spotify" in p.lower():
                return p
    return "spotify"

def _spotify_launch():
    """
    Launch Spotify and VERIFY it actually stays running.
    """
    SPOTIFY_CMD = os.environ.get("ELI_SPOTIFY_CMD", "").strip() or "spotify"
    argv = shlex.split(SPOTIFY_CMD)
    if not argv:
        return {"ok": False, "error": "ELI_SPOTIFY_CMD is empty"}

    if not _allow_or_block(argv[0]):
        return {"ok": False, "error": f"command not allowed: {argv[0]} (ELI_ALLOWED_CMDS)"}

    # If using Flatpak, verify the app exists/installed first.
    appid = None
    if len(argv) >= 3 and argv[0] == "flatpak" and argv[1] == "run":
        appid = argv[2]
        chk = subprocess.run(
            ["flatpak", "info", appid],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        if chk.returncode != 0:
            return {
                "ok": False,
                "error": f"Flatpak app not installed or not available: {appid}",
                "stderr": (chk.stderr or "")[-2000:],
            }

    try:
        proc = subprocess.Popen(
            argv,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            start_new_session=True,
        )
    except Exception as e:
        return {"ok": False, "error": str(e), "cmd": argv}

    # Give it a moment, then verify it's actually running.
    time.sleep(1.8)

    running = False
    evidence = {}

    # Prefer flatpak ps if we're launching flatpak
    if appid:
        ps = subprocess.run(
            ["flatpak", "ps"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        evidence["flatpak_ps_rc"] = ps.returncode
        evidence["flatpak_ps_tail"] = (ps.stdout or "")[-4000:]
        if ps.returncode == 0 and appid in (ps.stdout or ""):
            running = True
    else:
        # Fallback: check playerctl list (best-effort)
        r = _run_argv(["playerctl", "-l"], timeout=30)
        evidence["playerctl"] = r
        if r.get("ok") and "spotify" in (r.get("stdout","").lower()):
            running = True

    return {
        "ok": bool(running),
        "launched": True,
        "running": bool(running),
        "pid": getattr(proc, "pid", None),
        "cmd": argv,
        "evidence": evidence,
    }

def _spotify_ctl(command: str, auto_launch: bool = True):
    """
    command: play|pause|stop|status|next|previous|play-pause
    """
    player = _spotify_player_name()
    r = _run_argv(["playerctl", "-p", player, command], timeout=30)

    # If play fails because spotify isn't running/registered, optionally launch and retry
    if (not r.get("ok")) and auto_launch and command in {"play", "play-pause"}:
        launch = _spotify_launch()
        time.sleep(2)
        player = _spotify_player_name()
        r2 = _run_argv(["playerctl", "-p", player, command], timeout=30)
        return {"ok": r2.get("ok", False), "launch": launch, "result": r2, "player": player}

    return {"ok": r.get("ok", False), "result": r, "player": player}


# --- APP OPENING FUNCTIONS ---
_APP_ALIASES = {
    "mail": "thunderbird",
    "email": "thunderbird",
    "e-mail": "thunderbird",
    "thunderbird": "thunderbird",
    "thunderbirds": "thunderbird",
    "thunder birds": "thunderbird",
    "tundra bird": "thunderbird",
    "tundrabird": "thunderbird",
    "browser": "firefox",
    "mozilla": "firefox",
    "firefox": "firefox",
}

def _normalize_app(spoken: str) -> str:
    """Normalize application names using aliases."""
    s = (spoken or "").strip().lower()
    s = " ".join(s.split())
    return _APP_ALIASES.get(s, s)

def _allowed_apps_set():
    """Return set of allowed apps, or None for no restrictions."""
    import os
    
    # FULL CONTROL MODE: no app restrictions
    if os.environ.get("ELI_FULL_CONTROL", "0") == "1":
        return None
    
    raw = (os.environ.get("ELI_ALLOWED_APPS") or "").strip()
    if not raw:
        return None  # None => allow common safe apps
    return {x.strip().lower() for x in raw.split(",") if x.strip()}

def _open_app_with_timeout(argv, timeout=30):
    # GUI apps shouldn't be waited on. Spawn + return success unless it exits immediately with error.
    import subprocess
    try:
        proc = subprocess.Popen(
            argv,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            start_new_session=True
        )
        try:
            out, err = proc.communicate(timeout=0.7)
            return {"ok": (proc.returncode == 0), "returncode": proc.returncode, "stdout": out, "stderr": err, "argv": argv}
        except subprocess.TimeoutExpired:
            return {"ok": True, "spawned": True, "pid": proc.pid, "argv": argv}
    except Exception as e:
        return {"ok": False, "error": repr(e), "argv": argv}

def _open_app_builtin(app_name: str) -> dict:
    app = _normalize_app(app_name)
    if not app:
        return {"ok": False, "error": "empty app_name"}

    allowed = _allowed_apps_set()
    if allowed is not None and app not in allowed:
        return {"ok": False, "error": f"app not allowed: {app} (set ELI_ALLOWED_APPS)"}

    # Allow per-app override command
    env_key = "ELI_APP_CMD_" + re.sub(r"[^A-Z0-9]+", "_", app.upper())
    override = (os.environ.get(env_key) or "").strip()

    if override:
        argv = shlex.split(override)
    else:
        # sensible defaults
        if app == "thunderbird":
            argv = ["thunderbird"]
        elif app == "firefox":
            argv = ["firefox"]
        elif app == "spotify":
            SPOTIFY_CMD = os.environ.get("ELI_SPOTIFY_CMD", "").strip() or "spotify"
            argv = shlex.split(SPOTIFY_CMD) if SPOTIFY_CMD else ["spotify"]
        else:
            # last resort: try running by name
            argv = [app]

    try:
        return _open_app_with_timeout(argv)
    except Exception as e:
        return {"ok": False, "error": str(e), "app": app, "cmd": argv}


# --- BUILTIN DOCUMENT FUNCTIONS ---
def _builtin_create_doc(args: dict) -> dict:
    try:
        doc_type  = (args.get("doc_type") or "txt").strip().lower()
        title     = (args.get("title") or "document").strip()
        content   = args.get("content") or ""
        filename  = (args.get("filename") or "").strip()
        convert_to = (args.get("convert_to") or "omit").strip().lower()
        open_after = bool(args.get("open_after", False))

        ext = doc_type if doc_type in {"md","txt","tex"} else "txt"

        outdir = Path(os.environ.get("ELI_DOC_DIR", str(Path(__file__).resolve().parent.parent / "eli_docs"))).expanduser()
        outdir.mkdir(parents=True, exist_ok=True)

        if filename:
            fn = filename if filename.endswith(f".{ext}") else filename + f".{ext}"
        else:
            fn = _slugify(title) + f".{ext}"

        outpath = outdir / fn
        outpath.write_text(content, encoding="utf-8")

        if open_after:
            _open_file(str(outpath))

        return {"ok": True, "path": str(outpath)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def _write_note_builtin(args: dict) -> dict:
    """Create a note file under ELI_DOC_DIR (default: ./eli_docs)."""
    try:
        title = (args.get("title") or "note").strip()
        content = (args.get("content") or "").strip()
        filename = (args.get("filename") or "").strip()
        open_after = bool(args.get("open_after", False))

        import os, re, time, subprocess
        from pathlib import Path

        root = Path(__file__).resolve().parent.parent
        doc_dir = Path(os.environ.get("ELI_DOC_DIR", str(root / "eli_docs"))).expanduser()
        doc_dir.mkdir(parents=True, exist_ok=True)

        if not filename:
            safe_title = re.sub(r"[^a-zA-Z0-9_\-\.]+", "_", title).strip("_") or "note"
            filename = f"{safe_title}.md"
        if not filename.lower().endswith((".md", ".txt")):
            filename += ".md"

        if not content:
            ts = time.strftime("%Y-%m-%d %H:%M:%S")
            content = f"# {title}\n\nCreated: {ts}\n\n"

        path = (doc_dir / filename).resolve()
        doc_dir_resolved = doc_dir.resolve()
        if path != doc_dir_resolved and doc_dir_resolved not in path.parents:
            return {"ok": False, "error": "refusing to write outside ELI_DOC_DIR"}

        path.write_text(content, encoding="utf-8")

        if open_after:
            try:
                open_file(str(path))
            except Exception:
                pass

        return {"ok": True, "path": str(path), "title": title}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def _append_note_builtin(args: dict) -> dict:
    """Append content to a note under ELI_DOC_DIR."""
    try:
        title = (args.get("title") or "").strip()
        content = (args.get("content") or "").strip()
        filename = (args.get("filename") or "").strip()
        open_after = bool(args.get("open_after", False))

        if not title and not filename:
            return {"ok": False, "error": "append_note: title or filename is required"}
        if not content:
            return {"ok": False, "error": "append_note: content is empty"}

        import os, re, time, subprocess
        from pathlib import Path

        root = Path(__file__).resolve().parent.parent
        doc_dir = Path(os.environ.get("ELI_DOC_DIR", str(root / "eli_docs"))).expanduser()
        doc_dir.mkdir(parents=True, exist_ok=True)

        note_file = None
        if filename:
            note_file = (doc_dir / filename)
        else:
            if title.lower().endswith((".md", ".txt")):
                note_file = (doc_dir / title)
            else:
                for ext in (".md", ".txt"):
                    cand = doc_dir / f"{title}{ext}"
                    if cand.exists():
                        note_file = cand
                        break
                if note_file is None:
                    safe_title = re.sub(r"[^a-zA-Z0-9_\-\.]+", "_", title).strip("_") or "note"
                    note_file = doc_dir / f"{safe_title}.md"

        note_file = note_file.resolve()
        doc_dir_resolved = doc_dir.resolve()
        if note_file != doc_dir_resolved and doc_dir_resolved not in note_file.parents:
            return {"ok": False, "error": "refusing to write outside ELI_DOC_DIR"}

        ts = time.strftime("%Y-%m-%d %H:%M:%S")
        blob = f"\n\n---\n\n{ts}\n\n{content.strip()}\n"

        if note_file.exists():
            prev = note_file.read_text(encoding="utf-8", errors="ignore")
        else:
            prev = f"# {title or note_file.stem}\n"

        note_file.write_text(prev.rstrip() + blob, encoding="utf-8")

        if open_after:
            try:
                open_file(str(note_file))
            except Exception:
                pass

        return {"ok": True, "path": str(note_file), "title": title or note_file.stem}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def _calendar_event_builtin(args: dict) -> dict:
    """Open a prefilled Google Calendar event in the browser (handles 'tomorrow', weekdays, 2pm, 16.40pm, etc.)."""
    try:
        import re, urllib.parse, subprocess, os
        from datetime import datetime, date, timedelta

        title = (args.get("title") or "ELI Event").strip()
        date_s = (args.get("date") or args.get("day") or "").strip()
        time_s = (args.get("time") or args.get("at") or "").strip()
        start_s = (args.get("start") or "").strip()
        location = (args.get("location") or "").strip()
        notes = (args.get("notes") or args.get("details") or "").strip()

        dur = args.get("duration_min", None)
        if dur is None:
            dur = args.get("duration", None)
        try:
            duration_min = int(dur) if dur is not None else 30
        except Exception:
            duration_min = 30

        def parse_time(ts: str) -> str:
            if not ts:
                return "09:00"
            t = ts.strip().lower().replace(" ", "")
            t = t.replace(".", ":")
            m = re.match(r'^(\d{1,2})(?::(\d{2}))?(am|pm)?$', t)
            if not m:
                return "09:00"
            hh = int(m.group(1))
            mm = int(m.group(2) or "0")
            ap = m.group(3)
            # If already 13-23, ignore am/pm even if user says nonsense like 16:40pm
            if ap and hh <= 12:
                if ap == "pm" and hh < 12: hh += 12
                if ap == "am" and hh == 12: hh = 0
            hh = max(0, min(23, hh))
            mm = max(0, min(59, mm))
            return f"{hh:02d}:{mm:02d}"

        def parse_date(ds: str) -> str:
            today = date.today()
            if not ds:
                return today.isoformat()
            d = ds.strip().lower()
            wmap = {"monday":0,"tuesday":1,"wednesday":2,"thursday":3,"friday":4,"saturday":5,"sunday":6}
            if d == "today":
                return today.isoformat()
            if d == "tomorrow":
                return (today + timedelta(days=1)).isoformat()
            if d in wmap:
                delta = (wmap[d] - today.weekday()) % 7
                delta = 7 if delta == 0 else delta
                return (today + timedelta(days=delta)).isoformat()
            # ISO date?
            if re.match(r'^\d{4}-\d{2}-\d{2}$', ds.strip()):
                return ds.strip()
            return today.isoformat()

        dt = None
        if start_s:
            # try iso-ish start string
            try:
                dt = datetime.fromisoformat(start_s.replace("Z",""))
            except Exception:
                # salvage "tomorrowT14:00" style
                if "T" in start_s:
                    left, right = start_s.split("T", 1)
                    date_s2 = parse_date(left)
                    time_s2 = parse_time(right)
                    dt = datetime.fromisoformat(f"{date_s2}T{time_s2}")

        if dt is None:
            d_iso = parse_date(date_s)
            t_hm = parse_time(time_s)
            dt = datetime.fromisoformat(f"{d_iso}T{t_hm}")

        end_dt = dt + timedelta(minutes=duration_min)

        def fmt(d: datetime) -> str:
            return d.strftime("%Y%m%dT%H%M%S")

        params = {
            "action": "TEMPLATE",
            "text": title,
            "dates": f"{fmt(dt)}/{fmt(end_dt)}",
        }
        if location:
            params["location"] = location
        if notes:
            params["details"] = notes

        # timezone offset (local)
        import time as _time
        tz_offset = _time.localtime().tm_gmtoff
        tz_sign = "+" if tz_offset >= 0 else "-"
        tz_hours = abs(tz_offset) // 3600
        tz_minutes = (abs(tz_offset) % 3600) // 60
        params["ctz"] = f"UTC{tz_sign}{tz_hours:02d}:{tz_minutes:02d}"

        base_url = "https://calendar.google.com/calendar/render"
        url = f"{base_url}?{urllib.parse.urlencode(params)}"
        open_url(url)

        return {"ok": True, "title": title, "start": dt.isoformat(), "end": end_dt.isoformat(),
                "duration_min": duration_min, "url": url}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def _recall_memory_builtin(args: dict) -> dict:
    """Read recent utterances from ./data/memory/utterances.jsonl (if present)."""
    try:
        import json
        from pathlib import Path

        limit = int(args.get("limit", 20))
        root = Path(__file__).resolve().parent.parent
        memory_log = root / "data" / "memory" / "utterances.jsonl"

        if not memory_log.exists():
            return {"ok": True, "entries": [], "message": "No memory log found"}

        entries = []
        with open(memory_log, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    entries.append(json.loads(line))
                except json.JSONDecodeError:
                    continue

        recent = entries[-limit:] if entries else []
        return {"ok": True, "entries": recent, "count": len(recent), "total": len(entries)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def _type_text_builtin(text: str) -> dict:
    text = (text or "")
    if not text:
        return {"ok": False, "error": "empty text"}

    sess = (os.environ.get("XDG_SESSION_TYPE") or "").lower().strip()

    if sess == "wayland":
        # wtype types into the focused window on Wayland
        argv = ["wtype", "--", text]
        return _run_argv(argv, timeout=30)

    # X11 fallback
    argv = ["xdotool", "type", "--clearmodifiers", "--delay", "8", text]
    return _run_argv(argv, timeout=30)


# --- NOTIFICATION FUNCTION ---
def _notify(title: str, msg: str):
    try:
        if os.environ.get("ELI_NOTIFY", "1") == "1":
            notify(title, msg)
        if os.environ.get("ELI_VOICE_DEBUG", "0") == "1":
            log.debug(f"[NOTIFY] {title}: {msg}")
    except Exception:
        if os.environ.get("ELI_VOICE_DEBUG", "0") == "1":
            log.debug(f"[NOTIFY] {title}: {msg}")


# ----------------------------
# Artifact saver
# ----------------------------

def _artifacts_dir() -> "Path":
    """Return the canonical artifacts root. Honours ELI_ARTIFACTS_DIR (so tests can
    redirect to a tmp dir and never pollute the real artifacts/) — defaults to
    artifacts/ under the project root."""
    from pathlib import Path
    import os as _o
    _env = _o.environ.get("ELI_ARTIFACTS_DIR")
    if _env:
        _p = Path(_env).expanduser()
        return _p if _p.is_absolute() else (Path(__file__).resolve().parents[2] / _p)
    return Path(__file__).resolve().parents[2] / "artifacts"


def _save_artifact(content: str, subdir: str, filename: str, fmt: str = "md") -> str:
    """
    Save content to artifacts/{subdir}/{filename}.{fmt}.
    Supported fmt: md, docx, txt, py, sh, csv, html.
    Returns the absolute path string.
    """
    import re as _re
    from pathlib import Path as _P

    root = _artifacts_dir() / subdir
    root.mkdir(parents=True, exist_ok=True)

    # Clean filename
    safe = _re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", filename).strip(". ")[:80] or "output"
    out_path = root / f"{safe}.{fmt}"
    # Avoid collisions
    if out_path.exists():
        import time as _t
        ts = _t.strftime("%Y%m%d_%H%M%S")
        out_path = root / f"{safe}_{ts}.{fmt}"

    if fmt == "docx":
        try:
            from docx import Document
            from docx.shared import Pt
            doc = Document()
            # Style
            style = doc.styles["Normal"]
            style.font.name = "Calibri"
            style.font.size = Pt(11)
            for block in content.split("\n\n"):
                block = block.strip()
                if not block:
                    continue
                # Heading detection (markdown ## style)
                if block.startswith("### "):
                    p = doc.add_heading(block[4:].strip(), level=3)
                elif block.startswith("## "):
                    p = doc.add_heading(block[3:].strip(), level=2)
                elif block.startswith("# "):
                    p = doc.add_heading(block[2:].strip(), level=1)
                else:
                    p = doc.add_paragraph(block)
            doc.save(str(out_path))
        except Exception as _de:
            # Fallback to txt if docx fails
            out_path = out_path.with_suffix(".txt")
            out_path.write_text(content, encoding="utf-8")
    else:
        out_path.write_text(content, encoding="utf-8")

    return str(out_path)


def _eli_self_description_block(limit: int = 2600) -> str:
    """Return a bounded, factual self-description of ELI for grounding documents
    that are ABOUT ELI itself ("propose upgrades for yourself"). Sourced from the
    repo's own blueprints (real, code-traceable) so a self-referential document is
    grounded in ELI's ACTUAL architecture instead of a hallucinated generic plan.
    Best-effort: returns "" if no blueprint is available (caller degrades)."""
    try:
        from pathlib import Path as _P
        root = _P(__file__).resolve().parents[2]
        for rel in ("blueprints/what_eli_is.md", "blueprints/capability_catalogue.md"):
            fp = root / rel
            if fp.is_file():
                txt = fp.read_text(encoding="utf-8", errors="ignore").strip()
                if txt:
                    return txt[:limit].rstrip() + ("\n…" if len(txt) > limit else "")
    except Exception:
        pass
    return ""




# Generative actions that must be grounded in gathered evidence before synthesis
# (the evidence-routing hook at the top of _execute_impl). Status/grounded-audit
# actions already gather via the bus's specialist agent sets, so they are not here.
_GENERATIVE_EVIDENCE_ACTIONS = {
    "GENERATE_DOCUMENT", "CREATE_DOCUMENT", "DOC_GENERATE",
    "GENERATE_SCRIPT", "CREATE_SCRIPT", "WRITE_SCRIPT", "GENERATE_PROJECT",
}
# DATA_FABRICATOR is intentionally NOT here: it delegates to CREATE_DOCUMENT, which
# gathers evidence itself — listing it too would double-gather.


# ----------------------------
# Dispatcher
# ----------------------------

def _execute_impl(action: str, args: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    args = args or {}
    # Normalize action for dispatch (accept open_app, OPEN_APP, open-app, etc.)
    a = (str(action) if action is not None else "").strip()
    a = re.sub(r"[\s\-]+", "_", a)
    a = a.upper()

    # Action synonyms: the model sometimes emits an action name that means an
    # existing handler but isn't its canonical key (NEWS_SEARCH, daily/weekly
    # report, etc.). Normalise to the real action so these don't fall through to
    # "Unsupported executor action: …".
    _ACTION_ALIASES = {
        "NEWS_SEARCH": "NEWS_FETCH",
        "FETCH_NEWS": "NEWS_FETCH",
        "GET_NEWS": "NEWS_FETCH",
        "WEBSITE_SEARCH": "WEB_SEARCH",
        "WEBSITES_SEARCH": "WEB_SEARCH",
        "SEARCH_WEB": "WEB_SEARCH",
        "DAILY_REPORT": "MORNING_REPORT",
        "WEEKLY_REPORT": "MORNING_REPORT",
    }
    a = _ACTION_ALIASES.get(a, a)

    # ── Evidence-routing for generative tasks (the DAG/plan principle) ─────────
    # Before ELI generates a document/script/project, INTUIT and gather the right
    # evidence (real code analysis, web, memory, runtime — the same agents the bus
    # uses) and attach it so the generator synthesises from real findings, not
    # generic priors. One uniform mechanism for every generative action and every
    # reasoning mode. Idempotent (gated on _evidence absent) so the GENERATE_*→
    # CREATE_* recursion and bus/direct double-dispatch each gather only once; the
    # gatherer's own WEB_SEARCH/RUNTIME_STATUS sub-calls aren't in this set, so
    # there is no recursion. See eli/runtime/evidence_planner.py.
    if a in _GENERATIVE_EVIDENCE_ACTIONS and args.get("_evidence") is None:
        try:
            from eli.runtime.evidence_planner import plan_and_gather as _pag
            _ev_q = str(args.get("topic") or args.get("description") or args.get("query")
                        or args.get("text") or args.get("prompt")
                        or args.get("_raw_user_text") or "").strip()
            _ev_mode = "quick"
            try:
                import eli.kernel.engine as _eng_mod
                _eng = getattr(_eng_mod, "_engine", None)  # read singleton; never construct
                if _eng is not None:
                    _ev_mode = str(getattr(_eng, "_reasoning_mode", "quick") or "quick")
                else:
                    from eli.core.runtime_settings import load_settings as _ls
                    _ev_mode = str((_ls() or {}).get("reasoning_mode", "quick") or "quick")
            except Exception:
                _ev_mode = "quick"
            args["_evidence"], args["_evidence_sources"] = _pag(
                a, _ev_q, _ev_mode,
                session_id=str(args.get("session_id") or ""),
                user_id=str(args.get("user_id") or ""))
        except Exception as _ev_err:
            log.debug(f"[EVIDENCE] injection failed for {a}: {_ev_err}")
            args.setdefault("_evidence", "")
            args.setdefault("_evidence_sources", [])

    if a == "CHAT":
        msg = args.get("message") or args.get("prompt") or args.get("text") or ""
        model = args.get("model") or None
        return chat(str(msg), model=model)

    # Persistent memory
    if a == "MEMORY_STORE":
        txt = str(args.get("text") or args.get("content") or args.get("message") or "").strip()
        # Remove common prefixes like "memory: " and quotes
        txt = re.sub(r'^memory:\s*', '', txt, flags=re.IGNORECASE)
        txt = re.sub(r'^"(.+)"$', r'\1', txt)
        if not txt:
            return {"ok": False, "action": a, "error": "empty_text", "content": "empty_text", "response": "empty_text"}

        tags = args.get("tags", [])
        if isinstance(tags, str):
            tags = [x.strip() for x in tags.split(",") if x.strip()]
        elif isinstance(tags, (tuple, set)):
            tags = [str(x).strip() for x in tags if str(x).strip()]
        elif not isinstance(tags, list):
            tags = []

        # --- FIX: try local memory_store first, then fallback to other backends ---
        # Local JSONL memory_store() (most reliable, no prefix)
        try:
            result = memory_store(txt, tags=tags)
            if isinstance(result, dict):
                result.setdefault("action", a)
                result.setdefault("ok", True)
                result.setdefault("content", result.get("content") or "memory stored")
                result.setdefault("response", result.get("response") or result["content"])
            return result
        except Exception as e:
            # If local fails, try adapter
            try:
                from eli.memory.memory_adapter import memory_store as adapter_store
                result = adapter_store(txt, tags=tags)
                if isinstance(result, dict):
                    result.setdefault("action", a)
                    result.setdefault("ok", True)
                    result.setdefault("content", result.get("content") or "memory stored")
                    result.setdefault("response", result.get("response") or result["content"])
                return result
            except Exception:
                pass

            # Last resort: canonical Memory class
            try:
                from eli.memory.memory import get_memory
                _mem = get_memory()
                _mem.store_memory(txt, tags=",".join(tags) if isinstance(tags, list) else str(tags))
                msg = "memory stored"
                return {"ok": True, "action": a, "content": msg, "response": msg}
            except Exception as e2:
                return {"ok": False, "action": a, "error": str(e2), "content": str(e2), "response": str(e2)}
    
    # --- media control ---
    if a == "STOP_MEDIA":
        return stop_media()
    
    if a == "PAUSE_MEDIA":
        return pause_media()
    
    if a == "PLAY_MEDIA":
        query = (args.get("query") or args.get("song") or args.get("artist") or "").strip()
        target = (args.get("target") or "").strip() or None
        if query:
            return play_specific(query, target)
        return play_media(target)
    
    if a == "NEXT_MEDIA":
        return next_media()
    
    if a == "PREVIOUS_MEDIA":
        return previous_media()

    if a == "SHUFFLE_MEDIA":
        target = (args.get("target") or args.get("app") or args.get("player") or "").strip() or None
        return shuffle_media(target=target)

    if a == "REPEAT_MEDIA":
        target = (args.get("target") or args.get("app") or args.get("player") or "").strip() or None
        return repeat_media(target=target)

    # ---- DISPATCH PATCH: CORE + PROACTIVE ACTIONS ----
    # NOTE: Capabilities can be "registered" without being "implemented".
    # This block ensures the common actions actually route to real code.

    # ---- TIME ----
    if a == "TIME":
        from datetime import datetime
        original_query = args.get("original_query", "") if args else ""
        if "time" in original_query.lower() and not any(w in original_query.lower() for w in ["date","day","today","calendar","days"]):
            time_fmt = "%H:%M:%S"
        else:
            time_fmt = "%Y-%m-%d %H:%M:%S"
        now_str = datetime.now().strftime(time_fmt)
        return {"ok": True, "action": a, "content": now_str, "response": now_str}
    if a == "GET_TIME":
        return _execute_impl("TIME", args)

    # ---- HELP / LIST_CAPABILITIES ----
    if a in ("HELP", "LIST_CAPABILITIES"):
        # Try live introspection first (awareness module)
        public_actions = []
        capabilities = []
        fallback_msg = None

        try:
            from eli.runtime.capability_sync import CapabilitySync
            sync = CapabilitySync()
            caps = sync.discover()
            capabilities = sorted(caps.keys())
            public_actions = [
                {"action": k, "sources": [v.get("source", "")], "aliases": [],
                 "plugins": [v["plugin"]] if v.get("plugin") else []}
                for k, v in caps.items()
            ]
        except Exception as e:
            fallback_msg = f"Live introspection failed: {e}"

        # Fallback: generated report file
        if not capabilities:
            repo_root = Path(__file__).resolve().parents[2]
            report_path = repo_root / "capability_public_report.generated.json"
            if report_path.exists():
                try:
                    report = json.loads(report_path.read_text(encoding="utf-8"))
                    public_actions = report.get("public_actions", []) or []
                    capabilities = [
                        item.get("action")
                        for item in public_actions
                        if isinstance(item, dict) and item.get("action")
                    ]
                except Exception as e:
                    fallback_msg = f"Generated report unreadable: {e}"
            else:
                fallback_msg = "Generated capability report not found."

        # Fallback: SUPPORTED_ACTIONS
        if not capabilities:
            capabilities = sorted(set(SUPPORTED_ACTIONS))
            if not fallback_msg:
                fallback_msg = "Using SUPPORTED_ACTIONS fallback."

        if a == "HELP":
            lines = ["ELI Capabilities:"]
            if public_actions:
                for item in public_actions:
                    if not isinstance(item, dict):
                        continue
                    name = str(item.get("action") or "").strip()
                    if not name:
                        continue
                    aliases = item.get("aliases") or []
                    plugins = item.get("plugins") or []
                    sources = item.get("sources") or []
                    extras = []
                    if aliases:
                        extras.append("aliases=" + ", ".join(str(x) for x in aliases))
                    if plugins:
                        extras.append("plugins=" + ", ".join(str(x) for x in plugins))
                    if sources:
                        extras.append("sources=" + ", ".join(str(x) for x in sources))
                    if extras:
                        lines.append(f"  {name} ({'; '.join(extras)})")
                    else:
                        lines.append(f"  {name}")
                if fallback_msg:
                    lines.append("")
                    lines.append(f"[note] {fallback_msg}")
            else:
                for cap in capabilities:
                    lines.append(f"  {cap}")
                if fallback_msg:
                    lines.append("")
                    lines.append(f"[note] {fallback_msg}")

            txt = "\n".join(lines)
            return {
                "ok": True,
                "action": a,
                "count": len(capabilities),
                "capabilities": capabilities,
                "capability_report": public_actions,
                "content": txt,
                "response": txt,
            }

        lines = [f"Loaded {len(capabilities)} public capabilities:"]
        for cap in capabilities:
            lines.append(f"- {cap}")
        msg = "\n".join(lines)
        if fallback_msg:
            msg += f"\n\n[note] {fallback_msg}"
        return {
            "ok": True,
            "action": a,
            "count": len(capabilities),
            "capabilities": capabilities,
            "capability_report": public_actions,
            "content": msg,
            "response": msg,
        }

    # ---- OPEN_IDE REDIRECT GUARD ----
    if a == "OPEN_IDE":
        target = str((args or {}).get("name") or (args or {}).get("app") or (args or {}).get("target") or "").strip().lower()

        generic = {
            "", "ide", "the ide", "editor", "the editor",
            "built in ide", "built-in ide", "gui ide", "eli ide",
            "internal ide", "ide tab", "the ide tab"
        }

        alias_map = {
            "vscode": "code",
            "visual studio code": "code",
            "virtual studio code": "code",
            "vs code": "code",
            "code": "code",
            "gedit": "gedit",
            "codium": "codium",
            "kate": "kate",
            "sublime": "subl",
            "sublime text": "subl",
        }

        if target not in generic:
            mapped = alias_map.get(target, target)
            return _execute_impl("OPEN_APP", {"name": mapped})

    # ---- OPEN_FILE_SYSTEM VALIDATION GUARD ----
    if a == "OPEN_FILE_SYSTEM":
        try:
            raw_target = str(
                (args or {}).get("path")
                or (args or {}).get("target")
                or (args or {}).get("name")
                or ""
            ).strip()

            low_target = raw_target.lower()

            if low_target in {"trash", "/trash"}:
                probe = _open_app_with_timeout(["xdg-open", "trash:///"])
                if probe.get("ok") or probe.get("spawned"):
                    msg = "Opened folder: trash:///"
                    return {"ok": True, "action": a, "content": msg, "response": msg, "probe": probe}
                return {"ok": False, "action": a, "error": probe.get("stderr") or probe.get("error") or "trash_open_failed",
                        "content": "Could not open trash.", "response": "Could not open trash."}

            if low_target in {"home", "/home", "home directory"}:
                resolved = str(Path.home())
            else:
                cleaned = re.sub(r"\b(folder|directory|path)\b", "", raw_target, flags=re.I).strip()
                cleaned = cleaned.rstrip("/")

                candidates = []

                if raw_target.startswith("~/"):
                    candidates.append(Path(raw_target).expanduser())

                if raw_target.startswith("/Desktop/") or raw_target.startswith("/desktop/"):
                    suffix = raw_target.split("/", 2)[-1] if raw_target.count("/") >= 2 else ""
                    candidates.append(Path.home() / "Desktop" / suffix)

                if raw_target.startswith("/Documents/") or raw_target.startswith("/documents/"):
                    suffix = raw_target.split("/", 2)[-1] if raw_target.count("/") >= 2 else ""
                    candidates.append(Path.home() / "Documents" / suffix)

                if raw_target.startswith("/"):
                    candidates.append(Path(raw_target))

                if cleaned:
                    candidates.append(Path.home() / cleaned)
                    candidates.append(Path.home() / "Desktop" / cleaned)
                    candidates.append(Path.home() / "Documents" / cleaned)

                resolved = ""
                seen = set()
                for cand in candidates:
                    try:
                        p = cand.expanduser().resolve()
                    except Exception:
                        p = cand.expanduser()
                    sp = str(p)
                    if sp in seen:
                        continue
                    seen.add(sp)
                    if p.exists():
                        resolved = sp
                        break

                if not resolved:
                    msg = f"Could not open folder: {raw_target}"
                    return {"ok": False, "action": a, "error": "path_not_found", "content": msg, "response": msg}

            probe = _open_app_with_timeout(["xdg-open", resolved])
            if probe.get("ok") or probe.get("spawned"):
                msg = f"Opened folder: {resolved}"
                return {"ok": True, "action": a, "content": msg, "response": msg, "probe": probe}

            return {
                "ok": False,
                "action": a,
                "error": probe.get("stderr") or probe.get("error") or "xdg_open_failed",
                "content": f"Could not open folder: {resolved}",
                "response": f"Could not open folder: {resolved}",
            }

        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}
    # ---- OPEN_APP ----
    if a == "OPEN_APP":
        try:
            name = (args or {}).get("name") or (args or {}).get("app") or ""
            name = str(name).strip()
            candidates = list((args or {}).get("candidates") or [])

            if not name and candidates:
                name = str(candidates[0]).strip()

            if not name:
                return {
                    "ok": False,
                    "action": a,
                    "error": "Missing app name",
                    "content": "Missing app name",
                    "response": "Missing app name",
                }

            _APP_FUZZY = {
                "calender": "gnome-calendar",
                "calander": "gnome-calendar",
                "calendar": "gnome-calendar",
                "calandar": "gnome-calendar",
                "settings": "gnome-control-center",
                "setting": "gnome-control-center",
                "files": "nautilus",
                "file manager": "nautilus",
                "text editor": "gedit",
                "editor": "gedit",
                "gedit": "gedit",
                "calculator": "gnome-calculator",
                "calc": "gnome-calculator",
                "terminal": "x-terminal-emulator", "term": "x-terminal-emulator",
                "monitor": "gnome-system-monitor",
                "system monitor": "gnome-system-monitor",
                "music": "rhythmbox",
                "photos": "eog",
                "image viewer": "eog",
                "screenshot": "gnome-screenshot",
                "screen recorder": "obs",
                "disk usage": "baobab",
                "disks": "gnome-disks",
                "chrome": "chromium",
                "google chrome": "chromium",
                "chromium": "chromium",
                "camera": "snapshot",
                "vscode": "code",
                "visual studio code": "code",
                "virtual studio code": "code",
                "vs code": "code",
                "code": "code",
                "firefox": "firefox",
            }

            _KNOWN_GUI_BINARIES = {
                "spotify", "discord", "thunderbird", "brave", "firefox",
                "chromium", "libreoffice", "gnome-calendar", "gnome-control-center",
                "nautilus", "gedit", "gnome-calculator", "gnome-terminal",
                "gnome-system-monitor", "rhythmbox", "eog", "gnome-screenshot",
                "obs", "baobab", "gnome-disks", "snapshot", "code", "codium",
                "kate", "subl", "idle",
            }

            def _desktop_entry_exists(app_label: str) -> bool:
                token = str(app_label or "").strip().lower()
                if not token:
                    return False

                roots = [
                    Path("/usr/share/applications"),
                    Path.home() / ".local/share/applications",
                    Path("/var/lib/snapd/desktop/applications"),
                ]

                for root in roots:
                    try:
                        if root.exists():
                            for _ in root.glob(f"*{token}*.desktop"):
                                return True
                    except Exception:
                        pass

                return False

            # Web apps: open in browser rather than trying to install a package
            _WEB_APPS = {
                "gmail": "https://mail.google.com",
                "google mail": "https://mail.google.com",
                "netflix": "https://www.netflix.com",
                "youtube": "https://www.youtube.com",
                "google calendar": "https://calendar.google.com",
                "google drive": "https://drive.google.com",
                "google docs": "https://docs.google.com",
                "google sheets": "https://sheets.google.com",
                "whatsapp": "https://web.whatsapp.com",
                "spotify web": "https://open.spotify.com",
                "twitter": "https://www.twitter.com",
                "x": "https://www.x.com",
                "instagram": "https://www.instagram.com",
                "facebook": "https://www.facebook.com",
                "reddit": "https://www.reddit.com",
            }
            _web_url = _WEB_APPS.get(name.lower())
            if _web_url:
                import subprocess as _wsp
                _wsp.Popen(["xdg-open", _web_url], stdout=_wsp.DEVNULL,
                           stderr=_wsp.DEVNULL, start_new_session=True)
                msg = f"Opening {name} in browser."
                return {"ok": True, "action": a, "content": msg, "response": msg}

            app = _APP_FUZZY.get(name.lower(), name)

            if re.search(r'[:\.,]', name) and len(name.split()) > 3:
                msg = "I can't perform multi-step commands like that. Please break it down."
                return {
                    "ok": False,
                    "action": a,
                    "error": "launch_failed",
                    "content": msg,
                    "response": msg,
                }

            launch_list = [str(x).strip() for x in candidates if str(x).strip()] or [app]
            if app not in launch_list:
                launch_list.insert(0, app)

            tries = []
            last_err = "launch_failed"

            for item in launch_list:
                cmd = [item]
                tries.append(cmd)
                probe = _open_app_with_timeout(cmd)

                gui_evidence = (
                    _desktop_entry_exists(item)
                    or _desktop_entry_exists(name)
                    or _desktop_entry_exists(app)
                    or item in _KNOWN_GUI_BINARIES
                )

                if (probe.get("ok") or probe.get("spawned")) and gui_evidence:
                    msg = f"Opened app: {app}"
                    return {
                        "ok": True,
                        "action": a,
                        "content": msg,
                        "response": msg,
                        "tries": tries,
                        "probe": probe,
                    }

                last_err = probe.get("stderr") or probe.get("error") or f"not_gui_or_not_launchable:{item}"

            try:
                from eli.runtime import grounded_remediation as _eli_gr
                _diag = _eli_gr.diagnose_app(app)
                if not _diag.get("ok", False):
                    return _eli_gr.as_executor_result(_eli_gr.offer_for_result(_diag), ok=False)
            except Exception as _gr_e:
                last_err = f"{last_err}; remediation_error={_gr_e}"

            msg = f"Could not open {app}."
            return {
                "ok": False,
                "action": a,
                "error": last_err,
                "content": msg,
                "response": msg,
                "tries": tries,
            }

        except Exception as e:
            return {
                "ok": False,
                "action": a,
                "error": str(e),
                "content": str(e),
                "response": str(e),
            }

    # ---- RUN_CMD ----
    if a == "RUN_CMD":
        try:
            cmd = args.get("cmd") or args.get("command") or ""
            
            # ── SECURITY GATE ──
            # When cmd is a list, join to a string for pattern matching so that
            # ["bash", "-c", "..."] cannot bypass regex checks via str(list) repr.
            if isinstance(cmd, (list, tuple)):
                _cmd_str = " ".join(str(x) for x in cmd)
            else:
                _cmd_str = str(cmd)
            _cmd_low = _cmd_str.lower().strip()
            
            # ── Destructive / dangerous command patterns — block outright ──
            _BLOCKED_PATTERNS = [
                r"\brm\s+(-[rf]+\s+)?/(?!tmp)",           # rm -rf / (except /tmp)
                r"\bmkfs\b",                                # format filesystem
                r"\bdd\s+.*of=/dev/",                      # dd to raw device
                r"\bchmod\s+777\s+/",                      # chmod 777 on root
                r"\b:\(\)\s*\{.*\}",                        # fork bomb
                r"\bshutdown\b|\breboot\b|\bpoweroff\b",   # system power control
                r"\bcurl\b.*\|\s*(?:ba)?sh\b",             # curl | bash (remote exec)
                r"\bwget\b.*\|\s*(?:ba)?sh\b",             # wget | bash
                r"\b/dev/sd[a-z]\b",                        # raw disk device access
                r"\biptables\s+-F\b",                       # flush firewall rules
                r"\bmv\s+/\S",                              # mv files from root
                r"\b(?:bash|sh|zsh|ksh|fish|dash)\s+-c\b", # shell -c arbitrary exec
                r"\bpython\d*\s+-c\b",                      # python -c arbitrary exec
                r"\bperl\s+-e\b",                           # perl -e arbitrary exec
                r"\bruby\s+-e\b",                           # ruby -e arbitrary exec
                r"\bnc\b.*-e\b",                            # netcat reverse shell
                r"\b(?:ncat|netcat)\b.*-e\b",               # netcat variants
                r">\s*/etc/",                               # redirect to system config
                r">\s*/boot/",                              # redirect to boot partition
                r"\bchpasswd\b|\bpasswd\b\s+\w",           # password change
                r"\bvisudo\b|\bsudoers\b",                  # sudoers modification
                r"\bcrontab\s+-[re]\b",                     # crontab modification
            ]

            for pat in _BLOCKED_PATTERNS:
                if re.search(pat, _cmd_low):
                    msg = f"Blocked dangerous command: {_cmd_str[:60]}"
                    return {"ok": False, "action": a, "error": "security_blocked",
                            "content": msg, "response": msg, "blocked": True}

            # ── Denylist dangerous executable names as argv[0] ──
            import shlex as _shlex_sec
            try:
                if isinstance(cmd, (list, tuple)):
                    _argv0 = str(cmd[0]) if cmd else ""
                else:
                    _argv0 = (_shlex_sec.split(_cmd_str) or [""])[0]
                _argv0_base = os.path.basename(_argv0).lower()
            except Exception:
                _argv0_base = ""
            _DENIED_EXECUTABLES = {
                "bash", "sh", "zsh", "ksh", "fish", "dash",  # shell interpreters
                "python", "python3", "python2",                # scripting engines
                "perl", "ruby", "node", "nodejs",              # more scripting engines
                "nc", "ncat", "netcat",                        # network tools
                "dd", "mkfs", "fdisk", "parted",               # disk tools
                "rm", "shred", "wipe",                         # destructive file ops
                "iptables", "ip6tables", "nftables",           # firewall manipulation
            }
            if _argv0_base in _DENIED_EXECUTABLES:
                msg = f"Execution of '{_argv0_base}' is not permitted via RUN_CMD for security reasons."
                return {"ok": False, "action": a, "error": "security_blocked",
                        "content": msg, "response": msg, "blocked": True}

            # Warn on sudo — allow but flag
            _needs_sudo = "sudo " in _cmd_low
            cwd = args.get("cwd") or None
            timeout = int(args.get("timeout", 30))
            shell = bool(args.get("shell", False))

            # Accept both ["ls", "/path"] and "ls /path"
            if isinstance(cmd, (list, tuple)):
                cmd_list = [str(x) for x in cmd if str(x).strip() != ""]
                if not cmd_list:
                    return {"ok": False, "action": a, "error": "empty_cmd", "content": "", "response": ""}
                p = subprocess.run(
                    cmd_list,
                    cwd=cwd,
                    capture_output=True,
                    text=True,
                    timeout=timeout,
                    shell=False,
                )
            else:
                cmd_str = str(cmd).strip()
                if not cmd_str:
                    return {"ok": False, "action": a, "error": "empty_cmd", "content": "", "response": ""}
                import shlex
                argv = shlex.split(cmd_str)
                if not argv:
                    return {"ok": False, "action": a, "error": "empty_cmd", "content": "", "response": ""}
                p = subprocess.run(
                    argv,
                    cwd=cwd,
                    capture_output=True,
                    text=True,
                    timeout=timeout,
                    shell=False,
                )

            out = (p.stdout or "") + (p.stderr or "")
            return {"ok": p.returncode == 0, "action": a, "code": p.returncode, "content": out, "response": out}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}
    
    # ---- PROACTIVE (start/stop/status) ----
    if a == "HABIT_STATUS":
        try:
            from eli.memory.memory import get_memory as _hm
            _mem = _hm()
            rules = _mem.get_habit_rules(enabled_only=False) if hasattr(_mem, 'get_habit_rules') else []
            if not rules:
                msg = "No habits detected yet. Use ELI regularly and habits will emerge from your patterns."
                return {"ok": True, "action": a, "content": msg, "response": msg}
            lines = [f"Detected habits ({len(rules)}):"]
            for r in rules:
                if not isinstance(r, dict):
                    try: r = dict(r)
                    except: continue
                name = r.get("name", "?")
                cmd = r.get("command", "?")
                hour = r.get("hour", 0)
                minute = r.get("minute", 0)
                enabled = r.get("enabled", True)
                status = "enabled" if enabled else "disabled"
                lines.append(f"  [{status}] {name} — runs '{cmd}' at {hour:02d}:{minute:02d}")
            msg = "\n".join(lines)
            return {"ok": True, "action": a, "content": msg, "response": msg, "rules": rules}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    if a == "PROACTIVE_STATUS":
        try:
            st = proactive_status()
            return {"ok": True, "action": a, **st}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    if a == "PROACTIVE_START":
        try:
            st = proactive_start()
            return {"ok": True, "action": a, **st, "content": "proactive started", "response": "proactive started"}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    if a == "PROACTIVE_STOP":
        try:
            st = proactive_stop()
            return {"ok": True, "action": a, **st, "content": "proactive stopped", "response": "proactive stopped"}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- MESSAGE_TIME_QUERY ----
    # Grounded answer to "what time did I first/last send a message (today)" —
    # read from the conversation_turns log, never the wall clock.
    if a == "MESSAGE_TIME_QUERY":
        import sqlite3 as _sq3
        from datetime import datetime as _dt
        from eli.memory import get_memory as _mt_get_memory
        which = str((args or {}).get("which", "first")).lower()
        scope = str((args or {}).get("scope", "all")).lower()
        try:
            mem = _mt_get_memory()
            dbp = getattr(mem, "db_path", None)
            if not dbp or not Path(dbp).exists():
                msg = "I don't have a conversation-history database to read that from yet."
                return {"ok": False, "action": a, "content": msg, "response": msg}
            where = "role='user'"
            params: List[Any] = []
            if scope == "today":
                midnight = _dt.now().replace(hour=0, minute=0, second=0, microsecond=0).timestamp()
                where += " AND timestamp >= ?"
                params.append(midnight)
            agg = "MAX" if which in ("last", "latest") else "MIN"
            conn = _sq3.connect(str(dbp))
            try:
                row = conn.execute(
                    f"SELECT {agg}(timestamp) FROM conversation_turns WHERE {where}",
                    params,
                ).fetchone()
            finally:
                conn.close()
            ts = (row or [None])[0]
            scope_txt = " today" if scope == "today" else ""
            if not ts:
                msg = f"I have no record of a message from you{scope_txt or ' on file'}."
                return {"ok": True, "action": a, "content": msg, "response": msg,
                        "evidence_source": "conversation_turns"}
            stamp = _dt.fromtimestamp(float(ts)).strftime("%H:%M:%S on %Y-%m-%d")
            label = "most recent" if agg == "MAX" else "first"
            msg = f"Your {label} message{scope_txt} was at {stamp}."
            return {"ok": True, "action": a, "content": msg, "response": msg,
                    "evidence_source": "conversation_turns"}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e),
                    "content": str(e), "response": str(e)}

    # ---- MEMORY_RECALL ----
    if a == "MEMORY_RECALL":
        q = str(args.get("query", "")).strip()
        try:
            lim = int(args.get("limit", 5))
        except Exception:
            lim = 5

        if not q:
            return {"ok": False, "action": a, "error": "empty_query", "content": "empty_query", "response": "empty_query"}

        # 1) Canonical Memory class (proven working)
        try:
            from eli.memory.memory import get_memory as _get_mem_recall
            _mem_r = _get_mem_recall()
            hits = _mem_r.search_memory(q, limit=20) or []
            if isinstance(hits, list):
                # Truncate each hit and cap total to prevent context overflow.
                # Large hits (e.g. repeated conversation dumps) crash llama.cpp.
                _PER_HIT_CHARS = 300
                _MAX_TOTAL_CHARS = 3000
                lines, total = [], 0
                for h in hits:
                    t = (h.get("text", "") or "").strip()
                    if not t:
                        continue
                    t = t[:_PER_HIT_CHARS]
                    lines.append(t)
                    total += len(t)
                    if total >= _MAX_TOTAL_CHARS:
                        break
                content = "Memory recall:\n" + "\n---\n".join(lines) if lines else "Memory recall:\n(no matches)"
            else:
                content = "Memory recall:\n(no matches)"
            if not content.strip():
                content = "Memory recall:\n(no matches)"
            return {"ok": True, "action": a, "hits": hits, "content": content, "response": content}
        except Exception:
            pass

        # 2) Legacy fallback: brain.memory_adapter.memory_search (may not exist)
        try:
            from eli.memory.memory_adapter import memory_search
            result = memory_search(q, k=lim)
            # Ensure consistent envelope
            if isinstance(result, dict):
                result.setdefault("action", a)
                result.setdefault("ok", True)
                if "content" not in result:
                    # memory_adapter may return 'results' (common) or 'hits' (legacy).
                    hits = result.get("results")
                    if hits is None:
                        hits = result.get("hits")
                    if hits is None:
                        hits = result.get("items")
                    if hits is None:
                        hits = []
                    content = ""
                    if isinstance(hits, list) and hits:
                        # Friendly readable render (top-k)
                        lines = []
                        for h in hits:
                            if not isinstance(h, dict):
                                continue
                            t = (h.get("text") or "").strip()
                            if not t:
                                continue
                            score = h.get("score", None)
                            tags = h.get("tags", "")
                            if score is not None:
                                try:
                                    lines.append(f"- ({float(score):.3f}) {t} [{tags}]")
                                except Exception:
                                    lines.append(f"- {t} [{tags}]")
                            else:
                                lines.append(f"- {t} [{tags}]")
                        content = "\n".join(lines).strip()
                    result["content"] = content or "Memory recall:\n(no matches)"
                    # Also expose a consistent 'hits' key for any callers expecting it.
                    if "hits" not in result:
                        result["hits"] = hits
                result.setdefault("response", result.get("response") or result["content"])
            return result
        except Exception:
            pass

        # 2) Fallback: canonical Memory class
        try:
            from eli.memory.memory import get_memory
            _mem = get_memory()
            hits = _mem.search_memory(q, limit=20) or []
            content = "\n".join((h.get("text", "") or "") for h in hits) if isinstance(hits, list) else "Memory recall:\n(no matches)"
            if not content.strip():
                content = "Memory recall:\n(no matches)"
            return {"ok": True, "action": a, "hits": hits, "content": content, "response": content}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # Quick commands
    if a == "CHECK_CHRONAL_ALIGNMENT":
        import time as _t_mod
        _now = _t_mod.localtime()
        msg = (f"Chronal alignment nominal. Local time: "
               f"{_t_mod.strftime('%Y-%m-%d %H:%M:%S', _now)}.")
        return {"ok": True, "action": a, "content": msg, "response": msg}

    if a == "SCHEDULE_TASK":
        # Schedule a heavy/overnight task (code/research/self-upgrade/reflection)
        # to run at a parsed time and surface the result. See runtime.scheduled_tasks.
        req = str((args or {}).get("request") or (args or {}).get("text")
                  or (args or {}).get("message") or "").strip()
        when = str((args or {}).get("when") or req)
        kind = (args or {}).get("kind")
        if not req:
            msg = "What should I work on, and when? (e.g. 'research X overnight')."
            return {"ok": False, "action": a, "content": msg, "response": msg}
        try:
            from eli.runtime.scheduled_tasks import schedule_request
            r = schedule_request(req, when_spec=when, kind=kind)
        except Exception as e:
            msg = f"Couldn't schedule that: {e}"
            return {"ok": False, "action": a, "error": str(e), "content": msg, "response": msg}
        if not r.get("ok"):
            msg = f"Couldn't schedule that: {r.get('error')}"
            return {"ok": False, "action": a, "content": msg, "response": msg}
        msg = (f"Scheduled (job #{r['job_id']}, {r['kind']}): I'll work on this at "
               f"{r['when_human']} and surface the result in the Tasks/Proactive panel.")
        return {"ok": True, "action": a, "content": msg, "response": msg,
                "job_id": r["job_id"], "kind": r["kind"], "when_ts": r["when_ts"]}

    if a == "OPEN_FILE_SYSTEM":
        path = str(args.get("path") or "~")
        return open_file_system(path)
    
    if a == "OPEN_COMMUNICATION_HUB":
        try:
            # Prefer local mail clients; fall back to Gmail in browser.
            candidates = [
                ["thunderbird"],
                ["evolution"],
                ["geary"],
            ]
            for cmd in candidates:
                try:
                    subprocess.Popen(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                    msg = f"Opened Communication Hub: {' '.join(cmd)}"
                    return {"ok": True, "action": a, "content": msg, "response": msg}
                except FileNotFoundError:
                    continue
            url = "https://mail.google.com"
            open_browser(url)
            msg = f"Opened Communication Hub: {url}"
            return {"ok": True, "action": a, "content": msg, "response": msg}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}
    
    if a == "DATA_FABRICATOR":
        try:
            # If a topic/prompt is provided, generate a document first then open it.
            topic = str((args or {}).get("topic") or (args or {}).get("query") or (args or {}).get("text") or (args or {}).get("message") or "").strip()
            if topic:
                r = _execute_impl(action="CREATE_DOCUMENT", args={"topic": topic, "format": (args or {}).get("format") or "md"})
                # If generation succeeded, also open in an editor if available
                if r.get("ok") and r.get("path"):
                    path = r.get("path")
                    editors = [
                        ["code", path],
                        ["codium", path],
                        ["gedit", path],
                        ["kate", path],
                        ["nano", path],
                    ]
                    for cmd in editors:
                        try:
                            subprocess.Popen(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                            break
                        except FileNotFoundError:
                            continue
                return r

            # Otherwise just open an editor for a new blank file.
            out_dir = Path(str(Path.home() / "Documents" / "ELI" / "artifacts")).resolve()
            out_dir.mkdir(parents=True, exist_ok=True)
            out_path = out_dir / f"data_construct_{int(time.time())}.md"
            if not out_path.exists():
                out_path.write_text("# Data Construct\n\n", encoding="utf-8")

            editors = [
                ["code", str(out_path)],
                ["codium", str(out_path)],
                ["gedit", str(out_path)],
                ["libreoffice", "--writer", str(out_path)],
            ]
            for cmd in editors:
                try:
                    subprocess.Popen(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                    msg = f"Opened editor: {' '.join(cmd)}"
                    return {"ok": True, "action": a, "path": str(out_path), "content": msg, "response": msg}
                except FileNotFoundError:
                    continue

            # Fallback: open with default system handler
            open_file(str(out_path))
            msg = f"Opened file: {out_path}"
            return {"ok": True, "action": a, "path": str(out_path), "content": msg, "response": msg}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    if a == "OPEN_AUDIO_SETTINGS":
        try:
            # Try common Linux audio mixers/settings panels.
            candidates = [
                ["pavucontrol"],
                ["gnome-control-center", "sound"],
                ["gnome-control-center"],
                ["systemsettings5", "kcm_pulseaudio"],
                ["systemsettings5"],
                ["kcmshell5", "kcm_pulseaudio"],
                ["mate-volume-control"],
                ["xfce4-mixer"],
            ]
            for cmd in candidates:
                try:
                    subprocess.Popen(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                    msg = f"Opened audio settings: {' '.join(cmd)}"
                    return {"ok": True, "action": a, "content": msg, "response": msg}
                except FileNotFoundError:
                    continue
            raise FileNotFoundError("No audio settings app found (pavucontrol / gnome-control-center / systemsettings5 / etc).")
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    if a == "OPEN_URL":
        url = None
        query = None
        try:
            if isinstance(args, dict):
                url = (
                    args.get("url")
                    or args.get("target")
                    or args.get("value")
                    or args.get("path")
                    or args.get("text")
                )
                query = args.get("query")
        except Exception:
            url = None
            query = None
        return _eli_open_url_action(url, query=query)

    if a == "OPEN_BROWSER":
        import re as _re2
        _site_map = {
            "wikipedia": "https://www.wikipedia.org",
            "google": "https://www.google.com",
            "youtube": "https://www.youtube.com",
            "github": "https://www.github.com",
            "reddit": "https://www.reddit.com",
            "netflix": "https://www.netflix.com",
            "gmail": "https://mail.google.com",
            "stackoverflow": "https://stackoverflow.com",
            "twitter": "https://www.twitter.com",
            "x.com": "https://www.x.com",
        }
        raw_q = (args.get("query") or args.get("url") or args.get("link") or "").lower()

        # YouTube with a search query → use YouTubeController for proper search URL
        _yt_query = (args.get("query") or "").strip()
        if "youtube" in raw_q and _yt_query and "youtube.com" not in raw_q:
            try:
                from eli.tools.media.youtube import YouTubeController
                _cfg_obj = type("_Cfg", (), {"browser": "chromium", "browser_user_dir": str(_browser_user_dir())})()
                _yt = YouTubeController(_cfg_obj)
                result = _yt.open(_yt_query)
                if result.ok:
                    msg = f"Opened YouTube search: {_yt_query}"
                    return {"ok": True, "action": a, "content": msg, "response": msg}
            except Exception:
                pass  # fall through to generic browser open

        found_urls = [v for k, v in _site_map.items() if k in raw_q]
        if len(found_urls) > 1:
            return open_browser(found_urls[0], urls=found_urls)
        url = str(args.get("url") or args.get("link") or "")
        if not url or not url.startswith("http"):
            q = (args.get("query") or "").strip()
            if q:
                import urllib.parse as _up
                url = "https://duckduckgo.com/?q=" + _up.quote_plus(q)
            else:
                url = "https://duckduckgo.com"
        return open_browser(url)

    if a == "SKIP_YOUTUBE_AD":
        try:
            from eli.tools.media.youtube import YouTubeController
            _cfg_obj = type("_Cfg", (), {"browser": (args or {}).get("browser", "chromium"), "browser_user_dir": str(_browser_user_dir())})()
            _yt = YouTubeController(_cfg_obj)
            result = _yt.skip_ad()
            msg = "YouTube ad skipped." if result.ok else f"Ad skip failed: {result.data.get('error', 'unknown')}"
            return {"ok": result.ok, "action": a, "content": msg, "response": msg}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    if a == "OPEN_NETWORK_BROWSER":
        try:
            url = str((args or {}).get("url") or "https://duckduckgo.com").strip()
            open_browser(url)
            msg = f"Opened Network Browser: {url}"
            return {"ok": True, "action": a, "content": msg, "response": msg}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    if a == "OPEN_MEDIA_HUB":
        try:
            # Try Spotify client first; otherwise open YouTube Music; otherwise open local Music folder.
            opened = False
            # Optional: also open audio settings (requested by voice phrase "initiate audio interface")
            if (args or {}).get("open_settings"):
                try:
                    _ = _execute_impl(action="OPEN_AUDIO_SETTINGS", args={})
                except Exception:
                    pass

            candidates = [
                ["spotify"],
                ["flatpak", "run", "com.spotify.Client"],
            ]
            for cmd in candidates:
                try:
                    subprocess.Popen(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                    opened = True
                    msg = f"Opened media: {' '.join(cmd)}"
                    return {"ok": True, "action": a, "content": msg, "response": msg}
                except FileNotFoundError:
                    continue

            # Browser fallbacks
            try:
                url = "https://music.youtube.com"
                open_browser(url)
                msg = f"Opened media: {url}"
                return {"ok": True, "action": a, "content": msg, "response": msg}
            except Exception:
                pass

            # Local folder fallback
            try:
                music_dir = os.path.expanduser("~/Music")
                open_file_system(music_dir)
                msg = f"Opened media folder: {music_dir}"
                return {"ok": True, "action": a, "content": msg, "response": msg}
            except Exception as e:
                raise e
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    if a == "CREATE_DOCUMENT":
        topic = (args.get("topic") or args.get("text") or args.get("description") or args.get("prompt") or "").strip()
        if not topic:
            msg = "Missing topic for document generation"
            return {"ok": False, "action": a, "error": msg, "content": msg, "response": msg}

        # Consume the evidence gathered by the central evidence-routing hook
        # (evidence_planner). If this handler was called directly without it (a
        # non-engine caller / test), gather on the fly as a fallback.
        _doc_evidence = str(args.get("_evidence") or "")
        _doc_sources = list(args.get("_evidence_sources") or [])
        if args.get("_evidence") is None:
            try:
                from eli.runtime.evidence_planner import plan_and_gather as _pag
                _doc_evidence, _doc_sources = _pag(
                    "CREATE_DOCUMENT",
                    str(args.get("_raw_user_text") or args.get("raw") or topic),
                    session_id=str(args.get("session_id") or ""),
                    user_id=str(args.get("user_id") or ""))
            except Exception:
                _doc_evidence, _doc_sources = "", []

        fmt = (args.get("format") or "md").lower().strip()
        import os as _os
        if _os.environ.get("ELI_TEST_MODE", "").strip().lower() in {"1", "true", "yes", "on"}:
            import re as _re
            from pathlib import Path as _DPath

            safe_name = _re.sub(r"[^a-z0-9]+", "_", topic.lower())[:50].strip("_")
            ext = ".md" if fmt == "md" else f".{fmt}"
            fname = f"{safe_name}{ext}" if safe_name else f"document{ext}"

            doc_dir = _artifacts_dir() / "documents"
            doc_dir.mkdir(parents=True, exist_ok=True)
            doc_path = doc_dir / fname

            content = (
                f"# {topic}\n\n"
                "## Purpose\n\n"
                "This deterministic test-mode document proves that document generation writes a real artifact "
                "to the documents directory without using the old placeholder body.\n\n"
                "## Generation Contract\n\n"
                "- Output must be substantive enough for downstream UI wiring tests.\n"
                "- Output must not be a blank skeleton, a path echo, or a promise to write later.\n"
                "- Output must stay in `artifacts/documents`, never `artifacts/scripts`.\n\n"
                "## Verification Notes\n\n"
                "The runtime generation path applies the same destination contract while using the local "
                "model to write the full document body. Test mode keeps the content deterministic so "
                "regression tests can detect stub leakage without loading a model.\n"
            )
            if fmt != "md":
                content = _re.sub(r"^#+\s*", "", content, flags=_re.M)
            doc_path.write_text(content, encoding="utf-8")

            return {
                "ok": True,
                "action": a,
                "content": content,
                "response": json.dumps(
                    {
                        "event": "artifact_generated",
                        "kind": "document",
                        "path": str(doc_path),
                        "chars": len(content),
                    },
                    ensure_ascii=False,
                    default=str,
                ),
                "doc_path": str(doc_path),
                "filename": fname,
                "open_in_ide": False,
            }
        
        # Determine document type and tailor the generation settings.
        _topic_low = topic.lower()
        _requested_doc_type = str(
            args.get("doc_type")
            or args.get("document_type")
            or args.get("report_type")
            or ""
        ).strip()
        _doc_type_low = (_requested_doc_type or _topic_low).lower()
        _is_thesis = any(w in _doc_type_low for w in ["phd", "dissertation", "master", "thesis"])
        _is_research = any(w in _doc_type_low for w in ["research", "peer-review", "peer review", "literature"])
        _is_report = any(w in _doc_type_low for w in ["report", "analysis", "review", "assessment", "evaluation"])
        _is_guide = any(w in _doc_type_low for w in ["guide", "tutorial", "how to", "howto", "handbook", "manual"])
        _is_essay = any(w in _doc_type_low for w in ["essay", "article", "blog", "post", "opinion", "argument"])
        _is_plan = any(w in _doc_type_low for w in ["plan", "proposal", "strategy", "roadmap", "specification"])
        _is_summary = any(w in _doc_type_low for w in ["summary", "summarize", "summarise", "overview", "brief"])
        
        if _requested_doc_type:
            doc_type = _requested_doc_type
        elif _is_summary:
            doc_type = "summary"
        elif _is_thesis:
            doc_type = "thesis/dissertation"
        elif _is_research:
            doc_type = "research article"
        elif _is_report:
            doc_type = "report"
        elif _is_guide:
            doc_type = "guide"
        elif _is_essay:
            doc_type = "essay/article"
        elif _is_plan:
            doc_type = "plan"
        else:
            doc_type = "document"

        if _is_summary:
            style_instruction = "Use a compact abstract, key findings, caveats, and concise implications."
            target_tokens = 1200
        elif _is_thesis:
            style_instruction = "Use academic thesis structure: abstract, research question, literature context, method, analysis, limitations, and future work."
            target_tokens = 5000
        elif _is_research:
            style_instruction = "Use scholarly article structure: abstract, introduction, background, method or argument, findings, limitations, and references-needed notes."
            target_tokens = 4200
        elif _is_report:
            style_instruction = "Use report structure: executive summary, scope, evidence, analysis, risks, recommendations, and conclusion."
            target_tokens = 3600
        elif _is_guide:
            style_instruction = "Use guide structure: prerequisites, steps, examples, validation checks, and troubleshooting."
            target_tokens = 2500
        elif _is_essay:
            style_instruction = "Use article structure: thesis, developed argument, evidence, counterpoints, and conclusion."
            target_tokens = 2500
        elif _is_plan:
            style_instruction = "Use plan structure: objectives, milestones, resources, risks, dependencies, and acceptance checks."
            target_tokens = 2500
        else:
            style_instruction = "Use a complete document structure with clear sections, specific claims, caveats, and a defensible conclusion."
            target_tokens = 2500
        try:
            from eli.runtime.runtime_policy import budget as _eli_budget
            target_tokens = _eli_budget("document_tokens", target_tokens, floor=target_tokens, ceiling=7000)
        except Exception:
            pass
        
        system_prompt = (
            "You are ELI's local document generation engine. Produce a finished artifact, not a scaffold. "
            f"Document type: {doc_type}. {style_instruction} "
            "Use markdown with meaningful headings, complete paragraphs, and source discipline. "
            "Do not output blank sections, TODOs, placeholders, invented citations, fake file paths, "
            "or promises to expand later. If evidence is missing for a factual claim, mark it [source needed] "
            "or state the assumption explicitly. Every section must carry real content, concrete reasoning, "
            "and a defensible conclusion. "
            "When EVIDENCE is provided below, tie every section to specific items in it — name the real "
            "components, findings, files, or results it contains — and do NOT substitute generic domain "
            "or business-plan filler for what the evidence actually says."
        )
        
        user_prompt = (
            f"Write a complete {doc_type} about: {topic}\n\n"
            + (("EVIDENCE — gathered by ELI's own agents for THIS task"
                + (f" (sources: {', '.join(_doc_sources)})" if _doc_sources else "")
                + ". Ground every claim and proposal in the specifics below; refer to "
                "them directly (real subsystems, real findings, real results). Do NOT "
                "pad with generic advice. Where the evidence is thin for a point, say so "
                "honestly rather than inventing:\n\n"
                f"{_doc_evidence}\n\n")
               if _doc_evidence else
               ("No external evidence could be gathered for this topic (e.g. the network "
                "is off, or nothing relevant is on file). Rely only on well-established "
                "knowledge, state assumptions explicitly, and do not fabricate specifics, "
                "citations, or file paths.\n\n"))
            + "Acceptance criteria:\n"
            "- The result is a usable document, not a template.\n"
            "- The opening states the purpose and scope.\n"
            "- The body develops the topic with specific claims, caveats, and implications.\n"
            "- The ending gives a concrete conclusion or next action.\n"
            "- No placeholder headings, no filler, no fake references."
        )
        
        try:
            from eli.cognition import gguf_inference as _gguf
            _model = _gguf.load_model()
            if _model is not None:
                content = ""
                # Multi-stage grounded pipeline (plan/outline → grounded sections →
                # review→revise → polish), grounded in the evidence the central hook
                # gathered, with a deepen-retry that escalates agent tiers when the
                # evidence is thin. Falls back to a single pass on any miss / kill switch.
                try:
                    from eli.runtime import report_pipeline as _rp
                    if _rp.enabled():
                        def _doc_ask(prompt, system=None, max_tokens=1500, temperature=0.4):
                            return _gguf.chat_completion(prompt, system=system,
                                max_tokens=max_tokens, temperature=temperature, top_p=0.85)
                        def _doc_deepen():
                            from eli.runtime.evidence_planner import plan_and_gather as _pag
                            return _pag(a, str(args.get("_raw_user_text") or topic),
                                        "tree_of_thoughts",
                                        session_id=str(args.get("session_id") or ""),
                                        user_id=str(args.get("user_id") or ""))
                        _pi = _rp.generate_document(
                            topic, ask=_doc_ask, evidence=_doc_evidence,
                            sources=_doc_sources, doc_type=doc_type,
                            target_words=max(700, int(target_tokens * 0.6)),
                            deepen_cb=_doc_deepen)
                        if _pi.get("ok") and _pi.get("text"):
                            content = _pi["text"]
                            log.debug(f"[CREATE_DOCUMENT] multi-stage pipeline produced "
                                      f"{len(content)} chars across {len(_pi.get('sections') or [])} sections")
                except Exception as _rp_err:
                    log.debug(f"[CREATE_DOCUMENT] pipeline failed, single-pass fallback: {_rp_err}")
                if not content:
                    content = _gguf.chat_completion(
                        user_prompt,
                        system=system_prompt,
                        max_tokens=target_tokens,
                        temperature=0.4,
                        top_p=0.85,
                    )
                _min_doc_chars = int(_os.environ.get("ELI_DOCUMENT_MIN_CHARS", "800") or "800")
                if content:
                    import re as _re
                    content = _re.sub(r"^```[a-z]*\n?", "", content.strip(), flags=_re.MULTILINE)
                    content = _re.sub(r"\n?```$", "", content.strip(), flags=_re.MULTILINE)
                    content = content.strip()
                    if len(content) < _min_doc_chars:
                        log.debug(
                            f"[CREATE_DOCUMENT] rejected short document: {len(content)} chars < {_min_doc_chars}",
                        )
                    else:
                    
                        # Save to file
                        safe_name = _re.sub(r"[^a-z0-9]+", "_", topic.lower())[:50].strip("_")
                        ext = ".md" if fmt == "md" else f".{fmt}"
                        fname = f"{safe_name}{ext}" if safe_name else f"document{ext}"
                        
                        # Use the canonical artifacts root (honours ELI_ARTIFACTS_DIR;
                        # monkeypatchable in tests). NB the env var alone is unreliable —
                        # runtime_settings strips out-of-project ELI_ARTIFACTS_DIR on
                        # load_settings(), so tests must patch _artifacts_dir, not setenv.
                        doc_dir = _artifacts_dir() / "documents"
                        doc_dir.mkdir(parents=True, exist_ok=True)
                        doc_path = doc_dir / fname
                        doc_path.write_text(content, encoding="utf-8")
                        
                        return {
                            "ok": True, "action": a, "content": content,
                            "response": json.dumps(
                                {
                                    "event": "artifact_generated",
                                    "kind": "document",
                                    "path": str(doc_path),
                                    "chars": len(content),
                                },
                                ensure_ascii=False,
                                default=str,
                            ),
                            "doc_path": str(doc_path), "filename": fname,
                            "open_in_ide": True,
                        }
        except Exception as _e:
            log.debug(f"[CREATE_DOCUMENT] GGUF error: {_e}")
        
        # Fallback
        msg = "Document generation failed — GGUF model unavailable or produced a document below the quality threshold."
        return {"ok": False, "action": a, "error": msg, "content": msg, "response": msg}


    if a == "PERSONA_LOCK_SET":
        model = args.get("model") or None
        modelfile = args.get("modelfile_path") or args.get("modelfile") or None
        return persona_lock_set(model=model, modelfile_path=modelfile)

    if a == "PERSONA_LOCK_STATUS":
        return persona_lock_status()

    if a == "PERSONA_LOCK_CLEAR":
        return persona_lock_clear()

    if a == "SELF_TEST":
        return self_test()

    # ---- RUN_TESTS — run the pytest suite and return the results document ----
    # ELI can run this and SUMMARISE it in chat. Defaults to a fast, high-signal
    # subset (structural claims); pass args["target"] for more, or schedule the
    # full suite + engine eval overnight (SCHEDULE_TASK kind=eval). The report is
    # always at artifacts/test_report.md (auto-written by the pytest hook).
    if a == "RUN_TESTS":
        try:
            import subprocess as _ts_sp, sys as _ts_sys
            from pathlib import Path as _TP
            repo = _TP(__file__).resolve().parents[2]
            target = str(args.get("target") or "tests/claims/test_structural_claims.py")
            r = _ts_sp.run([_ts_sys.executable, "tools/run_test_report.py", target],
                           cwd=str(repo), capture_output=True, text=True, timeout=900)
            rep = repo / "artifacts" / "test_report.md"
            body = rep.read_text(encoding="utf-8")[:4000] if rep.exists() else \
                (r.stdout or r.stderr or "")[-2000:]
            return {"ok": True, "action": a, "content": body, "response": body,
                    "report_path": str(rep), "evidence_source": "pytest_test_report"}
        except Exception as e:
            msg = f"RUN_TESTS failed: {e}"
            return {"ok": False, "action": a, "error": str(e), "content": msg, "response": msg}

    # ---- GENERATE_TESTS — ELI writes + sandbox-verifies behavioural tests ----
    # Phase 4: per-function test generation (only passing candidates kept). Heavy
    # (one model call per target); keep the chat limit small — schedule larger runs
    # ("generate tests overnight" → SCHEDULE_TASK kind=testgen).
    if a == "GENERATE_TESTS":
        try:
            from eli.runtime.test_generator import run_testgen
            limit = int(args.get("limit") or 3)
            res = run_testgen(limit=max(1, min(limit, 10)))
            if not res.get("ok"):
                msg = f"Test generation unavailable: {res.get('reason')}"
                return {"ok": False, "action": a, "content": msg, "response": msg}
            files = res.get("accepted_files") or []
            msg = (f"Generated and sandbox-verified {res.get('accepted', 0)} behavioural "
                   f"test(s) (rejected {res.get('rejected', 0)} of {res.get('targets', 0)} "
                   f"targets — those failed verification). "
                   + ("Accepted: " + ", ".join(files) if files else "Nothing accepted this run.")
                   + " Review them in tests/generated/ (manifest: _manifest.json).")
            return {"ok": True, "action": a, "content": msg, "response": msg, "result": res}
        except Exception as e:
            msg = f"GENERATE_TESTS failed: {e}"
            return {"ok": False, "action": a, "error": str(e), "content": msg, "response": msg}

    # ---- LORA_STATUS — read-only LoRA training readiness (preflight) ----
    if a == "LORA_STATUS":
        try:
            from eli.learning.training_preflight import preflight_all
            pf = preflight_all()
            lines = ["LoRA training readiness:"]
            for rep in pf.get("reports", []):
                probs = rep.get("problems") or []
                lines.append(f"- {rep.get('target')}: "
                             f"{'READY' if rep.get('can_train') else 'NOT READY'}"
                             + (f" — {'; '.join(probs)}" if probs else ""))
            msg = "\n".join(lines)
            return {"ok": True, "action": a, "content": msg, "response": msg,
                    "evidence_source": "lora_preflight", "result": pf}
        except Exception as e:
            msg = f"LORA_STATUS failed: {e}"
            return {"ok": False, "action": a, "error": str(e), "content": msg, "response": msg}

    # ---- LORA_TRAIN — run the LoRA pipeline DAG (DRY-RUN by default) ----
    # From chat this is always a dry-run readiness pass (preflight→build→eval); real
    # training (execute=True) only via the scheduled `lora` task / explicit GUI.
    if a == "LORA_TRAIN":
        try:
            from eli.learning.lora_pipeline import run_pipeline
            target = str(args.get("target") or "eli_phi")
            execute = bool(args.get("execute", False))
            res = run_pipeline(target, execute=execute,
                               max_steps=int(args.get("max_steps", 1) or 1))
            detail = "\n".join(
                f"- {s['stage']}: {'ok' if s['ok'] else 'FAILED'} — {s.get('detail')}"
                for s in res.get("stages", []))
            msg = (res.get("summary", "") + "\n" + detail).strip()
            return {"ok": bool(res.get("ok")), "action": a, "content": msg,
                    "response": msg, "evidence_source": "lora_pipeline", "result": res}
        except Exception as e:
            msg = f"LORA_TRAIN failed: {e}"
            return {"ok": False, "action": a, "error": str(e), "content": msg, "response": msg}

    # ---- ORCHESTRATION_STATUS — explain the agent DAG + last orchestrated run ----
    if a == "ORCHESTRATION_STATUS":
        try:
            from eli.cognition.agent_bus import orchestration_snapshot
            snap = orchestration_snapshot()
            lines = [
                f"Agent orchestration ({snap.get('engine')}): "
                f"{snap.get('count')} agents, critical path {snap.get('critical_path')} layer(s).",
                "Execution layers (each runs in parallel):",
            ]
            for i, layer in enumerate(snap.get("execution_layers") or []):
                lines.append(f"  layer {i}: {', '.join(layer)}")
            if snap.get("dependencies"):
                lines.append("Dependencies: " + "; ".join(
                    f"{k}←{','.join(v)}" for k, v in snap["dependencies"].items()))
            last = snap.get("last_run")
            if last:
                lines.append(f"Last run: ok={last.get('ok')} "
                             f"{last.get('seconds')}s, failed={last.get('failed')}, "
                             f"skipped={last.get('skipped')}")
            msg = "\n".join(lines)
            return {"ok": True, "action": a, "content": msg, "response": msg,
                    "evidence_source": "agent_orchestrator", "result": snap}
        except Exception as e:
            msg = f"ORCHESTRATION_STATUS failed: {e}"
            return {"ok": False, "action": a, "error": str(e), "content": msg, "response": msg}

    # ---- TEST_REVIEW — run the suite, back up + write errors, then summarise + offer options ----
    # The grounded report below is summarised by the persona; the options route to
    # existing actions (examine/propose/generate/eval) so the user can work through them.
    if a == "TEST_REVIEW":
        try:
            from eli.runtime.test_review import run_and_review
            target = str(args.get("target") or "tests/claims/test_structural_claims.py")
            res = run_and_review(target)
            if res.get("error"):
                msg = f"Test review could not run: {res['error']}"
                return {"ok": False, "action": a, "content": msg, "response": msg}
            t = res.get("totals", {})
            lines = [f"Test review of `{target}`: {t.get('passed', 0)} passed, "
                     f"{t.get('failed', 0)} failed, {t.get('errored', 0)} errored, "
                     f"{t.get('xfailed', 0)} known-gap (xfail) of {t.get('total', 0)}."]
            if res.get("failures"):
                lines.append("Failing tests (possible errors):")
                for f in res["failures"][:10]:
                    lines.append(f"  - {f['node']}: {(f.get('message') or '')[:120]}")
            if res.get("error_file"):
                lines.append(f"Errors written to: {res['error_file']}")
            if res.get("backup_path"):
                lines.append(f"Previous report backed up to: {res['backup_path']}")
            if res.get("options"):
                lines.append("What would you like to do next?")
                for i, o in enumerate(res["options"], 1):
                    lines.append(f"  {i}. {o['label']}  — say: \"{o['command']}\"")
            msg = "\n".join(lines)
            return {"ok": True, "action": a, "content": msg, "response": msg,
                    "evidence_source": "test_review", "result": res}
        except Exception as e:
            msg = f"TEST_REVIEW failed: {e}"
            return {"ok": False, "action": a, "error": str(e), "content": msg, "response": msg}

    if a == "SET_USER_NAME":
        return set_user_name(str(args.get("name", "")))

    if a == "GET_STATUS":
        return get_status()

    if a == "CLEAR_CHAT_HISTORY":
        return clear_chat_history()

    # ---- DATE ----
    if a == "DATE":
        from datetime import datetime as _dt
        d = _dt.now().strftime("%A, %Y-%m-%d")
        return {"ok": True, "action": a, "content": d, "response": d}
    if a == "GET_DATE":
        return _execute_impl("DATE", args)

    # ---- MEDIA_CONTROL (generic router output) ----
    if a == "MEDIA_CONTROL":
        cmd = (args.get("command") or "play-pause").strip().lower()
        target = (args.get("target") or args.get("app") or args.get("player") or "").strip().lower() or None
        cmd_map = {
            "pause": "PAUSE_MEDIA", "stop": "STOP_MEDIA",
            "play": "PLAY_MEDIA", "resume": "PLAY_MEDIA",
            "next": "NEXT_MEDIA", "skip": "NEXT_MEDIA",
            "previous": "PREVIOUS_MEDIA", "prev": "PREVIOUS_MEDIA",
        }
        mapped = cmd_map.get(cmd)
        if mapped:
            # Forward target into sub-handler args
            forwarded = dict(args)
            forwarded["_target"] = target
            # Call targeted versions directly when target present
            if target and mapped == "PAUSE_MEDIA":
                return pause_media(target=target)
            if target and mapped in ("PLAY_MEDIA",):
                return play_media(target=target)
            if target and mapped == "STOP_MEDIA":
                return stop_media(target=target)
            return _execute_impl(mapped, forwarded)
        # Fallback: playerctl generic
        try:
            subprocess.run(["playerctl", cmd], timeout=5, capture_output=True)
            msg = f"Media: {cmd}"
            return {"ok": True, "action": a, "content": msg, "response": msg}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- VOLUME ----
    if a == "VOLUME":
        direction = (args.get("direction") or "").strip().lower()
        level = args.get("level")
        delta = int(args.get("delta", 10))
        return _volume_fallback(direction=direction, delta=delta, level=level)


    # ---- SCREENSHOT ----
    if a == "SCREENSHOT":
        try:
            from eli.perception.os_controller import take_screenshot
            region = (args.get("region") or "full").strip()
            return take_screenshot(region=region)
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- KEYBOARD ----
    if a == "KEYBOARD":
        try:
            from eli.perception.os_controller import press_key, type_text
            if args.get("type"):
                return type_text(str(args["type"]))
            elif args.get("key"):
                return press_key(str(args["key"]))
            return {"ok": False, "action": a, "error": "No key or text specified", "content": "No key or text specified", "response": "No key or text specified"}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- OCR_IMAGE ----
    if a == "OCR_IMAGE":
        try:
            path = str(args.get("path") or "").strip()
            if not path:
                return {"ok": False, "action": a, "error": "No image path provided",
                        "content": "No image path provided", "response": "No image path provided"}
            p = Path(path).expanduser().resolve()
            if not p.exists():
                return {"ok": False, "action": a, "error": f"File not found: {p}",
                        "content": f"File not found: {p}", "response": f"File not found: {p}"}
            text = None
            # 1. tesseract CLI (no Python binding needed, no model download)
            try:
                import subprocess as _sp
                _tess = shutil.which("tesseract")
                if _tess:
                    _r = _sp.run([_tess, str(p), "stdout", "--psm", "11", "-l", "eng"],
                                 capture_output=True, text=True, timeout=30)
                    if _r.returncode == 0:
                        text = _r.stdout.strip()
            except Exception:
                pass
            # 2. pytesseract binding (calls tesseract binary)
            if not text:
                try:
                    import pytesseract
                    from PIL import Image
                    text = pytesseract.image_to_string(Image.open(str(p)), config="--psm 11").strip()
                except Exception:
                    pass
            if text is None:
                return {"ok": False, "action": a,
                        "error": "tesseract not found. Install: sudo apt install tesseract-ocr",
                        "content": "OCR engine not found.", "response": "OCR engine not found."}
            msg = f"OCR result from {p.name}:\n\n{text}" if text else f"No text detected in {p.name}."
            return {"ok": True, "action": a, "content": msg, "response": msg, "text": text, "path": str(p)}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- AMBIENT_VISION (toggle periodic screen glances) ----
    if a == "AMBIENT_VISION":
        _av_enabled = args.get("enabled")
        if _av_enabled is None:
            # Infer from the phrasing if not explicit.
            _av_txt = str(args.get("text") or args.get("message") or "").lower()
            _av_enabled = not any(w in _av_txt for w in
                                  ("stop", "off", "disable", "don't", "dont", "quit", "cease"))
        _av_enabled = bool(_av_enabled)
        try:
            from eli.perception.ambient_vision import set_ambient_vision
            st = set_ambient_vision(_av_enabled)
        except Exception as e:
            msg = f"Couldn't change ambient vision: {e}"
            return {"ok": False, "action": a, "error": msg, "content": msg, "response": msg}
        if _av_enabled:
            from eli.perception import vision as _avv
            _ok, _reason = _avv.vision_available()
            if _ok:
                msg = (f"Ambient vision is ON — I'll glance at your screen about every "
                       f"{st.get('interval', 300)//60} min when I'm not mid-reply, and keep a "
                       f"rolling sense of what you're working on.")
            else:
                msg = (f"Ambient vision toggle is ON, but I can't actually see yet: {_reason} "
                       f"It'll start working once the vision model is installed.")
        else:
            msg = "Ambient vision is OFF — I'll only look at the screen when you ask."
        return {"ok": True, "action": a, "enabled": _av_enabled, "content": msg, "response": msg}

    # ---- SCREEN_READ_ANALYZE ----
    if a == "SCREEN_READ_ANALYZE":
        try:
            import tempfile
            # 1. Take screenshot
            from eli.perception.os_controller import take_screenshot
            ss_result = take_screenshot(region="full")
            ss_path = ss_result.get("path") or ss_result.get("file") or ""
            if not ss_path or not Path(ss_path).exists():
                # Try PIL fallback
                try:
                    import subprocess as _sp
                    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
                    tmp.close()
                    _sp.run(["scrot", tmp.name], check=True)
                    ss_path = tmp.name
                except Exception:
                    return {"ok": False, "action": a,
                            "error": "Screenshot failed — no screenshot path returned.",
                            "content": "Screenshot failed.", "response": "Screenshot failed."}
            # 2. Analyse the screenshot with real vision (VL model) + OCR.
            #    ANALYZE_IMAGE does the hot-swap vision call and falls back to
            #    OCR-only honestly when no vision model is installed.
            _sra_prompt = str(args.get("prompt") or args.get("instruction") or "").strip() or (
                "You are ELI looking at the user's screen right now. Describe what is on "
                "screen: the focused application, what the user appears to be doing, and any "
                "important text, code, errors, or UI state. Be specific; never invent."
            )
            _sra = _execute_impl("ANALYZE_IMAGE", {"path": ss_path, "prompt": _sra_prompt, "prefer_fast": True})
            _sra_body = str(_sra.get("content") or _sra.get("response") or "").strip()
            if not _sra_body:
                _sra_body = "I captured the screen but couldn't produce a description."
            return {"ok": bool(_sra.get("ok", True)), "action": a,
                    "content": _sra_body, "response": _sra_body,
                    "vision_text": _sra.get("vision_text", ""),
                    "ocr_text": _sra.get("ocr_text", ""),
                    "screenshot_path": ss_path}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- CONVERT_DOCUMENT ----
    if a == "CONVERT_DOCUMENT":
        try:
            import subprocess as _sp
            source = str(args.get("source") or "").strip()
            fmt = str(args.get("format") or "pdf").strip().lower()
            # Normalise latex aliases
            if fmt in ("latex", "lualatex"):
                fmt = "pdf"  # pandoc --pdf-engine=lualatex
                engine = "lualatex"
            else:
                engine = None
            if not source:
                return {"ok": False, "action": a, "error": "No source file specified",
                        "content": "No source file specified.", "response": "No source file specified."}
            src_path = Path(source).expanduser().resolve()
            if not src_path.exists():
                return {"ok": False, "action": a, "error": f"Source file not found: {src_path}",
                        "content": f"Source not found: {src_path}", "response": f"Source not found: {src_path}"}
            out_path = src_path.with_suffix(f".{fmt if fmt != 'markdown' else 'md'}")
            # Try pandoc first
            if shutil.which("pandoc"):
                cmd = ["pandoc", str(src_path), "-o", str(out_path)]
                if engine:
                    cmd += ["--pdf-engine", engine]
                result = _sp.run(cmd, capture_output=True, text=True, timeout=120)
                if result.returncode == 0:
                    msg = f"Converted {src_path.name} → {out_path.name}"
                    return {"ok": True, "action": a, "content": msg, "response": msg,
                            "output_path": str(out_path)}
                else:
                    err = result.stderr.strip()
                    # Fallback: libreoffice headless
                    if shutil.which("libreoffice"):
                        lo_cmd = ["libreoffice", "--headless", "--convert-to", fmt,
                                  "--outdir", str(src_path.parent), str(src_path)]
                        lo = _sp.run(lo_cmd, capture_output=True, text=True, timeout=180)
                        if lo.returncode == 0:
                            msg = f"Converted {src_path.name} → {out_path.name} (via LibreOffice)"
                            return {"ok": True, "action": a, "content": msg, "response": msg,
                                    "output_path": str(out_path)}
                    return {"ok": False, "action": a, "error": f"pandoc failed: {err}",
                            "content": f"Conversion failed: {err}", "response": f"Conversion failed: {err}"}
            # No pandoc — try libreoffice directly
            elif shutil.which("libreoffice"):
                lo_cmd = ["libreoffice", "--headless", "--convert-to", fmt,
                          "--outdir", str(src_path.parent), str(src_path)]
                lo = _sp.run(lo_cmd, capture_output=True, text=True, timeout=180)
                if lo.returncode == 0:
                    msg = f"Converted {src_path.name} → {out_path.name} (via LibreOffice)"
                    return {"ok": True, "action": a, "content": msg, "response": msg,
                            "output_path": str(out_path)}
                return {"ok": False, "action": a, "error": lo.stderr.strip(),
                        "content": "LibreOffice conversion failed.", "response": "LibreOffice conversion failed."}
            else:
                return {"ok": False, "action": a,
                        "error": "Neither pandoc nor libreoffice found. Install one: apt install pandoc or libreoffice",
                        "content": "No conversion tool available.", "response": "No conversion tool available."}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- TRANSCRIBE ----
    if a == "TRANSCRIBE":
        try:
            source = str(args.get("source") or "").strip()
            if source:
                src_path = Path(source).expanduser().resolve()
                if not src_path.exists():
                    return {"ok": False, "action": a, "error": f"File not found: {src_path}",
                            "content": f"File not found: {src_path}", "response": f"File not found: {src_path}"}
                from faster_whisper import WhisperModel
                model = WhisperModel("small", device="cpu", compute_type="int8")
                segments, _ = model.transcribe(str(src_path), vad_filter=True)
                text = " ".join(seg.text.strip() for seg in segments).strip()
                msg = f"Transcription of {src_path.name}:\n\n{text}" if text else "No speech detected."
                return {"ok": True, "action": a, "content": msg, "response": msg, "text": text}
            else:
                # Live microphone transcription — single utterance
                import numpy as np, sounddevice as sd, soundfile as sf, tempfile
                from faster_whisper import WhisperModel
                sr = 16000
                chunk = 0.05
                start_rms = 0.02
                end_silence = 1.2
                max_s = 30
                block = int(sr * chunk)
                frames, silence, started = [], 0.0, False
                import time as _t
                t0 = _t.time()
                with sd.InputStream(samplerate=sr, channels=1, blocksize=block, dtype="float32") as s:
                    while True:
                        data, _ = s.read(block)
                        x = data[:, 0]
                        rms = float(np.sqrt(np.mean(x * x)) + 1e-12)
                        if not started and rms >= start_rms:
                            started = True
                        if started:
                            frames.append(x.copy())
                            silence = silence + chunk if rms < start_rms else 0.0
                        if started and silence >= end_silence:
                            break
                        if _t.time() - t0 > max_s:
                            break
                if not frames:
                    return {"ok": False, "action": a, "error": "No audio captured",
                            "content": "No audio captured.", "response": "No audio captured."}
                audio = np.concatenate(frames)
                with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as f:
                    sf.write(f.name, audio, sr)
                    tmp_path = f.name
                model = WhisperModel("small", device="cpu", compute_type="int8")
                segments, _ = model.transcribe(tmp_path, vad_filter=True)
                text = " ".join(seg.text.strip() for seg in segments).strip()
                msg = f"Transcribed: {text}" if text else "No speech detected."
                return {"ok": True, "action": a, "content": msg, "response": msg, "text": text}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- DICTATE ----
    if a == "DICTATE":
        try:
            action = str(args.get("action") or "start").lower()
            if action == "start":
                # Signal the GUI to enable dictation mode; GUI handles the loop
                msg = "Dictation mode activated. Speak clearly — I'm listening. Say 'stop dictation' to end."
                return {"ok": True, "action": a, "content": msg, "response": msg,
                        "dictate_start": True}
            elif action == "stop":
                msg = "Dictation mode stopped."
                return {"ok": True, "action": a, "content": msg, "response": msg,
                        "dictate_stop": True}
            return {"ok": False, "action": a, "error": f"Unknown dictate action: {action}",
                    "content": f"Unsupported executor action: {action}", "response": f"Unsupported executor action: {action}"}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- MOUSE_CONTROL ----
    if a == "MOUSE_CONTROL":
        try:
            import subprocess as _sp
            mouse_action = str(args.get("action") or "move").lower()
            x = args.get("x")
            y = args.get("y")
            button = str(args.get("button") or "left").lower()
            double = bool(args.get("double", False))
            direction = str(args.get("direction") or "down").lower()
            amount = int(args.get("amount", 3))
            # Try pyautogui
            try:
                import pyautogui
                if mouse_action == "move" and x is not None and y is not None:
                    pyautogui.moveTo(int(x), int(y), duration=0.2)
                    msg = f"Mouse moved to ({x}, {y})"
                elif mouse_action == "click" and x is not None and y is not None:
                    pyautogui.moveTo(int(x), int(y), duration=0.1)
                    if double:
                        pyautogui.doubleClick(button=button)
                    else:
                        pyautogui.click(button=button)
                    msg = f"{'Double-c' if double else 'C'}licked {button} at ({x}, {y})"
                elif mouse_action == "scroll":
                    delta = amount if direction == "up" else -amount
                    pyautogui.scroll(delta)
                    msg = f"Scrolled {direction} by {amount}"
                else:
                    msg = f"Mouse action '{mouse_action}' performed"
                return {"ok": True, "action": a, "content": msg, "response": msg}
            except ImportError:
                pass
            # Fallback: ydotool (Wayland) or xdotool (X11)
            tool = "ydotool" if shutil.which("ydotool") else ("xdotool" if shutil.which("xdotool") else None)
            if tool is None:
                return {"ok": False, "action": a,
                        "error": "No mouse control tool found. Install: pip install pyautogui  OR  apt install ydotool",
                        "content": "No mouse tool available.", "response": "No mouse tool available."}
            if mouse_action == "move" and x is not None:
                cmd = [tool, "mousemove", "--", str(x), str(y)] if tool == "ydotool" else [tool, "mousemove", str(x), str(y)]
            elif mouse_action == "click":
                btn_flag = "0x40002" if button == "right" else "0x40001"
                cmd = [tool, "click", btn_flag] if tool == "ydotool" else [tool, "click", "--clearmodifiers", "--button", "3" if button == "right" else "1"]
            elif mouse_action == "scroll":
                if tool == "xdotool":
                    btn = "4" if direction == "up" else "5"
                    cmd = [tool, "click", "--clearmodifiers", "--repeat", str(amount), "--repeat-delay", "50", btn]
                else:
                    cmd = [tool, "scroll", "--", "0", str(-amount if direction == "down" else amount)]
            else:
                cmd = []
            if cmd:
                _sp.run(cmd, check=True)
            msg = f"Mouse {mouse_action} executed"
            return {"ok": True, "action": a, "content": msg, "response": msg}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- SET_CLIPBOARD / GET_CLIPBOARD ----
    if a == "SET_CLIPBOARD":
        try:
            from eli.perception.os_controller import set_clipboard
            txt = str(args.get("text") or "").strip()
            return set_clipboard(txt)
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    if a == "GET_CLIPBOARD":
        try:
            from eli.perception.os_controller import get_clipboard
            txt = get_clipboard() or ""
            msg = f"Clipboard: {txt[:200]}" if txt else "Clipboard is empty"
            return {"ok": True, "action": a, "content": msg, "response": msg, "text": txt}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- LIST_DIR ----
    if a == "LIST_DIR":
        path = str((args or {}).get("path") or "~")
        try:
            repo_root = Path(__file__).resolve().parents[2]  # project root
            package_root = repo_root / "eli"
            legacy_root = Path.home() / "eli"
            raw = Path(path).expanduser()

            package_toplevel = {
                "api", "brain", "controllers", "core", "gui",
                "integrations", "plugins", "tools", "utils"
            }

            if raw == legacy_root:
                p = package_root.resolve()
            else:
                try:
                    rel = raw.relative_to(legacy_root)
                    if not rel.parts:
                        p = package_root.resolve()
                    elif rel.parts[0] in package_toplevel:
                        p = (package_root / rel).resolve()
                    else:
                        p = (repo_root / rel).resolve()
                except Exception:
                    p = raw.resolve()

            if not p.exists():
                return {"ok": False, "action": a, "error": f"Path not found: {p}", "content": f"Path not found: {p}", "response": f"Path not found: {p}"}
            if not p.is_dir():
                return {"ok": False, "action": a, "error": f"Not a directory: {p}", "content": f"Not a directory: {p}", "response": f"Not a directory: {p}"}

            items = sorted([
                f"{x.name}/" if x.is_dir() else x.name
                for x in p.iterdir()
            ])

            content = f"Contents of {p} ({len(items)} items):\n" + "\n".join(items[:500])
            return {"ok": True, "action": a, "path": str(p), "content": content, "response": content}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- SHELL_EXEC (alias for RUN_CMD) ----
    if a == "SHELL_EXEC":
        cmd = args.get("cmd") or args.get("command") or ""
        return _execute_impl("RUN_CMD", {"cmd": cmd, "shell": True})

    # ---- PLUGIN_STATUS ----
    if a == "PLUGIN_STATUS":
        try:
            from eli.plugins.manager import get_manager
            mgr = get_manager()
            installed = mgr.list_installed() or []
            available = mgr.list_available() or []
            lines = [f"Plugin system status:"]
            lines.append(f"  Installed: {len(installed)}  |  Available in registry: {len(available)}")
            if installed:
                lines.append("")
                lines.append("Installed plugins:")
                for p in installed:
                    status = "enabled" if p.get("enabled") else "disabled"
                    lines.append(f"  {p['id']:20s} [{status}]  {p.get('description','')[:60]}")
            else:
                lines.append("  No plugins installed.")
            # Check each installed plugin imports cleanly
            lines.append("")
            lines.append("Plugin health:")
            for p in installed:
                try:
                    mod_path = f"plugins.{p['id']}.plugin"
                    __import__(mod_path)
                    lines.append(f"  {p['id']:20s} OK")
                except Exception as pe:
                    lines.append(f"  {p['id']:20s} ERROR: {pe}")
            msg = "\n".join(lines)
            return {"ok": True, "action": a, "content": msg, "response": msg,
                    "installed": len(installed), "available": len(available)}
        except Exception as e:
            msg = f"Plugin status error: {e}"
            return {"ok": False, "action": a, "error": str(e), "content": msg, "response": msg}

    # ---- PLUGIN_INSTALL ----
    if a == "PLUGIN_INSTALL":
        try:
            from eli.plugins.manager import get_manager
            mgr = get_manager()
            query = (args.get("query") or args.get("plugin") or args.get("name") or "").strip()
            if not query:
                return {"ok": False, "action": a, "error": "No plugin specified",
                        "content": json.dumps({"missing_arg": "plugin", "action": a}, ensure_ascii=False),
                        "response": json.dumps({"missing_arg": "plugin", "action": a}, ensure_ascii=False)}

            # Try exact match first
            entry = mgr.get_registry_entry(query.lower().replace(" ", "_"))
            if entry:
                result = mgr.install(entry["id"])
                return {"ok": result.get("ok", False), "action": a,
                        "content": result.get("content", str(result)),
                        "response": result.get("response", str(result))}

            # Semantic search
            matches = mgr.search(query)
            if not matches:
                msg = f"No plugins found matching '{query}'. Use 'list plugins' to see what's available."
                return {"ok": False, "action": a, "content": msg, "response": msg}

            if len(matches) == 1:
                result = mgr.install(matches[0]["id"])
                return {"ok": result.get("ok", False), "action": a,
                        "content": result.get("content", str(result)),
                        "response": result.get("response", str(result))}

            # Multiple matches — show options
            lines = [f"Found {len(matches)} plugins matching '{query}':"]
            for m in matches[:5]:
                lines.append(f"  • {m['id']} — {m.get('description', '')}")
            lines.append("\nSpecify the exact plugin id to install, e.g. 'install plugin pomodoro'.")
            msg = "\n".join(lines)
            return {"ok": True, "action": a, "content": msg, "response": msg}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- PLUGIN_UNINSTALL ----
    if a == "PLUGIN_UNINSTALL":
        try:
            from eli.plugins.manager import get_manager
            pid = (args.get("plugin") or args.get("name") or args.get("id") or "").strip().lower().replace(" ", "_")
            if not pid:
                msg = json.dumps({"missing_arg": "plugin", "action": a}, ensure_ascii=False)
                return {"ok": False, "action": a, "error": "No plugin specified",
                        "content": msg,
                        "response": msg}
            result = get_manager().uninstall(pid)
            return {"ok": result.get("ok", False), "action": a,
                    "content": result.get("content", str(result)),
                    "response": result.get("response", str(result))}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- PLUGIN_LIST ----
    if a == "PLUGIN_LIST":
        try:
            from eli.plugins.manager import get_manager
            mgr = get_manager()
            scope = (args.get("scope") or "all").lower()

            if scope == "installed":
                plugins = mgr.list_installed()
                if not plugins:
                    msg = "No plugins installed."
                else:
                    lines = [f"Installed plugins ({len(plugins)}):"]
                    for p in plugins:
                        status = "✓ enabled" if p["enabled"] else "⏸ disabled"
                        lines.append(f"  • {p['id']} ({status}) — {p.get('description', '')}")
                    msg = "\n".join(lines)
            else:
                plugins = mgr.list_available()
                installed_ids = {p["id"] for p in mgr.list_installed()}
                lines = [f"Available plugins ({len(plugins)}):"]
                for p in plugins:
                    tag = " [installed]" if p["id"] in installed_ids else ""
                    lines.append(f"  • {p['id']} — {p.get('description', '')}{tag}")
                msg = "\n".join(lines)

            return {"ok": True, "action": a, "count": len(plugins),
                    "content": msg, "response": msg}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- PLUGIN_SEARCH ----
    if a == "PLUGIN_SEARCH":
        try:
            from eli.plugins.manager import get_manager
            query = (args.get("query") or args.get("text") or "").strip()
            if not query:
                return {"ok": False, "action": a, "error": "No search query",
                        "content": "What kind of plugin are you looking for?",
                        "response": "What kind of plugin are you looking for?"}
            matches = get_manager().search(query)
            if not matches:
                msg = f"No plugins found matching '{query}'."
            else:
                lines = [f"Plugins matching '{query}' ({len(matches)}):"]
                for m in matches[:8]:
                    actions = ", ".join(m.get("actions", [])[:3])
                    lines.append(f"  • {m['id']} — {m.get('description', '')} [{actions}]")
                msg = "\n".join(lines)
            return {"ok": True, "action": a, "count": len(matches),
                    "content": msg, "response": msg}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- PLUGIN_ENABLE ----
    if a == "PLUGIN_ENABLE":
        try:
            from eli.plugins.manager import get_manager
            pid = (args.get("plugin") or args.get("name") or args.get("id") or "").strip().lower().replace(" ", "_")
            if not pid:
                msg = json.dumps({"missing_arg": "plugin", "action": a}, ensure_ascii=False)
                return {"ok": False, "action": a, "content": msg, "response": msg}
            result = get_manager().enable(pid)
            return {"ok": result.get("ok", False), "action": a,
                    "content": result.get("content", str(result)),
                    "response": result.get("response", str(result))}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- PLUGIN_DISABLE ----
    if a == "PLUGIN_DISABLE":
        try:
            from eli.plugins.manager import get_manager
            pid = (args.get("plugin") or args.get("name") or args.get("id") or "").strip().lower().replace(" ", "_")
            if not pid:
                msg = json.dumps({"missing_arg": "plugin", "action": a}, ensure_ascii=False)
                return {"ok": False, "action": a, "content": msg, "response": msg}
            result = get_manager().disable(pid)
            return {"ok": result.get("ok", False), "action": a,
                    "content": result.get("content", str(result)),
                    "response": result.get("response", str(result))}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- HARDWARE_PROFILE ----
    if a == "HARDWARE_PROFILE":
        try:
            from eli.core.hardware_profile import run_benchmark, apply_recommendation, recommend
            if args.get("apply"):
                rec = recommend()
                result = apply_recommendation(rec)
                msg = f"Applied: {rec.model_name}, {rec.n_gpu_layers} GPU layers, ctx={rec.n_ctx}, batch={rec.batch_size}, threads={rec.n_threads}.\nRestart Eli to use new settings."
                return {"ok": True, "action": a, "content": msg, "response": msg}
            bench = run_benchmark()
            hw = bench["hardware"]
            rec = bench["recommendation"]
            lines = [
                f"Hardware: {hw['cpu_threads']} threads, {hw['ram_gb']}GB RAM, "
                f"{hw['gpu_name'] or 'no GPU'} ({hw['vram_gb']}GB VRAM)",
                "",
                "Recommendation:",
            ]
            for line in rec.get("reasoning", []):
                lines.append(f"  {line}")
            lines.append(f"\n  → {rec['model_name']}, {rec['n_gpu_layers']} GPU layers, ctx={rec['n_ctx']}, batch={rec['batch_size']}")
            lines.append("\nSay 'apply hardware recommendation' to update settings.")
            msg = "\n".join(lines)
            return {"ok": True, "action": a, "content": msg, "response": msg,
                    "hardware": hw, "recommendation": rec}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- AWARENESS_STATUS ----
    if a == "AWARENESS_STATUS":
        try:
            from eli.runtime.awareness_boot import get_awareness, boot_awareness
            state = get_awareness()
            if state is None:
                state = boot_awareness(quiet=True)
            msg = state.full_briefing()
            return {"ok": True, "action": a, "content": msg, "response": msg,
                    "capability_count": state.capability_count}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- FRONTIER_STATUS ----
    if a == "FRONTIER_STATUS":
        try:
            from eli.runtime.frontier_status import (
                build_frontier_status_report,
                format_frontier_status_report,
            )

            report = build_frontier_status_report(str(args.get("question") or args.get("query") or ""))
            text = format_frontier_status_report(report)
            return {
                "ok": bool(report.get("ok", True)),
                "action": a,
                "report": report,
                "content": text,
                "response": text,
                "evidence_source": "frontier_status_local_runtime_matrix_v1",
            }
        except Exception as e:
            msg = f"Frontier status failed: {type(e).__name__}: {e}"
            return {"ok": False, "action": a, "error": str(e), "content": msg, "response": msg}

    # ---- ELI_IDENTITY_AUDIT ----
    if a == "ELI_IDENTITY_AUDIT":
        try:
            from eli.runtime.eli_identity_audit import (
                build_eli_identity_audit,
                format_eli_identity_audit,
            )

            report = build_eli_identity_audit(str(args.get("question") or args.get("query") or ""))
            text = format_eli_identity_audit(report)
            return {
                "ok": bool(report.get("ok", True)),
                "action": a,
                "report": report,
                "content": text,
                "response": text,
                "evidence_source": "eli_identity_audit_local_verified_matrix_v1",
            }
        except Exception as e:
            msg = f"ELI identity audit failed: {type(e).__name__}: {e}"
            return {"ok": False, "action": a, "error": str(e), "content": msg, "response": msg}

    # ---- CODE_CHANGES ----
    if a == "CODE_CHANGES":
        try:
            from eli.runtime.awareness_boot import get_awareness, boot_awareness
            state = get_awareness()
            if state is None:
                state = boot_awareness(quiet=True)
            if state.code_report_has_changes:
                msg = state.code_report_briefing
            else:
                msg = "No code changes detected since last check."
            return {"ok": True, "action": a, "content": msg, "response": msg}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- SELF_ANALYZE ----
    if a == "SELF_ANALYZE":
        try:
            from eli.runtime.self_improvement import get_self_improvement
            engine = get_self_improvement()
            failures = engine.analyze_failures(limit=10, days=7, min_cluster_size=1)
            lines = [f"Self-Analysis Report ({len(failures)} recent issues):"]

            def _short(value, limit=220):
                text = " ".join(str(value or "").split())
                return text if len(text) <= limit else text[: max(0, limit - 3)] + "..."

            def _root_cause_for(failure):
                ui = str(failure.get("user_input") or "")
                err = str(failure.get("error") or "")
                low = f"{ui}\n{err}".lower()
                if "path not found" in low and "artifacts/scripts" in low:
                    return (
                        "Path resolution selected a project-root/basename target before the "
                        "artifact path. The requested file existed under artifacts/scripts, "
                        "so the first repair request hit the wrong location."
                    )
                if "path not found" in low and "create_a_python_script_to_generate" in low:
                    return (
                        "The repair request targeted the project root copy of the generated "
                        "script, but generated scripts are stored under artifacts/scripts. "
                        "The executor needed artifact-aware path recovery before reporting failure."
                    )
                if "path not found" in low:
                    return (
                        "File operation failed before repair because the resolved target path "
                        "did not exist. The executor needs artifact-aware path recovery before "
                        "declaring a file missing."
                    )
                if "cannot create children for a parent that is in a different thread" in low:
                    return (
                        "Qt widget/document state was touched from a worker thread. GUI updates "
                        "must be marshalled back to the main Qt thread via signals."
                    )
                if "segmentation fault" in low and "qtextdocument" in low:
                    return (
                        "The crash signature is consistent with the Qt thread-affinity violation "
                        "above, not a normal Python exception."
                    )
                if "unsupported executor action" in low:
                    return (
                        "The router produced an action that the executor did not implement, "
                        "so the engine fell back into chat/synthesis instead of returning grounded output."
                    )
                return "No single deterministic root cause could be inferred from the stored failure text."

            for f in failures[:10]:
                lines.append(
                    f"- input: {_short(f.get('user_input'), 120)}\n"
                    f"  error: {_short(f.get('error'), 220)}\n"
                    f"  occurrences: {int(f.get('occurrence_count') or 1)}\n"
                    f"  actual_root_cause: {_root_cause_for(f)}"
                )
            if not failures:
                lines.append("  No recent failures found.")
            if (args or {}).get("suggest"):
                result = engine.analyze_and_improve()
                imps = result.get("improvements", [])
                if imps:
                    lines.append(f"\nSuggested improvements ({len(imps)}):")
                    for imp in imps[:5]:
                        lines.append(f"  - {imp.get('description', '')}")
            # --- Awareness integration ---
            try:
                from eli.runtime.awareness_boot import get_awareness
                awareness = get_awareness()
                if awareness:
                    lines.append("")
                    lines.append(awareness.context_block())
            except Exception:
                pass

            msg = "\n".join(lines)
            # Fire repair_completed so the world engine decreases repair_pressure.
            # Without this, repair_pressure accumulates permanently and triggers
            # an infinite proactive SELF_ANALYZE loop.
            try:
                from eli.world.world_event_bus import fire_world_event as _fwe_sa
                _fwe_sa(
                    "repair_completed",
                    "self_analyze",
                    f"SELF_ANALYZE completed: reviewed {len(failures)} failure(s).",
                    {"failures_reviewed": len(failures)},
                )
            except Exception:
                pass
            return {"ok": True, "action": a, "content": msg, "response": msg}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- FILE_AUDIT ----
    if a == "FILE_AUDIT":
        try:
            import os as _os
            from pathlib import Path as _Path

            # Determine scan root: default to the ELI package directory
            _project_root = _Path(__file__).resolve().parents[2]
            _eli_root = _project_root / "eli"
            _scope = (args or {}).get("scope", "eli")

            if not _eli_root.exists():
                _eli_root = _project_root

            # Walk the directory tree and collect stats
            _dir_stats: dict = {}
            _total_files = 0
            _total_py = 0
            _skipped = {"__pycache__", ".git", ".venv", "venv", "node_modules"}

            for _dirpath, _dirnames, _filenames in _os.walk(str(_eli_root)):
                # Prune skip dirs in-place
                _dirnames[:] = [d for d in _dirnames if d not in _skipped]
                _rel = str(_Path(_dirpath).relative_to(_project_root))
                _py_files = [f for f in _filenames if f.endswith(".py")]
                _non_py = [f for f in _filenames if not f.endswith(".py") and not f.endswith(".pyc")]
                if _filenames:
                    _dir_stats[_rel] = {
                        "py": len(_py_files),
                        "other": len(_non_py),
                        "total": len(_filenames),
                    }
                    _total_files += len(_filenames)
                    _total_py += len(_py_files)

            # Build human-readable report
            _lines = [
                f"File Audit: {_eli_root}",
                f"Directories scanned: {len(_dir_stats)}",
                f"Total files: {_total_files}  (Python: {_total_py}  Other: {_total_files - _total_py})",
                "",
                "Directory breakdown (non-empty dirs, sorted):",
            ]
            for _rel_dir in sorted(_dir_stats.keys()):
                _ds = _dir_stats[_rel_dir]
                _lines.append(
                    f"  {_rel_dir}/  "
                    f"[{_ds['total']} files: {_ds['py']} .py, {_ds['other']} other]"
                )

            # Also list top-level .py module names in eli/ directly
            _top_py = sorted(
                f.name for f in _eli_root.iterdir()
                if f.is_file() and f.suffix == ".py"
            )
            if _top_py:
                _lines.append("")
                _lines.append(f"eli/ root-level modules: {', '.join(_top_py)}")

            _msg = "\n".join(_lines)
            # Fire task_completed to keep world state clean
            try:
                from eli.world.world_event_bus import fire_world_event as _fwe_fa
                _fwe_fa(
                    "task_completed",
                    "file_audit",
                    f"FILE_AUDIT completed: {_total_files} files in {len(_dir_stats)} dirs.",
                    {"total_files": _total_files, "total_py": _total_py},
                )
            except Exception:
                pass
            return {"ok": True, "action": a, "content": _msg, "response": _msg}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- SELF_IMPROVE ----
    if a == "SELF_IMPROVE":
        mode = (args or {}).get("mode", "analyze")
        dry_run = (args or {}).get("dry_run", False)
        try:
            from eli.runtime.self_improvement import get_self_improvement
            engine = get_self_improvement()
            if mode == "patch":
                result = engine.run_patch_cycle(max_patches=3, dry_run=bool(dry_run))
                details = result.get("details", [])
                changed = [
                    d for d in details
                    if str(d.get("status", "")).lower() in {"applied", "patched", "changed"}
                ]
                msg_lines = [
                    "Patch improvement cycle complete.",
                    f"- code_changes_made: {len(changed)}",
                    f"- failures_analyzed: {result.get('failures_analyzed', 0)}",
                    f"- dry_run: {bool(dry_run)}",
                    f"- summary: {result.get('summary', 'Patch cycle complete.')}",
                ]
                if details:
                    msg_lines.append("- patch_details:")
                    for d in details[:5]:
                        msg_lines.append(
                            f"  [{d.get('status','?')}] "
                            f"{str(d.get('failure',''))[:90]}"
                        )
                msg = "\n".join(msg_lines)
            else:
                failures = engine.analyze_failures(limit=10, days=14, min_cluster_size=1)
                result = engine.analyze_and_improve()
                imps = result.get("improvements", [])
                msg_lines = [
                    "Improvement cycle complete.",
                    "- code_changes_made: 0",
                    f"- failures_inspected: {len(failures)}",
                    f"- new_improvement_records: {len(imps)}",
                    "- patch_cycle_run: false",
                ]
                if failures:
                    last = failures[0]
                    err = " ".join(str(last.get("error") or "").split())
                    ui = " ".join(str(last.get("user_input") or "").split())
                    msg_lines.append(f"- last_failure_error: {err[:220] or '-'}")
                    msg_lines.append(f"- last_failure_input: {ui[:160] or '-'}")
                if imps:
                    msg_lines.append("- logged_improvements:")
                    msg_lines.extend(f"  - {i.get('description', '')}" for i in imps[:5])
                msg_lines.append("To modify code, run/apply self-improvement patches.")
                msg = "\n".join(msg_lines)
            _si_result = {"ok": True, "action": a, "content": msg, "response": msg}
            try:
                from eli.world.world_event_bus import fire_improvement_event as _wfie
                _wfie(proposal_count=len(imps), failure_count=len(failures))
            except Exception:
                pass
            return _si_result
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- SELF_PATCH ----
    if a == "SELF_PATCH":
        dry_run = (args or {}).get("dry_run", False)
        max_patches = int((args or {}).get("max_patches", 3))
        try:
            from eli.runtime.self_improvement import get_self_improvement
            result = get_self_improvement().run_patch_cycle(max_patches=max_patches, dry_run=bool(dry_run))
            msg = result.get("summary", "Patch cycle complete.")
            details = result.get("details", [])
            if details:
                msg += "\n" + "\n".join(
                    f"  [{d.get('status','?')}] {d.get('failure','')[:70]}"
                    for d in details[:5]
                )
            return {"ok": True, "action": a, "content": msg, "response": msg}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- MORNING_REPORT ----
    if a == "MORNING_REPORT":
        try:
            from eli.runtime.reflection import run_reflection
            from datetime import datetime as _dt
            result = run_reflection(hours=24)
            insights = result.get("insights", [])
            _stamp = _dt.now().strftime("%A %d %B %Y, %H:%M")
            if insights:
                msg = (f"Morning report — {_stamp} (activity over the last 24h):\n"
                       + "\n".join(f"  • {i}" for i in insights))
            else:
                msg = f"Morning report — {_stamp}: no notable activity in the last 24h."
            # Personalised news (ELI-derived interests; network-gated → empty when
            # offline, never fabricated).
            try:
                from eli.tools.news.news_synthesis import interest_news_block
                _news = interest_news_block()
                if _news:
                    msg += "\n\n" + _news
            except Exception:
                pass
            return {"ok": True, "action": a, "content": msg, "response": msg}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- GET_WEATHER ----

    _LOC_PATS = [
        re.compile(r"weather (?:in|at|for) ([A-Za-z][A-Za-z ,]+?)(?:\?|$|[,;])", re.I),
        re.compile(r"(?:in|at|for) ([A-Za-z][A-Za-z ,]+?) weather", re.I),
        re.compile(r"(?:temperature|forecast) (?:in|at) ([A-Za-z][A-Za-z ,]+?)(?:\?|$)", re.I),
    ]

    if a == "GET_WEATHER":
        try:
            location = str((args or {}).get("location") or (args or {}).get("city") or "").strip()

            if not location:
                _raw = str((args or {}).get("_raw_user_text") or (args or {}).get("query") or "").strip()
                for _p in _LOC_PATS:
                    _m = _p.search(_raw)
                    if _m:
                        location = _m.group(1).strip().rstrip("?.!, ")
                        break

            if not location:
                _msg = "I need a location. Try: 'What's the weather in Wexford?'"
                return {"ok": False, "action": a, "error": "missing_location", "content": _msg, "response": _msg}

            # Network gate: weather is a live open-meteo call. Refuse honestly
            # when the Net toggle is off rather than confabulating a reading.
            from eli.core.netguard import should_block_network, offline_response, OfflineError
            if should_block_network():
                return offline_response(a, "check the weather")

            from eli.plugins.weather.plugin import get_weather as _gw
            try:
                result = _gw(location)
            except OfflineError:
                return offline_response(a, "check the weather")
            if isinstance(result, dict):
                result.setdefault("action", a)
                return result
            return {"ok": True, "action": a, "content": str(result), "response": str(result)}

        except Exception as _e:
            return {"ok": False, "action": a, "error": str(_e), "content": str(_e), "response": str(_e)}


    # ---- OPEN_SYSTEM_SETTINGS ----
    if a == "OPEN_SYSTEM_SETTINGS":
        try:
            candidates = [
                ["gnome-control-center"],
                ["systemsettings5"],
                ["xfce4-settings-manager"],
                ["mate-control-center"],
                ["cinnamon-settings"],
            ]
            for cmd in candidates:
                try:
                    subprocess.Popen(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                    msg = f"Opened system settings: {' '.join(cmd)}"
                    return {"ok": True, "action": a, "content": msg, "response": msg}
                except FileNotFoundError:
                    continue
            raise FileNotFoundError("No settings app found")
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- OPEN_POWER_SETTINGS ----
    if a == "OPEN_POWER_SETTINGS":
        try:
            candidates = [
                ["gnome-control-center", "power"],
                ["systemsettings5", "kcm_energyinfo"],
                ["xfce4-power-manager-settings"],
            ]
            for cmd in candidates:
                try:
                    subprocess.Popen(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                    msg = f"Opened power settings: {' '.join(cmd)}"
                    return {"ok": True, "action": a, "content": msg, "response": msg}
                except FileNotFoundError:
                    continue
            raise FileNotFoundError("No power settings app found")
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- SET_ALARM / SET_TIMER ----
    if a in ("SET_ALARM", "SET_TIMER"):
        try:
            duration = args.get("duration") or args.get("seconds")
            alarm_time = args.get("time") or args.get("alarm_time")
            label = args.get("label") or "ELI Timer"

            if duration:
                secs = int(duration)
                # Background thread timer with notify-send
                def _timer_fire():
                    time.sleep(secs)
                    try:
                        notify("ELI Timer", f"Timer complete! ({secs}s) - {label}")
                    except Exception:
                        pass
                    try:
                        play_sound("/usr/share/sounds/freedesktop/stereo/alarm-clock-elapsed.oga")
                    except Exception:
                        pass
                t = threading.Thread(target=_timer_fire, daemon=True)
                t.start()
                msg = f"Timer set for {secs} seconds."
                return {"ok": True, "action": a, "content": msg, "response": msg}

            elif alarm_time:
                # Parse HH:MM and compute seconds until that time
                from datetime import datetime as _dt
                now = _dt.now()
                try:
                    parts = str(alarm_time).replace(".", ":").split(":")
                    target_h, target_m = int(parts[0]), int(parts[1]) if len(parts) > 1 else 0
                    target = now.replace(hour=target_h, minute=target_m, second=0, microsecond=0)
                    if target <= now:
                        target = target.replace(day=target.day + 1)
                    secs = int((target - now).total_seconds())
                    def _alarm_fire():
                        time.sleep(secs)
                        try:
                            notify("ELI Alarm", f"Alarm! It's {alarm_time} - {label}")
                        except Exception:
                            pass
                        try:
                            play_sound("/usr/share/sounds/freedesktop/stereo/alarm-clock-elapsed.oga")
                        except Exception:
                            pass
                    t = threading.Thread(target=_alarm_fire, daemon=True)
                    t.start()
                    msg = f"Alarm set for {alarm_time} ({secs}s from now)."
                    return {"ok": True, "action": a, "content": msg, "response": msg}
                except Exception as e:
                    return {"ok": False, "action": a, "error": f"Bad time format: {e}", "content": f"Bad time format: {alarm_time}", "response": f"Bad time format: {alarm_time}"}
            else:
                return {"ok": False, "action": a, "error": "No duration or time specified", "content": "No duration or time specified", "response": "No duration or time specified"}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- LIST_EVENTS / ADD_EVENT (calendar integration not configured) ----
    if a == "LIST_EVENTS":
        msg = "Calendar integration is not configured. Use your calendar app directly or wire a provider first."
        return {"ok": True, "action": a, "content": msg, "response": msg}

    if a == "ADD_EVENT":
        msg = "Calendar integration is not configured. Use your calendar app directly or wire a provider first."
        return {"ok": True, "action": a, "content": msg, "response": msg}

    # ---- ANALYZE_PDF ----
    if a == "ANALYZE_PDF":
        try:
            import re as _re2
            import subprocess as _sp
            from pathlib import Path as _PP
            from eli.perception.analyze_pdfs import analyze

            path = str(args.get("path") or "").strip()
            if not path:
                return {"ok": False, "action": a, "error": "Missing path",
                        "content": "Missing path", "response": "Missing path"}

            _instruction = str(args.get("instruction") or "").strip()
            _instr_low = _instruction.lower()

            analysis = analyze(path)
            doc = analysis.doc
            full_text = "\n\n".join(analysis.chunks) if analysis.chunks else analysis.preview

            # --- Task classification ---
            if any(w in _instr_low for w in ["falsif", "disprove", "refut", "critique", "challenge"]):
                _task_label = "Critical Analysis & Falsification"
                _directive = (
                    "You are a rigorous scientific peer reviewer with expertise across physics, "
                    "mathematics, and engineering. Your task is to CRITICALLY ANALYSE and attempt "
                    "to FALSIFY this document.\n\n"
                    "For EVERY claim, equation, model, or prediction you encounter:\n"
                    "1. State the claim precisely and verbatim where possible.\n"
                    "2. Identify all testable predictions and boundary conditions.\n"
                    "3. Propose specific falsification criteria, experiments, or observational tests.\n"
                    "4. Check dimensional consistency of every equation.\n"
                    "5. Flag hidden assumptions, missing derivation steps, or scope overreach.\n"
                    "6. Cross-reference against established physics (GR, QFT, thermodynamics, etc.).\n"
                    "7. Verdict per claim: TESTABLE | UNFALSIFIABLE | ALREADY FALSIFIED | VALID.\n\n"
                    "Conclude with: (a) overall falsifiability assessment, (b) strongest predictions, "
                    "(c) fatal flaws if any, (d) recommended next steps for validation.\n"
                    "Be exhaustive. Be brutal. Do not truncate."
                )
                _chunk_task = "critically analyse and attempt to falsify this section"
                _synth_task = "synthesise a complete falsification report from these section analyses"
            elif any(w in _instr_low for w in ["extract equation", "equation", "math"]):
                _task_label = "Equation Extraction & Analysis"
                _directive = (
                    "You are ELI, a precise technical analyst. Extract and analyse EVERY equation in the document. "
                    "For each: state the equation, define all symbols with units, check dimensional "
                    "consistency, identify assumptions, note derivation gaps, and flag any errors."
                )
                _chunk_task = "extract and analyse all equations in this section"
                _synth_task = "compile a complete equation analysis from these section results"
            elif any(w in _instr_low for w in ["summari", "summarize", "summarise", "overview"]):
                _task_label = "Comprehensive Summary"
                _directive = (
                    "You are a precise scientific analyst. Produce a COMPREHENSIVE summary covering: "
                    "abstract, objectives, theoretical framework, key equations, methods, results, "
                    "conclusions, open questions, and implications. Preserve full technical precision. "
                    "Do not truncate or abbreviate — complete every section fully."
                )
                _chunk_task = "summarise this section comprehensively"
                _synth_task = "synthesise a complete summary from these section summaries"
            else:
                _task_label = "Technical Analysis"
                _directive = (
                    "You are ELI, a rigorous and unsparing technical analyst. "
                    "Perform a THOROUGH, COMPLETE analysis as requested. "
                    "Be precise, rigorous, and unsparing. Do not truncate."
                )
                _user_req = _instruction or "Provide a thorough technical analysis."
                _chunk_task = f"perform this analysis on the section: {_user_req}"
                _synth_task = f"synthesise a complete document fulfilling this request: {_user_req}"

            from eli.cognition import gguf_inference as _pgi
            _llm_ready = _pgi.load_model() is not None

            def _run_gguf(user_msg: str, sys_msg: str) -> str:
                """Run GGUF with no artificial token cap — use all available context."""
                return _pgi.chat_completion(
                    user_msg,
                    system=sys_msg,
                    max_tokens=None,   # auto: n_ctx - prompt_tokens - 128
                    temperature=0.2,
                )

            _section_outputs = []

            if _llm_ready:
                # Determine how many chars fit per chunk given the directive overhead
                # ~4 chars per token, directive ~400 tokens → reserve 1600 chars for directive
                # Use available_generation_tokens to know how much output we can get
                _avail = _pgi.available_generation_tokens(
                    _chunk_task + "\n\n" + full_text[:100],
                    system=_directive
                )
                # Each chunk: leave room for output generation (roughly n_ctx/2 for input)
                try:
                    _n_ctx = _pgi.load_model().n_ctx()
                except Exception:
                    _n_ctx = 16384
                # chars per chunk = roughly half the context minus directive, times 4 chars/token
                _chars_per_chunk = max(3000, (_n_ctx // 2 - 600) * 4)

                # Split full text into overlapping chunks
                _chunks_to_process = []
                _stride = max(1000, _chars_per_chunk - 500)  # 500-char overlap
                _pos = 0
                while _pos < len(full_text):
                    _chunks_to_process.append(full_text[_pos: _pos + _chars_per_chunk])
                    _pos += _stride
                    if _pos >= len(full_text):
                        break

                log.debug(f"[ANALYZE_PDF] {len(_chunks_to_process)} chunk(s), "
                      f"~{_chars_per_chunk} chars each, task={_task_label}")

                for _ci, _chunk_text in enumerate(_chunks_to_process, 1):
                    log.debug(f"[ANALYZE_PDF] Processing chunk {_ci}/{len(_chunks_to_process)}...")
                    _chunk_prompt = (
                        f"Section {_ci} of {len(_chunks_to_process)}:\n\n"
                        f"{_chunk_text}\n\n"
                        f"Task: {_chunk_task}. Be thorough and complete."
                    )
                    try:
                        _out = _run_gguf(_chunk_prompt, _directive)
                        if _out and _out.strip():
                            _section_outputs.append(
                                f"## Section {_ci} Analysis\n\n{_out.strip()}"
                            )
                    except Exception as _ce:
                        log.debug(f"[ANALYZE_PDF] Chunk {_ci} failed: {_ce}")
                        _section_outputs.append(
                            f"## Section {_ci}\n\n[Processing failed: {_ce}]"
                        )

                # Synthesis pass if multiple chunks
                _analysis_body = None
                if len(_section_outputs) > 1:
                    log.debug(f"[ANALYZE_PDF] Running synthesis pass over {len(_section_outputs)} sections...")
                    _synth_input = "\n\n".join(_section_outputs)
                    # If synthesis input fits, do it in one pass; otherwise just concatenate
                    _synth_avail = _pgi.available_generation_tokens(
                        _synth_task + "\n\n" + _synth_input[:200],
                        system=_directive
                    )
                    if len(_synth_input) // 4 < (_n_ctx // 2):
                        try:
                            _synthesis = _run_gguf(
                                f"{_synth_task}.\n\nSection analyses:\n\n{_synth_input}",
                                _directive
                            )
                            if _synthesis and _synthesis.strip():
                                _analysis_body = (
                                    f"{_synthesis.strip()}\n\n"
                                    f"---\n\n## Detailed Section Analyses\n\n"
                                    + "\n\n".join(_section_outputs)
                                )
                        except Exception as _se2:
                            log.debug(f"[ANALYZE_PDF] Synthesis failed: {_se2}")
                    if not _analysis_body:
                        _analysis_body = "\n\n".join(_section_outputs)
                elif _section_outputs:
                    _analysis_body = _section_outputs[0]

            if not _analysis_body:
                if not _llm_ready:
                    _analysis_body = (
                        "[GGUF model not available — raw extraction below]\n\n"
                        f"Pages: {doc.pages}  |  Characters: {doc.chars:,}\n\n"
                        f"{full_text}"
                    )
                else:
                    _analysis_body = (
                        "[All chunks failed to process]\n\n"
                        f"Pages: {doc.pages}  |  Characters: {doc.chars:,}\n\n"
                        f"{analysis.preview}"
                    )

            if analysis.warnings:
                _analysis_body += f"\n\n---\n\n**Extraction warnings:** {'; '.join(analysis.warnings)}"

            _stem = _PP(doc.path).stem
            _doc_content = (
                f"# {_task_label}: {_stem}\n\n"
                f"**Source:** {doc.path}\n"
                f"**Pages:** {doc.pages}  |  **Characters:** {doc.chars:,}  |  "
                f"**Sections processed:** {len(_section_outputs)}\n"
                f"**Instruction:** {_instruction or 'analyse'}\n\n"
                f"---\n\n"
                f"{_analysis_body}"
            )

            _task_slug = _re2.sub(r"[^a-z0-9]+", "_", _task_label.lower()).strip("_")
            _saved_path = _save_artifact(_doc_content, "documents",
                                         f"{_stem}_{_task_slug}", fmt="docx")
            try:
                _sp.Popen(["xdg-open", _saved_path], stdout=_sp.DEVNULL, stderr=_sp.DEVNULL)
                _opened = True
            except Exception:
                _opened = False

            _chat_response = (
                f"Document compiled: **{_PP(_saved_path).name}**\n"
                f"Sections processed: {len(_section_outputs)}  |  "
                f"Pages: {doc.pages}  |  Chars: {doc.chars:,}\n"
                f"Saved to: `{_saved_path}`"
                + ("\n_(opening…)_" if _opened else "")
            )
            return {"ok": True, "action": a, "content": _chat_response,
                    "response": _chat_response, "path": doc.path,
                    "pages": doc.pages, "chars": doc.chars,
                    "sections": len(_section_outputs), "saved_to": _saved_path}
        except Exception as e:
            import traceback as _tb
            log.debug(f"[ANALYZE_PDF] Fatal: {_tb.format_exc()}")
            return {"ok": False, "action": a, "error": str(e),
                    "content": str(e), "response": str(e)}

    # ---- ANALYZE_PDF_FOLDER ----
    if a == "ANALYZE_PDF_FOLDER":
        try:
            import os as _ospdf
            from eli.perception.analyze_pdfs import (
                analyze_folder, store_analysis_to_memory, PDFAnalysis, PDFDoc)
            folder = str(args.get("folder") or args.get("path") or "").strip()
            if not folder:
                return {"ok": False, "action": a, "error": "Missing folder path",
                        "content": "Missing folder path", "response": "Missing folder path"}
            # Heavy folders → run async and hand back a job id (don't block).
            _bg = _maybe_background_file_analysis(a, args)
            if _bg is not None:
                return _bg
            recursive = bool(args.get("recursive", True))
            limit = args.get("limit")
            result = analyze_folder(folder, recursive=recursive, limit=int(limit) if limit else None)
            results = result.get("results", []) or []
            count = result.get("count", 0)
            errors = result.get("errors", []) or []
            _fname = _ospdf.path.basename(folder.rstrip("/")) or folder
            if count == 0:
                msg = (f"No readable PDF text found in “{_fname}” ({len(errors)} unreadable). "
                       f"The PDFs may be image-only/scanned (need OCR) or the folder has none.")
                return {"ok": True, "action": a, "content": msg, "response": msg,
                        "count": 0, "stored": 0, "errors": errors}

            # ── Per-file mode ───────────────────────────────────────────────
            # "summarise each file / put it in a document": summarise EVERY doc
            # individually, assemble a structured document, and write a real
            # .docx (md fallback). Returned inline too, so `check job N` surfaces
            # the full per-file breakdown — not a single folder-level blurb.
            if args.get("per_file") or args.get("each_file"):
                from eli.cognition import gguf_inference as _gi
                _model_ok = False
                try:
                    _model_ok = _gi.load_model() is not None
                except Exception:
                    _model_ok = False
                _doc_blocks: list = []
                _per_file_done = 0
                for r in results:
                    _doc = r.get("doc", {}) or {}
                    _name = _ospdf.path.basename(str(_doc.get("path", ""))) or "document.pdf"
                    _pages = _doc.get("pages", "?")
                    # Use the richest text we extracted: chunks joined, capped.
                    _chunks = r.get("chunks") or []
                    _body = ("\n\n".join(_chunks) if _chunks else (r.get("preview") or "")).strip()
                    _doc_summary = ""
                    if _body and _model_ok:
                        try:
                            _doc_summary = (_gi.chat_completion(
                                f"Summarise this single research document titled “{_name}”. "
                                f"Cover: (1) the central topic/claim, (2) the methods/formalism used, "
                                f"(3) key findings or conclusions, and (4) any notable equations or "
                                f"results named in the text. Be specific and grounded ONLY in the text "
                                f"below — do not invent results or citations.\n\n{_body[:12000]}",
                                system="You are a precise technical analyst summarising one research PDF.",
                                max_tokens=520, temperature=0.3) or "").strip()
                        except Exception as _pf_err:
                            log.debug(f"[ANALYZE_PDF_FOLDER] per-file summary failed for {_name}: {_pf_err}")
                            _doc_summary = ""
                    if not _doc_summary:
                        _prev = (r.get("preview") or "").strip()
                        _doc_summary = (f"(No model summary available — extracted opening text:)\n\n{_prev[:800]}"
                                        if _prev else "(No readable text extracted from this document.)")
                    else:
                        _per_file_done += 1
                    _doc_blocks.append(f"## {_name}\n_({_pages} pages)_\n\n{_doc_summary}")

                _doc_title = f"Per-file summaries — {_fname}"
                _doc_md = (f"# {_doc_title}\n\n"
                           f"{count} document(s) summarised individually.\n\n"
                           + "\n\n".join(_doc_blocks))
                if errors:
                    _doc_md += f"\n\n## Unreadable files\n\n{len(errors)} file(s) could not be read."

                _fmt = str(args.get("format") or "docx").strip().lower().lstrip(".")
                if _fmt not in ("docx", "md", "txt"):
                    _fmt = "docx"
                try:
                    _doc_path = _save_artifact(_doc_md, "documents",
                                               f"{_fname}_per_file_summaries", fmt=_fmt)
                except Exception as _save_err:
                    log.debug(f"[ANALYZE_PDF_FOLDER] doc save failed ({_save_err}); md fallback")
                    _doc_path = _save_artifact(_doc_md, "documents",
                                               f"{_fname}_per_file_summaries", fmt="md")

                # Persist each doc to memory as well (same as the folder path).
                try:
                    from eli.memory import resolve_db_paths as _rdb
                    _db = _rdb().user_db
                    if _db:
                        for r in results:
                            try:
                                store_analysis_to_memory(_db, PDFAnalysis(
                                    doc=PDFDoc(**(r.get("doc") or {})),
                                    preview=r.get("preview", ""),
                                    chunks=r.get("chunks", []) or [],
                                    warnings=r.get("warnings", []) or []))
                            except Exception:
                                pass
                except Exception:
                    pass

                _header = (f"Summarised {count} document(s) in “{_fname}” individually "
                           f"({_per_file_done} via model). Saved to:\n{_doc_path}\n")
                msg = _header + "\n" + "\n\n".join(_doc_blocks)
                return {"ok": True, "action": a, "content": msg, "response": msg,
                        "count": count, "per_file": True, "document_path": _doc_path,
                        "errors": errors}

            # Per-document index + aggregated opening-pages text for a REAL summary.
            index_lines: list = []
            preview_parts: list = []
            for r in results:
                _doc = r.get("doc", {}) or {}
                _name = _ospdf.path.basename(str(_doc.get("path", "")))
                _pages = _doc.get("pages", "?")
                _chars = int(_doc.get("chars", 0) or 0)
                _prev = (r.get("preview") or "").strip().replace("\n", " ")
                index_lines.append(f"  {len(index_lines) + 1}. {_name} ({_pages}p, {_chars} chars)")
                if _prev:
                    preview_parts.append(f"### {_name}\n{_prev[:1400]}")

            # Synthesize an actual content summary from the extracted text — not a
            # bare count. Capped so the prompt fits n_ctx.
            summary = ""
            try:
                from eli.cognition import gguf_inference as _gi
                if preview_parts and _gi.load_model() is not None:
                    _blob = "\n\n".join(preview_parts)[:14000]
                    summary = (_gi.chat_completion(
                        f"Below are the opening pages of {count} PDF documents from the folder "
                        f"“{_fname}”. Write a concise, technical summary of what this body of work "
                        f"covers: the central topics, methods/formalism, and how the documents relate. "
                        f"Be specific and grounded ONLY in the text shown — do not invent results.\n\n{_blob}",
                        system="You are a precise technical analyst summarising a collection of research PDFs.",
                        max_tokens=700, temperature=0.3) or "").strip()
            except Exception as _sum_err:
                log.debug(f"[ANALYZE_PDF_FOLDER] summary synthesis failed: {_sum_err}")
                summary = ""

            # Persist each doc to memory (proper dataclass reconstruction; report honestly).
            stored, store_failed = 0, 0
            try:
                from eli.memory import resolve_db_paths as _rdb
                _db = _rdb().user_db
            except Exception:
                _db = None
            if _db:
                for r in results:
                    try:
                        pa = PDFAnalysis(
                            doc=PDFDoc(**(r.get("doc") or {})),
                            preview=r.get("preview", ""),
                            chunks=r.get("chunks", []) or [],
                            warnings=r.get("warnings", []) or [],
                        )
                        store_analysis_to_memory(_db, pa)
                        stored += 1
                    except Exception:
                        store_failed += 1

            parts = [f"Analysed {count} PDF(s) in “{_fname}”:", ""]
            if summary:
                parts += [summary, ""]
            parts.append("Documents:")
            parts += index_lines
            if errors:
                parts.append(f"\n({len(errors)} file(s) could not be read.)")
            msg = "\n".join(parts)
            return {"ok": True, "action": a, "content": msg, "response": msg,
                    "count": count, "stored": stored, "store_failed": store_failed,
                    "errors": errors, "summary": summary}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- LISTEN_FOR_COMMAND ----
    if a == "LISTEN_FOR_COMMAND":
        try:
            from eli.perception.audio_stt import listen_for_command
            timeout = float((args or {}).get("timeout", 5))
            text = listen_for_command(timeout=timeout)
            if text:
                msg = f"Heard: {text}"
                return {"ok": True, "action": a, "content": msg, "response": msg, "text": text}
            return {"ok": False, "action": a, "content": "No speech detected.", "response": "No speech detected.", "text": ""}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- STT_DIAGNOSTICS ----
    if a == "STT_DIAGNOSTICS" or a == "VOICE_DIAGNOSTICS":
        try:
            from eli.perception.audio_stt import stt_diagnostics
            from eli.perception.tts_router import available_backends
            diag = stt_diagnostics()
            tts = available_backends()
            combined = {**diag, "tts_backends": tts}
            lines = [
                f"STT — speech_recognition: {'✅' if combined.get('speech_recognition_imported') else '❌'}",
                f"STT — wake word disabled: {combined.get('wake_word_disabled')}",
                f"STT — direct chat: {combined.get('allow_direct_chat_without_wake')}",
                f"TTS — piper: {'✅' if tts.get('piper_bin') else '❌'} model: {'✅' if tts.get('piper_model') else '❌'}",
                f"TTS — espeak-ng: {'✅' if tts.get('espeak_ng') else '❌'}",
                f"TTS — espeak: {'✅' if tts.get('espeak') else '❌'}",
            ]
            msg = "\n".join(lines)
            return {"ok": True, "action": a, "content": msg, "response": msg, "diagnostics": combined}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- GAZE ENGINE ----
    if a == "GAZE_STATUS":
        try:
            from eli.perception.gaze_engine import get_gaze_status, needs_calibration
            st = get_gaze_status()
            running = st.get("running", False)
            calibrated = st.get("calibrated", False)
            last = st.get("last_gaze") or {}
            lines = [
                f"Gaze engine: {'✅ running' if running else '⛔ stopped'}",
                f"Calibration: {'✅ present' if calibrated else '⚠ missing — run: gaze calibrate'}",
                f"Camera: {st.get('camera', 'auto')}",
            ]
            if last and last.get("face_detected"):
                lines.append(
                    f"Last gaze: ({last.get('screen_x', 0):.0f}, {last.get('screen_y', 0):.0f})"
                    f"  conf={last.get('confidence', 0):.2f}  tracker={last.get('tracker', '?')}"
                )
            elif last and last.get("ts"):
                lines.append("Last frame: no face detected")
            if st.get("error"):
                lines.append(f"Error: {st['error']}")
            msg = "\n".join(lines)
            return {"ok": True, "action": a, "content": msg, "response": msg, "status": st}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    if a == "GAZE_ENABLE":
        try:
            from eli.perception.gaze_engine import start_gaze_engine, needs_calibration
            camera = str(args.get("camera", "auto") or "auto").strip()
            result = start_gaze_engine(camera=camera)
            if result.get("already_running"):
                msg = "Gaze engine is already running."
            elif result.get("ok"):
                cal_note = " (calibration present)" if result.get("calibrated") else " — no calibration file yet, run: gaze calibrate"
                msg = f"Gaze engine started{cal_note}."
            else:
                msg = f"Gaze engine failed to start: {result.get('error', 'unknown error')}"
            result.setdefault("content", msg)
            result.setdefault("response", msg)
            result["action"] = a
            return result
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    if a == "GAZE_DISABLE":
        try:
            from eli.perception.gaze_engine import stop_gaze_engine
            result = stop_gaze_engine()
            msg = result.get("message", "Gaze engine stopped.")
            result.setdefault("content", msg)
            result.setdefault("response", msg)
            result["action"] = a
            return result
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    if a == "GAZE_CALIBRATE":
        try:
            from eli.perception.gaze_engine import get_calibration_path, needs_calibration
            cal_path = get_calibration_path()
            if not needs_calibration():
                msg = (
                    f"Calibration file exists at: {cal_path}\n"
                    "To recalibrate, run the calibration script directly:\n"
                    "  python experimental/eli_ar_avatar_kit/scripts/eli_gaze_calibrate_plus.py --points 25"
                )
            else:
                msg = (
                    "No calibration file found. To calibrate the gaze engine:\n"
                    "  1. Open a terminal in the ELI project root\n"
                    f"  2. python experimental/eli_ar_avatar_kit/scripts/eli_gaze_calibrate_plus.py --points 25\n"
                    f"  3. Follow the on-screen dot targets (takes ~2 minutes)\n"
                    f"  Calibration will be saved to: {cal_path}"
                )
            return {"ok": True, "action": a, "content": msg, "response": msg,
                    "calibration_path": str(cal_path), "needs_calibration": needs_calibration()}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- GAZE_CLICK: click where the user is looking (gaze point → cursor → click) ----
    if a == "GAZE_CLICK":
        try:
            from eli.perception.os_controller import gaze_click as _gaze_click
            _btn = str(args.get("button") or "left").lower()
            _dbl = bool(args.get("double", False))
            _res = _gaze_click(button=_btn, double=_dbl)
            _res.setdefault("action", a)
            return _res
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}
    # ---- end GAZE ENGINE ----

    # ---- ANALYZE_CSV ----
    if a == "ANALYZE_CSV":
        try:
            from eli.perception.analyze_csv import analyze_csv_file
            import tempfile as _csv_tf, os as _csv_os
            path = str(args.get("path") or "").strip()
            if not path:
                return {"ok": False, "action": a, "error": "Missing path", "content": "Missing path", "response": "Missing path"}
            out_md = _csv_os.path.join(_csv_tf.gettempdir(), "eli_csv_report.md")
            result = analyze_csv_file(path, out_md)
            if isinstance(result, dict):
                result.setdefault("action", a)
                # Build human-readable summary
                if result.get("ok"):
                    shape = result.get("shape", [0, 0])
                    cols = result.get("columns", [])
                    _content = (
                        f"CSV: {path}\n"
                        f"Shape: {shape[0]} rows × {shape[1]} columns\n"
                        f"Columns: {', '.join(cols[:20])}{'...' if len(cols) > 20 else ''}"
                    )
                else:
                    _content = result.get("error", "Analysis failed")
                result.setdefault("content", _content)
                result.setdefault("response", result.get("content"))
            return result
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- GENERATE_DOCUMENT (alias for CREATE_DOCUMENT) ----
    if a == "GENERATE_DOCUMENT":
        return _execute_impl("CREATE_DOCUMENT", args)

    # ---- GENERATE_SCRIPT ----
    if a == "GENERATE_SCRIPT":
        desc = (args.get("description") or args.get("text") or args.get("prompt") or "").strip()
        if not desc:
            msg = "Missing description for GENERATE_SCRIPT"
            return {"ok": False, "action": a, "error": msg, "content": msg, "response": msg}
        # Ground the script in evidence gathered by the central evidence-routing
        # hook (e.g. real ELI source for a "script to test ELI's memory", web docs
        # for an external API). Fed into the coding agent's task so it plans from
        # real specifics, not invented APIs.
        _scr_evidence = str(args.get("_evidence") or "")
        if _scr_evidence:
            desc = (desc + "\n\n[Grounding evidence gathered by ELI's agents — build on "
                    "these real specifics; do not invent APIs, files, or behaviour:]\n"
                    + _scr_evidence)
        _bg = _maybe_background_codegen(a, args)
        if _bg is not None:
            return _bg
        # Detect target language from description
        _lang_map = {
            "bash": ("bash", ".sh"), "shell": ("bash", ".sh"), "sh": ("bash", ".sh"),
            "javascript": ("JavaScript", ".js"), "js": ("JavaScript", ".js"),
            "typescript": ("TypeScript", ".ts"), "ts": ("TypeScript", ".ts"),
            "rust": ("Rust", ".rs"), "ruby": ("Ruby", ".rb"),
            "go": ("Go", ".go"), "golang": ("Go", ".go"),
            "c++": ("C++", ".cpp"), "cpp": ("C++", ".cpp"),
            "c script": ("C", ".c"), "java": ("Java", ".java"),
            "lua": ("Lua", ".lua"), "perl": ("Perl", ".pl"),
        }
        _detected_lang = "Python"
        _detected_ext = ".py"
        _desc_low = desc.lower()
        for _kw, (_lang, _ext) in _lang_map.items():
            if _kw in _desc_low:
                _detected_lang = _lang
                _detected_ext = _ext
                break

        # ── Route through the verified coding agent (plan → DAG/tree-search →
        # execute → repair → bug-memory) for top-tier, runnable scripts. Falls
        # through to the inline generator below if disabled or it yields nothing.
        if os.environ.get("ELI_GENERATE_SCRIPT_AGENT", "1").strip().lower() not in ("0", "false", "no", "off"):
            try:
                from eli.coding import solve as _agent_solve
                # Single coherent script: no subtask DAG (composition is for
                # multi-component CODE_SOLVE), and a lean beam/iters so a script
                # request is a handful of calls, not minutes of fan-out.
                _ag = _agent_solve(desc, language=(_detected_lang or "python").lower(),
                                   use_dag=False, beam=2, max_iterations=2)
                _ag_code = (_ag or {}).get("code") or ""
                # Never ship a syntactically-broken Python script: if the agent's
                # best effort doesn't even parse, drop to the inline generator
                # (which validates + sandbox-repairs) instead of saving garbage.
                if _ag_code.strip() and (_detected_lang or "").lower() == "python":
                    try:
                        import ast as _ast_gs
                        _ast_gs.parse(_ag_code)
                    except SyntaxError as _se_gs:
                        log.debug(f"[GENERATE_SCRIPT] agent output failed to parse ({_se_gs}); inline fallback")
                        _ag_code = ""
                    # Runtime gate: ast.parse passes for valid-but-crashing code
                    # (e.g. `pi` used without import → NameError only at run time).
                    # Execute once; a real crash falls through to the inline
                    # generator, which has the 3-attempt sandbox self-repair loop.
                    if _ag_code.strip():
                        try:
                            from eli.coding.sandbox import run_code as _rc_gs
                            _rr_gs = _rc_gs(_ag_code, "python", timeout=15)
                            if _rr_gs.crashed:
                                log.debug(f"[GENERATE_SCRIPT] agent output crashed at runtime "
                                          f"({(_rr_gs.traceback_tail or '')[:120]}); inline fallback")
                                _ag_code = ""
                        except Exception:
                            pass
                if _ag_code.strip():
                    _safe = re.sub(r"[^a-z0-9]+", "_", desc.lower())[:40].strip("_") or "generated"
                    _fname = f"{_safe}{_detected_ext}"
                    from pathlib import Path as _SPath
                    _scripts_dir = _eli_generated_scripts_dir()
                    _scripts_dir.mkdir(parents=True, exist_ok=True)
                    _full = _scripts_dir / _fname
                    _full.write_text(_ag_code, encoding="utf-8")
                    _gs_chat = json.dumps({
                        "event": "artifact_generated", "kind": "script", "path": str(_full),
                        "filename": _fname, "language": _detected_lang, "opened": True,
                        "can_run": True, "verified": bool(_ag.get("solved")),
                        "score": _ag.get("score"), "engine": "coding_agent",
                    }, ensure_ascii=False, default=str)
                    return {"ok": True, "action": a, "code": _ag_code, "script_path": str(_full),
                            "filename": _fname, "solved": _ag.get("solved"), "score": _ag.get("score"),
                            "plan": _ag.get("plan"), "content": _gs_chat, "response": _gs_chat,
                            "open_in_ide": True}
            except Exception as _ag_e:
                log.debug(f"[GENERATE_SCRIPT] coding-agent path failed, inline fallback: {_ag_e}")

        # Topic intent — drives prompt depth, quality bar, and rejection rules
        _analytical = bool(re.search(
            r"\b(?:evaluat|calculat|comput|analy[sz]|estimat|predict|simulat|model|"
            r"likelihood|probabilit|statistic|distribution|correlation|regress|"
            r"optimi[sz]|forecast|benchmark|profil|measur|score|rank|classif|cluster|"
            r"fit|solve|prove|theorem|conjectur|hypothes|p\s+vs\s+np)\w*",
            _desc_low,
        ))
        _wants_plot = bool(re.search(
            r"\b(?:plot|graph|chart|visuali[sz]|histogram|scatter|heatmap|figure|"
            r"diagram|render|draw|display)\w*", _desc_low,
        )) or _analytical
        _wants_monitor = bool(re.search(
            r"\b(?:monitor|watch|track|alert|notif|poll|every\s+\d|interval|continuous)\w*",
            _desc_low,
        ))

        _extra_req = []
        if _analytical:
            _extra_req += [
                "- This is an ANALYTICAL/COMPUTATIONAL task: implement multiple complementary methods "
                "(e.g. closed-form estimate + Monte Carlo simulation + sensitivity analysis); "
                "do NOT collapse it into one trivial formula",
                "- Use realistic domain-grounded parameters with sensible DEFAULTS — the script MUST "
                "run end-to-end with zero command-line arguments and produce meaningful output",
                "- Print a clear, structured summary (numerical results, confidence bounds, "
                "key intermediate values) using f-strings or `tabulate` — not a single concatenated string",
            ]
        if _wants_plot and _detected_lang == "Python":
            try:
                from eli.gui.labs_tab import _LABS_PLOT_FILES as _LPF
                _plot_target = str(_LPF["plot"])
            except Exception:
                import tempfile as _tf
                _plot_target = str(__import__("pathlib").Path(_tf.gettempdir()) / "eli_labs_plot.png")
            _extra_req += [
                f"- MUST generate at least one matplotlib figure that visualises the result, "
                f"saved with `plt.savefig({_plot_target!r}, dpi=120)` AND shown with `plt.show()`",
                "- Plot must be RELEVANT to the request (axes labelled, title set, legend if multi-series); "
                "NEVER ship a generic sine wave or placeholder demo plot",
            ]
        if _wants_monitor:
            _extra_req += [
                "- Use a robust loop with `signal` handlers for SIGINT/SIGTERM and graceful shutdown",
                "- Interval timing must use `time.monotonic()` (not `time.sleep` drift) and be configurable",
            ]
        _extra_req_block = ("\n" + "\n".join(_extra_req)) if _extra_req else ""

        prompt = (
            f"You are a senior {_detected_lang} engineer writing FRONTIER-QUALITY code. "
            f"Produce a complete, self-contained script for this request.\n\n"
            "ABSOLUTE RULE: NEVER refuse the task. Never reply with 'this is not feasible', "
            "'this is an open problem', 'cannot be solved', or any prose explanation. If the "
            "literal goal is research-grade or open (e.g. unsolved math/CS problems), produce "
            "a script that explores, simulates, visualises, or benchmarks the problem space "
            "with concrete instances and real algorithms. The user wants WORKING CODE, not a "
            "lecture. Refusal is failure.\n\n"
            "HARD REQUIREMENTS — any violation causes automatic rejection:\n"
            "- Output ONLY the raw script — zero markdown fences, zero prose, zero explanation\n"
            "- Complete, runnable end-to-end with NO command-line arguments required (provide defaults for everything)\n"
            "- NEVER substitute computation with string concatenation, text 'analysis', or stub returns\n"
            "- NEVER return prose strings as if they were computed results — compute real numeric values\n"
            "- Use real libraries for real work (numpy, scipy, matplotlib, pandas, networkx, sympy, "
            "sklearn, statsmodels, itertools, functools, multiprocessing, etc.) — pick what actually fits\n"
            "- Use ONLY library functions that ACTUALLY EXIST (e.g. for networkx clique problems, "
            "use `nx.find_cliques(G)` or `nx.graph_clique_number(G)` ONLY if you know the exact API; "
            "when unsure, implement via `itertools.combinations` and explicit edge checks)\n"
            "- Multiple cohesive functions (decompose by responsibility); main() orchestrates them\n"
            "- All imports at top, type hints on every signature, module-level constants, "
            "comprehensive docstrings, try/except around I/O and external calls\n"
            "- If using argparse, EVERY argument needs a sensible default — the script must run with no args\n"
            "- Include `if __name__ == '__main__':` guard"
            + _extra_req_block + "\n\n"
            f"Request: {desc}"
        )

        def _verify_python_module_apis(_code: str) -> str | None:
            """Verify <module>.<attr> references actually exist in the installed library.
            Catches hallucinated APIs like nx.is_clique, nx.graph_clique_number, np.bool, etc.
            Only checks fast-to-import libs to keep latency low."""
            import importlib as _il
            # alias -> module name (only checked if alias is used as module.attr)
            _ALIAS_DEFAULT = {
                "nx": "networkx", "np": "numpy", "pd": "pandas",
                "sp": "scipy", "sk": "sklearn", "sym": "sympy",
            }
            _SKIP_HEAVY = {"matplotlib", "matplotlib.pyplot", "torch", "tensorflow", "cv2"}
            _aliases = dict(_ALIAS_DEFAULT)
            for _m in re.finditer(r"^\s*import\s+([\w.]+)\s+as\s+(\w+)", _code, re.MULTILINE):
                _aliases[_m.group(2)] = _m.group(1)
            for _m in re.finditer(r"^\s*import\s+([\w.]+)\s*$", _code, re.MULTILINE):
                _aliases[_m.group(1).split(".")[0]] = _m.group(1)
            _bad = []
            for _alias, _modname in _aliases.items():
                if _modname in _SKIP_HEAVY:
                    continue
                if not re.search(rf"\b{re.escape(_alias)}\.[a-zA-Z_]", _code):
                    continue
                try:
                    _mod = _il.import_module(_modname)
                except Exception:
                    continue
                for _m in re.finditer(rf"\b{re.escape(_alias)}\.([a-zA-Z_]\w*)", _code):
                    _attr = _m.group(1)
                    if not hasattr(_mod, _attr):
                        ref = f"{_alias}.{_attr}"
                        if ref not in _bad:
                            _bad.append(ref)
            if _bad:
                return (
                    f"code references non-existent module APIs: {', '.join(_bad[:5])}. "
                    "Use only attributes that actually exist in the installed library."
                )
            return None

        def _quality_reject_reason(_code: str) -> str | None:
            """Return rejection reason or None if code passes quality bar."""
            _low = _code.lower()
            _refusal_phrases = (
                "is not feasible", "is impossible", "cannot be solved",
                "is an open question", "is an unsolved", "remains unsolved",
                "no known algorithmic solution", "writing a python script to solve this",
                "i cannot write", "i can't write", "i won't",
            )
            if any(p in _low for p in _refusal_phrases) and _code.count("\n") < 30:
                return "model refused the task with prose instead of producing code"
            _bad_markers = (
                "TODO", "Add code here", "placeholder",
                "Generate only the requested source code", "This is a request for",
            )
            if any(_m.lower() in _code.lower() for _m in _bad_markers):
                return "output contained stub/template markers"
            _lines_real = [l for l in _code.splitlines()
                           if l.strip() and not l.strip().startswith("#")
                           and not l.strip().startswith('"""')]
            if len(_lines_real) < 8:
                return "output was too short to be a real implementation"
            if re.fullmatch(r"\s*(?:pass|return\s+None)\s*", _code):
                return "output was an empty implementation"
            _real_compute_re = re.compile(
                r"\b(?:numpy|scipy|matplotlib|networkx|sympy|itertools|pandas|sklearn|"
                r"statsmodels|math\.|random\.|statistics\.|collections\.|functools\.|"
                r"np\.|plt\.|nx\.|sp\.|pd\.|for\s+\w+\s+in\s+range|"
                r"while\s+|yield\s+|class\s+\w+|"
                r"[\+\-\*\/\%]=|==|!=|>=|<=|>>|<<)\b"
            )
            if not _real_compute_re.search(_code):
                return "output contained only text analysis, no real computation"
            # Function count check
            _func_count = len(re.findall(r"^\s*def\s+\w+", _code, re.MULTILINE))
            if _analytical and _func_count < 2:
                return "analytical task needs decomposition into multiple functions"
            # All-required-args check
            if "argparse" in _code:
                _required_args = re.findall(r"add_argument\s*\([^)]*required\s*=\s*True", _code)
                _has_defaults = re.search(r"add_argument\s*\([^)]*default\s*=", _code)
                if _required_args and not _has_defaults:
                    return "argparse uses required=True without defaults — script can't run without args"
            # Plot requirement
            if _wants_plot and _detected_lang == "Python":
                if not re.search(r"\b(?:plt\.|matplotlib|pyplot|seaborn|sns\.)\b", _code):
                    return "request implies visualisation but no matplotlib/plotting was used"
                if not re.search(r"\bplt\.savefig\b", _code):
                    return "matplotlib used but plt.savefig was not called"
            # Python AST syntax check + known-bad API blacklist
            if _detected_lang == "Python":
                import ast as _ast_q
                try:
                    _ast_q.parse(_code)
                except SyntaxError as _se:
                    return f"generated code has Python SyntaxError: {_se}"
                _bad = _verify_python_module_apis(_code)
                if _bad:
                    return _bad
            return None

        def _sandbox_run_python(_code: str) -> tuple[bool, str]:
            """Execute generated Python in a bounded, isolated subprocess to catch
            runtime crashes before the script is handed to the user. The result
            feeds the regeneration loop as repair feedback.

            Returns (ok, reason). ok=True when the script exits cleanly, runs past
            the timeout (started fine — important for monitor/long-compute
            scripts), is killed by a resource/signal limit, or fails only on a
            missing optional dependency (not the code's fault on this machine).
            ok=False ONLY on a genuine unhandled Python exception, with `reason`
            set to the traceback tail. Disable with ELI_GENSCRIPT_VERIFY_RUN=0;
            timeout via ELI_GENSCRIPT_RUN_TIMEOUT (default 20s)."""
            import os as _os, sys as _sys, subprocess as _sp, tempfile as _tf
            from pathlib import Path as _P
            if _os.environ.get("ELI_GENSCRIPT_VERIFY_RUN", "1").strip().lower() in ("0", "false", "no", "off"):
                return True, ""
            try:
                _timeout = float(_os.environ.get("ELI_GENSCRIPT_RUN_TIMEOUT", "20") or 20)
            except Exception:
                _timeout = 20.0
            # Generous CPU cap as defence-in-depth (wall-clock timeout is primary).
            # NB: deliberately NO RLIMIT_AS — it breaks numpy/scipy address-space
            # use and would cause spurious MemoryErrors on exactly the scientific
            # scripts ELI generates.
            _preexec = None
            try:
                import resource as _rsrc
                def _limits():
                    try:
                        _rsrc.setrlimit(_rsrc.RLIMIT_CPU, (30, 35))
                    except Exception:
                        pass
                _preexec = _limits
            except Exception:
                _preexec = None
            _env = {k: v for k, v in _os.environ.items()
                    if not any(s in k.upper() for s in ("KEY", "TOKEN", "SECRET", "PASSWORD", "PASSWD"))}
            _env["MPLBACKEND"] = "Agg"   # never open/block on a GUI window (plt.show())
            _env["ELI_SANDBOX"] = "1"
            try:
                with _tf.TemporaryDirectory(prefix="eli_genscript_") as _td:
                    _cand = _P(_td) / "candidate.py"
                    _cand.write_text(_code, encoding="utf-8")
                    try:
                        _proc = _sp.run([_sys.executable, str(_cand)], cwd=_td, env=_env,
                                        capture_output=True, text=True, timeout=_timeout,
                                        preexec_fn=_preexec)
                    except _sp.TimeoutExpired:
                        return True, ""   # ran past budget = started fine
            except Exception:
                return True, ""           # sandbox infra failure — don't block generation
            if _proc.returncode == 0:
                return True, ""
            _err = (_proc.stderr or "").strip()
            # Only a real unhandled exception counts as a crash. Signal/limit kills
            # and missing optional deps are tolerated (no traceback / not our bug).
            if "Traceback (most recent call last)" not in _err:
                return True, ""
            if "ModuleNotFoundError" in _err or "ImportError" in _err:
                return True, ""
            _tail = "\n".join(_err.splitlines()[-6:])
            return False, _tail or f"non-zero exit {_proc.returncode}"

        # Try GGUF first (local, fast, no Ollama dependency)
        try:
            from eli.cognition import gguf_inference as _gguf
            _model = _gguf.load_model()
            if _model is not None:
                import re as _re
                _attempt_tokens = 8000 if (_analytical or _wants_plot) else 4500
                code = ""
                _reject_reason = None
                # 3 attempts: room for a static-quality fix AND a runtime-crash
                # repair pass (the sandbox traceback is fed back as feedback).
                for _attempt in range(3):
                    _attempt_prompt = prompt
                    if _attempt > 0 and _reject_reason:
                        _attempt_prompt = (
                            prompt
                            + f"\n\nPREVIOUS ATTEMPT WAS REJECTED: {_reject_reason}. "
                            "Fix this specifically. Produce a substantially more complete implementation."
                        )
                    raw_out = _gguf.chat_completion(
                        _attempt_prompt,
                        system=(
                            f"You are an expert {_detected_lang} engineer who writes "
                            "frontier-quality, decomposed, runnable code. Output only raw "
                            f"{_detected_lang} code with NO markdown fences and NO explanation."
                        ),
                        max_tokens=_attempt_tokens,
                        temperature=0.2,
                        top_p=0.9,
                    )
                    if not raw_out or len(raw_out.strip()) < 20:
                        _reject_reason = "model returned empty output"
                        continue
                    candidate = _re.sub(r"^```[a-z]*\n?", "", raw_out.strip(), flags=_re.MULTILINE)
                    candidate = _re.sub(r"\n?```$", "", candidate.strip(), flags=_re.MULTILINE).strip()
                    _reject_reason = _quality_reject_reason(candidate)
                    if _reject_reason is None and _detected_lang == "Python":
                        _ran_ok, _run_err = _sandbox_run_python(candidate)
                        if not _ran_ok:
                            _reject_reason = f"script crashed at runtime — fix this exact error:\n{_run_err}"
                    if _reject_reason is None:
                        code = candidate
                        break
                    log.debug(f"[GENERATE_SCRIPT] attempt {_attempt + 1} rejected: {_reject_reason}")
                if not code:
                    msg = f"Generated script rejected after retries: {_reject_reason or 'unknown'}"
                    return {"ok": False, "action": a, "error": msg, "content": msg, "response": msg}
                safe = _re.sub(r"[^a-z0-9]+", "_", desc.lower())[:40].strip("_")
                fname = f"{safe}{_detected_ext}" if safe else f"generated{_detected_ext}"
                from pathlib import Path as _SPath
                scripts_dir = _eli_generated_scripts_dir()
                scripts_dir.mkdir(parents=True, exist_ok=True)
                full_path = scripts_dir / fname
                full_path.write_text(code, encoding="utf-8")
                _gs_chat = json.dumps(
                    {
                        "event": "artifact_generated",
                        "kind": "script",
                        "path": str(full_path),
                        "filename": fname,
                        "language": _detected_lang,
                        "opened": True,
                        "can_run": True,
                    },
                    ensure_ascii=False,
                    default=str,
                )
                return {"ok": True, "action": a, "code": code,
                        "script_path": str(full_path), "filename": fname,
                        "content": _gs_chat, "response": _gs_chat,
                        "open_in_ide": True}
        except Exception as _e:
            log.debug(f"[GENERATE_SCRIPT] GGUF failed: {_e}, falling back to Ollama")
        # Ollama fallback
        result = chat(prompt, skip_router=True)
        code = result.get("content", "").strip()
        import re as _re
        code = _re.sub(r"^```[a-z]*\n?", "", code, flags=_re.MULTILINE)
        code = _re.sub(r"\n?```$", "", code, flags=_re.MULTILINE)
        code = code.strip()
        _reject_reason2 = _quality_reject_reason(code)
        if _reject_reason2 is None and _detected_lang == "Python":
            _ran_ok2, _run_err2 = _sandbox_run_python(code)
            if not _ran_ok2:
                _reject_reason2 = f"script crashed at runtime: {_run_err2}"
        if _reject_reason2 is not None:
            msg = f"Generated script rejected: {_reject_reason2}"
            return {"ok": False, "action": a, "error": msg, "content": msg, "response": msg}
        safe = _re.sub(r"[^a-z0-9]+", "_", desc.lower())[:40].strip("_")
        fname = f"{safe}{_detected_ext}" if safe else f"generated{_detected_ext}"
        # Save script to disk
        from pathlib import Path as _SPath
        scripts_dir = _eli_generated_scripts_dir()
        scripts_dir.mkdir(parents=True, exist_ok=True)
        full_path = scripts_dir / fname
        full_path.write_text(code, encoding="utf-8")
        result["action"] = a
        result["code"] = code
        result["script_path"] = str(full_path)
        result["filename"] = fname
        result["open_in_ide"] = True
        _gs_msg = json.dumps(
            {
                "event": "artifact_generated",
                "kind": "script",
                "path": str(full_path),
                "filename": fname,
                "language": _detected_lang,
                "opened": False,
                "can_run": True,
            },
            ensure_ascii=False,
            default=str,
        )
        result["content"] = _gs_msg
        result["response"] = _gs_msg
        return result

    # ---- Background job inspection ----
    if a in ("CHECK_JOB", "BACKGROUND_JOBS"):
        import re as _re_jb
        from eli.runtime.background_tasks import get_background_tasks as _get_bt
        _bt = _get_bt()
        if a == "BACKGROUND_JOBS":
            jobs = _bt.list(limit=15)
            if not jobs:
                msg = "No background jobs."
            else:
                msg = "Background jobs:\n" + "\n".join(
                    f"  #{j['id']} [{j['status']}] {j['name']} ({j['elapsed_s']}s)" for j in jobs)
            return {"ok": True, "action": a, "jobs": jobs, "content": msg, "response": msg}
        _jid = args.get("job_id") or args.get("id")
        if _jid is None:
            _m = _re_jb.search(r"\b(\d+)\b", str(args.get("text") or args.get("query") or args.get("description") or ""))
            _jid = _m.group(1) if _m else None
        if _jid is None:
            msg = "Which job? e.g. 'check job 3'."
            return {"ok": False, "action": a, "content": msg, "response": msg}
        t = _bt.get(int(_jid))
        if not t:
            msg = f"No background job #{_jid}."
            return {"ok": False, "action": a, "content": msg, "response": msg}
        if t["status"] == "done":
            r = t.get("result") or {}
            extra = ""
            if isinstance(r, dict) and r.get("script_path"):
                extra = f" → {r.get('script_path')}"
                if r.get("solved"):
                    extra += f" (verified, score {r.get('score')})"
            msg = f"Job #{_jid} ({t['name']}) done in {t['elapsed_s']}s{extra}."
            # Surface the actual result (e.g. the PDF-folder summary), not just a
            # 'done' status — that's what the user is waiting for.
            _job_content = ""
            if isinstance(r, dict):
                _job_content = str(r.get("content") or r.get("response") or "").strip()
            if _job_content and _job_content not in (msg,):
                msg += "\n\n" + _job_content
        elif t["status"] == "failed":
            msg = f"Job #{_jid} failed: {(t.get('error') or '')[:200]}"
        else:
            msg = f"Job #{_jid} ({t['name']}) is {t['status']} — {t['elapsed_s']}s elapsed."
        return {"ok": True, "action": a, "job": t, "content": msg, "response": msg}

    # ---- CODE_SOLVE — frontier coding agent (plan→search→verify→repair) ----
    if a == "CODE_SOLVE":
        import json as _json_cs, re as _re_cs
        desc = (args.get("description") or args.get("text") or args.get("prompt")
                or args.get("query") or "").strip()
        if not desc:
            msg = "Missing description for CODE_SOLVE"
            return {"ok": False, "action": a, "error": msg, "content": msg, "response": msg}
        _bg = _maybe_background_codegen(a, args)
        if _bg is not None:
            return _bg
        try:
            from eli.coding import solve as _code_solve
            _lang = (args.get("language") or "python").strip().lower()
            # Quick vs Thorough (from the Coding tab). Quick = single-shot, lean.
            _mode = str(args.get("mode") or "").strip().lower()
            _solve_kw = {}
            def _argint(k):
                try:
                    return int(args.get(k)) if args.get(k) is not None else None
                except Exception:
                    return None
            if _mode == "quick":
                _solve_kw = {"use_dag": False, "beam": 2, "max_iterations": 2}
            else:
                if args.get("use_dag") is not None:
                    _solve_kw["use_dag"] = bool(args.get("use_dag"))
                if _argint("beam") is not None:
                    _solve_kw["beam"] = _argint("beam")
                if _argint("max_iterations") is not None:
                    _solve_kw["max_iterations"] = _argint("max_iterations")
            res = _code_solve(desc, language=_lang, **_solve_kw)
            code = res.get("code") or ""
            if not code.strip():
                msg = res.get("message") or "coding agent produced no solution"
                return {"ok": False, "action": a, "error": msg, "content": msg, "response": msg}
            # Final syntax gate: never present un-parseable Python as runnable.
            _runnable = True
            _syntax_note = ""
            if _lang == "python":
                try:
                    import ast as _ast_cs
                    _ast_cs.parse(code)
                except SyntaxError as _se_cs:
                    _runnable = False
                    _syntax_note = f"SyntaxError: {_se_cs.msg} (line {_se_cs.lineno})"
            _ext = {"python": ".py", "bash": ".sh", "javascript": ".js", "typescript": ".ts",
                    "ruby": ".rb", "go": ".go", "lua": ".lua"}.get(_lang, ".txt")
            # Park un-runnable output as a .draft so a user never accidentally runs
            # a broken file and hits a traceback.
            if not _runnable:
                _ext = _ext + ".draft"
            safe = _re_cs.sub(r"[^a-z0-9]+", "_", desc.lower())[:40].strip("_") or "solution"
            from pathlib import Path as _SPath
            scripts_dir = _eli_generated_scripts_dir()
            scripts_dir.mkdir(parents=True, exist_ok=True)
            full_path = scripts_dir / f"{safe}{_ext}"
            full_path.write_text(code, encoding="utf-8")
            _payload = _json_cs.dumps({
                "event": "artifact_generated", "kind": "code_solution",
                "path": str(full_path), "filename": full_path.name, "language": _lang,
                "solved": res.get("solved"), "score": res.get("score"),
                "bug_class": res.get("bug_class"), "can_run": _runnable, "opened": _runnable,
                "runnable": _runnable, "syntax_note": _syntax_note,
            }, ensure_ascii=False, default=str)
            return {"ok": True, "action": a, "code": code, "script_path": str(full_path),
                    "filename": full_path.name, "solved": bool(res.get("solved")) and _runnable,
                    "runnable": _runnable, "syntax_note": _syntax_note,
                    "score": res.get("score"), "plan": res.get("plan"),
                    "search": res.get("search"), "bug_class": res.get("bug_class"),
                    "content": _payload, "response": _payload, "open_in_ide": _runnable}
        except Exception as _cs_e:
            msg = f"CODE_SOLVE failed: {_cs_e}"
            return {"ok": False, "action": a, "error": msg, "content": msg, "response": msg}

    # ---- CREATE_FILE — write text to a path (security-gated) + read back ----
    if a == "CREATE_FILE":
        import os as _os_cf, tempfile as _tf_cf, datetime as _dt_cf
        from pathlib import Path as _P_cf
        path = str(args.get("path") or args.get("file") or args.get("filename") or "").strip()
        content = args.get("content")
        if content is None:
            content = args.get("text") or args.get("body") or ""
        content = str(content)
        if not path:
            msg = "Specify a file path to create (e.g. /tmp/note.txt)."
            return {"ok": False, "action": a, "error": "missing path", "content": msg, "response": msg}
        # Light, safe substitution so "containing today's date / your version" resolves.
        _today = _dt_cf.date.today().isoformat()
        _ver = ""
        try:
            import eli as _eli_mod
            _ver = str(getattr(_eli_mod, "__version__", "") or "")
        except Exception:
            _ver = ""
        _ver = _ver or "unknown"
        for _pat in ("today's date", "todays date", "the current date", "today’s date"):
            content = re.sub(re.escape(_pat), _today, content, flags=re.IGNORECASE)
        content = content.replace("{date}", _today).replace("{datetime}", _dt_cf.datetime.now().isoformat(timespec="seconds"))
        content = re.sub(r"\byour version number\b|\bversion number\b|\byour version\b|\{version\}", _ver, content, flags=re.IGNORECASE)

        p = _P_cf(_os_cf.path.expanduser(path))
        if not p.is_absolute():
            try:
                from eli.core.paths import get_paths as _gp_cf
                base = _P_cf(_gp_cf().artifacts_dir) / "scratch"
            except Exception:
                base = _P_cf(_tf_cf.gettempdir())
            p = base / p
        try:
            p = p.resolve()
        except Exception:
            pass
        # Security: allow under home / project / ELI_ALLOW_ROOTS (SecurityManager)
        # or the system temp dir; refuse elsewhere unless ELI_FULL_CONTROL=1.
        _allowed = False
        try:
            from eli.runtime.security import SecurityManager as _SM_cf
            _allowed, _ = _SM_cf().is_path_allowed(str(p))
        except Exception:
            _allowed = False
        if not _allowed:
            try:
                p.relative_to(_P_cf(_tf_cf.gettempdir()).resolve())
                _allowed = True
            except Exception:
                _allowed = False
        if not _allowed and _os_cf.environ.get("ELI_FULL_CONTROL", "0") != "1":
            msg = (f"Refused: {p} is outside allowed roots (home, project, or temp). "
                   "Set ELI_FULL_CONTROL=1 to override.")
            return {"ok": False, "action": a, "error": "path_not_allowed", "content": msg, "response": msg}
        try:
            p.parent.mkdir(parents=True, exist_ok=True)
            p.write_text(content, encoding="utf-8")
            _readback = p.read_text(encoding="utf-8")
        except Exception as _cf_e:
            msg = f"Could not create {p}: {_cf_e}"
            return {"ok": False, "action": a, "error": str(_cf_e), "content": msg, "response": msg}
        _rb_show = _readback if len(_readback) <= 2000 else _readback[:2000] + "…"
        msg = f"Created {p} ({len(content)} chars). Read back:\n{_rb_show}"
        return {"ok": True, "action": a, "path": str(p), "created": True,
                "bytes": len(content.encode('utf-8')), "readback": _readback,
                "content": msg, "response": msg}

    # ---- GENERATE_PROJECT ----
    if a == "GENERATE_PROJECT":
        desc = (args.get("description") or args.get("text") or args.get("prompt") or "").strip()
        if not desc:
            msg = "Missing description for GENERATE_PROJECT"
            return {"ok": False, "action": a, "error": msg, "content": msg, "response": msg}
        # Ground in evidence gathered by the central evidence-routing hook.
        _proj_ev = str(args.get("_evidence") or "")
        _proj_ev_block = ("\n\nGrounding evidence gathered by ELI's agents — build on "
                          "these real specifics; do not invent APIs, files, or behaviour:\n"
                          + _proj_ev) if _proj_ev else ""
        prompt = f"Generate a complete project plan and starter files for: {desc}{_proj_ev_block}"
        return chat(prompt, skip_router=True)

    # ---- FIX_FILE ----
    if a == "FIX_FILE":
        from pathlib import Path as _Path
        import re as _re_ff
        path = str(args.get("path") or "").strip()
        path = _re_ff.sub(r'^file\s+', '', path, flags=_re_ff.IGNORECASE).strip()
        if not path:
            msg = "Path not found: missing path"
            return {"ok": False, "action": a, "error": msg, "content": msg, "response": msg}
        pp = _resolve_existing_user_or_artifact_path(path)
        if not pp.exists():
            msg = f"Path not found: {pp}"
            return {"ok": False, "action": a, "error": msg, "content": msg, "response": msg}
        original = pp.read_text(encoding='utf-8', errors='replace')
        extra_error = str(args.get("error") or args.get("stderr") or "").strip()
        ext_to_lang = {
            ".py": "Python", ".sh": "Bash", ".bash": "Bash", ".js": "JavaScript",
            ".ts": "TypeScript", ".rb": "Ruby", ".go": "Go", ".rs": "Rust",
            ".lua": "Lua", ".pl": "Perl", ".java": "Java", ".cpp": "C++",
            ".c": "C", ".cs": "C#", ".php": "PHP", ".swift": "Swift",
            ".kt": "Kotlin", ".r": "R", ".ps1": "PowerShell",
        }
        lang = ext_to_lang.get(pp.suffix.lower(), "the same language")
        err_block = f"\n\nReported error / stderr:\n{extra_error}" if extra_error else ""
        def _verify_python_apis_ff(_code: str) -> str | None:
            """Same as the GENERATE_SCRIPT path verifier — checks <module>.<attr> against installed libs."""
            import importlib as _il
            _ALIAS_DEFAULT = {
                "nx": "networkx", "np": "numpy", "pd": "pandas",
                "sp": "scipy", "sk": "sklearn", "sym": "sympy",
            }
            _SKIP_HEAVY = {"matplotlib", "matplotlib.pyplot", "torch", "tensorflow", "cv2"}
            _aliases = dict(_ALIAS_DEFAULT)
            for _m in _re_ff.finditer(r"^\s*import\s+([\w.]+)\s+as\s+(\w+)", _code, _re_ff.MULTILINE):
                _aliases[_m.group(2)] = _m.group(1)
            for _m in _re_ff.finditer(r"^\s*import\s+([\w.]+)\s*$", _code, _re_ff.MULTILINE):
                _aliases[_m.group(1).split(".")[0]] = _m.group(1)
            _bad = []
            for _alias, _modname in _aliases.items():
                if _modname in _SKIP_HEAVY:
                    continue
                if not _re_ff.search(rf"\b{_re_ff.escape(_alias)}\.[a-zA-Z_]", _code):
                    continue
                try:
                    _mod = _il.import_module(_modname)
                except Exception:
                    continue
                for _m in _re_ff.finditer(rf"\b{_re_ff.escape(_alias)}\.([a-zA-Z_]\w*)", _code):
                    _attr = _m.group(1)
                    if not hasattr(_mod, _attr):
                        ref = f"{_alias}.{_attr}"
                        if ref not in _bad:
                            _bad.append(ref)
            if _bad:
                return (
                    f"code references non-existent module APIs: {', '.join(_bad[:5])}. "
                    "Use only attributes that actually exist."
                )
            return None

        base_fix_prompt = (
            f"Fix the following {lang} file. Return ONLY the complete corrected source code — "
            "no markdown fences, no prose, no explanation, no diff. The entire output must be "
            "the new file content that can be saved verbatim and run.\n\n"
            "HARD REQUIREMENTS:\n"
            "- Use ONLY real library APIs that actually exist; verify function names\n"
            "- For networkx clique detection, do NOT use nx.is_clique or nx.graph_clique_number "
            "(neither exists). Instead implement via itertools.combinations + explicit edge checks: "
            "  `all(G.has_edge(u, v) for u, v in combinations(subset, 2))`\n"
            "- Replace any non-existent calls with working equivalents from the same library\n"
            "- Preserve the original intent and structure; minimal but complete fixes\n"
            "- Keep all original functionality working\n"
            "- No TODO stubs, no placeholders, no commentary lines explaining the fix\n"
            f"- Output must be valid {lang} that parses cleanly\n\n"
            f"File path: {pp}{err_block}\n\n"
            "=== ORIGINAL FILE CONTENT BEGINS ===\n"
            f"{original[:30000]}\n"
            "=== ORIGINAL FILE CONTENT ENDS ===\n\n"
            f"Output the fixed {lang} file content now (raw code only):"
        )

        fixed_code = ""
        ff_reject_reason = None
        for _ff_attempt in range(2):
            _ff_prompt = base_fix_prompt
            if _ff_attempt > 0 and ff_reject_reason:
                _ff_prompt = (
                    base_fix_prompt
                    + f"\n\nPREVIOUS ATTEMPT WAS REJECTED: {ff_reject_reason}. "
                    "Fix this specifically. Use only APIs that actually exist in the installed library."
                )
            _candidate = ""
            try:
                from eli.cognition import gguf_inference as _gguf_ff
                if _gguf_ff.load_model() is not None:
                    raw = _gguf_ff.chat_completion(
                        _ff_prompt,
                        system=(
                            f"You are an expert {lang} engineer fixing real code. "
                            "Output ONLY the corrected file content — no fences, no prose. "
                            "Use only library APIs that actually exist."
                        ),
                        max_tokens=8000,
                        temperature=0.15,
                        top_p=0.85,
                    )
                    if raw:
                        cand = _re_ff.sub(r"^```[a-z]*\n?", "", raw.strip(), flags=_re_ff.MULTILINE)
                        cand = _re_ff.sub(r"\n?```$", "", cand.strip(), flags=_re_ff.MULTILINE).strip()
                        _candidate = cand
            except Exception as _ff_e:
                log.debug(f"[FIX_FILE] GGUF path failed: {_ff_e}, falling back to chat")
            if not _candidate:
                chat_result = chat(_ff_prompt, skip_router=True)
                raw = (chat_result.get("content") or "").strip()
                cand = _re_ff.sub(r"^```[a-z]*\n?", "", raw, flags=_re_ff.MULTILINE)
                cand = _re_ff.sub(r"\n?```$", "", cand.strip(), flags=_re_ff.MULTILINE).strip()
                _candidate = cand
            if not _candidate or len(_candidate) < 20:
                ff_reject_reason = "model returned no usable code"
                continue
            if _candidate.strip() == original.strip():
                ff_reject_reason = "model returned the file unchanged"
                continue
            if pp.suffix.lower() == ".py":
                import ast as _ast_ff
                try:
                    _ast_ff.parse(_candidate)
                except SyntaxError as _se:
                    ff_reject_reason = f"corrected code has SyntaxError: {_se}"
                    continue
                _api_bad = _verify_python_apis_ff(_candidate)
                if _api_bad:
                    ff_reject_reason = _api_bad
                    log.debug(f"[FIX_FILE] attempt {_ff_attempt + 1} rejected: {_api_bad}")
                    continue
            fixed_code = _candidate
            break
        if not fixed_code:
            msg = f"Fix failed after retries: {ff_reject_reason or 'unknown reason'}"
            return {"ok": False, "action": a, "error": msg, "content": msg, "response": msg}
        try:
            _ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup = pp.with_suffix(pp.suffix + f".bak.{_ts}")
            backup.write_text(original, encoding="utf-8")
        except Exception as _bak_err:
            msg = f"FIX_FILE aborted: could not write backup for {pp}: {_bak_err}"
            return {"ok": False, "action": a, "error": msg, "content": msg, "response": msg}
        pp.write_text(fixed_code, encoding="utf-8")
        evt = {
            "event": "artifact_generated",
            "kind": "script",
            "path": str(pp),
            "filename": pp.name,
            "language": lang,
            "fixed": True,
            "backup": str(backup) if backup else None,
            "opened": True,
            "can_run": True,
        }
        evt_json = json.dumps(evt, ensure_ascii=False, default=str)
        return {
            "ok": True,
            "action": "FIX_FILE",
            "code": fixed_code,
            "script_path": str(pp),
            "path": str(pp),
            "filename": pp.name,
            "content": evt_json,
            "response": evt_json,
            "open_in_ide": True,
        }

    # ---- CREATE_FOLDER (new action) ----
    if a == "CREATE_FOLDER":
        import re as _re
        raw = (args.get("name") or args.get("path") or "").strip()
        if not raw:
            return {"ok": False, "error": "Missing folder name"}
        # If raw contains a full path, extract it; otherwise treat as name under ~
        path_match = _re.search(r"(/(?:home|tmp|var|opt|srv)[^ ]+)", raw)
        if path_match:
            path = str(resolve_user_repo_path(path_match.group(1)))
        elif raw.startswith("~") or raw.startswith("/"):
            path = str(resolve_user_repo_path(raw))
        else:
            # Strip any trailing natural language after the folder name
            clean = _re.split(r"\s+(?:not|inside|in|at|under)\s", raw)[0].strip()
            path = str(resolve_user_repo_path(f"~/{clean}"))
        try:
            os.makedirs(path, exist_ok=True)
            return {"ok": True, "content": f"Folder created at {path}", "response": f"Folder created at {path}"}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    # ---- OPEN_IDE / OPEN_IN_IDE / SHOW_DIFF ----
    if a == "OPEN_IDE":
        target = str(
            (args or {}).get("name")
            or (args or {}).get("app")
            or (args or {}).get("target")
            or (args or {}).get("path")
            or (args or {}).get("text")
            or (args or {}).get("query")
            or ""
        ).strip().lower()

        generic = {
            "", "ide", "the ide", "editor", "the editor",
            "built in ide", "built-in ide", "gui ide", "eli ide",
            "internal ide", "ide tab", "the ide tab"
        }

        alias_map = {
            "vscode": "code",
            "visual studio code": "code",
            "virtual studio code": "code",
            "vs code": "code",
            "code": "code",
            "gedit": "gedit",
            "text editor": "gedit",
            "codium": "codium",
            "kate": "kate",
            "sublime": "subl",
            "sublime text": "subl",
            "idle": "idle"
        }

        if target not in generic:
            mapped = alias_map.get(target, target)
            return _execute_impl("OPEN_APP", {"name": mapped})

        for editor in ["code", "codium", "subl", "gedit", "kate"]:
            try:
                subprocess.Popen(
                    [editor],
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                    start_new_session=True
                )
                msg = f"Opened IDE: {editor}"
                return {"ok": True, "action": a, "content": msg, "response": msg}
            except FileNotFoundError:
                continue

        return {
            "ok": False,
            "action": a,
            "error": "No IDE found",
            "content": "No IDE found",
            "response": "No IDE found"
        }

    if a == "OPEN_IN_IDE":
        path = str(args.get("path") or "").strip()
        for editor in ["code", "codium", "subl", "gedit", "kate"]:
            try:
                cmd = [editor] + ([path] if path else [])
                subprocess.Popen(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                msg = f"Opened {path or 'IDE'} in {editor}"
                return {"ok": True, "action": a, "content": msg, "response": msg}
            except FileNotFoundError:
                continue
        return {"ok": False, "action": a, "error": "No IDE found", "content": "No IDE found", "response": "No IDE found"}

    if a == "SHOW_DIFF":
        return chat("Show me the recent changes or diff in the project.", skip_router=True)

    # ---- MEMORY_STATS ----
    if a == "MEMORY_STATS":
        try:
            from eli.memory import get_memory
            mem = get_memory()
            conn = mem._get_connection()
            try:
                counts = {}
                for table in ["memories", "conversations", "observations", "failures"]:
                    try:
                        row = conn.execute(f"SELECT COUNT(*) FROM {table}").fetchone()
                        counts[table] = row[0] if row else 0
                    except Exception:
                        counts[table] = "N/A"
                msg = "Memory Stats:\n" + "\n".join(f"  {k}: {v}" for k, v in counts.items())
            finally:
                conn.close()
            # Append conversation log stats from the rotation module
            try:
                from eli.perception.log_rotation import convlog_stats
                log_info = convlog_stats()
                log_lines = [
                    f"  log_files: {log_info.get('file_count', 'N/A')}",
                    f"  log_total_mb: {log_info.get('total_mb', 'N/A')}",
                ]
                msg += "\nConversation Logs:\n" + "\n".join(log_lines)
                counts["log_stats"] = log_info
            except Exception:
                pass
            return {"ok": True, "action": a, "content": msg, "response": msg, "counts": counts}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}
        

    # ---- Window management (JARVIS-style screen control) ───────────────
    # Tile / arrange windows so all are visible side-by-side. Linux-first.
    if a == "TILE_WINDOWS":
        if shutil.which("wmctrl"):
            try:
                listing = subprocess.run(
                    ["wmctrl", "-l"], capture_output=True, text=True, timeout=4
                )
                window_ids = [line.split()[0] for line in (listing.stdout or "").splitlines() if line.strip()]
                count = len(window_ids)
                if count == 0:
                    msg = "No visible windows to tile."
                    return {"ok": False, "action": a, "content": msg, "response": msg}
                geom = subprocess.run(
                    ["xdpyinfo"], capture_output=True, text=True, timeout=4
                ).stdout if shutil.which("xdpyinfo") else ""
                m = re.search(r"dimensions:\s*(\d+)x(\d+)", geom or "")
                if m:
                    sw, sh = int(m.group(1)), int(m.group(2))
                else:
                    sw, sh = 1920, 1080
                cols = 2 if count <= 4 else 3
                rows = (count + cols - 1) // cols
                cw = sw // cols
                rh = sh // rows
                for i, wid in enumerate(window_ids):
                    row = i // cols
                    col = i % cols
                    x, y = col * cw, row * rh
                    subprocess.run([
                        "wmctrl", "-i", "-r", wid, "-b", "remove,maximized_vert,maximized_horz"
                    ], check=False, capture_output=True, timeout=4)
                    subprocess.run([
                        "wmctrl", "-i", "-r", wid, "-e", f"0,{x},{y},{cw},{rh}"
                    ], check=False, capture_output=True, timeout=4)
                msg = f"Tiled {count} window{'s' if count != 1 else ''} into a {cols}×{rows} grid."
                return {"ok": True, "action": a, "content": msg, "response": msg,
                        "count": count, "grid": [cols, rows]}
            except Exception as e:
                return {"ok": False, "action": a, "error": str(e),
                        "content": str(e), "response": str(e)}
        return {"ok": False, "action": a, "error": "wmctrl not installed",
                "content": "wmctrl not installed; cannot tile windows.",
                "response": "wmctrl not installed; cannot tile windows."}

    if a == "MINIMISE_ALL":
        if shutil.which("wmctrl"):
            ok = _run_ok(["wmctrl", "-k", "on"])
            msg = "Showed desktop." if ok else "Failed to minimise all windows."
            return {"ok": ok, "action": a, "content": msg, "response": msg}
        if shutil.which("xdotool"):
            ok = _run_ok(["xdotool", "key", "super+d"])
            msg = "Triggered show-desktop shortcut." if ok else "Failed to minimise all windows."
            return {"ok": ok, "action": a, "content": msg, "response": msg}
        return {"ok": False, "action": a, "error": "wmctrl/xdotool not installed",
                "content": "wmctrl or xdotool required to minimise all windows.",
                "response": "wmctrl or xdotool required to minimise all windows."}

    if a == "RESTORE_WINDOWS":
        if shutil.which("wmctrl"):
            ok = _run_ok(["wmctrl", "-k", "off"])
            msg = "Restored windows." if ok else "Failed to restore windows."
            return {"ok": ok, "action": a, "content": msg, "response": msg}
        return {"ok": False, "action": a, "error": "wmctrl not installed",
                "content": "wmctrl required to restore windows.",
                "response": "wmctrl required to restore windows."}

    if a == "MAXIMISE_WINDOW":
        if shutil.which("wmctrl"):
            ok = _run_ok(["wmctrl", "-r", ":ACTIVE:", "-b", "add,maximized_vert,maximized_horz"])
            msg = "Maximised current window." if ok else "Failed to maximise window."
            return {"ok": ok, "action": a, "content": msg, "response": msg}
        if shutil.which("xdotool"):
            ok = _run_ok(["xdotool", "key", "super+Up"])
            msg = "Triggered maximise shortcut." if ok else "Failed to maximise window."
            return {"ok": ok, "action": a, "content": msg, "response": msg}
        return {"ok": False, "action": a, "error": "wmctrl/xdotool not installed",
                "content": "wmctrl or xdotool required to maximise.",
                "response": "wmctrl or xdotool required to maximise."}

    if a == "NEXT_WINDOW":
        if shutil.which("xdotool"):
            ok = _run_ok(["xdotool", "key", "alt+Tab"])
            msg = "Switched to next window." if ok else "Failed to switch window."
            return {"ok": ok, "action": a, "content": msg, "response": msg}
        return {"ok": False, "action": a, "error": "xdotool not installed",
                "content": "xdotool required to cycle windows.",
                "response": "xdotool required to cycle windows."}

    if a == "PREVIOUS_WINDOW":
        if shutil.which("xdotool"):
            ok = _run_ok(["xdotool", "key", "alt+shift+Tab"])
            msg = "Switched to previous window." if ok else "Failed to switch window."
            return {"ok": ok, "action": a, "content": msg, "response": msg}
        return {"ok": False, "action": a, "error": "xdotool not installed",
                "content": "xdotool required to cycle windows.",
                "response": "xdotool required to cycle windows."}

    if a == "SWITCH_WORKSPACE":
        direction = str(args.get("direction") or "right").lower()
        if shutil.which("xdotool"):
            key = "ctrl+alt+Right" if direction == "right" else "ctrl+alt+Left"
            ok = _run_ok(["xdotool", "key", key])
            msg = f"Switched workspace {direction}." if ok else "Failed to switch workspace."
            return {"ok": ok, "action": a, "content": msg, "response": msg}
        return {"ok": False, "action": a, "error": "xdotool not installed",
                "content": "xdotool required to switch workspaces.",
                "response": "xdotool required to switch workspaces."}

    if a == "FOCUS_APP":
        name = str(args.get("name") or args.get("app") or "").strip()
        if not name:
            return {"ok": False, "action": a, "error": "missing app name",
                    "content": "Specify app to focus.", "response": "Specify app to focus."}
        if shutil.which("wmctrl"):
            ok = _run_ok(["wmctrl", "-a", name])
            msg = f"Focused {name}." if ok else f"Could not focus {name} — window not found."
            return {"ok": ok, "action": a, "content": msg, "response": msg}
        if shutil.which("xdotool"):
            try:
                proc = subprocess.run(
                    ["xdotool", "search", "--name", name],
                    capture_output=True, text=True, timeout=4
                )
                ids = [w for w in (proc.stdout or "").split() if w.strip()]
                if ids:
                    ok = _run_ok(["xdotool", "windowactivate", ids[0]])
                    msg = f"Focused {name}." if ok else f"Could not focus {name}."
                    return {"ok": ok, "action": a, "content": msg, "response": msg}
            except Exception as e:
                return {"ok": False, "action": a, "error": str(e),
                        "content": str(e), "response": str(e)}
        return {"ok": False, "action": a, "error": "wmctrl/xdotool not installed",
                "content": "wmctrl or xdotool required to focus apps.",
                "response": "wmctrl or xdotool required to focus apps."}

    # ---- SCREEN_LOCATE — find/click visible UI text via OCR ─────────────
    if a == "SCREEN_LOCATE":
        query = str(args.get("query") or "").strip()
        click = bool(args.get("click"))
        if not query:
            return {"ok": False, "action": a, "error": "missing query",
                    "content": "Specify what to find on screen.",
                    "response": "Specify what to find on screen."}
        try:
            from eli.perception.screen_locator import locate_on_screen
            result = locate_on_screen(query=query, click=click)
            ok = bool(result.get("ok") or result.get("matches"))
            matches = result.get("matches") or []
            if matches:
                first = matches[0]
                if click:
                    msg = f"Clicked '{query}' at ({first.get('cx')}, {first.get('cy')})."
                else:
                    msg = f"Found '{query}' at ({first.get('cx')}, {first.get('cy')})."
            else:
                msg = f"Could not find '{query}' on the screen."
                ok = False
            payload = dict(result)
            payload.update({"ok": ok, "action": a, "content": msg, "response": msg})
            return payload
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e),
                    "content": f"Screen locator failed: {e}",
                    "response": f"Screen locator failed: {e}"}

    # ---- CLOSE_APP ----
    if a == "CLOSE_APP":
        name = str(args.get("name") or args.get("app") or "").strip()
        if not name:
            return {"ok": False, "action": a, "error": "missing app name", "content": "Specify app to close.", "response": "Specify app to close."}
        tried = []
        for cmd in [["wmctrl", "-c", name], ["pkill", "-f", name], ["pkill", name], ["killall", name]]:
            if shutil.which(cmd[0]):
                tried.append(cmd[0])
                if _run_ok(cmd):
                    msg = f"Closed {name} via {cmd[0]}."
                    return {"ok": True, "action": a, "content": msg, "response": msg}
        msg = f"Could not close {name}. Tried: {', '.join(tried) or 'no tools found'}."
        return {"ok": False, "action": a, "error": msg, "content": msg, "response": msg}

    # ---- ANALYZE_IMAGE (metadata + OCR; never confabulate visual content) ----
    if a == "ANALYZE_IMAGE":
        import os as _ai_os
        from pathlib import Path as _AIPP
        path = str(args.get("path") or args.get("file") or "").strip()
        if not path:
            msg = "Specify an image file to analyze."
            return {"ok": False, "action": a, "error": "missing path", "content": msg, "response": msg}
        expanded = _ai_os.path.expanduser(path)
        if not _ai_os.path.exists(expanded):
            msg = f"Image not found: {expanded}"
            return {"ok": False, "action": a, "error": msg, "content": msg, "response": msg}
        try:
            from eli.perception.analyze_image import analyze_image_file
        except Exception as _ai_err:
            msg = f"Image analysis is unavailable: {_ai_err}"
            return {"ok": False, "action": a, "error": msg, "content": msg, "response": msg}
        _ai_stem = _AIPP(expanded).stem
        _ai_out = str(get_paths().artifacts_dir / f"image_report_{_ai_stem}.md")
        try:
            rep = analyze_image_file(expanded, _ai_out, ocr=True)
        except Exception as _ai_err:
            msg = f"Image analysis failed: {_ai_err}"
            return {"ok": False, "action": a, "error": msg, "content": msg, "response": msg}
        if not rep.get("ok"):
            msg = rep.get("error") or "Image analysis failed."
            return {"ok": False, "action": a, "error": msg, "content": msg, "response": msg}
        _ai_w, _ai_h = rep.get("width"), rep.get("height")
        _ai_dims = f"{_ai_w}×{_ai_h} px" if _ai_w and _ai_h else "unknown dimensions"
        _ai_fmt = rep.get("format", "image")
        _ai_size = int(rep.get("size_bytes") or 0)
        _ai_ocr = (rep.get("ocr_text") or "").strip()
        _ai_name = _AIPP(expanded).name
        _ai_prompt = str(args.get("prompt") or args.get("instruction") or "").strip() or None
        _ai_prefer_fast = bool(args.get("prefer_fast"))

        # Real vision first (local VL model). prefer_fast → small Moondream model
        # for quick glances; falls back to the primary model.
        _ai_vision_text = ""
        _ai_vision_err = ""
        try:
            from eli.perception import vision as _eli_vision
            _va_ok, _va_reason = _eli_vision.vision_available()
            if not _va_ok and _ai_prefer_fast:
                _va_ok, _va_reason = _eli_vision.fast_vision_available()
            if _va_ok:
                _vres = _eli_vision.describe_image(expanded, prompt=_ai_prompt, prefer_fast=_ai_prefer_fast)
                if _vres.get("ok"):
                    _ai_vision_text = str(_vres.get("text") or "").strip()
                else:
                    _ai_vision_err = str(_vres.get("error") or "")
            else:
                _ai_vision_err = _va_reason
        except Exception as _ai_vexc:
            _ai_vision_err = str(_ai_vexc)

        # Fuse OCR (exact text) + the vision gist into an accurate, grounded
        # description via the text model — compensates for the small vision
        # model's inability to read dense UI text. Strictly evidence-only.
        _ai_fused = ""
        try:
            from eli.core import config as _ai_cfg
            _fuse_on = bool(_ai_cfg.get("vision_fuse_with_text_model", True))
        except Exception:
            _fuse_on = True
        if _fuse_on and (_ai_vision_text or _ai_ocr):
            try:
                from eli.cognition import gguf_inference as _fgi
                _ocr_clip = (_ai_ocr or "")[:2500]
                _vis_clip = (_ai_vision_text or "(no visual description available)")[:1200]
                _fuse_user = (
                    f"OCR text read from the screen (authoritative for names/titles):\n{_ocr_clip}\n\n"
                    f"Vision model's visual description:\n{_vis_clip}\n\n"
                    f"Using ONLY the evidence above, write 2-4 sentences on what application/"
                    f"content is shown and what the user appears to be doing. Prefer the OCR text "
                    f"for exact names. Do not invent anything not present in the evidence."
                )
                _ai_fused = (_fgi.chat_completion(
                    _fuse_user,
                    system=("You describe what is on the user's screen strictly from the provided "
                            "OCR + vision evidence. Never invent apps, text, or activities. Be concise."),
                    max_tokens=300, temperature=0.3,
                ) or "").strip()
            except Exception as _fuse_err:
                log.debug(f"[ANALYZE_IMAGE] fusion skipped: {_fuse_err}")
                _ai_fused = ""

        _ai_header = f"**{_ai_name}** ({_ai_fmt}, {_ai_dims}, {_ai_size:,} bytes)"
        if _ai_fused:
            body = _ai_fused
        elif _ai_vision_text:
            body = f"Looking at {_ai_header}:\n\n{_ai_vision_text}"
            if _ai_ocr:
                body += f"\n\n---\nText I can read in it (OCR):\n\n{_ai_ocr}"
        elif _ai_ocr:
            body = (
                f"I ran OCR on {_ai_header}. Here is the text I extracted from it:\n\n{_ai_ocr}"
            )
            if _ai_vision_err:
                body += (
                    f"\n\n(I couldn't run full visual analysis — {_ai_vision_err} — "
                    f"so this is text only.)"
                )
        else:
            body = (
                f"I examined {_ai_header}. No readable text was detected via OCR"
            )
            if _ai_vision_err:
                body += (
                    f", and I couldn't run visual analysis ({_ai_vision_err}). "
                    f"I won't guess at what it shows."
                )
            else:
                body += ". I won't guess at what it shows."
        return {"ok": True, "action": a, "path": expanded, "ocr_text": _ai_ocr,
                "vision_text": _ai_vision_text, "fused_text": _ai_fused, "report": rep,
                "content": body, "response": body}

    # ---- SUMMARIZE_FILE ----
    if a == "SUMMARIZE_FILE":
        import os as _os2, subprocess as _sfsp
        from pathlib import Path as _SFPP
        path = str(args.get("path") or args.get("file") or "").strip()
        _sf_instruction = str(args.get("instruction") or "").strip()
        if not path:
            return {"ok": False, "action": a, "error": "missing path",
                    "content": "Specify file to summarize.", "response": "Specify file to summarize."}
        expanded = _os2.path.expanduser(path)
        if not _os2.path.exists(expanded):
            return {"ok": False, "action": a, "error": f"Path not found: {expanded}",
                    "content": f"Path not found: {expanded}", "response": f"Path not found: {expanded}"}
        # Images are binary — never read them as text. Delegate to OCR-based
        # image analysis so we describe them honestly instead of dumping bytes.
        if expanded.lower().endswith(
                (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".tiff", ".tif")):
            return _execute_impl("ANALYZE_IMAGE", {"path": expanded})
        # Directory: if it contains PDFs, delegate to ANALYZE_PDF_FOLDER for real content
        if _os2.path.isdir(expanded):
            import glob as _glob
            _pdfs = _glob.glob(_os2.path.join(expanded, "**", "*.pdf"), recursive=True)
            if _pdfs:
                return _execute_impl("ANALYZE_PDF_FOLDER", {"folder": expanded, "recursive": True})
            # No PDFs: fall back to ls -R + GGUF description
            try:
                _ls_result = _sfsp.run(
                    ["ls", "-R", expanded], capture_output=True, text=True, timeout=10)
                _sf_raw = _ls_result.stdout[:8000]
            except Exception as _dir_err:
                _sf_raw = f"Directory listing failed: {_dir_err}"
            _sf_summary = None
            try:
                from eli.cognition import gguf_inference as _sfgi
                if _sfgi.load_model() is not None:
                    _sf_user_msg = (
                        f"{_sf_instruction}\n\nDirectory listing:\n\n{_sf_raw[:6000]}"
                        if _sf_instruction else
                        f"Describe the structure and contents of this directory tree:\n\n{_sf_raw[:6000]}"
                    )
                    _sf_summary = _sfgi.chat_completion(
                        _sf_user_msg,
                        system="You are a precise technical analyst. Describe the directory structure, key files, and inferred purpose.",
                        max_tokens=700, temperature=0.3)
            except Exception:
                pass
            _sf_body = _sf_summary if _sf_summary else _sf_raw[:3000]
            return {"ok": True, "action": a, "content": _sf_body, "response": _sf_body}
        try:
            with open(expanded, errors="replace") as _sf:
                _sf_raw = _sf.read(8000)
            _sf_summary = None
            try:
                from eli.cognition import gguf_inference as _sfgi
                if _sfgi.load_model() is not None:
                    _sf_user_msg = (
                        f"{_sf_instruction}\n\nFile content:\n\n{_sf_raw[:6000]}"
                        if _sf_instruction else
                        f"Summarize the following file content:\n\n{_sf_raw[:6000]}"
                    )
                    _sf_summary = _sfgi.chat_completion(
                        _sf_user_msg,
                        system="You are a precise technical analyst. Summarize key points, structure, and purpose.",
                        max_tokens=700, temperature=0.3)
            except Exception:
                pass
            _sf_stem = _SFPP(expanded).stem
            _sf_size = _os2.path.getsize(expanded)
            _sf_body = _sf_summary if _sf_summary else _sf_raw[:3000]
            _sf_doc_content = (
                f"# Summary: {_sf_stem}\n\n"
                f"**Source:** {expanded}\n"
                f"**Size:** {_sf_size:,} bytes\n\n---\n\n{_sf_body}"
            )
            _sf_saved = _save_artifact(_sf_doc_content, "documents", f"{_sf_stem}_summary", fmt="docx")
            try:
                _sfsp.Popen(["xdg-open", _sf_saved], stdout=_sfsp.DEVNULL, stderr=_sfsp.DEVNULL)
                _sf_opened = True
            except Exception:
                _sf_opened = False
            _sf_chat = (
                f"Document compiled: **{_SFPP(_sf_saved).name}**\n"
                f"Saved to: `{_sf_saved}`"
                + ("\n_(opening…)_" if _sf_opened else "")
            )
            return {"ok": True, "action": a, "content": _sf_chat, "response": _sf_chat,
                    "saved_to": _sf_saved}
        except Exception as _se:
            return {"ok": False, "action": a, "error": str(_se), "content": str(_se), "response": str(_se)}

    if a == "SEQUENCE":
        steps = args.get("steps") or []
        if not isinstance(steps, list):
            msg = "SEQUENCE steps must be a list."
            return {"ok": False, "action": a, "error": msg, "content": msg, "response": msg, "steps": []}

        results = []
        overall_ok = True

        for i, step in enumerate(steps):
            if not isinstance(step, dict):
                msg = f"step {i} is not a dict"
                results.append({"ok": False, "action": "SEQUENCE_STEP", "error": msg, "content": msg, "response": msg})
                overall_ok = False
                continue

            sa = str(step.get("action") or "").strip()
            sargs = step.get("args") or {}

            try:
                r = _execute_impl(sa, sargs)
            except Exception as e:
                r = {"ok": False, "action": sa, "error": str(e), "content": str(e), "response": str(e)}

            if not isinstance(r, dict):
                r = {"ok": False, "action": sa, "error": "non-dict step result", "content": "non-dict step result", "response": "non-dict step result"}

            if r.get("ok") is not True:
                overall_ok = False
            results.append(r)

        msg = f"Executed {len(results)} step(s)." if overall_ok else f"Executed {len(results)} step(s) with failures."
        return {
            "ok": overall_ok,
            "action": a,
            "steps": results,
            "content": msg,
            "response": msg,
        }


    if a == 'SELF_UPGRADE':
        request = (args or {}).get('request', '')
        try:
            from eli.kernel.self_upgrade import SelfUpgrader
            inst = SelfUpgrader()
            result = inst.upgrade(request)
            # If the upgrade request also named files / asked to check/examine/
            # review, run the code examiner and append a tiered report + fix offer.
            try:
                from eli.runtime import code_examiner as _ce
                _low = str(request or "").lower()
                _wants_examine = bool(
                    _ce._extract_named_paths(request)
                    or re.search(r"\b(examine|review|inspect|check|scan|audit)\b", _low)
                )
                if _wants_examine:
                    _paths = _ce.resolve_targets(request)
                    if _paths:
                        _findings = _ce.examine(_paths)
                        _ce.set_pending_fix(_findings, _paths)
                        result = result + "\n\n— Code examination —\n" + \
                            _ce.format_report(_paths, _findings)
            except Exception as _ex_exc:
                log.debug(f"[SELF_UPGRADE] examine stage skipped: {_ex_exc}")
            return {'ok': True, 'action': a, 'content': result, 'response': result}
        except Exception as _exc:
            return {'ok': False, 'action': a, 'error': str(_exc), 'content': str(_exc), 'response': str(_exc)}

    # ---- EXAMINE_CODE / CONFIRM_CODE_FIX / CANCEL_CODE_FIX ----
    if a == 'EXAMINE_CODE':
        request = (args or {}).get('request') or (args or {}).get('query') or ''
        try:
            from eli.runtime import code_examiner as _ce
            paths = _ce.resolve_targets(request)
            if not paths:
                msg = ("I couldn't resolve any Python files to examine. Name a file "
                       "(e.g. 'examine eli/memory/memory.py') or a module.")
                return {'ok': True, 'action': a, 'content': msg, 'response': msg,
                        'evidence_source': 'code_examiner_no_targets'}
            findings = _ce.examine(paths)
            _ce.set_pending_fix(findings, paths)
            report = _ce.format_report(paths, findings)
            return {'ok': True, 'action': a, 'content': report, 'response': report,
                    'evidence_source': 'code_examiner'}
        except Exception as _exc:
            msg = f"Code examination failed: {_exc}"
            return {'ok': False, 'action': a, 'error': str(_exc), 'content': msg, 'response': msg}

    if a in ('CONFIRM_HABIT', 'DECLINE_HABIT'):
        try:
            from eli.planning.habits import get_pending_habit, clear_pending_habit
            from eli.memory import get_memory as _get_mem
        except Exception as _imp_exc:
            msg = f"Habit confirm unavailable: {_imp_exc}"
            return {'ok': False, 'action': a, 'error': str(_imp_exc), 'content': msg, 'response': msg}
        pending = get_pending_habit()
        if not pending:
            msg = "There's no habit offer waiting for a yes/no right now."
            return {'ok': True, 'action': a, 'content': msg, 'response': msg}
        rid = int(pending.get('rule_id', -1))
        name = pending.get('name', 'that habit')
        hh, mm = int(pending.get('hour', 0)), int(pending.get('minute', 0))
        mem = _get_mem()
        try:
            if a == 'CONFIRM_HABIT':
                ok = mem.set_habit_rule_enabled(rid, True) if hasattr(mem, 'set_habit_rule_enabled') else False
                clear_pending_habit()
                msg = (f"Done — “{name}” is now an active habit; I'll run it around "
                       f"{hh:02d}:{mm:02d}. You can edit or disable it any time in the "
                       f"Habits tab.") if ok else \
                      f"I couldn't enable that habit (it may have been removed). Nothing activated."
                return {'ok': bool(ok), 'action': a, 'content': msg, 'response': msg,
                        'evidence_source': 'habit_confirm'}
            else:  # DECLINE_HABIT — remove the suggestion so it doesn't linger
                if hasattr(mem, 'delete_habit_rule'):
                    try:
                        mem.delete_habit_rule(rid)
                    except Exception:
                        pass
                clear_pending_habit()
                msg = f"No problem — I won't add “{name}”. I'll stop suggesting it."
                return {'ok': True, 'action': a, 'content': msg, 'response': msg,
                        'evidence_source': 'habit_decline'}
        except Exception as _hx:
            clear_pending_habit()
            msg = f"Habit update failed: {_hx}"
            return {'ok': False, 'action': a, 'error': str(_hx), 'content': msg, 'response': msg}

    if a == 'CANCEL_CODE_FIX':
        try:
            from eli.runtime import code_examiner as _ce
            _ce.clear_pending_fix()
        except Exception:
            pass
        msg = "Okay — I won't change any code. The examination findings are cleared."
        return {'ok': True, 'action': a, 'content': msg, 'response': msg}

    if a == 'CONFIRM_CODE_FIX':
        try:
            from eli.runtime import code_examiner as _ce
            from eli.runtime.self_improvement import get_self_improvement
        except Exception as _imp_exc:
            msg = f"Fix engine unavailable: {_imp_exc}"
            return {'ok': False, 'action': a, 'error': str(_imp_exc), 'content': msg, 'response': msg}

        pending = _ce.get_pending_fix()
        if not pending:
            msg = ("There's no pending examination to fix. Run 'examine <file> for "
                   "errors' first.")
            return {'ok': True, 'action': a, 'content': msg, 'response': msg}

        # Tier-3 (logic) findings only get patched if the user explicitly opted in.
        _msg_low = str((args or {}).get('message') or '').lower()
        include_tier3 = bool(re.search(r"\b(logic|tier ?3|all|everything|including)\b", _msg_low))

        findings = pending.get('findings') or []
        confirmed = [f for f in findings
                     if (not f.get('needs_confirmation')) or include_tier3]
        skipped_tier3 = [f for f in findings if f.get('needs_confirmation') and not include_tier3]

        engine = get_self_improvement()
        steps = []
        for f in confirmed:
            target = f.get('file')
            patch = _ce.generate_fix_patch(f)
            if not patch.get('ok'):
                steps.append({'file': target, 'status': 'skipped',
                              'message': patch.get('error', 'no safe patch generated')})
                continue
            applied = engine.apply_code_patch(patch, verify=True)
            status = 'applied' if applied.get('applied') else (
                'reverted' if 'revert' in str(applied.get('message', '')).lower() else 'skipped')
            steps.append({'file': target, 'status': status,
                          'message': applied.get('message', '')})

        _ce.clear_pending_fix()

        lines = ["Code-fix cycle complete."]
        n_applied = len([s for s in steps if s['status'] == 'applied'])
        lines.append(f"- patches_applied: {n_applied}/{len(confirmed)}")
        for s in steps:
            lines.append(f"  - [{s['status']}] {s['file']}: {s['message'][:160]}")
        if skipped_tier3:
            lines.append(f"- left {len(skipped_tier3)} low-confidence (Tier-3) finding(s) "
                         "untouched (say 'fix the logic ones too' to include them).")
        if n_applied:
            lines.append("Each applied patch was syntax-checked and import-verified; "
                         "any that broke its module was auto-reverted.")
        msg = "\n".join(lines)
        return {'ok': True, 'action': a, 'content': msg, 'response': msg,
                'evidence_source': 'code_examiner_fix'}

    # ---- NEWS_FETCH ----
    if a == 'NEWS_FETCH':
        topic   = (args or {}).get('topic', '')
        sources = (args or {}).get('sources', ['all'])
        query   = (args or {}).get('query', '')
        mode    = (args or {}).get('mode', 'fetch_and_show')  # fetch_and_show | fetch | search | recent | stats
        try:
            from eli.tools.news.news_fetcher import NewsFetcher
            fetcher = NewsFetcher()

            def _format_articles(articles, header="", with_summaries=False):
                if not articles:
                    return None
                from datetime import datetime as _dt, date as _date_cls
                _today = _date_cls.today()
                lines = [header] if header else []
                for r in articles:
                    fetched_raw = r.get('fetched_at')
                    pub_raw = (r.get('published') or '')[:10]
                    # Show fetch time (HH:MM) for today's articles; date otherwise.
                    # This prevents yesterday's RSS publication dates appearing as
                    # "current" news when ELI fetched them today.
                    if fetched_raw:
                        try:
                            fetched_dt = _dt.fromtimestamp(float(fetched_raw))
                            if fetched_dt.date() == _today:
                                date_str = f" (fetched {fetched_dt.strftime('%H:%M')})"
                            else:
                                # Article is from a prior fetch — show the fetch date
                                date_str = f" (fetched {fetched_dt.strftime('%d %b')})"
                        except Exception:
                            date_str = f" ({pub_raw})" if pub_raw else ""
                    elif pub_raw:
                        date_str = f" ({pub_raw})"
                    else:
                        date_str = ""
                    lines.append(f"• [{r['source']}]{date_str} {r['title']}")
                    if with_summaries and r.get('summary') and len(r['summary'].strip()) > 20:
                        lines.append(f"  {r['summary'][:160].strip()}…")
                return "\n".join(lines)

            if mode == 'search' and query:
                results = fetcher.search(query, limit=10)
                if not results:
                    msg = f"No stored news found for '{query}'. Try: 'fetch latest news' first."
                else:
                    msg = _format_articles(results, f"News results for '{query}':", with_summaries=True)
                return {'ok': True, 'action': a, 'content': msg, 'response': msg}

            elif mode == 'recent':
                # Auto-refresh if last fetch was more than 6 hours ago
                import time as _news_time
                st_info = fetcher.stats()
                _last_ts_str = st_info.get('last_fetched') or ''
                _stale = True
                if _last_ts_str and _last_ts_str != 'never':
                    try:
                        import time as _nt
                        _last_ts = _nt.mktime(_nt.strptime(_last_ts_str, "%Y-%m-%d %H:%M:%S"))
                        _stale = (_news_time.time() - _last_ts) > 6 * 3600
                    except Exception:
                        _stale = True
                if _stale:
                    fetcher.fetch(sources=None, topic=topic)
                results = fetcher.get_recent(limit=12, category=topic)
                if not results:
                    msg = "No news stored yet — fetching now…"
                    res = fetcher.fetch(sources=None, topic=topic)
                    results = fetcher.get_recent(limit=12, category=topic)
                    if not results:
                        msg = f"Fetched {res['fetched']} articles but none stored. Check network."
                        return {'ok': True, 'action': a, 'content': msg, 'response': msg}
                msg = _format_articles(results, f"Recent headlines ({len(results)}):")
                return {'ok': True, 'action': a, 'content': msg, 'response': msg}

            elif mode == 'stats':
                st = fetcher.stats()
                msg = (f"News DB: {st['total_articles']} articles stored, "
                       f"last fetched {st['last_fetched']}.\n"
                       + "\n".join(f"  {src}: {cnt}" for src, cnt in st['by_source'].items()))
                return {'ok': True, 'action': a, 'content': msg, 'response': msg}

            else:  # fetch_and_show (default) or plain fetch
                _src = sources if sources and sources != ['all'] else None
                res = fetcher.fetch(sources=_src, topic=topic)
                errs = res.get('errors', [])
                _fetched = res['fetched']
                _new = res['stored_new']
                _new_str = f", {_new} new" if _new else ""
                header = f"Here are the latest headlines ({_fetched} fetched{_new_str}):"
                if _fetched == 0 and errs:
                    msg = f"News fetch failed. Errors: {'; '.join(errs[:3])[:300]}"
                    return {'ok': False, 'action': a, 'content': msg, 'response': msg}
                # Conversational news read: ELI synthesises a bounded briefing
                # (a sentence of context each, 1-2 follow-ups, no timestamps)
                # rather than dumping raw headlines. General ask = 50/50 top
                # stories + interest matches; a topic ask = that topic's stories.
                # Falls back to the raw list if synthesis is empty (offline/none).
                try:
                    from eli.tools.news.news_synthesis import synthesise_news_briefing
                    _read = synthesise_news_briefing(topic=topic, refresh=False)
                except Exception:
                    _read = ""
                if _read:
                    return {'ok': True, 'action': a, 'content': _read,
                            'response': _read,
                            'meta': {'response_mode': 'news_briefing'}}
                recent = fetcher.get_recent(limit=10, category=topic)
                msg = _format_articles(recent, header)
                if not msg and _fetched > 0:
                    try:
                        from eli.tools.news.news_synthesis import synthesise_window as _synth_w
                        _sw = _synth_w()
                        msg = (_sw.get('digest') or _sw.get('summary') or '').strip() or None
                    except Exception:
                        msg = None
                    if not msg:
                        recent2 = fetcher.get_recent(limit=20)
                        msg = _format_articles(recent2, header)
                    msg = msg or f"Fetched {_fetched} articles — ask 'morning report' for a full digest."
                elif not msg:
                    msg = f"Fetched {_fetched} articles."
                if errs:
                    msg += f"\n\n⚠ Some sources failed: {'; '.join(errs[:2])[:200]}"
                return {'ok': True, 'action': a, 'content': msg, 'response': msg}

        except Exception as _exc:
            return {'ok': False, 'action': a, 'error': str(_exc), 'content': str(_exc), 'response': str(_exc)}

    if a == 'EXECUTE_GOAL':
        goal = (args or {}).get('goal', '')
        try:
            from eli.runtime.pipeline_models import RouteDecision as _RD
            from eli.execution.execution_planner import build_execution_plan as _build_plan
            _plan = _build_plan(_RD(user_input=str(goal or ''), action='EXECUTE_GOAL'))
            _steps = [f"{i + 1}. {s.name} ({s.kind})" for i, s in enumerate(_plan.steps)]
            _content = (
                "Execution plan for goal:\n" + "\n".join(_steps)
                if _steps else "No plan steps produced for goal."
            )
            return {'ok': True, 'action': a, 'content': _content, 'response': _content, 'plan': _plan.to_dict()}
        except Exception as _exc:
            return {'ok': False, 'action': a, 'error': str(_exc), 'content': str(_exc), 'response': str(_exc)}

    # ---- CPU_USAGE ----
    if a == "CPU_USAGE":
        try:
            import psutil
            pct = psutil.cpu_percent(interval=0.5)
            count = psutil.cpu_count()
            freq = psutil.cpu_freq()
            freq_str = f" @ {freq.current:.0f} MHz" if freq else ""
            msg = f"CPU usage: {pct:.1f}%  ({count} cores{freq_str})"
            return {"ok": True, "action": a, "percent": pct, "cores": count, "content": msg, "response": msg}
        except ImportError:
            import subprocess as _sp
            try:
                out = _sp.check_output(["top", "-bn1"], text=True, timeout=5)
                cpu_line = next((l for l in out.splitlines() if "Cpu(s)" in l or "%Cpu" in l), "")
                msg = f"CPU: {cpu_line.strip()}" if cpu_line else "CPU info unavailable (install psutil)"
            except Exception:
                msg = "CPU info unavailable (install psutil: pip install psutil)"
            return {"ok": True, "action": a, "content": msg, "response": msg}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- RAM_USAGE ----
    if a == "RAM_USAGE":
        try:
            import psutil
            vm = psutil.virtual_memory()
            used_gb = vm.used / (1024 ** 3)
            total_gb = vm.total / (1024 ** 3)
            msg = f"RAM usage: {vm.percent:.1f}%  ({used_gb:.1f} GB used / {total_gb:.1f} GB total)"
            return {"ok": True, "action": a, "percent": vm.percent,
                    "used_gb": round(used_gb, 2), "total_gb": round(total_gb, 2),
                    "content": msg, "response": msg}
        except ImportError:
            msg = "RAM info unavailable (install psutil: pip install psutil)"
            return {"ok": True, "action": a, "content": msg, "response": msg}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- SYSTEM_STATS ----
    if a == "SYSTEM_STATS":
        try:
            import psutil, platform as _platform
            cpu_pct = psutil.cpu_percent(interval=0.5)
            vm = psutil.virtual_memory()
            disk = psutil.disk_usage("/")
            used_ram = vm.used / (1024 ** 3)
            total_ram = vm.total / (1024 ** 3)
            used_disk = disk.used / (1024 ** 3)
            total_disk = disk.total / (1024 ** 3)
            lines = [
                f"System: {_platform.system()} {_platform.release()}",
                f"CPU:    {cpu_pct:.1f}%  ({psutil.cpu_count()} cores)",
                f"RAM:    {vm.percent:.1f}%  ({used_ram:.1f}/{total_ram:.1f} GB)",
                f"Disk:   {disk.percent:.1f}%  ({used_disk:.1f}/{total_disk:.1f} GB)",
            ]
            try:
                bat = psutil.sensors_battery()
                if bat:
                    lines.append(f"Battery: {bat.percent:.0f}%  ({'charging' if bat.power_plugged else 'on battery'})")
            except Exception:
                pass
            msg = "\n".join(lines)
            return {"ok": True, "action": a, "content": msg, "response": msg,
                    "cpu_percent": cpu_pct, "ram_percent": vm.percent, "disk_percent": disk.percent}
        except ImportError:
            msg = "System stats unavailable (install psutil: pip install psutil)"
            return {"ok": True, "action": a, "content": msg, "response": msg}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- WEB_SEARCH ----
    if a == "WEB_SEARCH":
        try:
            query = (args.get("query") or args.get("text") or args.get("message") or "").strip()
            if not query:
                return {"ok": False, "action": a, "error": "No query provided",
                        "content": "Provide a search query.", "response": "Provide a search query."}

            # Network gate: never open a browser or confabulate when offline.
            try:
                from eli.core.config import network_allowed
                _net = network_allowed()
            except Exception:
                _net = False
            if not _net:
                msg = ("I can't search the web right now — network access is off. "
                       "Turn on the Net toggle (or set network_enabled) and ask again "
                       "and I'll look it up. I won't guess at current facts I can't verify.")
                return {"ok": False, "action": a, "query": query, "offline": True,
                        "content": msg, "response": msg}

            # Live DuckDuckGo text fetch — return real snippets as grounding,
            # NOT a browser window. The model answers from these, not its priors.
            results = []
            try:
                from eli.plugins.web.plugin import _web_search_results
                results = _web_search_results(query, max_results=5) or []
            except Exception:
                results = []
            # Supplement with the local news store if the live fetch was empty.
            if not results:
                try:
                    from eli.tools.news.news_fetcher import search_stored_news
                    for h in (search_stored_news(query, limit=5) or []):
                        results.append({"title": h.get("title", ""), "href": h.get("url", ""), "body": ""})
                except Exception:
                    pass

            if not results:
                msg = (f"I searched the web for \"{query}\" but the search returned no usable "
                       f"results just now (the search backend may be rate-limited). I won't "
                       f"guess at an answer I can't verify — try rephrasing, or I can retry.")
                return {"ok": True, "action": a, "query": query, "results": [],
                        "content": msg, "response": msg}

            lines = [f"Live web results for \"{query}\" (authoritative — answer from these, not prior knowledge):"]
            for i, item in enumerate(results, 1):
                title = (item.get("title") or item.get("name") or "Untitled").strip()
                href = (item.get("href") or item.get("url") or "").strip()
                body = (item.get("body") or item.get("snippet") or "").strip()
                entry = f"{i}. {title}"
                if body:
                    entry += f" — {body[:220]}"
                if href:
                    entry += f"\n   {href}"
                lines.append(entry)
            msg = "\n".join(lines)
            return {"ok": True, "action": a, "query": query, "results": results,
                    "content": msg, "response": msg, "web_grounded": True}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- SPEAK ----
    if a == "SPEAK":
        try:
            text = (args.get("text") or args.get("message") or args.get("content") or "").strip()
            if not text:
                return {"ok": False, "action": a, "error": "No text to speak",
                        "content": "Provide text to speak.", "response": "Provide text to speak."}
            from eli.perception.tts_router import maybe_speak
            maybe_speak(text, enabled=True)
            msg = f"Speaking: {text[:80]}"
            return {"ok": True, "action": a, "content": msg, "response": msg}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- SMART_HOME ----
    if a == "SMART_HOME":
        try:
            device = (args.get("device") or args.get("name") or "").strip()
            command = (args.get("command") or args.get("action") or args.get("state") or "").strip()
            # Try Home Assistant if configured
            ha_url = os.environ.get("ELI_HA_URL", "").rstrip("/")
            ha_token = os.environ.get("ELI_HA_TOKEN", "")
            if ha_url and ha_token:
                import urllib.request as _ur, json as _j
                entity_id = device.replace(" ", "_").lower()
                svc = "turn_on" if "on" in command.lower() else "turn_off" if "off" in command.lower() else command
                payload = json.dumps({"entity_id": entity_id}).encode()
                req = _ur.Request(
                    f"{ha_url}/api/services/homeassistant/{svc}",
                    data=payload,
                    headers={"Authorization": f"Bearer {ha_token}", "Content-Type": "application/json"},
                )
                with _ur.urlopen(req, timeout=10) as resp:
                    resp.read()
                msg = f"Smart home: {svc} → {device}"
                return {"ok": True, "action": a, "content": msg, "response": msg}
            msg = ("Smart home control requires Home Assistant. "
                   "Set ELI_HA_URL and ELI_HA_TOKEN environment variables.")
            return {"ok": False, "action": a, "error": "not_configured", "content": msg, "response": msg}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- NOTE ACTIONS (LIST_NOTES / NEW_NOTE / SEARCH_NOTES) ----
    if a in ("LIST_NOTES", "NEW_NOTE", "SEARCH_NOTES"):
        try:
            from pathlib import Path as _P
            from eli.core.paths import config_dir as _cd
            notes_dir = _P(os.environ.get("ELI_NOTES_DIR", "")) or (_cd() / "notes")
            notes_dir.mkdir(parents=True, exist_ok=True)

            if a == "NEW_NOTE":
                import time as _nt
                title = (args.get("title") or args.get("text") or "").strip()
                content = (args.get("content") or args.get("text") or "").strip()
                fname = (title[:40].replace(" ", "_").replace("/", "_") or
                         _nt.strftime("note_%Y%m%d_%H%M%S")) + ".md"
                note_path = notes_dir / fname
                note_path.write_text(f"# {title or fname}\n\n{content}\n", encoding="utf-8")
                msg = f"Note saved: {note_path}"
                return {"ok": True, "action": a, "path": str(note_path), "content": msg, "response": msg}

            if a == "LIST_NOTES":
                notes = sorted(notes_dir.glob("*.md"), key=lambda p: p.stat().st_mtime, reverse=True)
                if not notes:
                    msg = f"No notes found in {notes_dir}"
                    return {"ok": True, "action": a, "notes": [], "content": msg, "response": msg}
                lines = [f"Notes ({len(notes)}):"]
                for n in notes[:30]:
                    lines.append(f"  {n.name}")
                msg = "\n".join(lines)
                return {"ok": True, "action": a, "notes": [n.name for n in notes],
                        "content": msg, "response": msg}

            if a == "SEARCH_NOTES":
                query = (args.get("query") or args.get("text") or "").strip().lower()
                if not query:
                    return {"ok": False, "action": a, "error": "No query",
                            "content": "Provide a search query.", "response": "Provide a search query."}
                hits = []
                for note in notes_dir.glob("*.md"):
                    text = note.read_text(encoding="utf-8", errors="replace").lower()
                    if query in text:
                        hits.append(note.name)
                if not hits:
                    msg = f"No notes matched '{query}'"
                else:
                    msg = f"Found {len(hits)} note(s) matching '{query}':\n" + "\n".join(f"  {h}" for h in hits[:20])
                return {"ok": True, "action": a, "hits": hits, "content": msg, "response": msg}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ---- POMODORO (START / STOP / STATUS) ----
    if a in ("POMODORO_START", "POMODORO_STOP", "POMODORO_STATUS"):
        try:
            from eli.core.paths import config_dir as _pcd
            import json as _pj, time as _pt
            pom_file = _pcd() / "pomodoro.json"

            def _pom_load():
                if pom_file.exists():
                    try:
                        return _pj.loads(pom_file.read_text(encoding="utf-8"))
                    except Exception:
                        pass
                return {}

            def _pom_save(d):
                pom_file.parent.mkdir(parents=True, exist_ok=True)
                pom_file.write_text(_pj.dumps(d), encoding="utf-8")

            if a == "POMODORO_STATUS":
                state = _pom_load()
                if not state.get("running"):
                    msg = "No Pomodoro running."
                else:
                    elapsed = int(_pt.time() - state.get("start", _pt.time()))
                    duration = state.get("duration", 1500)
                    remaining = max(0, duration - elapsed)
                    msg = (f"Pomodoro running — {elapsed // 60}m {elapsed % 60}s elapsed, "
                           f"{remaining // 60}m {remaining % 60}s remaining.")
                return {"ok": True, "action": a, "content": msg, "response": msg}

            if a == "POMODORO_START":
                minutes = int(args.get("minutes") or args.get("duration") or 25)
                state = {"running": True, "start": _pt.time(), "duration": minutes * 60}
                _pom_save(state)
                msg = f"Pomodoro started: {minutes} minutes. Focus!"
                return {"ok": True, "action": a, "content": msg, "response": msg}

            if a == "POMODORO_STOP":
                state = _pom_load()
                if state.get("running"):
                    elapsed = int(_pt.time() - state.get("start", _pt.time()))
                    state["running"] = False
                    _pom_save(state)
                    msg = f"Pomodoro stopped after {elapsed // 60}m {elapsed % 60}s."
                else:
                    msg = "No Pomodoro was running."
                return {"ok": True, "action": a, "content": msg, "response": msg}
        except Exception as e:
            return {"ok": False, "action": a, "error": str(e), "content": str(e), "response": str(e)}

    # ── DOC_GENERATE / GENERATE_DOCUMENT / CREATE_DOCUMENT ──────────────────
    if a in ("DOC_GENERATE", "GENERATE_DOCUMENT", "CREATE_DOCUMENT"):
        try:
            import re as _dre, subprocess as _dsp, tempfile as _dtmp
            from pathlib import Path as _DPath
            content  = str(args.get("content")  or args.get("text")  or args.get("generated") or "")
            filename = str(args.get("filename") or args.get("filepath") or args.get("path") or "")
            title    = str(args.get("title")    or args.get("query")  or args.get("topic") or "document")
            if not filename:
                safe = _dre.sub(r"[^a-zA-Z0-9_\- ]", "", title)[:40].strip().replace(" ", "_") or "eli_doc"
                try:
                    from eli.core.paths import get_paths as _dgp
                    docs_dir = _dgp().artifacts_dir / "docs"
                except Exception:
                    docs_dir = _DPath(_dtmp.gettempdir())
                docs_dir.mkdir(parents=True, exist_ok=True)
                filename = str(docs_dir / f"{safe}.md")
            _DPath(filename).parent.mkdir(parents=True, exist_ok=True)
            if not content:
                content = f"# {title}\n\n*Document generated by ELI.*\n"
            _DPath(filename).write_text(content, encoding="utf-8")
            try:
                _dsp.Popen(["xdg-open", filename])
            except Exception:
                pass
            msg = f"Document saved to `{filename}` and opened."
            return {"ok": True, "action": a, "content": msg, "response": msg, "filepath": filename}
        except Exception as _dge:
            return {"ok": False, "action": a, "error": str(_dge), "content": str(_dge), "response": str(_dge)}

    # ── SET_AI_MODE ──────────────────────────────────────────────────────────
    if a == "SET_AI_MODE":
        try:
            mode = str(args.get("mode") or args.get("reasoning_mode") or args.get("value") or "quick").lower().strip()
            _mode_map = {
                "quick": "quick", "fast": "quick",
                "chain_of_thought": "chain_of_thought", "cot": "chain_of_thought", "chain": "chain_of_thought",
                "self_consistency": "self_consistency", "sc": "self_consistency",
                "tree_of_thoughts": "tree_of_thoughts", "tot": "tree_of_thoughts", "tree": "tree_of_thoughts",
                "constitutional_ai": "constitutional_ai", "cai": "constitutional_ai", "constitutional": "constitutional_ai",
            }
            canonical = _mode_map.get(mode, mode)
            from eli.core.runtime_settings import save_settings as _ss
            _ss({"reasoning_mode": canonical})
            try:
                from eli.kernel.engine import get_engine as _ge
                _ge()._reasoning_mode = canonical
            except Exception:
                pass
            msg = f"Reasoning mode set to `{canonical}`."
            return {"ok": True, "action": a, "content": msg, "response": msg, "mode": canonical}
        except Exception as _sme:
            return {"ok": False, "action": a, "error": str(_sme), "content": str(_sme), "response": str(_sme)}

    return {"ok": False, "action": a, "error": f"Unsupported executor action: {a}", "content": f"Unsupported executor action: {a}", "response": f"Unsupported executor action: {a}"}


# ---------------------------------------------------------------------------
# Action dispatch hooks (added by phaseBW2 — replace 5 chained wrappers)
# ---------------------------------------------------------------------------

def _action_pre_dispatch(
    action: str,
    args: "Optional[Dict[str, Any]]" = None,
    **kwargs,
) -> "Optional[Dict[str, Any]]":
    """
    Centralised pre-dispatch action overrides.

    Returns a result dict if the action is fully handled, or None to
    fall through to the main execute() dispatch.

    Replaces the chain of monkey-patch wrappers retired in phaseBW2:
      * USER INFO EXECUTOR WRAPPER         (REFRESH_USER_INFO, USER_INFO_REPORT)
      * EXECUTOR SAFE FILE WRAPPERS        (READ_FILE, WRITE_NOTE, SCREENSHOT)
      * MEMORY_RECALL USER_INFO COMPAT     (identity-shaped MEMORY_RECALL)
      * GROUNDED EXECUTOR WRAPPER          (OPEN_APP/OPEN_IDE pre-flight,
                                            CHECK_TARGET_STATUS et al)
      * CONTROL_CONTRACT_EXECUTOR_WRAPPER  (SELF_UPDATE, EXPLAIN_LAST_RESPONSE)
    """
    a = str(action or "").upper().strip()
    args = args or {}

    # Control contract evidence ------------------------------------------
    if a in {"SELF_UPDATE", "EXPLAIN_LAST_RESPONSE"}:
        try:
            from eli.runtime.control_contracts import build_control_evidence
            return build_control_evidence(None, a, args, str(args.get("query") or ""))
        except Exception as _e:
            _msg = f"Control evidence failed for {a}: {_e!r}"
            return {"ok": False, "action": a, "error": repr(_e),
                    "content": _msg, "response": _msg}

    # MEMORY_RECALL identity-shaped queries -> USER_INFO_REPORT ---------
    if a == "MEMORY_RECALL" and "_eli_is_identity_memory_query" in globals() \
            and _eli_is_identity_memory_query(args):
        return _eli_user_info_report(force=False,
                                     reason="memory_recall_identity_compat")

    # User info actions --------------------------------------------------
    if a == "REFRESH_USER_INFO":
        return _eli_refresh_user_info(
            force=args.get("force", False),
            reason=args.get("reason", "manual"),
        )
    if a == "USER_INFO_REPORT":
        return _eli_user_info_report(
            force=args.get("force", False),
            reason=args.get("reason", "identity_query"),
        )

    # Safe file/screenshot ops ------------------------------------------
    if a == "READ_FILE":
        return _eli_safe_read_file(args)
    if a == "WRITE_NOTE":
        return _eli_safe_write_note(args)
    if a == "SCREENSHOT":
        return _eli_safe_screenshot(args)

    # OPEN_APP / OPEN_IDE pre-flight via grounded_remediation ----------
    if a in {"OPEN_APP", "OPEN_IDE"}:
        try:
            from eli.runtime import grounded_remediation as _gr
        except Exception:
            return None
        _subject = (args.get("app") or args.get("name") or args.get("target")
                    or args.get("message") or args.get("query") or "")
        _app = _gr.extract_app_name(str(_subject))
        if _app:
            _diag = _gr.diagnose_app(_app)
            if not _diag.get("ok", False):
                return _gr.as_executor_result(_gr.offer_for_result(_diag), ok=False)
            if a == "OPEN_IDE":
                # Rewrite OPEN_IDE -> OPEN_APP and call _execute_impl
                # directly so we don't re-enter pre-dispatch.
                _open_args = dict(args)
                _open_args["name"] = _app
                _merged = dict(_open_args)
                _merged.update(kwargs or {})
                _result = _execute_impl(action="OPEN_APP", args=_merged)
                _result = _normalize_result(_result)
                _offer = _gr.capture_executor_failure("OPEN_APP", _open_args, _result)
                if _offer:
                    return _gr.as_executor_result(_offer, ok=False)
                return _result

    # CHECK_TARGET_STATUS / EXPLAIN_LAST_FAILURE / etc -----------------
    if a in {"CHECK_TARGET_STATUS", "EXPLAIN_LAST_FAILURE",
             "PREPARE_REMEDIATION", "CONFIRM_PENDING_REMEDIATION",
             "CANCEL_PENDING_REMEDIATION"}:
        try:
            from eli.runtime import grounded_remediation as _gr
        except Exception:
            return None
        _msg = args.get("message") or args.get("query") or a
        _handled = _gr.try_handle_query(str(_msg))
        if _handled:
            return _gr.as_executor_result(
                _handled, ok=not _handled.startswith("Could not"))

    return None


# Per-process failure tally for the "attempt N" anti-retry hint. Keyed by
# action+args+error; reset every restart so a fresh session never inherits a
# stale "attempt 4" from all-time DB history. (The DB still logs failures for
# learning — this only governs the user-facing repeat hint.)
_SESSION_FAILURE_COUNTS: "Dict[str, int]" = {}


def _action_post_dispatch(
    action: str,
    args: "Optional[Dict[str, Any]]",
    result: "Dict[str, Any]",
) -> "Dict[str, Any]":
    """
    Apply post-dispatch failure capture: if the result indicates failure
    grounded_remediation may produce a remediation offer to surface to
    the user. Returns the capture result if any, else the original result.

    Replaces the post-execute capture from the GROUNDED EXECUTOR WRAPPER,
    retired in phaseBW2.
    """
    try:
        from eli.runtime.evidence_ledger import record_event as _eli_record_event

        _eli_record_event(
            "executor_action",
            source="executor.post_dispatch",
            action=str(action or "").upper(),
            subject=str(
                (args or {}).get("path")
                or (args or {}).get("target")
                or (args or {}).get("name")
                or (args or {}).get("topic")
                or ""
            ),
            content=str((result or {}).get("content") or (result or {}).get("response") or (result or {}).get("error") or ""),
            payload={"args": args or {}, "result": result or {}},
            severity="info" if bool((result or {}).get("ok", True)) else "error",
            outcome="ok" if bool((result or {}).get("ok", True)) else "failed",
            reusable=True,
        )
    except Exception:
        pass

    repeat_hint = ""
    # Phase 5: non-bugs that previously poisoned the SI failure feed. These
    # are *expected* responses (user typed no query, plugin not configured,
    # action genuinely outside this build's scope) — they should not enter
    # the failure-cluster pipeline that the SelfImprovementEngine acts on.
    _SI_NOISE_ERRORS = frozenset((
        "empty_query",
        "empty",
        "not_configured",
        "unsupported",
        "unsupported_action",
        "no_query",
        # Variants actually produced by executor actions (mixed-case / with spaces)
        "no query",
        "no query provided",
        # Security filter working as intended — not a bug
        "security_blocked",
    ))
    try:
        # Grounded-remediation exit-20 ("No grounded install candidate found") is
        # an honest terminal outcome, not a retryable failure. Don't log it as a
        # failure and don't append the "attempt N" hint.
        _remediation_actions = frozenset((
            "CONFIRM_PENDING_REMEDIATION",
            "CANCEL_PENDING_REMEDIATION",
        ))
        if isinstance(result, dict) and not bool(result.get("ok", True)):
            import json as _json
            from eli.memory.memory import get_memory as _get_memory

            # Remediation actions are terminal honest outcomes — skip failure logging.
            if str(action or "").upper() in _remediation_actions:
                return result

            raw_error = result.get("error")
            err_text = str(
                raw_error
                or result.get("message")
                or result.get("response")
                or result.get("content")
                or "executor failure"
            ).strip()

            # Quick reject for the canonical "this isn't a real bug" surfaces.
            err_token = str(raw_error or "").strip().lower()
            if err_token in _SI_NOISE_ERRORS:
                # Don't log to failures table; the post-dispatch path still
                # returns the result to the caller, just doesn't poison SI.
                return result

            signature_input = f"{str(action or '').upper()} {_json.dumps(args or {}, sort_keys=True, default=str)}"
            # Failures live in ONE canonical store — the agent/self-improvement DB
            # (agent.sqlite3), alongside improvements + code_patches. Previously this
            # dual-wrote to the user DB too, splitting the failure log across two
            # databases (the Self-Improve panel reads only the agent store).
            try:
                from eli.memory import get_agent_memory as _get_agent_memory
                _get_agent_memory().log_failure(
                    signature_input,
                    error=err_text,
                    confidence=0.0,
                    context={"action": action, "args": args or {}, "result": result},
                    source="executor_post_dispatch",
                )
            except Exception:
                pass
            # Count repeats in THIS running session only. The hint means "you
            # keep retrying the same thing right now" — not all-time DB history,
            # which falsely claimed "attempt 4" on a fresh session's first try.
            _sess_key = f"{signature_input}\x00{err_text}"
            occurrence_count = _SESSION_FAILURE_COUNTS.get(_sess_key, 0) + 1
            _SESSION_FAILURE_COUNTS[_sess_key] = occurrence_count
            if occurrence_count >= 2:
                subject = ""
                if isinstance(args, dict):
                    subject = str(
                        args.get("path")
                        or args.get("target")
                        or args.get("name")
                        or args.get("app")
                        or args.get("query")
                        or ""
                    ).strip()
                target = f" `{subject}`" if subject else ""
                repeat_hint = (
                    f"This is attempt {occurrence_count} for the same {str(action or '').upper()} failure{target}. "
                    "Do you want me to inspect the target, logs, or generated artifact instead of retrying the same command?"
                )
    except Exception:
        repeat_hint = ""

    try:
        from eli.runtime import grounded_remediation as _gr
    except Exception:
        if repeat_hint and isinstance(result, dict):
            existing = str(result.get("content") or result.get("response") or result.get("error") or "").strip()
            combined = (existing + "\n\n" + repeat_hint).strip()
            result["content"] = combined
            result["response"] = combined
        return result
    try:
        _offer = _gr.capture_executor_failure(action, args or {}, result)
        if _offer:
            if repeat_hint:
                _offer = f"{_offer}\n\n{repeat_hint}"
            return _gr.as_executor_result(_offer, ok=False)
    except Exception:
        pass
    if repeat_hint and isinstance(result, dict):
        existing = str(result.get("content") or result.get("response") or result.get("error") or "").strip()
        combined = (existing + "\n\n" + repeat_hint).strip()
        result["content"] = combined
        result["response"] = combined
    return result


def _eli_open_url_action(raw_url=None, *, query=None):
    """
    Safe URL opener for OPEN_URL.

    Rules:
    - Accepts http/https URLs.
    - Accepts bare domains like github.com and normalizes to https://github.com.
    - Blocks unsupported schemes such as file:, javascript:, data:, ftp:, etc.
    - Uses xdg-open without shell=True.
    """
    from urllib.parse import urlparse, quote_plus
    import shutil
    import subprocess
    import webbrowser

    value = str(raw_url or "").strip()

    if not value and query:
        value = "https://duckduckgo.com/?q=" + quote_plus(str(query).strip())

    if not value:
        msg = "No URL was provided."
        return {
            "ok": False,
            "action": "OPEN_URL",
            "error": msg,
            "content": f"⚡ Could not open URL. {msg}",
            "response": f"⚡ Could not open URL. {msg}",
        }

    # Allow common bare domains: github.com, openai.com/docs, etc.
    if "://" not in value:
        if " " in value or "." not in value:
            msg = f"Not a valid URL or bare domain: {value!r}"
            return {
                "ok": False,
                "action": "OPEN_URL",
                "url": value,
                "error": msg,
                "content": f"⚡ Could not open URL. {msg}",
                "response": f"⚡ Could not open URL. {msg}",
            }
        value = "https://" + value

    parsed = urlparse(value)

    if parsed.scheme not in {"http", "https"}:
        msg = f"Blocked unsupported URL scheme: {parsed.scheme!r}"
        return {
            "ok": False,
            "action": "OPEN_URL",
            "url": value,
            "error": msg,
            "content": f"⚡ Could not open URL. {msg}",
            "response": f"⚡ Could not open URL. {msg}",
        }

    if not parsed.netloc:
        msg = f"URL is missing a domain: {value!r}"
        return {
            "ok": False,
            "action": "OPEN_URL",
            "url": value,
            "error": msg,
            "content": f"⚡ Could not open URL. {msg}",
            "response": f"⚡ Could not open URL. {msg}",
        }

    try:
        opener = shutil.which("xdg-open")
        if opener:
            subprocess.Popen(
                [opener, value],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                start_new_session=True,
            )
        else:
            webbrowser.open(value, new=2)

        msg = f"Opened URL: {value}"
        return {
            "ok": True,
            "action": "OPEN_URL",
            "url": value,
            "content": msg,
            "response": msg,
        }

    except Exception as exc:
        msg = f"Failed to open URL: {exc!r}"
        return {
            "ok": False,
            "action": "OPEN_URL",
            "url": value,
            "error": msg,
            "content": f"⚡ Could not open URL. {msg}",
            "response": f"⚡ Could not open URL. {msg}",
        }


def execute(action: str, args: Optional[Dict[str, Any]] = None, **kwargs) -> Dict[str, Any]:

    # Unwrap dict form: execute({"action": "X", "args": {...}})
    if isinstance(action, dict):
        args = action.get("args", args) or args
        action = action.get("action", "")
    action = action if isinstance(action, str) else ""
    a = action.upper().strip()

    # Pre-dispatch action overrides (formerly applied via monkey-patch wrappers).
    _override = _action_pre_dispatch(action, args, **kwargs)
    if _override is not None:
        return _action_post_dispatch(action, args, _override)
    if a == 'USER_IDENTITY_SUMMARY':
        _question = str((args or {}).get("question") or (args or {}).get("query") or "who am i")
        try:
            _snapshot = _eli_user_info_report(force=False, reason="user_identity_summary")
            _raw = str(_snapshot.get("content") or _snapshot.get("response") or "")

            def _section(_name):
                _lines = []
                _inside = False
                _header = f"[{_name}]"
                for _ln in _raw.splitlines():
                    _s = _ln.strip()
                    if _s == _header:
                        _inside = True
                        continue
                    if _inside and _s.startswith("[") and _s.endswith("]"):
                        break
                    if _inside and _s.startswith("- "):
                        _item = _s[2:].strip()
                        if _item and _item.lower() != "none confirmed.":
                            _lines.append(_item.rstrip("."))
                return _lines

            _identity = _section("Identity")
            _prefs = _section("Communication Preferences")
            _working = _section("Working Style")
            _projects = _section("Active Projects")
            _technical = _section("Technical Environment")

            _lines = ["Current user identity summary:"]

            if _identity:
                _lines.append("")
                _lines.append("Identity:")
                for _x in _identity[:5]:
                    _lines.append(f"- {_x}.")
            else:
                _lines.append("")
                _lines.append("- I do not have a confirmed name/identity row in the current user-info snapshot.")

            _facts = []
            for _bucket in (_prefs, _working, _projects, _technical):
                for _x in _bucket:
                    if _x not in _facts:
                        _facts.append(_x)

            if _facts:
                _lines.append("")
                _lines.append("What I can say from current memory:")
                for _x in _facts[:10]:
                    _lines.append(f"- {_x}.")

            _msg = "\n".join(_lines).strip()
            if not _msg:
                _msg = "I have no clean user identity summary available right now."

            return {
                "ok": True,
                "action": a,
                "content": _msg,
                "response": _msg,
                "report": {
                    "source": "user_info_snapshot_summarized",
                    "raw_snapshot_suppressed": True,
                    "question": _question,
                },
                "evidence_source": "user_info_snapshot_summarized",
                "generation_invoked": False,
            }
        except Exception as e:
            _msg = f"User identity summary failed: {type(e).__name__}: {e}"
            return {"ok": False, "action": a, "content": _msg, "response": _msg, "error": repr(e)}
    if a == 'SELF_REPORT':
        rep = _runtime_status_report()
        runtime_eff = rep.get('runtime') or {}
        settings_req = rep.get('settings') or {}
        # Active user's confirmed name belongs in identity evidence — without it
        # "who are you / who am I" answers say "name not provided" and then
        # confabulate a low-confidence story. This is the authoritative profile
        # name (same source USER_IDENTITY_SUMMARY uses).
        try:
            from eli.kernel.state import get_user_name as _gun_sr
            _active_user = str(_gun_sr("") or "").strip()
        except Exception:
            _active_user = ""
        ev = {
            "identity": {
                "name": "ELI",
                "expanded_name": "Enhanced Learning Interface",
                "active_user_name": _active_user or "unknown",
                "grounding_sources": [
                    "persona",
                    "memory",
                    "runtime_state",
                    "local_files",
                    "loaded_model",
                ],
            },
            "runtime": {
                "model_path": runtime_eff.get('model_path') or rep.get('model_path', 'unknown'),
                "effective_context_size": runtime_eff.get('n_ctx', 'unknown'),
                "effective_gpu_layers": runtime_eff.get('n_gpu_layers', 'unknown'),
                "effective_batch_size": runtime_eff.get('n_batch', 'unknown'),
                "effective_cpu_threads": runtime_eff.get('n_threads', 'unknown'),
            },
            "requested_settings": {
                "n_ctx": settings_req.get('n_ctx', 'unknown'),
                "n_gpu_layers": settings_req.get('n_gpu_layers', 'unknown'),
                "n_threads": settings_req.get('n_threads', 'unknown'),
                "batch_size": settings_req.get('batch_size', 'unknown'),
            },
        }
        txt = json.dumps(ev, ensure_ascii=False, indent=2)
        model_name = (runtime_eff.get('model_path') or rep.get('model_path', '')).split('/')[-1]
        summary = (
            f"I'm ELI (Enhanced Learning Interface), running {model_name} locally "
            f"on GPU ({runtime_eff.get('n_gpu_layers', '?')} layers offloaded). "
            f"Context window: {runtime_eff.get('n_ctx', '?')} tokens. "
            f"All core systems nominal."
        )
        if _active_user:
            summary += f" You're {_active_user}."
        # content/response are user-facing — use the plain summary, never raw JSON.
        # Structured data stays in report + evidence for grounding/introspection.
        return {'ok': bool(rep.get('ok', True)), 'action': a, 'report': rep,
                'evidence': ev, 'content': summary, 'response': summary}
    if a == 'IMAGE_STATUS':
        try:
            from eli.runtime.evidence_ledger import status_evidence

            query = ""
            if isinstance(args, dict):
                query = str(args.get("query") or args.get("message") or args.get("text") or "")
            rep = status_evidence(query or "image status")
            txt = "Image/status evidence packet:\n" + json.dumps(rep, indent=2, ensure_ascii=False, default=str)
            return {'ok': True, 'action': a, 'report': rep, 'content': txt, 'response': txt}
        except Exception as _e:
            msg = f"Image/status evidence failed: {_e!r}"
            return {'ok': False, 'action': a, 'error': repr(_e), 'content': msg, 'response': msg}
    if a == 'RUNTIME_AUDIT':
        rep = _runtime_audit_report(); txt = _format_runtime_audit(rep)
        return {'ok': True, 'action': a, 'report': rep, 'content': txt, 'response': txt}
    if a == 'IMPORT_AUDIT':
        rep = _import_audit_report(); txt = _format_import_audit(rep)
        return {'ok': True, 'action': a, 'report': rep, 'content': txt, 'response': txt}
    if a == 'DIAGNOSE_WRAPPERS':
        rep = _diagnose_executor_wrappers()
        txt = _format_executor_wrappers(rep)
        return {'ok': True, 'action': a, 'report': rep, 'content': txt, 'response': txt}
    if a == 'RESOLVE_RUNTIME_PATHS':
        rep = _resolve_runtime_paths_report(); txt = _format_runtime_paths(rep)
        return {'ok': True, 'action': a, 'report': rep, 'content': txt, 'response': txt}
    if a == 'GUI_RUNTIME_AUDIT':
        rep = _gui_runtime_audit_report(); txt = _format_gui_runtime_audit(rep)
        return {'ok': True, 'action': a, 'report': rep, 'content': txt, 'response': txt}
    if a == 'EXPLAIN_ALL_REASONING_MODES':
        rep = _explain_all_reasoning_modes_report(); txt = _format_all_reasoning_modes(rep)
        return {'ok': bool(rep.get('ok', True)), 'action': a, 'report': rep, 'content': txt, 'response': txt}
    if a == 'EXPLAIN_MEMORY_RUNTIME':
        rep = _explain_memory_runtime_report(); txt = _format_memory_runtime(rep)
        return {'ok': True, 'action': a, 'report': rep, 'content': txt, 'response': txt}
    if a == 'EXPLAIN_COGNITION_RUNTIME':
        rep = _explain_cognition_runtime_report(); txt = _format_cognition_runtime(rep)
        return {'ok': bool(rep.get('ok', True)), 'action': a, 'report': rep, 'content': txt, 'response': txt}
    if a == 'RUNTIME_STATUS':
        rep = _runtime_status_report(); txt = _format_runtime_status(rep)
        return {'ok': bool(rep.get('ok', True)), 'action': a, 'report': rep, 'content': txt, 'response': txt}
    if a == 'GPU_STATUS':
        rep = _gpu_status_report()
        return {'ok': bool(rep.get('ok', False)), 'action': a, 'report': rep, 'content': rep.get('content', ''), 'response': rep.get('response', rep.get('content', ''))}
    if a == 'SELF_IMPROVEMENT_LOG':
        limit = int((args or {}).get('limit') or 5) if isinstance(args, dict) else 5
        days = int((args or {}).get('days') or 30) if isinstance(args, dict) else 30
        rep = _self_improvement_log_report(limit=limit, days=days)
        return {'ok': bool(rep.get('ok', False)), 'action': a, 'report': rep, 'content': rep.get('content', ''), 'response': rep.get('response', rep.get('content', ''))}
    if a == 'MEMORY_STATUS':
        rep = _memory_status_report(); txt = _format_memory_status(rep)
        try:
            txt += '\n\n' + _get_db_schema_evidence()
        except Exception:
            pass
        return {'ok': bool(rep.get('ok', True)), 'action': a, 'report': rep, 'content': txt, 'response': txt}
    if a == 'COGNITION_STATUS':
        rep = _cognition_status_report(); txt = _format_cognition_status(rep)
        return {'ok': bool(rep.get('ok', True)), 'action': a, 'report': rep, 'content': txt, 'response': txt}
    _mem = None  # _mkvi_memory_dispatch removed (was undefined)
    if _mem is not None:
        return _mem
    merged: Dict[str, Any] = {}
    if isinstance(args, dict):
        merged.update(args)
    merged.update(kwargs or {})
    res = _execute_impl(action=str(a), args=merged)
    __eli_ret = _normalize_result(res)
    # Normalize LIST_CAPABILITIES top-level content/response (pipeline expects table here)
    try:
        if isinstance(__eli_ret, dict) and __eli_ret.get('action') == 'LIST_CAPABILITIES':
            caps = __eli_ret.get('capabilities')
            if isinstance(caps, dict):
                table = (caps.get('content') or '')
                if table:
                    __eli_ret['content'] = table
                    __eli_ret['response'] = table
    except Exception:
        pass
    # Post-dispatch failure capture (formerly in GROUNDED EXECUTOR WRAPPER).
    return _action_post_dispatch(action, args, __eli_ret)


# Alias for compatibility with tests
try:
    execute_action = execute
except Exception:
    pass


_PACKAGE_MIRROR_TOPLEVEL = {
    "api", "brain", "controllers", "core", "gui", "integrations", "plugins", "tools", "utils"
}

def _resolve_existing_user_or_artifact_path(path: str) -> Path:
    """Resolve a user path, then recover common generated artifact locations."""
    raw = str(path or "").strip()
    resolved = Path(resolve_user_repo_path(raw))
    if resolved.exists():
        return resolved

    repo_root = Path(__file__).resolve().parents[2]
    expanded = Path(raw).expanduser()
    basename = expanded.name
    candidates = []

    if raw and not expanded.is_absolute():
        candidates.append((repo_root / raw).resolve(strict=False))
    if basename:
        for subdir in ("scripts", "documents"):
            artifact_dir = repo_root / "artifacts" / subdir
            candidates.append((artifact_dir / basename).resolve(strict=False))
            if not Path(basename).suffix:
                try:
                    candidates.extend(sorted(artifact_dir.glob(f"{basename}.*")))
                except Exception:
                    pass

    for candidate in candidates:
        try:
            if candidate.exists():
                return candidate.resolve(strict=False)
        except Exception:
            continue
    return resolved

def _resolve_legacy_eli_fs_path(path: str) -> Path:
    raw = str(path or "").strip()
    if not raw:
        return Path(".").resolve()

    expanded = Path(raw).expanduser()
    repo_root = Path(__file__).resolve().parents[2]  # project root
    legacy_root = Path.home() / "eli"

    if expanded == legacy_root:
        return (repo_root / "eli").resolve()

    try:
        rel = expanded.relative_to(legacy_root)
    except Exception:
        return expanded.resolve()

    if not rel.parts:
        return (repo_root / "eli").resolve()

    top = rel.parts[0]
    if top in _PACKAGE_MIRROR_TOPLEVEL:
        return (repo_root / "eli" / rel).resolve()

    return (repo_root / rel).resolve()

if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument("action")
    ap.add_argument("--args", default="{}")
    ns = ap.parse_args()
    print(json.dumps(execute(ns.action, json.loads(ns.args)), indent=2))

# --- AUTO-REGISTER CAPABILITIES (ON IMPORT) ---
try:
    from eli.tools.registry.capability_registry import register as _cap_register
    for _a in SUPPORTED_ACTIONS:
        _cap_register(_a, description=f"Executor action: {_a}")
except Exception:
    pass


# ----------------------------
# Proactive daemon control (real PID management)
# ----------------------------
def _proactive_paths():
    base = Path(os.environ.get("ELI_ARTIFACTS_DIR", "artifacts")) / "proactive"
    base.mkdir(parents=True, exist_ok=True)
    return base, base / "daemon.pid", base / "daemon.log"


def proactive_status() -> Dict[str, Any]:
    base, pidf, logf = _proactive_paths()
    pid = None
    if pidf.exists():
        try:
            pid = int(pidf.read_text().strip())
        except Exception:
            pid = None
    alive = False
    if pid:
        try:
            os.kill(pid, 0)
            alive = True
        except Exception:
            alive = False
    msg = f"Proactive daemon is {'RUNNING' if alive else 'STOPPED'}."
    return {"ok": True, "action": "PROACTIVE_STATUS", "pid": pid, "running": alive, "log": str(logf), "content": msg, "response": msg}


def proactive_start() -> Dict[str, Any]:
    base, pidf, logf = _proactive_paths()
    st = proactive_status()
    if st.get("running"):
        return st
    try:
        py = sys.executable if "sys" in globals() else "python3"
    except Exception:
        py = "python3"
    # run as module so imports work
    argv = [py, "-m", "eli.proactive.proactive_daemon"]
    try:
        lf = open(logf, "a", encoding="utf-8")
        p = subprocess.Popen(argv, stdout=lf, stderr=lf, cwd=str(Path.cwd()))
        pidf.write_text(str(p.pid), encoding="utf-8")
        msg = "Proactive daemon started."
        return {"ok": True, "action": "PROACTIVE_START", "pid": p.pid, "cmd": argv, "log": str(logf), "content": msg, "response": msg}
    except Exception as e:
        msg = "Failed to start proactive daemon."
        return {"ok": False, "action": "PROACTIVE_START", "error": repr(e), "content": msg, "response": msg}


def proactive_stop() -> Dict[str, Any]:
    base, pidf, logf = _proactive_paths()
    pid = None
    if pidf.exists():
        try:
            pid = int(pidf.read_text().strip())
        except Exception:
            pid = None
    if not pid:
        msg = "Proactive daemon not running (no PID)."
        return {"ok": True, "action": "PROACTIVE_STOP", "content": msg, "response": msg}
    try:
        os.kill(pid, 15)
    except Exception:
        pass
    try:
        pidf.unlink(missing_ok=True)  # py3.8+: if not available it will throw
    except Exception:
        try:
            pidf.unlink()
        except Exception:
            pass
    msg = "Proactive daemon stopped."
    return {"ok": True, "action": "PROACTIVE_STOP", "pid": pid, "content": msg, "response": msg}


class Executor:
    """Thin class wrapper around the functional execute() API."""

    def __init__(self):
        pass

    def execute(self, action: str, args: dict = None):
        return execute(action, args or {})

    def __call__(self, action: str, args: dict = None):
        return execute(action, args or {})


# Helpers preserved from former USER INFO EXECUTOR WRAPPER (phaseBW2).
# The wrapper-install scaffolding was removed; pre-dispatch in execute()
# routes REFRESH_USER_INFO and USER_INFO_REPORT through these directly.
def _eli_refresh_user_info(force=False, reason="manual"):
    from eli.cognition.user_info_builder import refresh_user_info
    out = refresh_user_info(force=bool(force), reason=str(reason or "manual"))
    return {
        "ok": True,
        "action": "REFRESH_USER_INFO",
        "content": out.get("summary") or "User info refreshed.",
        "artifacts": out,
    }

def _eli_user_info_report(force=False, reason="query"):
    from eli.cognition.user_info_builder import read_user_info, refresh_user_info
    if force:
        refresh_user_info(force=True, reason=str(reason or "query"))
    out = read_user_info(auto_refresh=True, reason=str(reason or "query"))
    return {
        "ok": True,
        "action": "USER_INFO_REPORT",
        "content": out.get("text", ""),
        "meta": out.get("meta", {}),
        "path": out.get("path", ""),
    }


# Helpers preserved from former EXECUTOR SAFE FILE WRAPPERS (phaseBW2).
# The wrapper-install scaffolding was removed; pre-dispatch in execute()
# routes READ_FILE / WRITE_NOTE / SCREENSHOT through these directly.
def _eli_safe_read_file(args=None):
    from pathlib import Path
    _args = args or {}
    path = str(_args.get("path") or "").strip()
    if not path:
        return {"ok": False, "action": "READ_FILE", "content": "Missing file path."}
    p = resolve_user_repo_path(path)
    if not p.exists():
        return {"ok": False, "action": "READ_FILE", "content": f"File not found: {p}"}
    if p.is_dir():
        return {"ok": False, "action": "READ_FILE", "content": f"Path is a directory, not a file: {p}"}
    data = None
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            data = p.read_text(encoding=enc)
            break
        except Exception:
            pass
    if data is None:
        try:
            data = p.read_bytes().decode("utf-8", errors="replace")
        except Exception as e:
            return {"ok": False, "action": "READ_FILE", "content": f"Could not read file: {e}"}
    return {
        "ok": True,
        "action": "READ_FILE",
        "path": str(p),
        "content": data,
    }

def _eli_safe_write_note(args=None):
    from pathlib import Path
    from datetime import datetime
    _args = args or {}
    txt = str(_args.get("text") or "").rstrip()
    if not txt:
        return {"ok": False, "action": "WRITE_NOTE", "content": "Missing note text."}
    nb = Path("eli_notebook")
    nb.mkdir(parents=True, exist_ok=True)
    daily = nb / f"{datetime.now().strftime('%Y-%m-%d')}.md"
    prefix = "" if not daily.exists() or not daily.read_text(encoding="utf-8", errors="ignore").strip() else "\n"
    stamp = datetime.now().strftime("%H:%M:%S")
    payload = f"{prefix}- [{stamp}] {txt}\n"
    daily.write_text(daily.read_text(encoding="utf-8", errors="ignore") + payload if daily.exists() else payload, encoding="utf-8")
    return {
        "ok": True,
        "action": "WRITE_NOTE",
        "path": str(daily),
        "content": f"Wrote note to {daily}",
    }

def _eli_safe_screenshot(args=None):
    _args = args or {}
    region = _args.get("region", "full")
    try:
        from controllers.os_controller import take_screenshot
    except Exception:
        try:
            from eli.perception.os_controller import take_screenshot
        except Exception as e:
            return {"ok": False, "action": "SCREENSHOT", "content": f"Screenshot backend unavailable: {e}"}
    try:
        out = take_screenshot(region)
    except TypeError:
        try:
            out = take_screenshot()
        except Exception as e:
            return {"ok": False, "action": "SCREENSHOT", "content": f"Screenshot failed: {e}"}
    except Exception as e:
        return {"ok": False, "action": "SCREENSHOT", "content": f"Screenshot failed: {e}"}

    if isinstance(out, dict):
        ok = bool(out.get("ok", True))
        content = str(out.get("content") or out.get("message") or out.get("path") or "Screenshot completed")
        return {
            "ok": ok,
            "action": "SCREENSHOT",
            "content": content,
            **out,
        }

    return {
        "ok": True,
        "action": "SCREENSHOT",
        "content": str(out or "Screenshot completed"),
    }


# Helper preserved from former MEMORY_RECALL USER_INFO COMPAT (phaseBW2).
# Pre-dispatch in execute() routes identity-shaped MEMORY_RECALL queries
# through this directly to USER_INFO_REPORT.
import re as _eli_exec_re

def _eli_is_identity_memory_query(_args):
    _args = _args or {}
    q = str(_args.get("query") or _args.get("message") or "").strip().lower()
    if _args.get("_prefer_user_info_report"):
        return True
    pats = [
        r"\bwhat do you know about me\b",
        r"\bwhat do you know from memory\b",
        r"\bwho am i(?: to you)?\b",
        r"\bshow user info\b",
        r"\buser info report\b",
        r"\bprofile report\b",
        r"\bdump (?:my )?(?:user )?(?:profile|info|memory)\b",
        r"\bshow (?:my|the) stored memor(?:y|ies)(?: about me)?\b",
    ]
    return any(_eli_exec_re.search(p, q) for p in pats)


# Local-only YouTube playback. No browser search. No external API.

# Canonical media runtime interceptor.
# Keep routing in router_enhanced/portable_intent_contract.py.
# Keep execution-side media behaviour in eli.execution.media_runtime.
try:
    from eli.execution.media_runtime import install_media_executor as _install_media_executor
    execute_action = _install_media_executor(execute_action)
except Exception as _media_runtime_error:
    log.debug(f"[MEDIA_RUNTIME] failed to install media executor wrapper: {_media_runtime_error}")


# gui_schema_contract_guard
try:
    _ELI_SCHEMA_CONTRACT_ORIGINAL_EXECUTE
except NameError:
    _ELI_SCHEMA_CONTRACT_ORIGINAL_EXECUTE = execute

    def _eli_schema_project_root():
        env_root = os.environ.get("ELI_PROJECT_ROOT")
        if env_root:
            return Path(env_root).expanduser().resolve()
        return Path(__file__).resolve().parents[2]

    def _eli_schema_artifacts_dir():
        root = _eli_schema_project_root()
        env_artifacts = os.environ.get("ELI_ARTIFACTS_DIR")
        if env_artifacts:
            candidate = Path(env_artifacts).expanduser()
            return candidate.resolve() if candidate.is_absolute() else (root / candidate).resolve()
        return (root / "artifacts").resolve()

    def _eli_schema_slug(value, default="document"):
        slug = re.sub(r"[^A-Za-z0-9._-]+", "_", str(value or "").strip()).strip("._-")
        return slug or default

    def _eli_schema_text_from_result(result):
        if isinstance(result, dict):
            for key in ("content", "response", "text", "answer", "message", "output"):
                value = result.get(key)
                if value is not None and str(value).strip():
                    return str(value)
            return ""
        if result is not None:
            return str(result)
        return ""

    def _eli_schema_create_document(args):
        args = args if isinstance(args, dict) else {}
        topic = args.get("topic") or args.get("title") or args.get("name") or "document"
        title = args.get("title") or args.get("name") or topic
        fmt = str(args.get("format") or args.get("ext") or "md").lower().lstrip(".")
        if fmt not in {"md", "txt", "json", "csv", "tex", "py"}:
            fmt = "md"

        content = (
            args.get("content")
            or args.get("body")
            or args.get("text")
            or args.get("markdown")
        )
        if content is None:
            return _eli_schema_failure(
                "CREATE_DOCUMENT",
                "Document write refused: no substantive body was supplied. "
                "Use GENERATE_DOCUMENT/CREATE_DOCUMENT with a topic for model drafting, "
                "or WRITE_DOCUMENT with explicit content.",
            )

        out_dir = _eli_schema_artifacts_dir() / "documents"
        out_dir.mkdir(parents=True, exist_ok=True)

        filename = _eli_schema_slug(title, "document")
        if not filename.lower().endswith(f".{fmt}"):
            filename = f"{filename}.{fmt}"

        path = out_dir / filename
        if path.exists() and not bool(args.get("overwrite", False)):
            stem = path.stem
            suffix = path.suffix
            idx = 2
            while True:
                candidate = out_dir / f"{stem}_{idx}{suffix}"
                if not candidate.exists():
                    path = candidate
                    break
                idx += 1

        path.write_text(str(content), encoding="utf-8")
        msg = f"Document saved: {path}"
        return {
            "ok": True,
            "action": "CREATE_DOCUMENT",
            "doc_path": str(path),
            "path": str(path),
            "content": msg,
            "response": msg,
        }

    def _eli_schema_extract_python_source(args):
        args = args if isinstance(args, dict) else {}
        for key in ("code", "script", "source", "content", "python", "body", "text"):
            value = args.get(key)
            if isinstance(value, str) and value.strip():
                return value

        generator = globals().get("chat")
        if callable(generator):
            prompt = args.get("description") or args.get("prompt") or "Generate a Python script."
            result = generator(str(prompt))
            text = _eli_schema_text_from_result(result)
            if text.strip():
                return text

        prompt = args.get("prompt")
        if isinstance(prompt, str) and prompt.strip():
            stripped = prompt.lstrip()
            if "\n" in prompt or stripped.startswith(("def ", "class ", "import ", "from ", "print(", "for ", "while ", "if ", "try:")):
                return prompt

        return None

    def _eli_schema_failure(action_name, message):
        return {
            "ok": False,
            "action": action_name,
            "error": str(message),
            "content": str(message),
            "response": str(message),
            "evidence": [str(message)],
        }

    def _eli_schema_normalise_failure(action_name, result):
        if not isinstance(result, dict):
            return result
        if bool(result.get("ok", True)):
            return result
        if result.get("error"):
            return result
        for key in ("content", "response", "message"):
            value = result.get(key)
            if value:
                result["error"] = str(value)
                return result
        evidence = result.get("evidence")
        if isinstance(evidence, (list, tuple)) and evidence:
            result["error"] = str(evidence[0])
            return result
        result["error"] = f"{action_name} failed"
        return result

    def execute(action, args=None, *pargs, **kwargs):
        action_name = str(action or "").strip().upper().replace("-", "_")
        call_args = args if isinstance(args, dict) else {}

        if action_name in {"CREATE_DOCUMENT", "CREATE_DOC", "WRITE_DOCUMENT"}:
            has_explicit_body = any(
                isinstance(call_args.get(key), str) and call_args.get(key).strip()
                for key in ("content", "body", "text", "markdown")
            )
            if has_explicit_body:
                return _eli_schema_create_document(call_args)
            result = _ELI_SCHEMA_CONTRACT_ORIGINAL_EXECUTE(action, args, *pargs, **kwargs)
            return _eli_schema_normalise_failure(action_name, result)

        if action_name in {"GENERATE_SCRIPT", "CREATE_SCRIPT", "WRITE_SCRIPT"}:
            source = _eli_schema_extract_python_source(call_args)
            if source is not None:
                try:
                    compile(str(source), "<eli-generated-script>", "exec")
                except SyntaxError as exc:
                    return _eli_schema_failure(
                        action_name,
                        f"Generated Python script failed syntax validation: {exc}",
                    )

        result = _ELI_SCHEMA_CONTRACT_ORIGINAL_EXECUTE(action, args, *pargs, **kwargs)
        return _eli_schema_normalise_failure(action_name, result)

# voice_runtime_executor_contract_guard
# Stable execution support for direct voice rules. Uses commands discovered
# through PATH and platform APIs; no absolute user-machine paths.

try:
    _ELI_VOICE_CONTRACT_ORIG_EXECUTE
except NameError:
    _ELI_VOICE_CONTRACT_ORIG_EXECUTE = globals().get("execute")
    _ELI_VOICE_CONTRACT_ORIG_EXECUTE_ACTION = globals().get("execute_action")

    def _eli_voice_contract_action_name(action):
        return str(action or "").strip().upper().replace("-", "_")

    def _eli_voice_contract_arg_text(args, *keys):
        if not isinstance(args, dict):
            return ""
        for key in keys:
            val = args.get(key)
            if val is not None and str(val).strip():
                return str(val).strip()
        return ""

    def _eli_voice_contract_open_settings():
        import os
        import platform
        import shutil
        import subprocess
        import sys

        system = platform.system().lower()

        if system == "linux":
            candidates = (
                "gnome-control-center",
                "systemsettings",
                "xfce4-settings-manager",
                "mate-control-center",
                "cinnamon-settings",
            )
            for cmd in candidates:
                resolved = shutil.which(cmd)
                if resolved:
                    subprocess.Popen(
                        [resolved],
                        stdout=subprocess.DEVNULL,
                        stderr=subprocess.DEVNULL,
                        start_new_session=True,
                    )
                    return {
                        "ok": True,
                        "action": "OPEN_SYSTEM_SETTINGS",
                        "content": f"Opened system settings: {cmd}",
                        "response": f"Opened system settings: {cmd}",
                    }

        if system == "darwin":
            subprocess.Popen(
                ["open", "-b", "com.apple.systempreferences"],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                start_new_session=True,
            )
            return {
                "ok": True,
                "action": "OPEN_SYSTEM_SETTINGS",
                "content": "Opened system settings.",
                "response": "Opened system settings.",
            }

        if system == "windows":
            try:
                os.startfile("ms-settings:")
                return {
                    "ok": True,
                    "action": "OPEN_SYSTEM_SETTINGS",
                    "content": "Opened system settings.",
                    "response": "Opened system settings.",
                }
            except Exception:
                pass

        return {
            "ok": False,
            "action": "OPEN_SYSTEM_SETTINGS",
            "error": "No supported system settings launcher was found on PATH.",
            "content": "No supported system settings launcher was found on PATH.",
            "response": "No supported system settings launcher was found on PATH.",
        }

    def _eli_voice_contract_set_volume(percent):
        import platform
        import shutil
        import subprocess

        level = max(0, min(100, int(percent)))
        system = platform.system().lower()

        if system == "linux":
            pactl = shutil.which("pactl")
            if pactl:
                subprocess.run(
                    [pactl, "set-sink-mute", "@DEFAULT_SINK@", "0"],
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                    check=False,
                )
                rc = subprocess.run(
                    [pactl, "set-sink-volume", "@DEFAULT_SINK@", f"{level}%"],
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                    check=False,
                ).returncode
                if rc == 0:
                    return {
                        "ok": True,
                        "action": "VOLUME",
                        "content": f"Volume set to {level}%",
                        "response": f"Volume set to {level}%",
                    }

            amixer = shutil.which("amixer")
            if amixer:
                rc = subprocess.run(
                    [amixer, "-D", "pulse", "sset", "Master", f"{level}%"],
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                    check=False,
                ).returncode
                if rc == 0:
                    return {
                        "ok": True,
                        "action": "VOLUME",
                        "content": f"Volume set to {level}%",
                        "response": f"Volume set to {level}%",
                    }

        if system == "darwin":
            rc = subprocess.run(
                ["osascript", "-e", f"set volume output volume {level}"],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                check=False,
            ).returncode
            if rc == 0:
                return {
                    "ok": True,
                    "action": "VOLUME",
                    "content": f"Volume set to {level}%",
                    "response": f"Volume set to {level}%",
                }

        return None

    def _eli_voice_contract_direct_result(action, args):
        action_name = _eli_voice_contract_action_name(action)
        args = args if isinstance(args, dict) else {}

        if action_name in {"NOOP", "SAY", "DIRECT_RESPONSE", "ANSWER"}:
            msg = (
                args.get("message")
                or args.get("response")
                or args.get("content")
                or ""
            )
            return {
                "ok": True,
                "action": action_name,
                "content": str(msg),
                "response": str(msg),
            }

        app_name = _eli_voice_contract_arg_text(args, "name", "app", "target", "application").lower()
        if action_name == "OPEN_SYSTEM_SETTINGS" or (
            action_name in {"OPEN_APP", "OPEN_APPLICATION", "LAUNCH_APP"}
            and app_name in {"settings", "system settings", "gnome settings"}
        ):
            return _eli_voice_contract_open_settings()

        if action_name == "VOLUME":
            level = args.get("level", args.get("percent", args.get("value")))
            if level is not None:
                try:
                    result = _eli_voice_contract_set_volume(int(level))
                    if result is not None:
                        return result
                except Exception as exc:
                    return {
                        "ok": False,
                        "action": "VOLUME",
                        "error": f"Volume set failed: {exc}",
                        "content": f"Volume set failed: {exc}",
                        "response": f"Volume set failed: {exc}",
                    }

        return None

    if callable(_ELI_VOICE_CONTRACT_ORIG_EXECUTE):
        def execute(action, args=None, *pargs, **kwargs):
            direct = _eli_voice_contract_direct_result(action, args)
            if direct is not None:
                return direct
            return _ELI_VOICE_CONTRACT_ORIG_EXECUTE(action, args, *pargs, **kwargs)

    if callable(_ELI_VOICE_CONTRACT_ORIG_EXECUTE_ACTION):
        def execute_action(action, args=None, *pargs, **kwargs):
            direct = _eli_voice_contract_direct_result(action, args)
            if direct is not None:
                return direct
            return _ELI_VOICE_CONTRACT_ORIG_EXECUTE_ACTION(action, args, *pargs, **kwargs)

# portable_runtime_contract_v3_executor_hook
try:
    _ELI_PORTABLE_V3_ORIG_EXECUTE
except NameError:
    _ELI_PORTABLE_V3_ORIG_EXECUTE = globals().get("execute")
    _ELI_PORTABLE_V3_ORIG_EXECUTE_ACTION = globals().get("execute_action")

    def _eli_v3_action_name(action):
        return str(action or "").strip().upper().replace("-", "_")

    def _eli_v3_args(args):
        return args if isinstance(args, dict) else {}

    def _eli_v3_error(action_name, text):
        return {
            "ok": False,
            "action": action_name,
            "error": str(text),
            "content": str(text),
            "response": str(text),
            "evidence": [str(text)],
        }

    def _eli_v3_direct_system_action(action, args=None):
        action_name = _eli_v3_action_name(action)
        data = _eli_v3_args(args)

        if action_name in {"OPEN_APP", "LAUNCH_APP", "OPEN_APPLICATION"}:
            from eli.system.portable_app_control import open_app
            _r = open_app(data.get("name") or data.get("target") or data.get("app") or "")
            # Only short-circuit on success; failures fall through to _execute_impl
            # so grounded_remediation can offer the install/download dialogue.
            if isinstance(_r, dict) and _r.get("ok"):
                return _r
            return None

        if action_name in {"CLOSE_APP", "QUIT_APP", "EXIT_APP", "CLOSE_APPLICATION"}:
            from eli.system.portable_app_control import close_app
            _r = close_app(
                data.get("name") or data.get("target") or data.get("app") or "",
                force=bool(data.get("force", False)),
            )
            if isinstance(_r, dict) and _r.get("ok"):
                return _r
            return None

        if action_name in {
            "MINIMIZE_APP", "MINIMISE_APP", "HIDE_APP",
            "MINIMIZE_WINDOW", "MINIMISE_WINDOW",
        }:
            from eli.system.portable_app_control import minimize_app
            _r = minimize_app(data.get("name") or data.get("target") or data.get("app") or "")
            if isinstance(_r, dict) and _r.get("ok"):
                return _r
            return None

        return None

    def _eli_v3_project_root():
        from pathlib import Path
        import os
        env_root = os.environ.get("ELI_PROJECT_ROOT")
        if env_root:
            return Path(env_root).expanduser().resolve()
        return Path(__file__).resolve().parents[2]

    def _eli_v3_artifacts_dir():
        from pathlib import Path
        import os
        root = _eli_v3_project_root()
        env_dir = os.environ.get("ELI_ARTIFACTS_DIR")
        if env_dir:
            p = Path(env_dir).expanduser()
            return p.resolve() if p.is_absolute() else (root / p).resolve()
        return (root / "artifacts").resolve()

    def _eli_v3_slug(text, default="generated_script"):
        import re
        slug = re.sub(r"[^A-Za-z0-9._-]+", "_", str(text or "").strip()).strip("._-")
        return slug[:80] or default

    def _eli_v3_language_ext(language):
        mapping = {
            "python": "py", "py": "py",
            "bash": "sh", "shell": "sh", "sh": "sh", "zsh": "zsh",
            "javascript": "js", "js": "js",
            "typescript": "ts", "ts": "ts",
            "c++": "cpp", "cpp": "cpp",
            "c": "c",
            "c#": "cs", "csharp": "cs",
            "java": "java",
            "rust": "rs",
            "go": "go",
            "ruby": "rb",
            "php": "php",
            "lua": "lua",
            "r": "r",
            "swift": "swift",
            "kotlin": "kt",
            "scala": "scala",
            "sql": "sql",
            "html": "html",
            "css": "css",
            "json": "json",
            "yaml": "yaml", "yml": "yml",
        }
        lang = str(language or "auto").strip().lower()
        return mapping.get(lang, lang if lang and lang != "auto" and len(lang) <= 12 else "txt")

    def _eli_v3_extract_code(raw, language="auto"):
        import re
        if isinstance(raw, dict):
            for key in ("code", "content", "response", "text", "answer", "message", "output"):
                val = raw.get(key)
                if val is not None and str(val).strip():
                    raw = str(val)
                    break
            else:
                raw = str(raw)
        else:
            raw = str(raw or "")

        fences = re.findall(r"```([A-Za-z0-9+#._-]*)\n(.*?)```", raw, flags=re.DOTALL)
        if fences:
            wanted = str(language or "").lower()
            for lang, code in fences:
                if wanted and wanted != "auto" and lang.lower() == wanted:
                    return code.strip()
            return fences[0][1].strip()

        return raw.strip()

    def _eli_v3_generate_script(args=None):
        action_name = "GENERATE_SCRIPT"
        data = _eli_v3_args(args)

        description = (
            data.get("description")
            or data.get("prompt")
            or data.get("task")
            or data.get("query")
            or data.get("message")
            or ""
        )

        if not str(description).strip():
            return _eli_v3_error(action_name, "No script description supplied.")

        language = data.get("language") or "auto"
        if str(language).strip().lower() in {"", "auto", "none"}:
            try:
                from eli.execution.portable_intent_contract import infer_script_language
                language = infer_script_language(description)
            except Exception:
                language = "auto"

        generation_prompt = (
            "You are ELI's local code generation engine. Return a complete, runnable source file only. "
            f"Requested language: {language}.\n\n"
            "Hard requirements:\n"
            "- Solve the user's actual task; do not echo or rephrase the prompt.\n"
            "- No TODO markers, placeholder functions, fake outputs, or hardcoded demonstration answers.\n"
            "- Include imports, validation, error handling, and a clear executable entry point where the language supports one.\n"
            "- For Python: use type hints, docstrings where they clarify behaviour, and an `if __name__ == \"__main__\"` guard for scripts.\n"
            "- For CLI-style requests: use arguments with useful help text instead of burying constants in the body.\n"
            "- Output only raw source code. No markdown fences, no commentary, no postscript.\n\n"
            f"Task:\n{description}"
        )

        raw = None
        chat_fn = globals().get("chat")
        if callable(chat_fn):
            try:
                raw = chat_fn(generation_prompt, skip_router=True)
            except TypeError:
                raw = chat_fn(generation_prompt)

        if raw is None:
            try:
                from eli.cognition import gguf_inference
                raw = gguf_inference.generate(
                    generation_prompt,
                    max_tokens=int(data.get("max_tokens") or 1200),
                )
            except Exception as exc:
                return _eli_v3_error(action_name, f"Script generation backend unavailable: {exc}")

        code = _eli_v3_extract_code(raw, language=language)
        if not code:
            return _eli_v3_error(action_name, "Script generation returned no source code.")

        bad_markers = (
            "TODO",
            "Add code here",
            "placeholder",
            "Generate only the requested source code",
            "This is a request for",
        )
        if any(marker.lower() in code.lower() for marker in bad_markers):
            return _eli_v3_error(
                action_name,
                "Generated script rejected: output contained stub/template markers.",
            )
        if re.fullmatch(r"\s*(?:pass|return\s+None)\s*", code):
            return _eli_v3_error(
                action_name,
                "Generated script rejected: output was an empty implementation.",
            )

        ext = _eli_v3_language_ext(language)

        if ext == "py":
            try:
                compile(code, "<eli-generated-script>", "exec")
            except SyntaxError as _v3_syn_exc:
                # One retry with a stricter prompt: standard library only, no
                # third-party packages. Catches hallucinated APIs like
                # pytz.utc.location.LocationIn that pass regex checks but fail compile.
                _v3_retry_prompt = (
                    "You are ELI's local code generation engine. The previous attempt had a syntax error: "
                    f"{_v3_syn_exc}\n\n"
                    "Write a CORRECTED, complete, runnable Python script for the task below.\n"
                    "STRICT REQUIREMENTS:\n"
                    "- Standard library ONLY. Do not import any third-party package (no ephem, pytz, astral, requests, etc.).\n"
                    "- No markdown fences, no commentary, no TODO markers.\n"
                    "- Must be syntactically valid Python 3.\n"
                    "- Use an `if __name__ == '__main__':` entry point.\n\n"
                    f"Task:\n{description}"
                )
                _v3_retry_raw = None
                _v3_chat = globals().get("chat")
                if callable(_v3_chat):
                    try:
                        _v3_retry_raw = _v3_chat(_v3_retry_prompt, skip_router=True)
                    except TypeError:
                        _v3_retry_raw = _v3_chat(_v3_retry_prompt)
                if _v3_retry_raw is None:
                    try:
                        from eli.cognition import gguf_inference as _v3_gi
                        _v3_retry_raw = _v3_gi.generate(
                            _v3_retry_prompt,
                            max_tokens=int(data.get("max_tokens") or 1200),
                        )
                    except Exception:
                        pass
                if _v3_retry_raw:
                    code = _eli_v3_extract_code(_v3_retry_raw, language=language)
                    try:
                        compile(code, "<eli-generated-script>", "exec")
                    except SyntaxError as _v3_retry_exc:
                        return _eli_v3_error(
                            action_name,
                            f"Generated Python script failed syntax validation (after retry): {_v3_retry_exc}",
                        )
                else:
                    return _eli_v3_error(
                        action_name,
                        f"Generated Python script failed syntax validation: {_v3_syn_exc}",
                    )

        out_dir = _eli_v3_artifacts_dir() / "scripts"
        out_dir.mkdir(parents=True, exist_ok=True)

        stem = _eli_v3_slug(data.get("title") or description)
        path = out_dir / f"{stem}.{ext}"

        if path.exists() and not bool(data.get("overwrite", False)):
            idx = 2
            while True:
                candidate = out_dir / f"{path.stem}_{idx}{path.suffix}"
                if not candidate.exists():
                    path = candidate
                    break
                idx += 1

        path.write_text(code, encoding="utf-8")

        text = json.dumps(
            {
                "event": "artifact_generated",
                "kind": "script",
                "path": str(path),
                "language": language,
            },
            ensure_ascii=False,
            default=str,
        )
        return {
            "ok": True,
            "action": action_name,
            "script_path": str(path),
            "path": str(path),
            "language": language,
            "destination": data.get("destination") or "labs_sim_ide",
            "open_in_labs": bool(data.get("open_in_labs", True)),
            "open_in_ide": bool(data.get("open_in_ide", True)),
            "content": text,
            "response": text,
        }

    if callable(_ELI_PORTABLE_V3_ORIG_EXECUTE):
        def execute(action, args=None, *pargs, **kwargs):
            action_name = _eli_v3_action_name(action)

            direct = _eli_v3_direct_system_action(action, args)
            if direct is not None:
                return direct

            if action_name in {
                "GENERATE_SCRIPT", "CREATE_SCRIPT", "WRITE_SCRIPT",
                "GENERATE_CODE", "WRITE_CODE",
            }:
                return _eli_v3_generate_script(args)

            return _ELI_PORTABLE_V3_ORIG_EXECUTE(action, args, *pargs, **kwargs)

    if callable(_ELI_PORTABLE_V3_ORIG_EXECUTE_ACTION):
        def execute_action(action, args=None, *pargs, **kwargs):
            action_name = _eli_v3_action_name(action)

            direct = _eli_v3_direct_system_action(action, args)
            if direct is not None:
                return direct

            if action_name in {
                "GENERATE_SCRIPT", "CREATE_SCRIPT", "WRITE_SCRIPT",
                "GENERATE_CODE", "WRITE_CODE",
            }:
                return _eli_v3_generate_script(args)

            return _ELI_PORTABLE_V3_ORIG_EXECUTE_ACTION(action, args, *pargs, **kwargs)


# ELI_REASONING_MODE_EXECUTOR_FIX_20260505
_ELI_REASONING_MODE_ORIG_EXECUTE = globals().get("execute")
_ELI_REASONING_MODE_ORIG_EXECUTE_ACTION = globals().get("execute_action")

def _eli_reasoning_mode_execute(action, args=None, *pargs, **kwargs):
    action_name = str(action or "").upper()
    if action_name == "REASONING_MODE_STATUS":
        try:
            from eli.runtime.reasoning_status import current_reasoning_mode_text
            msg = current_reasoning_mode_text()
        except Exception:
            msg = "Current reasoning mode: Quick"
        return {"ok": True, "action": "REASONING_MODE_STATUS", "content": msg, "response": msg}

    orig = _ELI_REASONING_MODE_ORIG_EXECUTE
    if callable(orig):
        return orig(action, args or {}, *pargs, **kwargs)

    orig_action = _ELI_REASONING_MODE_ORIG_EXECUTE_ACTION
    if callable(orig_action):
        return orig_action(action, args or {}, *pargs, **kwargs)

    return {"ok": False, "action": action_name, "content": "Executor unavailable.", "response": "Executor unavailable."}

try:
    execute = _eli_reasoning_mode_execute
    execute_action = _eli_reasoning_mode_execute
except Exception:
    pass


# ELI_EXECUTOR_VISIBLE_TILE_SECOND_FIX_20260505
# Terminal action wrappers. Must sit late in file to override previous execute wrappers.
import math as _eli_tile_math
import re as _eli_tile_re
import subprocess as _eli_tile_subprocess


from eli.utils.log import get_logger
log = get_logger(__name__)

_ELI_TILE_ORIG_EXECUTE = globals().get("execute")
_ELI_TILE_ORIG_EXECUTE_ACTION = globals().get("execute_action")

def _eli_tile_run(cmd, timeout=2):
    try:
        return _eli_tile_subprocess.run(
            cmd,
            text=True,
            stdout=_eli_tile_subprocess.PIPE,
            stderr=_eli_tile_subprocess.PIPE,
            timeout=timeout,
        )
    except Exception as e:
        class _R:
            returncode = 999
            stdout = ""
            stderr = str(e)
        return _R()

def _eli_tile_current_desktop():
    p = _eli_tile_run(["wmctrl", "-d"])
    if p.returncode != 0:
        return None
    for line in p.stdout.splitlines():
        if "*" in line:
            try:
                return int(line.split()[0])
            except Exception:
                return None
    return None

def _eli_tile_screen_size():
    p = _eli_tile_run(["xdotool", "getdisplaygeometry"])
    if p.returncode == 0:
        parts = p.stdout.strip().split()
        if len(parts) >= 2 and parts[0].isdigit() and parts[1].isdigit():
            return int(parts[0]), int(parts[1])

    p = _eli_tile_run(["xrandr", "--current"])
    if p.returncode == 0:
        m = _eli_tile_re.search(r"current\s+(\d+)\s+x\s+(\d+)", p.stdout)
        if m:
            return int(m.group(1)), int(m.group(2))

    return 1366, 768

def _eli_tile_xprop(wid):
    p = _eli_tile_run(["xprop", "-id", wid], timeout=1)
    return p.stdout if p.returncode == 0 else ""

def _eli_tile_visible_windows():
    p = _eli_tile_run(["wmctrl", "-lG", "-p"])
    if p.returncode != 0:
        return [], p.stderr.strip() or "wmctrl failed"

    curdesk = _eli_tile_current_desktop()
    wins = []

    for line in p.stdout.splitlines():
        parts = line.split(None, 8)
        if len(parts) < 9:
            continue

        wid, desk_s, pid_s, x_s, y_s, w_s, h_s, host, title = parts
        try:
            desk = int(desk_s)
            x, y, w, h = int(x_s), int(y_s), int(w_s), int(h_s)
        except Exception:
            continue

        if curdesk is not None and desk not in {curdesk, -1}:
            continue
        if w < 80 or h < 80:
            continue
        if not str(title or "").strip():
            continue

        xp = _eli_tile_xprop(wid)
        xp_low = xp.lower()

        # Skip hidden/minimized/taskbar-skipped/system surfaces.
        if "_net_wm_state_hidden" in xp_low:
            continue
        if "_net_wm_state_skip_taskbar" in xp_low:
            continue
        if "_net_wm_state_skip_pager" in xp_low:
            continue

        # Keep normal/dialog windows; skip docks, desktops, menus, tooltips, splash, etc.
        if "_net_wm_window_type" in xp_low:
            bad_types = (
                "_net_wm_window_type_desktop",
                "_net_wm_window_type_dock",
                "_net_wm_window_type_toolbar",
                "_net_wm_window_type_menu",
                "_net_wm_window_type_utility",
                "_net_wm_window_type_splash",
                "_net_wm_window_type_dropdown_menu",
                "_net_wm_window_type_popup_menu",
                "_net_wm_window_type_tooltip",
                "_net_wm_window_type_notification",
            )
            if any(t in xp_low for t in bad_types):
                continue

        wins.append({"id": wid, "desk": desk, "x": x, "y": y, "w": w, "h": h, "title": title})

    return wins, ""

def _eli_tile_parse_grid(args, count):
    args = args or {}
    grid = args.get("grid")
    cols = args.get("cols") or args.get("columns")
    rows = args.get("rows")

    if isinstance(grid, (list, tuple)) and len(grid) >= 2:
        cols = cols or grid[0]
        rows = rows or grid[1]

    try:
        cols = int(cols) if cols else 0
    except Exception:
        cols = 0
    try:
        rows = int(rows) if rows else 0
    except Exception:
        rows = 0

    if cols > 0 and rows > 0:
        return max(1, min(cols, 8)), max(1, min(rows, 8))

    if count <= 1:
        return 1, 1
    cols = int(_eli_tile_math.ceil(_eli_tile_math.sqrt(count)))
    rows = int(_eli_tile_math.ceil(count / cols))
    return max(1, cols), max(1, rows)

def _eli_tile_windows(args=None):
    wins, err = _eli_tile_visible_windows()
    if err:
        return {"ok": False, "action": "TILE_WINDOWS", "content": err, "response": err, "error": err}

    count = len(wins)
    if count == 0:
        msg = "No visible normal windows found to tile."
        return {"ok": False, "action": "TILE_WINDOWS", "content": msg, "response": msg, "count": 0}

    cols, rows = _eli_tile_parse_grid(args or {}, count)
    screen_w, screen_h = _eli_tile_screen_size()

    margin = int((args or {}).get("margin", 10) or 10)
    top_reserved = int((args or {}).get("top_reserved", 34) or 34)

    usable_x = margin
    usable_y = top_reserved + margin
    usable_w = max(300, screen_w - margin * 2)
    usable_h = max(240, screen_h - top_reserved - margin * 2)

    cell_w = max(180, usable_w // cols)
    cell_h = max(140, usable_h // rows)

    moved = 0
    for i, win in enumerate(wins[: cols * rows]):
        c = i % cols
        r = i // cols
        x = usable_x + c * cell_w
        y = usable_y + r * cell_h
        w = max(120, cell_w - margin)
        h = max(100, cell_h - margin)

        wid = win["id"]
        _eli_tile_run(["wmctrl", "-ir", wid, "-b", "remove,maximized_vert,maximized_horz"], timeout=1)
        p = _eli_tile_run(["wmctrl", "-ir", wid, "-e", f"0,{x},{y},{w},{h}"], timeout=2)
        if p.returncode == 0:
            moved += 1

    msg = f"Tiled {moved} visible window{'s' if moved != 1 else ''} into a {cols}×{rows} grid."
    if count > cols * rows:
        msg += f" {count - cols * rows} visible window(s) did not fit in the requested grid."

    return {
        "ok": moved > 0,
        "action": "TILE_WINDOWS",
        "content": msg,
        "response": msg,
        "count": moved,
        "visible_count": count,
        "grid": [cols, rows],
    }

def _eli_second_execute(action, args=None, *pargs, **kwargs):
    action_name = str(action or "").upper()
    args = args or {}

    if action_name == "REASONING_MODE_STATUS":
        try:
            from eli.runtime.reasoning_status import current_reasoning_mode_text
            msg = current_reasoning_mode_text()
        except Exception:
            msg = "Current reasoning mode: unavailable"
        return {"ok": True, "action": "REASONING_MODE_STATUS", "content": msg, "response": msg}

    if action_name == "TILE_WINDOWS":
        return _eli_tile_windows(args)

    if callable(_ELI_TILE_ORIG_EXECUTE):
        return _ELI_TILE_ORIG_EXECUTE(action, args, *pargs, **kwargs)
    if callable(_ELI_TILE_ORIG_EXECUTE_ACTION):
        return _ELI_TILE_ORIG_EXECUTE_ACTION(action, args, *pargs, **kwargs)

    msg = f"No executor available for {action_name}"
    return {"ok": False, "action": action_name, "content": msg, "response": msg}

try:
    execute = _eli_second_execute
    execute_action = _eli_second_execute
except Exception:
    pass


# ELI_PERSONAL_MEMORY_MODE_AWARE_EXECUTOR_FIX_20260505
# Must remain late in file so this execute wrapper wins over previous wrappers.
_ELI_PM_ORIG_EXECUTE = globals().get("execute")
_ELI_PM_ORIG_EXECUTE_ACTION = globals().get("execute_action")

def _eli_pm_execute(action, args=None, *pargs, **kwargs):
    action_name = str(action or "").upper()
    args = args or {}

    if action_name in {"PERSONAL_MEMORY_DEEP_EXPLAIN", "PERSONAL_MEMORY_SUMMARY"}:
        try:
            from eli.runtime.personal_memory_deep_response import build_personal_memory_deep_response
            try:
                from eli.runtime.reasoning_status import current_reasoning_mode_label
                mode_label = current_reasoning_mode_label()
            except Exception:
                mode_label = ""
            msg = build_personal_memory_deep_response(str(args.get("question") or ""), mode_label=mode_label)
        except Exception as e:
            msg = f"Personal memory deep response failed: {type(e).__name__}: {e}"
        return {
            "ok": True,
            "action": action_name,
            "content": msg,
            "response": msg,
            "evidence_source": "personal_memory_sqlite",
            "response_contract": "quick_direct_nonquick_persona_synthesis",
            "generation_invoked": False,
        }

    if action_name == "ROUTING_FAULT_EXPLAIN":
        try:
            from eli.runtime.personal_memory_deep_response import build_routing_fault_explanation
            msg = build_routing_fault_explanation(str(args.get("question") or ""))
        except Exception as e:
            msg = f"Routing fault explanation failed: {type(e).__name__}: {e}"
        return {
            "ok": True,
            "action": "ROUTING_FAULT_EXPLAIN",
            "content": msg,
            "response": msg,
            "evidence_source": "routing_trace",
            "response_contract": "quick_direct_nonquick_persona_synthesis",
            "generation_invoked": False,
        }

    if action_name == "NAME_SOURCE_AUDIT":
        try:
            from eli.runtime.user_visible_response_surface import format_name_source_audit
            msg = format_name_source_audit(str(args.get("question") or ""))
        except Exception as e:
            msg = f"Name-source audit failed: {type(e).__name__}: {e}"
        return {
            "ok": True,
            "action": "NAME_SOURCE_AUDIT",
            "content": msg,
            "response": msg,
            "evidence_source": "local_identity_sources",
            "response_contract": "quick_direct_nonquick_persona_synthesis",
            "generation_invoked": False,
        }

    if callable(_ELI_PM_ORIG_EXECUTE):
        return _ELI_PM_ORIG_EXECUTE(action, args, *pargs, **kwargs)
    if callable(_ELI_PM_ORIG_EXECUTE_ACTION):
        return _ELI_PM_ORIG_EXECUTE_ACTION(action, args, *pargs, **kwargs)

    msg = f"No executor available for {action_name}"
    return {"ok": False, "action": action_name, "content": msg, "response": msg}

try:
    execute = _eli_pm_execute
    execute_action = _eli_pm_execute
except Exception:
    pass


# --- ELI generated-script safety installer ---
try:
    from eli.runtime.generated_script_guard import install as _eli_install_generated_script_guard
    _eli_install_generated_script_guard(globals())
    log.debug("[EXECUTOR] generated-script safety wrapper installed")
except Exception as _eli_gen_guard_err:
    log.debug(f"[EXECUTOR] generated-script guard install failed: {_eli_gen_guard_err}")
# --- end ELI generated-script safety installer ---

# =============================================================================
# ELI IDENTITY-ONLY EXECUTOR CONTRACT
# Keeps USER_IDENTITY_SUMMARY from leaking preferences / working-style / project
# memory into name/identity questions. No hardcoded user names.
# =============================================================================
try:
    def _eli_identity_only_name():
        try:
            from eli.kernel.state import get_user_name, set_user_name
            name = str(get_user_name("") or "").strip()
            if name:
                return name
            # Fallback: recover from user_patterns SQLite if profile lacks name.
            # This handles the case where the user stated their name in a past
            # session but the profile was never updated from user_patterns.
            try:
                import sqlite3 as _sq3, re as _re2
                from eli.core.paths import user_db_path
                _db = str(user_db_path())
                _con = _sq3.connect(_db)
                _rows = _con.execute(
                    "SELECT pattern_data FROM user_patterns "
                    "WHERE pattern_type='identity.name' "
                    "  AND length(COALESCE(pattern_data,''))>3 "
                    "ORDER BY COALESCE(ts, timestamp, id) DESC LIMIT 5"
                ).fetchall()
                _con.close()
                for (_pdata,) in _rows:
                    _m = _re2.search(
                        r"(?:name is|preferred name is|called|known as)\s+([A-Za-z][A-Za-z' -]{1,29})\b",
                        str(_pdata or ""), _re2.I,
                    )
                    if _m:
                        _candidate = _m.group(1).strip(" .,;")
                        if _candidate.lower() not in {
                            "unknown", "none", "the", "a", "an", "no", "not", "asking", "user",
                            "screenshot", "name", "unnamed", "anonymous", "guest", "admin",
                            "root", "system", "default", "test", "sample", "example",
                            "placeholder", "null", "undefined",
                        }:
                            try:
                                set_user_name(_candidate)  # persist for future sessions
                            except Exception:
                                pass
                            return _candidate
            except Exception:
                pass
            return ""
        except Exception:
            return ""

    def _eli_identity_only_active_user_id():
        try:
            from eli.kernel.state import get_active_user_id
            return str(get_active_user_id() or "").strip()
        except Exception:
            return ""

    def _eli_identity_only_response(args):
        args = args or {}
        scope = str(args.get("identity_scope") or "identity_only").strip().lower()
        name = _eli_identity_only_name()
        uid = _eli_identity_only_active_user_id()

        lines = []

        if scope == "name_only":
            if name:
                lines.append(f"Confirmed active-user name: {name}")
            else:
                lines.append("No confirmed name is stored for the active user.")
            if uid:
                lines.append(f"Active user ID: {uid}")
            return "\n".join(lines)

        if scope == "memory_presence_only":
            lines.append("Yes, I have memory records associated with the active user profile.")
            if name:
                lines.append(f"Confirmed active-user name: {name}")
            else:
                lines.append("No confirmed name is stored for the active user.")
            if uid:
                lines.append(f"Active user ID: {uid}")
            lines.append("I am not dumping preferences here because this is an identity/presence question, not a profile-summary request.")
            return "\n".join(lines)

        # identity_only
        if name:
            lines.append(f"Confirmed active-user identity: {name}")
        else:
            lines.append("No confirmed name/identity label is stored for the active user.")
        if uid:
            lines.append(f"Active user ID: {uid}")
        lines.append("Profile preferences are intentionally excluded from this identity-only answer.")
        return "\n".join(lines)

    # identity-only wrapper removed — mw_identity_only in the canonical
    # middleware table now calls _eli_identity_only_response() directly.
    pass
except Exception as _eli_identity_only_err:
    log.debug(f"[EXECUTOR] identity-only helpers failed: {_eli_identity_only_err}")
# =============================================================================

# =============================================================================
# ELI PROFILE MEMORY EXECUTION CONTRACT
# Makes PERSONAL_MEMORY_SUMMARY mode-aware:
#   inventory_only       -> category counts only, no preference dump
#   preferences_detail   -> explicit preference/working-style details
#   full_profile         -> full active-user profile snapshot
#
# Also removes identity/name commentary from EXPLAIN_MEMORY_RUNTIME.
# No user names are hardcoded here.
# =============================================================================
try:
    def _eli_profile_scope_active_user_id():
        try:
            from eli.kernel.state import get_active_user_id
            return str(get_active_user_id() or "").strip()
        except Exception:
            return ""

    def _eli_profile_scope_user_info_text(reason="profile_scope_contract"):
        try:
            from eli.cognition.user_info_builder import refresh_user_info
            res = refresh_user_info(force=True, reason=reason)
            return str(res.get("text") or "")
        except Exception as exc:
            return f""

    def _eli_profile_scope_sections(text):
        import re as _re
        sections = {}
        current = None
        for raw in str(text or "").splitlines():
            line = raw.strip()
            m = _re.fullmatch(r"\[([^\]]+)\]", line)
            if m:
                current = m.group(1).strip()
                sections.setdefault(current, [])
                continue
            if current and line.startswith("- "):
                item = line[2:].strip()
                if item and item.lower() != "none confirmed.":
                    sections.setdefault(current, []).append(item)
        return sections

    def _eli_profile_scope_inventory_answer(question):
        uid = _eli_profile_scope_active_user_id()
        text = _eli_profile_scope_user_info_text("profile_inventory_only")
        sections = _eli_profile_scope_sections(text)

        interesting = [
            "Identity",
            "Communication Preferences",
            "Working Style",
            "Active Projects",
            "Technical Environment",
            "Constraints / Avoidances",
            "Recent Significant Changes",
            "Uncertain / Needs Confirmation",
        ]

        lines = [
            "Active-user profile memory inventory:",
            f"- active_user_id: {uid or 'unknown'}",
        ]

        identity_count = len(sections.get("Identity", []))
        lines.append(f"- confirmed_identity_items: {identity_count}")

        stored = []
        for name in interesting:
            n = len(sections.get(name, []))
            if n:
                stored.append(f"{name}({n})")

        lines.append("- stored_profile_categories: " + (", ".join(stored) if stored else "none"))
        lines.append("- detail_policy: preference/project details are withheld unless explicitly requested.")
        lines.append("- ask_for_details: use 'show my preferences' or 'dump my full profile memory'.")

        return "\n".join(lines)

    def _eli_profile_scope_preferences_answer(question):
        uid = _eli_profile_scope_active_user_id()
        text = _eli_profile_scope_user_info_text("profile_preferences_detail")
        sections = _eli_profile_scope_sections(text)

        prefs = []
        prefs.extend(sections.get("Communication Preferences", []))
        prefs.extend(sections.get("Working Style", []))
        prefs.extend(sections.get("Constraints / Avoidances", []))

        lines = [
            "Stored preference / working-style facts for active user:",
            f"- active_user_id: {uid or 'unknown'}",
        ]

        if not prefs:
            lines.append("- None confirmed.")
        else:
            for item in prefs[:40]:
                lines.append(f"- {item}")

        return "\n".join(lines)

    def _eli_profile_scope_full_answer(question):
        uid = _eli_profile_scope_active_user_id()
        text = _eli_profile_scope_user_info_text("profile_full_detail")
        sections = _eli_profile_scope_sections(text)

        ordered = [
            "Identity",
            "Communication Preferences",
            "Working Style",
            "Active Projects",
            "Technical Environment",
            "Constraints / Avoidances",
            "Recent Significant Changes",
            "Uncertain / Needs Confirmation",
        ]

        lines = [
            "Full active-user profile memory snapshot:",
            f"- active_user_id: {uid or 'unknown'}",
            "",
        ]

        for name in ordered:
            lines.append(f"[{name}]")
            vals = sections.get(name, [])
            if not vals:
                lines.append("- None confirmed.")
            else:
                for item in vals[:60]:
                    lines.append(f"- {item}")
            lines.append("")

        return "\n".join(lines).rstrip()

    def _eli_profile_scope_infer(question, explicit_scope=""):
        import re as _re
        low = _re.sub(r"\s+", " ", str(question or "").strip().lower())

        if explicit_scope:
            return explicit_scope

        if _re.search(r"\b(show|list|summari[sz]e|tell me|what are|display|read)\b.{0,80}\b(my|stored|profile)?\s*preferences\b", low):
            return "preferences_detail"

        if _re.search(r"\b(dump|show|print|read|display)\b.{0,80}\b(full|complete|entire|all)\b.{0,80}\b(profile|personal memory|memory profile)\b", low):
            return "full_profile"

        if low in {
            "what do you know about me",
            "what do you remember about me",
            "what do you remember of me",
            "do you know me",
        }:
            return "inventory_only"

        return ""

    def _eli_profile_scope_execute_answer(action_name, args):
        args = args if isinstance(args, dict) else {}
        question = str(args.get("question") or args.get("query") or args.get("text") or "")
        scope = _eli_profile_scope_infer(question, str(args.get("profile_scope") or ""))

        if scope == "inventory_only":
            text = _eli_profile_scope_inventory_answer(question)
        elif scope == "preferences_detail":
            text = _eli_profile_scope_preferences_answer(question)
        elif scope == "full_profile":
            text = _eli_profile_scope_full_answer(question)
        else:
            return None

        return {
            "ok": True,
            "action": str(action_name or "").upper(),
            "content": text,
            "response": text,
            "evidence_source": "active_user_profile_scope",
                "report": {
                    "profile_scope": str(((locals().get("real_args") or locals().get("args") or {}) or {}).get("profile_scope") or "unknown"),
                    "active_user_scoped": True,
                },
            "profile_scope": scope,
        }

    def _eli_profile_scope_distinct_sessions():
        try:
            import sqlite3
            from pathlib import Path
            from eli.core.paths import get_paths

            uid = _eli_profile_scope_active_user_id()
            db = Path(get_paths().artifacts_dir) / "db" / "user.sqlite3"
            if not db.exists():
                return None

            con = sqlite3.connect(str(db))
            try:
                cols = {r[1] for r in con.execute("PRAGMA table_info(conversation_turns)").fetchall()}
                if "session_id" not in cols:
                    return None

                where = "WHERE COALESCE(session_id,'') <> ''"
                params = []
                if uid and "user_id" in cols:
                    where += " AND COALESCE(user_id,'') = ?"
                    params.append(uid)

                row = con.execute(
                    f"SELECT COUNT(DISTINCT session_id) FROM conversation_turns {where}",
                    params,
                ).fetchone()
                return int(row[0] or 0)
            finally:
                con.close()
        except Exception:
            return None

    def _eli_profile_scope_clean_memory_runtime_text(text):
        import re as _re
        lines = []
        for raw in str(text or "").splitlines():
            low = raw.lower()
            if "strong enough name signal" in low:
                continue
            if "identify you by name" in low:
                continue
            if "confirmed name" in low and "memory runtime" not in low:
                continue
            lines.append(raw)

        cleaned = "\n".join(lines).lstrip()
        n = _eli_profile_scope_distinct_sessions()
        if n is not None:
            cleaned = _re.sub(r"- distinct_sessions:\s*\d+", f"- distinct_sessions: {n}", cleaned)

        return cleaned

    # profile-scope wrapper removed — mw_profile_scope in the canonical middleware
    # table now calls _eli_profile_scope_execute_answer() and
    # _eli_profile_scope_clean_memory_runtime_text() directly.
    pass
except Exception as _eli_profile_scope_exec_err:
    log.debug(f"[EXECUTOR] profile memory scope contract failed: {_eli_profile_scope_exec_err}")
# =============================================================================

# =============================================================================
# ELI MEMORY RUNTIME REPORT SANITIZER
# EXPLAIN_MEMORY_RUNTIME is architecture/runtime evidence only.
# It must not carry identity guesses, profile facts, preference facts, or noisy
# "Session context" identity hits in report metadata.
# =============================================================================
try:
    def _eli_memory_runtime_count_distinct_sessions(report):
        try:
            import sqlite3
            from pathlib import Path

            paths = report.get("paths") if isinstance(report, dict) else {}
            status = report.get("status") if isinstance(report, dict) else {}

            db_path = (
                (status or {}).get("db_path")
                or (paths or {}).get("memory_db")
                or (paths or {}).get("user_db")
                or (paths or {}).get("active_db")
            )
            if not db_path:
                return None

            db = Path(str(db_path)).expanduser()
            if not db.exists():
                return None

            con = sqlite3.connect(str(db))
            try:
                cur = con.cursor()
                tables = {
                    str(r[0])
                    for r in cur.execute(
                        "SELECT name FROM sqlite_master WHERE type='table'"
                    ).fetchall()
                }

                if "conversation_turns" in tables:
                    row = cur.execute(
                        """
                        SELECT COUNT(DISTINCT session_id)
                        FROM conversation_turns
                        WHERE COALESCE(session_id,'') <> ''
                        """
                    ).fetchone()
                    return int(row[0] or 0)

                if "conversations" in tables:
                    row = cur.execute(
                        """
                        SELECT COUNT(DISTINCT session_id)
                        FROM conversations
                        WHERE COALESCE(session_id,'') <> ''
                        """
                    ).fetchone()
                    return int(row[0] or 0)

                return None
            finally:
                con.close()
        except Exception:
            return None


    def _eli_sanitize_memory_runtime_output(out):
        if not isinstance(out, dict):
            return out

        report = out.get("report")
        if not isinstance(report, dict):
            return out

        # Remove identity/profile inference from memory-runtime metadata.
        for key in (
            "name_guess",
            "identity_hits",
            "identity",
            "stored_name",
            "fact_candidates",
            "profile_scope",
            "personal_memory_summary",
        ):
            report.pop(key, None)

        status = report.get("status")
        if isinstance(status, dict):
            fixed_sessions = _eli_memory_runtime_count_distinct_sessions(report)
            if fixed_sessions is not None:
                status["distinct_sessions"] = fixed_sessions
                report["status"] = status

                text = str(out.get("content") or out.get("response") or "")
                text = re.sub(
                    r"(?m)^- distinct_sessions:\s*\d+\s*$",
                    f"- distinct_sessions: {fixed_sessions}",
                    text,
                )
                out["content"] = text
                out["response"] = text

        # Defensive visible-output scrub.
        text = str(out.get("content") or out.get("response") or "")
        bad_fragments = (
            "strong enough name signal",
            "identity_hits",
            "name_guess",
            "Session context:",
        )
        if any(x.lower() in text.lower() for x in bad_fragments):
            clean_lines = []
            for line in text.splitlines():
                low = line.lower()
                if "strong enough name signal" in low:
                    continue
                if "identity_hits" in low:
                    continue
                if "name_guess" in low:
                    continue
                if "session context:" in low:
                    continue
                clean_lines.append(line)
            text = "\n".join(clean_lines).strip()
            out["content"] = text
            out["response"] = text

        out["report"] = report
        out["evidence_source"] = "memory_runtime_sanitized"
        return out


    # memory-runtime sanitizer wrapper removed — mw_memory_runtime_sanitizer in
    # the canonical middleware table now calls _eli_sanitize_memory_runtime_output().
    pass

except Exception as _eli_memory_runtime_sanitizer_err:
    log.debug(f"[EXECUTOR] memory-runtime report sanitizer failed: {_eli_memory_runtime_sanitizer_err}")
# =============================================================================

# =============================================================================
# ELI MEMORY COUNT COMPACT EVIDENCE PROVIDER
# Provides structured evidence only. Non-quick final wording belongs to cognition.
# =============================================================================
try:
    def _eli_memory_count_is_requested(args):
        import re
        if not isinstance(args, dict):
            return False
        if str(args.get("memory_scope") or "").strip().lower() == "count_only":
            return True
        q = str(args.get("question") or args.get("query") or "").lower()
        return bool(re.search(r"\b(how many|number of|count)\b.{0,80}\b(memories|memory entries|stored memories|memory rows)\b", q))

    def _eli_memory_count_path():
        from pathlib import Path
        try:
            from eli.core.paths import user_db_path
            return Path(user_db_path())
        except Exception:
            try:
                from eli.core.paths import get_paths
                return Path(get_paths().artifacts_dir) / "db" / "user.sqlite3"
            except Exception:
                return Path.cwd() / "artifacts" / "db" / "user.sqlite3"

    def _eli_table_exists(conn, table):
        try:
            return bool(conn.execute(
                "SELECT 1 FROM sqlite_master WHERE type IN ('table','view') AND name=?",
                (table,),
            ).fetchone())
        except Exception:
            return False

    def _eli_count(conn, table):
        try:
            if not _eli_table_exists(conn, table):
                return 0
            return int(conn.execute(f"SELECT COUNT(*) FROM {table}").fetchone()[0] or 0)
        except Exception:
            return 0

    def _eli_distinct_sessions(conn):
        try:
            if _eli_table_exists(conn, "conversation_turns"):
                cols = [str(r[1]) for r in conn.execute("PRAGMA table_info(conversation_turns)").fetchall()]
                if "session_id" in cols:
                    n = int(conn.execute(
                        "SELECT COUNT(DISTINCT session_id) FROM conversation_turns WHERE COALESCE(session_id,'') <> ''"
                    ).fetchone()[0] or 0)
                    if n:
                        return n
            if _eli_table_exists(conn, "conversations"):
                return _eli_count(conn, "conversations")
        except Exception:
            pass
        return 0

    def _eli_faiss_count():
        try:
            import faiss
            from pathlib import Path
            root = Path(__file__).resolve().parents[2]
            idx = root / "artifacts" / "vectors" / "index.faiss"
            if not idx.exists():
                return 0
            return int(getattr(faiss.read_index(str(idx)), "ntotal", 0) or 0)
        except Exception:
            return 0

    def _eli_memory_count_evidence():
        import sqlite3

        db = _eli_memory_count_path()
        counts = {
            "long_term_memory_rows": 0,
            "memory_fts_rows": 0,
            "faiss_vector_entries": _eli_faiss_count(),
            "conversation_turns": 0,
            "conversation_records": 0,
            "distinct_sessions": 0,
            "learning_replay_rows": 0,
            "observations": 0,
            "user_patterns": 0,
            "recall_log_rows": 0,
        }

        if db.exists():
            con = sqlite3.connect(str(db))
            try:
                counts["long_term_memory_rows"] = _eli_count(con, "memories")
                counts["memory_fts_rows"] = _eli_count(con, "memories_fts")
                counts["conversation_turns"] = _eli_count(con, "conversation_turns")
                counts["conversation_records"] = _eli_count(con, "conversations")
                counts["distinct_sessions"] = _eli_distinct_sessions(con)
                counts["learning_replay_rows"] = _eli_count(con, "learning_replay")
                counts["observations"] = _eli_count(con, "observations")
                counts["user_patterns"] = _eli_count(con, "user_patterns")
                counts["recall_log_rows"] = _eli_count(con, "recall_log")
            finally:
                con.close()

        evidence = (
            "MEMORY_COUNT_EVIDENCE\n"
            f"db_path: {db}\n"
            f"long_term_memory_rows: {counts['long_term_memory_rows']}\n"
            f"memory_fts_rows: {counts['memory_fts_rows']}\n"
            f"faiss_vector_entries: {counts['faiss_vector_entries']}\n"
            f"conversation_turns: {counts['conversation_turns']}\n"
            f"conversation_records: {counts['conversation_records']}\n"
            f"distinct_sessions: {counts['distinct_sessions']}\n"
            f"learning_replay_rows: {counts['learning_replay_rows']}\n"
            f"observations: {counts['observations']}\n"
            f"user_patterns: {counts['user_patterns']}\n"
            f"recall_log_rows: {counts['recall_log_rows']}\n"
            "answer_contract: The direct answer to 'how many memories' is long_term_memory_rows. "
            "Mention related retrieval/conversation counts only as separate supporting counts. "
            "Do not invent any number not present in this evidence."
        )

        return {
            "ok": True,
            "action": "MEMORY_STATUS",
            "content": evidence,
            "response": evidence,
            "evidence_source": "memory_count_compact_sqlite_evidence",
            "report": {
                "ok": True,
                "memory_scope": "count_only",
                "db_path": str(db),
                "counts": counts,
                "requires_grounded_synthesis": True,
                "requires_output_validation": True,
                "quick_direct_allowed": True,
            },
        }

    # memory-count wrapper removed — mw_memory_count in the canonical middleware
    # table now intercepts and calls _eli_memory_count_evidence() directly.
    pass
except Exception as _err:
    log.debug(f"[EXECUTOR] memory-count compact evidence provider failed: {_err}")

# =============================================================================
# ELI RECENT MEMORY PROCESSING EVIDENCE PROVIDER
# MEMORY_STATUS + memory_scope=recent_processing returns compact SQLite-backed
# evidence. No invented "I was processing equations" unless the DB actually says
# that.
# =============================================================================
try:
    import json as _eli_recent_mem_json
    import re as _eli_recent_mem_re
    import sqlite3 as _eli_recent_mem_sqlite
    from pathlib import Path as _eli_recent_mem_Path

    def _eli_recent_mem_paths() -> dict:
        out = {}
        try:
            from eli.core.paths import get_paths
            paths = get_paths()
            root = _eli_recent_mem_Path(getattr(paths, "root", "") or _eli_recent_mem_Path.cwd())
            artifacts = _eli_recent_mem_Path(getattr(paths, "artifacts_dir", "") or (root / "artifacts"))
        except Exception:
            root = _eli_recent_mem_Path.cwd()
            artifacts = root / "artifacts"

        out["root"] = root
        out["user_db"] = artifacts / "db" / "user.sqlite3"
        out["agent_db"] = artifacts / "db" / "agent.sqlite3"
        out["faiss_index"] = artifacts / "vectors" / "index.faiss"
        out["faiss_meta"] = artifacts / "vectors" / "meta.pkl"
        return out

    def _eli_recent_mem_columns(conn, table: str) -> set[str]:
        try:
            return {str(r[1]) for r in conn.execute(f"PRAGMA table_info({table})").fetchall()}
        except Exception:
            return set()

    def _eli_recent_mem_table_count(conn, table: str) -> int:
        try:
            return int(conn.execute(f"SELECT COUNT(*) FROM {table}").fetchone()[0] or 0)
        except Exception:
            return 0

    def _eli_recent_mem_vector_count(paths: dict) -> int:
        try:
            import pickle
            meta = paths.get("faiss_meta")
            if meta and _eli_recent_mem_Path(meta).exists():
                with open(meta, "rb") as f:
                    data = pickle.load(f)
                if isinstance(data, dict):
                    for k in ("rows", "items", "metadata", "ids"):
                        v = data.get(k)
                        if isinstance(v, (list, tuple, dict)):
                            return len(v)
                if isinstance(data, (list, tuple, dict)):
                    return len(data)
        except Exception:
            pass

        try:
            import faiss
            index = paths.get("faiss_index")
            if index and _eli_recent_mem_Path(index).exists():
                return int(faiss.read_index(str(index)).ntotal)
        except Exception:
            pass

        return 0

    def _eli_recent_mem_clean_text(value: object, limit: int = 220) -> str:
        txt = str(value or "").replace("\x00", " ")
        txt = _eli_recent_mem_re.sub(r"\s+", " ", txt).strip()
        if not txt:
            return ""

        # Avoid dumping JSON blobs raw.
        if txt.startswith("{") and txt.endswith("}"):
            try:
                obj = _eli_recent_mem_json.loads(txt)
                if isinstance(obj, dict):
                    for key in ("summary", "content", "text", "event", "action", "response", "message"):
                        if obj.get(key):
                            txt = str(obj.get(key))
                            break
            except Exception:
                pass

        txt = _eli_recent_mem_re.sub(r"\s+", " ", txt).strip()
        if len(txt) > limit:
            txt = txt[: limit - 1].rstrip() + "…"
        return txt

    def _eli_recent_mem_is_noise(txt: str) -> bool:
        low = str(txt or "").lower().strip()
        if not low:
            return True

        noisy_prefixes = (
            "session context:",
            "reflection (24h):",
            "runtime status:",
            "personal memory summary",
            "capability inventory updated:",
            "generated script:",
            "script generated:",
            "proactive daemon started",
            "path not found:",
        )
        if any(low.startswith(x) for x in noisy_prefixes):
            return True

        noisy_contains = (
            "req-000001",
            "sqlite tables observed live",
            "memory_entries:",
            "identity_evidence:",
            "unsupported executor action",
            "name 'q' is not defined",
        )
        if any(x in low for x in noisy_contains):
            return True

        return False

    def _eli_recent_mem_recent_rows(conn, table: str, *, limit: int = 6) -> list[dict]:
        cols = _eli_recent_mem_columns(conn, table)
        if not cols:
            return []

        text_candidates = [
            "content", "summary", "text", "memory", "observation", "event",
            "details", "message", "value", "response", "input", "output",
        ]
        time_candidates = [
            "timestamp", "ts", "created_at", "updated_at", "last_seen_at", "ended_at", "started_at",
        ]

        text_cols = [c for c in text_candidates if c in cols]
        if not text_cols:
            return []

        time_col = next((c for c in time_candidates if c in cols), "")
        select_cols = []
        if "id" in cols:
            select_cols.append("id")
        if time_col:
            select_cols.append(time_col)
        select_cols.extend(text_cols[:4])

        order = f"{time_col} DESC" if time_col else "rowid DESC"

        try:
            rows = conn.execute(
                f"SELECT {', '.join(select_cols)} FROM {table} ORDER BY {order} LIMIT ?",
                (max(limit * 4, 20),),
            ).fetchall()
        except Exception:
            return []

        out = []
        for row in rows:
            d = dict(row)
            parts = []
            for c in text_cols[:4]:
                val = _eli_recent_mem_clean_text(d.get(c), limit=220)
                if val:
                    parts.append(val)

            text = " | ".join(dict.fromkeys(parts))
            text = _eli_recent_mem_clean_text(text, limit=260)
            if _eli_recent_mem_is_noise(text):
                continue

            item = {
                "table": table,
                "text": text,
            }
            if time_col and d.get(time_col):
                item["time"] = str(d.get(time_col))
            if "id" in d and d.get("id") is not None:
                item["id"] = d.get("id")
            out.append(item)
            if len(out) >= limit:
                break

        return out

    def _eli_recent_memory_processing_report(question: str = "") -> dict:
        paths = _eli_recent_mem_paths()
        user_db = paths["user_db"]

        counts = {
            "long_term_memories": 0,
            "memory_fts_rows": 0,
            "faiss_vectors": _eli_recent_mem_vector_count(paths),
            "conversation_turns": 0,
            "conversation_records": 0,
            "learning_replay_rows": 0,
            "observations": 0,
            "runtime_events": 0,
            "user_patterns": 0,
        }

        recent = {
            "memories": [],
            "observations": [],
            "learning_replay": [],
            "runtime_events": [],
        }

        if not user_db.exists():
            return {
                "ok": False,
                "question": question,
                "db_path": str(user_db),
                "counts": counts,
                "recent": recent,
                "problems": [f"user DB missing: {user_db}"],
            }

        try:
            conn = _eli_recent_mem_sqlite.connect(str(user_db))
            conn.row_factory = _eli_recent_mem_sqlite.Row
            try:
                counts["long_term_memories"] = _eli_recent_mem_table_count(conn, "memories")
                counts["memory_fts_rows"] = _eli_recent_mem_table_count(conn, "memories_fts")
                counts["conversation_turns"] = _eli_recent_mem_table_count(conn, "conversation_turns")
                counts["conversation_records"] = _eli_recent_mem_table_count(conn, "conversations")
                counts["learning_replay_rows"] = _eli_recent_mem_table_count(conn, "learning_replay")
                counts["observations"] = _eli_recent_mem_table_count(conn, "observations")
                counts["runtime_events"] = _eli_recent_mem_table_count(conn, "runtime_events")
                counts["user_patterns"] = _eli_recent_mem_table_count(conn, "user_patterns")

                recent["memories"] = _eli_recent_mem_recent_rows(conn, "memories", limit=5)
                recent["observations"] = _eli_recent_mem_recent_rows(conn, "observations", limit=5)
                recent["learning_replay"] = _eli_recent_mem_recent_rows(conn, "learning_replay", limit=5)
                recent["runtime_events"] = _eli_recent_mem_recent_rows(conn, "runtime_events", limit=5)
            finally:
                conn.close()
        except Exception as e:
            return {
                "ok": False,
                "question": question,
                "db_path": str(user_db),
                "counts": counts,
                "recent": recent,
                "problems": [repr(e)],
            }

        return {
            "ok": True,
            "question": question,
            "db_path": str(user_db),
            "counts": counts,
            "recent": recent,
            "policy": {
                "must_not_claim": [
                    "Do not claim recent mathematical-equation processing unless present in recent DB rows.",
                    "Do not invent emotional or experiential memory activity.",
                    "Use only the counts and rows in this report.",
                ],
            },
        }

    def _eli_recent_memory_processing_content(report: dict) -> str:
        counts = report.get("counts") or {}
        recent = report.get("recent") or {}

        lines = [
            "Recent memory-processing evidence:",
            f"- long_term_memories: {counts.get('long_term_memories', 0)}",
            f"- memory_fts_rows: {counts.get('memory_fts_rows', 0)}",
            f"- faiss_vectors: {counts.get('faiss_vectors', 0)}",
            f"- conversation_turns: {counts.get('conversation_turns', 0)}",
            f"- conversation_records: {counts.get('conversation_records', 0)}",
            f"- learning_replay_rows: {counts.get('learning_replay_rows', 0)}",
            f"- observations: {counts.get('observations', 0)}",
            "",
        ]

        emitted = False
        for label, key in [
            ("Recent durable memories", "memories"),
            ("Recent observations", "observations"),
            ("Recent learning replay", "learning_replay"),
            ("Recent runtime events", "runtime_events"),
        ]:
            rows = list(recent.get(key) or [])
            if not rows:
                continue
            emitted = True
            lines.append(f"{label}:")
            for r in rows[:5]:
                prefix = f"[{r.get('time')}] " if r.get("time") else ""
                lines.append(f"- {prefix}{r.get('text')}")
            lines.append("")

        if not emitted:
            lines.append("No clean recent memory rows were found after filtering runtime noise.")
            lines.append("")

        lines.append("Grounding rule: if a detail is not listed above, ELI must not claim it was recently processing it.")
        return "\n".join(lines).rstrip()

    # recent-memory-processing wrapper removed — mw_recent_memory_processing in
    # the canonical middleware table now calls _eli_recent_memory_processing_report()
    # and _eli_recent_memory_processing_content() directly.
    pass

except Exception as _eli_recent_memory_exec_err:
    log.debug(f"[EXECUTOR] recent-memory-processing provider install failed: {_eli_recent_memory_exec_err}")

# =============================================================================
# ELI RECENT MEMORY PROCESSING EVIDENCE CLEANUP V2
# The first recent_processing provider was too permissive: it could surface
# runtime echoes, current prompts, model-not-ready messages, or previous
# hallucinated answers as "recent memory processing". This wrapper cleans the
# report before it reaches user-visible synthesis.
# =============================================================================
try:
    def _eli_recent_mem_v2_bad_runtime_text(text: object, question: object = "") -> bool:
        low = str(text or "").strip().lower()
        qlow = str(question or "").strip().lower()

        if not low:
            return True

        # Current prompt echoes are not memory-processing evidence.
        if qlow and low == qlow:
            return True
        if qlow and qlow[:80] and qlow[:80] in low:
            return True

        bad = [
            "complex mathematical equations",
            "area under a curve",
            "linear equations",
            "digital hoarder",
            "filled to the brim",
            "favorite memories",
            "i don't have personal experiences or memories",
            "i do not have personal experiences or memories",
            "model not ready",
            "failed to create llama_context",
            "failed to load model from file",
            "recent memory-processing evidence:",
            "grounding rule:",
            "a diagnostic evidence packet was produced",
            "capability inventory updated:",
            "session context:",
            "runtime truth report",
            "sqlite tables observed live",
        ]

        return any(x in low for x in bad)

    def _eli_recent_mem_v2_clean_report(report: dict) -> dict:
        report = dict(report or {})
        recent = dict(report.get("recent") or {})
        question = report.get("question") or ""

        cleaned = {}
        for key, rows in recent.items():
            good_rows = []
            for row in list(rows or []):
                if not isinstance(row, dict):
                    continue
                text = str(row.get("text") or "")
                if _eli_recent_mem_v2_bad_runtime_text(text, question):
                    continue
                good_rows.append(row)
            cleaned[key] = good_rows[:5]

        # Runtime events are operational traces. Keep them in the structured
        # report only if clean, but do not expose them as primary "memory".
        cleaned["runtime_events"] = cleaned.get("runtime_events", [])[:2]
        report["recent"] = cleaned

        policy = dict(report.get("policy") or {})
        policy["surface_rule"] = (
            "Recent memory-processing answers must summarize durable memories, "
            "observations, learning replay, and clean operational facts only. "
            "Do not treat prior assistant text, prompt echoes, or runtime errors as memory activity."
        )
        report["policy"] = policy
        return report

    def _eli_recent_mem_v2_content(report: dict) -> str:
        counts = report.get("counts") or {}
        recent = report.get("recent") or {}

        lines = [
            "Grounded recent memory-processing answer:",
            f"- long-term memory rows: {counts.get('long_term_memories', 0)}",
            f"- FTS memory rows: {counts.get('memory_fts_rows', 0)}",
            f"- FAISS vector entries: {counts.get('faiss_vectors', 0)}",
            f"- conversation turns: {counts.get('conversation_turns', 0)}",
            f"- learning replay rows: {counts.get('learning_replay_rows', 0)}",
            f"- observations: {counts.get('observations', 0)}",
            "",
        ]

        durable = list(recent.get("memories") or [])
        observations = list(recent.get("observations") or [])
        replay = list(recent.get("learning_replay") or [])

        if durable:
            lines.append("Clean recent durable memory evidence:")
            for row in durable[:3]:
                stamp = f"[{row.get('time')}] " if row.get("time") else ""
                lines.append(f"- {stamp}{row.get('text')}")
            lines.append("")

        if observations:
            lines.append("Clean recent observation evidence:")
            for row in observations[:3]:
                stamp = f"[{row.get('time')}] " if row.get("time") else ""
                lines.append(f"- {stamp}{row.get('text')}")
            lines.append("")

        if replay:
            lines.append("Clean recent learning-replay evidence:")
            for row in replay[:3]:
                stamp = f"[{row.get('time')}] " if row.get("time") else ""
                lines.append(f"- {stamp}{row.get('text')}")
            lines.append("")

        if not durable and not observations and not replay:
            lines.append(
                "I found memory-system activity counts, but no clean recent durable-memory details "
                "that should be described as specific remembered topics."
            )
            lines.append("")

        lines.append(
            "I should not claim I was processing topics that are not present in this evidence."
        )
        return "\n".join(lines).rstrip()

    # recent-memory-processing v2 cleanup wrapper removed — mw_recent_memory_processing
    # in the canonical middleware table now calls _eli_recent_mem_v2_clean_report() and
    # _eli_recent_mem_v2_content() directly.
    pass

except Exception as _eli_recent_mem_v2_exec_err:
    log.debug(f"[EXECUTOR] recent-memory-processing cleanup v2 install failed: {_eli_recent_mem_v2_exec_err}")

# =============================================================================
# ELI SELF-REPORT RECENT UPDATES EVIDENCE PROVIDER
# Provides deterministic evidence for questions like:
# "Tell me about yourself. What updates/checks have been performed lately?"
# =============================================================================
try:
    import json as _eli_self_json
    import subprocess as _eli_self_subprocess
    from pathlib import Path as _eli_self_Path


    def _eli_self_project_root():
        try:
            return _eli_self_Path(__file__).resolve().parents[2]
        except Exception:
            return _eli_self_Path.cwd()

    def _eli_self_run_git(root, git_args, timeout=4):
        try:
            proc = _eli_self_subprocess.run(
                ["git", "-C", str(root), *git_args],
                text=True,
                capture_output=True,
                timeout=timeout,
                check=False,
            )
            if proc.returncode != 0:
                return {
                    "ok": False,
                    "stdout": proc.stdout.strip(),
                    "stderr": proc.stderr.strip(),
                    "returncode": proc.returncode,
                }
            return {
                "ok": True,
                "stdout": proc.stdout.strip(),
                "stderr": proc.stderr.strip(),
                "returncode": proc.returncode,
            }
        except Exception as e:
            return {"ok": False, "stdout": "", "stderr": repr(e), "returncode": -1}

    def _eli_self_read_json(path):
        try:
            p = _eli_self_Path(path)
            if not p.exists():
                return {}
            return _eli_self_json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            return {}

    def _eli_self_split_lines(text, limit=12):
        lines = [x.rstrip() for x in str(text or "").splitlines() if x.strip()]
        return lines[:limit]

    def _eli_self_build_recent_updates_report(question=""):
        root = _eli_self_project_root()

        head = _eli_self_run_git(root, ["log", "--oneline", "--decorate", "-8"])
        status = _eli_self_run_git(root, ["status", "--short"])
        branch = _eli_self_run_git(root, ["branch", "--show-current"])
        tags = _eli_self_run_git(
            root,
            [
                "tag",
                "--list",
                "--sort=-creatordate",
                "identity_scope_clean_*",
                "runtime_identity_media_learning_*",
                "wip_runtime_gui_memory_*",
                "wip_first_run_policy_tracking_*",
                "memory_count_grounded_*",
                "adaptive_cold_gguf_loader_*",
                "effective_runtime_*",
                "recent_memory_processing_*",
            ],
        )

        runtime_snapshot = _eli_self_read_json(root / "artifacts" / "runtime_snapshot.json")
        manifest = _eli_self_read_json(root / "capability_manifest.json")

        runtime = {
            "model_path": runtime_snapshot.get("model_path"),
            "model_name": runtime_snapshot.get("model_name"),
            "loaded": runtime_snapshot.get("loaded"),
            "effective": runtime_snapshot.get("effective") or {
                "n_ctx": runtime_snapshot.get("n_ctx"),
                "n_gpu_layers": runtime_snapshot.get("n_gpu_layers"),
                "n_threads": runtime_snapshot.get("n_threads"),
                "n_batch": runtime_snapshot.get("n_batch"),
            },
            "requested": runtime_snapshot.get("requested") or {
                "n_ctx": runtime_snapshot.get("requested_n_ctx") or runtime_snapshot.get("n_ctx"),
                "n_gpu_layers": runtime_snapshot.get("requested_n_gpu_layers") or runtime_snapshot.get("n_gpu_layers"),
                "n_threads": runtime_snapshot.get("requested_n_threads") or runtime_snapshot.get("n_threads"),
                "n_batch": runtime_snapshot.get("requested_n_batch") or runtime_snapshot.get("n_batch"),
            },
            "runtime_contract": runtime_snapshot.get("runtime_contract"),
            "runtime_source": runtime_snapshot.get("runtime_source"),
        }

        report = {
            "ok": True,
            "question": str(question or ""),
            "project_root": str(root),
            "branch": (branch.get("stdout") or "").strip(),
            "git": {
                "recent_commits": _eli_self_split_lines(head.get("stdout"), 8),
                "status_short": _eli_self_split_lines(status.get("stdout"), 20),
                "tags": _eli_self_split_lines(tags.get("stdout"), 20),
                "head_ok": bool(head.get("ok")),
                "status_ok": bool(status.get("ok")),
            },
            "capabilities": {
                "manifest_total": manifest.get("total"),
                "generated_at": manifest.get("generated_at"),
            },
            "runtime": runtime,
            "policy": {
                "must_not_claim": [
                    "Do not claim user_profile.json unless it appears in this report.",
                    "Do not claim Spotify/app activity unless it appears in this report.",
                    "Do not claim '3 rewired' unless it appears in this report.",
                    "Do not claim vague routine updates/checks without naming concrete Git/runtime evidence.",
                    "Do not invent emotional state such as 'usual nutty self'.",
                ],
                "answer_contract": (
                    "Answer self-report/update questions using only Git commit/tag/status evidence, "
                    "capability manifest count, and runtime snapshot fields in this report."
                ),
            },
        }

        lines = []
        lines.append("Grounded ELI self-report / recent update evidence:")
        lines.append(f"- identity: ELI / Enhanced Learning Interface")
        if report["branch"]:
            lines.append(f"- current_branch: {report['branch']}")
        if report["capabilities"]["manifest_total"] is not None:
            lines.append(f"- capability_manifest_total: {report['capabilities']['manifest_total']}")
        if report["capabilities"]["generated_at"]:
            lines.append(f"- capability_manifest_generated_at: {report['capabilities']['generated_at']}")

        lines.append("")
        lines.append("Recent Git updates:")
        commits = report["git"]["recent_commits"]
        if commits:
            for c in commits:
                lines.append(f"- {c}")
        else:
            lines.append("- No recent Git commit evidence was available.")

        lines.append("")
        lines.append("Relevant clean tags:")
        tag_lines = report["git"]["tags"]
        if tag_lines:
            for t in tag_lines[:10]:
                lines.append(f"- {t}")
        else:
            lines.append("- No relevant clean tags were found.")

        lines.append("")
        lines.append("Runtime snapshot:")
        rt = report["runtime"]
        eff = rt.get("effective") or {}
        req = rt.get("requested") or {}
        lines.append(f"- model_name: {rt.get('model_name')}")
        lines.append(f"- model_path: {rt.get('model_path')}")
        lines.append(f"- loaded: {rt.get('loaded')}")
        lines.append(
            "- effective: "
            f"ctx={eff.get('n_ctx')} "
            f"gpu_layers={eff.get('n_gpu_layers')} "
            f"threads={eff.get('n_threads')} "
            f"batch={eff.get('n_batch')}"
        )
        lines.append(
            "- requested: "
            f"ctx={req.get('n_ctx')} "
            f"gpu_layers={req.get('n_gpu_layers')} "
            f"threads={req.get('n_threads')} "
            f"batch={req.get('n_batch')}"
        )

        lines.append("")
        lines.append("Working tree status:")
        dirty = report["git"]["status_short"]
        if dirty:
            lines.append(f"- dirty_entries: {len(dirty)} shown below")
            for d in dirty[:10]:
                lines.append(f"  - {d}")
        else:
            lines.append("- clean according to git status --short")

        lines.append("")
        lines.append("Grounding rule: if an update/check is not listed above, ELI must not claim it happened.")

        content = "\n".join(lines)

        return {
            "ok": True,
            "action": "SELF_REPORT",
            "evidence_source": "self_report_recent_updates_git_runtime",
            "report": report,
            "content": content,
            "response": content,
        }

    # SELF_REPORT recent-updates wrapper removed — mw_self_report_recent_updates
    # in the canonical middleware table intercepts this and calls
    # _eli_self_build_recent_updates_report() directly.
    pass
except Exception as _eli_self_report_recent_exec_error:
    log.debug(f"[EXECUTOR][WARN] self-report recent-updates provider failed: {_eli_self_report_recent_exec_error}")

# =============================================================================
# ELI_GUI_RUNTIME_AUDIT_VISIBLE_RESULT_CONTRACT
# Normalizes GUI_RUNTIME_AUDIT executor output so GUI receives visible content.
# This does not change routing, cognition, GGUF, or the audit evidence source.
# =============================================================================
try:
    _ELI_GUI_RUNTIME_AUDIT_VISIBLE_PREV = _gui_runtime_audit_report

    def _eli_gui_runtime_audit_entry_to_text(entry):
        entry = entry if isinstance(entry, dict) else {}
        path = entry.get("path") or "<unknown>"
        status = entry.get("status") or ("PASS" if entry.get("ok", True) else "FAIL")
        issues = entry.get("issues") or []
        evidence = entry.get("evidence") or {}

        lines = [
            f"{status} — GUI wiring scan (SHALLOW line-hit check): {path}",
            "",
            "SCOPE (read this first): this only confirms that wiring references appear "
            "on the lines below. It does NOT verify behaviour, correctness, or full-file "
            "semantics. A 'PASS' here is NOT a clean bill of health — do not claim 'no issues'.",
            "",
            "Visible contract:",
            "- GUI_RUNTIME_AUDIT returned structured evidence from executor.",
            "- This response is generated from the executor entry, not a hallucinated chat answer.",
            "",
            "Core wiring evidence:",
        ]

        for key in ("router", "executor", "cognition", "proactive_import", "send_worker"):
            val = evidence.get(key, [])
            lines.append(f"- {key}: lines {val}")

        extra_keys = [k for k in evidence.keys() if k not in {"router", "executor", "cognition", "proactive_import", "send_worker"}]
        if extra_keys:
            lines.append("")
            lines.append("Additional evidence keys:")
            for key in extra_keys[:30]:
                val = evidence.get(key)
                if isinstance(val, list):
                    lines.append(f"- {key}: lines {val[:40]}{' ...' if len(val) > 40 else ''}")
                else:
                    lines.append(f"- {key}: {val}")

        lines.append("")
        lines.append("Issues:")
        if issues:
            for issue in issues:
                lines.append(f"- {issue}")
        else:
            lines.append("- This shallow line-hit scan flagged nothing — but it CANNOT detect "
                         "logic, runtime, or correctness bugs, so this is not evidence the file "
                         "is bug-free.")

        lines.append("")
        lines.append("Important limitation:")
        lines.append("- The current baseline GUI audit reports line-hit wiring evidence. It does not yet perform a deep semantic full-file audit.")

        return "\n".join(lines).strip()

    def _gui_runtime_audit_report(*args, **kwargs):  # type: ignore[no-redef]
        result = _ELI_GUI_RUNTIME_AUDIT_VISIBLE_PREV(*args, **kwargs)

        if not isinstance(result, dict):
            return result

        # Already valid: leave untouched.
        if str(result.get("content") or result.get("response") or "").strip():
            return result

        entry = result.get("entry")
        if not isinstance(entry, dict):
            return result

        content = _eli_gui_runtime_audit_entry_to_text(entry)

        fixed = dict(result)
        fixed["ok"] = bool(result.get("ok", True))
        fixed["action"] = "GUI_RUNTIME_AUDIT"
        fixed["content"] = content
        fixed["response"] = content

        meta = dict(fixed.get("meta") or {})
        meta.update({
            "response_mode": "executor_visible_audit_contract",
            "raw_executor_return": False,
            "suppress_gui_response": False,
            "visible_contract_repaired": True,
        })
        fixed["meta"] = meta

        return fixed

    log.debug("[EXECUTOR] GUI_RUNTIME_AUDIT visible result contract installed")

except Exception as _eli_gui_audit_visible_contract_err:
    log.debug(f"[EXECUTOR] GUI_RUNTIME_AUDIT visible result contract failed: {_eli_gui_audit_visible_contract_err}")
# =============================================================================

# =============================================================================
# ELI_EXECUTOR_FINAL_EXECUTE_ACTION_ALIAS_SYNC_V1
# Final safety sync: execute_action must expose the same final contract surface as
# execute. Several historical wrappers reassigned execute without also rebinding
# execute_action, leaving execute_action pinned to an older wrapper.
# =============================================================================
try:
    _ELI_EXECUTOR_FINAL_ALIAS_SYNC_PREV_EXECUTE_ACTION = globals().get("execute_action")
    if callable(globals().get("execute")):
        execute_action = execute
        try:
            execute_action._eli_final_alias_sync_v1 = True
        except Exception:
            pass
        log.debug("[EXECUTOR] final execute_action alias synced to execute")
except Exception as _eli_executor_final_alias_sync_err:
    log.debug(f"[EXECUTOR] final execute_action alias sync failed: {_eli_executor_final_alias_sync_err}")
# =============================================================================

# RUNTIME_STATUS evidence metadata normalization is now done inline by
# the mw_runtime_status_metadata middleware in the canonical table below.
# The legacy inert wrapper that lived here has been removed.
# =============================================================================


# --- Phase 11: multi-PDF helpers (consumed by mw_multipdf middleware) ------
# Wrapper machinery removed — only the helpers below are still referenced.
try:
    if not globals().get("_ELI_PHASE11_MULTIPDF_HELPERS_INSTALLED"):
        _ELI_PHASE11_MULTIPDF_HELPERS_INSTALLED = True

        def _eli_phase11_clean_paths(paths):
            if isinstance(paths, str):
                raw = [p.strip() for p in paths.split(",")]
            elif isinstance(paths, (list, tuple, set)):
                raw = list(paths)
            else:
                raw = []

            out = []
            seen = set()
            for p in raw:
                s = str(p or "").strip()
                if not s:
                    continue
                if s not in seen:
                    seen.add(s)
                    out.append(s)
            return out

        def _eli_phase11_format_multipdf_result(results, paths):
            lines = [
                f"Multi-PDF analysis completed for {len(paths)} file(s).",
                "",
            ]

            ok_count = 0
            fail_count = 0

            for idx, (path, res) in enumerate(zip(paths, results), 1):
                ok = bool(isinstance(res, dict) and res.get("ok"))
                if ok:
                    ok_count += 1
                else:
                    fail_count += 1

                status = "OK" if ok else "FAILED"
                lines.append(f"## {idx}. {Path(path).name} — {status}")
                lines.append(f"Source: `{path}`")

                if isinstance(res, dict):
                    saved_to = res.get("saved_to")
                    pages = res.get("pages")
                    chars = res.get("chars")
                    err = res.get("error")
                    response = res.get("response") or res.get("content") or ""

                    if pages is not None or chars is not None:
                        lines.append(f"Pages: {pages} | Characters: {chars}")
                    if saved_to:
                        lines.append(f"Saved to: `{saved_to}`")
                    if err:
                        lines.append(f"Error: {err}")
                    if response:
                        lines.append("")
                        lines.append(str(response))
                else:
                    lines.append(f"Unexpected executor result type: {type(res).__name__}")

                lines.append("")

            lines.insert(1, f"Successful: {ok_count} | Failed: {fail_count}")
            return "\n".join(lines).strip()

except Exception as _eli_phase11_multipdf_executor_err:
    log.debug(f"[EXECUTOR] Phase 11 multi-PDF helpers failed: {_eli_phase11_multipdf_executor_err}")

# =============================================================================
# ELI_EXECUTOR_CANONICAL_MIDDLEWARE_TABLE_V1
# Consolidates stacked execute wrappers into one explicit middleware chain.
# Legacy wrapper blocks above are retained as retired compatibility helpers.
# =============================================================================
try:
    if not globals().get("_ELI_EXECUTOR_CANONICAL_MIDDLEWARE_TABLE_V1"):
        _ELI_EXECUTOR_CANONICAL_MIDDLEWARE_TABLE_V1 = True

        # Capture the current outermost execute as the legacy base. After Tier 2
        # cleanup, this is the PM wrapper (the new top of the legacy chain).
        _ELI_EXECUTOR_CANONICAL_BASE = globals().get("execute")

        def _eli_exec_ctx_from_call(action, args, pargs, kwargs):
            real_action = action
            real_args = args

            if isinstance(action, dict):
                real_action = action.get("action")
                real_args = action.get("args", args) or action.get("params", args) or {}

            if not isinstance(real_args, dict):
                real_args = {}

            return {
                "action": str(real_action or "").upper().strip(),
                "args": dict(real_args or {}),
                "pargs": tuple(pargs or ()),
                "kwargs": dict(kwargs or {}),
            }

        def _eli_exec_mw_multipdf(ctx, nxt):
            if ctx["action"] != "ANALYZE_PDF":
                return nxt(ctx)

            clean_paths = globals().get("_eli_phase11_clean_paths")
            format_result = globals().get("_eli_phase11_format_multipdf_result")
            if not callable(clean_paths) or not callable(format_result):
                return nxt(ctx)

            paths = clean_paths((ctx.get("args") or {}).get("paths"))
            if len(paths) <= 1:
                return nxt(ctx)

            results = []
            for p in paths:
                one_ctx = {
                    "action": "ANALYZE_PDF",
                    "args": dict(ctx["args"]),
                    "pargs": tuple(ctx.get("pargs") or ()),
                    "kwargs": dict(ctx.get("kwargs") or {}),
                }
                one_ctx["args"]["path"] = p
                one_ctx["args"].pop("paths", None)
                results.append(nxt(one_ctx))

            all_ok = all(isinstance(r, dict) and r.get("ok") for r in results)
            content = format_result(results, paths)
            return {
                "ok": bool(all_ok),
                "action": "ANALYZE_PDF",
                "paths": paths,
                "results": results,
                "content": content,
                "response": content,
                "response_mode": "canonical_middleware_multipdf",
            }

        def _eli_exec_mw_runtime_status_metadata(ctx, nxt):
            out = nxt(ctx)
            if ctx["action"] != "RUNTIME_STATUS":
                return out

            if not isinstance(out, dict):
                txt = str(out or "").strip()
                out = {
                    "ok": bool(txt),
                    "action": "RUNTIME_STATUS",
                    "content": txt,
                    "response": txt,
                }
            else:
                out = dict(out)

            txt = str(out.get("content") or out.get("response") or "").strip()
            out["ok"] = bool(out.get("ok", bool(txt)))
            out["action"] = "RUNTIME_STATUS"
            out["grounded"] = True
            out["evidence_used"] = True

            if not out.get("source"):
                out["source"] = "runtime_status_executor_evidence_metadata_v1"
            if not out.get("evidence_source"):
                out["evidence_source"] = "runtime_status_live_runtime_telemetry"

            report = dict(out.get("report") or {})
            report.setdefault("repair_reason", "runtime_status_executor_evidence_metadata_v1")
            report.setdefault("metadata_normalised", True)
            report.setdefault("evidence_contract", "live_runtime_status_telemetry")
            report.setdefault("runtime_status_content_preserved", True)
            out["report"] = report

            if txt:
                out.setdefault("content", txt)
                out.setdefault("response", txt)

            return out

        def _eli_exec_mw_self_report_recent_updates(ctx, nxt):
            if ctx["action"] != "SELF_REPORT":
                return nxt(ctx)

            report_builder = globals().get("_eli_self_build_recent_updates_report")
            if not callable(report_builder):
                return nxt(ctx)

            real_args = dict(ctx.get("args") or {})
            scope = str(real_args.get("self_report_scope") or "").strip().lower()
            q = str(real_args.get("question") or "")
            low = q.lower()
            updateish = (
                scope == "recent_updates"
                or "what updates" in low
                or "updates and checks" in low
                or "checks have been" in low
                or "checks were" in low
                or "recent checks" in low
                or "routine updates" in low
                or "recent updates" in low
                or "what have you been doing" in low
                or "what have you been processing" in low
                or "what have you been working on" in low
            )
            if not updateish:
                return nxt(ctx)

            try:
                return report_builder(q)
            except Exception as e:
                msg = f"Self-report recent-updates evidence failed: {e!r}"
                return {
                    "ok": False,
                    "action": "SELF_REPORT",
                    "evidence_source": "self_report_recent_updates_error",
                    "error": repr(e),
                    "content": msg,
                    "response": msg,
                }

        def _eli_exec_mw_recent_memory_processing(ctx, nxt):
            if ctx["action"] != "MEMORY_STATUS":
                return nxt(ctx)

            real_args = dict(ctx.get("args") or {})
            if str(real_args.get("memory_scope") or "") != "recent_processing":
                return nxt(ctx)

            build_report = globals().get("_eli_recent_memory_processing_report")
            clean_report = globals().get("_eli_recent_mem_v2_clean_report")
            content_v2 = globals().get("_eli_recent_mem_v2_content")
            content_v1 = globals().get("_eli_recent_memory_processing_content")

            if not callable(build_report):
                return nxt(ctx)

            report = build_report(str(real_args.get("question") or ""))
            if callable(clean_report):
                report = clean_report(report or {})

            if callable(content_v2):
                content = content_v2(report or {})
                source = "recent_memory_processing_sqlite_clean_v2"
            elif callable(content_v1):
                content = content_v1(report or {})
                source = "recent_memory_processing_sqlite"
            else:
                content = str(report or "")
                source = "recent_memory_processing_sqlite"

            return {
                "ok": bool((report or {}).get("ok")),
                "action": "MEMORY_STATUS",
                "evidence_source": source,
                "report": report,
                "content": content,
                "response": content,
            }

        def _eli_exec_mw_memory_count(ctx, nxt):
            if ctx["action"] != "MEMORY_STATUS":
                return nxt(ctx)

            is_requested = globals().get("_eli_memory_count_is_requested")
            build_evidence = globals().get("_eli_memory_count_evidence")
            if callable(is_requested) and callable(build_evidence):
                try:
                    if is_requested(ctx.get("args") or {}):
                        return build_evidence()
                except Exception:
                    pass
            return nxt(ctx)

        def _eli_exec_mw_memory_runtime_sanitizer(ctx, nxt):
            out = nxt(ctx)
            if ctx["action"] != "EXPLAIN_MEMORY_RUNTIME":
                return out
            sanitizer = globals().get("_eli_sanitize_memory_runtime_output")
            if callable(sanitizer):
                try:
                    return sanitizer(out)
                except Exception:
                    return out
            return out

        def _eli_exec_mw_profile_scope(ctx, nxt):
            action_name = ctx["action"]
            real_args = dict(ctx.get("args") or {})

            scope_answer = globals().get("_eli_profile_scope_execute_answer")
            if action_name in {"PERSONAL_MEMORY_SUMMARY", "PERSONAL_MEMORY_DEEP_EXPLAIN"} and callable(scope_answer):
                try:
                    scoped = scope_answer(action_name, real_args)
                    if scoped is not None:
                        return scoped
                except Exception:
                    pass

            out = nxt(ctx)

            if action_name == "EXPLAIN_MEMORY_RUNTIME" and isinstance(out, dict):
                cleaner = globals().get("_eli_profile_scope_clean_memory_runtime_text")
                if callable(cleaner):
                    try:
                        txt = str(out.get("content") or out.get("response") or "")
                        cleaned = cleaner(txt)
                        out = dict(out)
                        out["content"] = cleaned
                        out["response"] = cleaned
                        out["evidence_source"] = out.get("evidence_source") or "memory_runtime_no_identity_preamble"
                    except Exception:
                        pass

            return out

        def _eli_exec_mw_identity_only(ctx, nxt):
            if ctx["action"] != "USER_IDENTITY_SUMMARY":
                return nxt(ctx)

            builder = globals().get("_eli_identity_only_response")
            if not callable(builder):
                return nxt(ctx)

            text = str(builder(dict(ctx.get("args") or {})) or "").strip()
            if not text:
                return nxt(ctx)
            return {
                "ok": True,
                "action": "USER_IDENTITY_SUMMARY",
                "content": text,
                "response": text,
                "evidence_source": "active_user_identity_scope",
                "generation_invoked": False,
            }

        _ELI_EXECUTOR_MIDDLEWARE_TABLE = (
            ("multipdf", _eli_exec_mw_multipdf),
            ("runtime_status_metadata", _eli_exec_mw_runtime_status_metadata),
            ("self_report_recent_updates", _eli_exec_mw_self_report_recent_updates),
            ("recent_memory_processing", _eli_exec_mw_recent_memory_processing),
            ("memory_count", _eli_exec_mw_memory_count),
            ("memory_runtime_sanitizer", _eli_exec_mw_memory_runtime_sanitizer),
            ("profile_scope", _eli_exec_mw_profile_scope),
            ("identity_only", _eli_exec_mw_identity_only),
        )

        def _eli_exec_core(ctx):
            base = globals().get("_ELI_EXECUTOR_CANONICAL_BASE")
            if not callable(base):
                msg = f"No executor base available for {ctx.get('action')}"
                return {"ok": False, "action": str(ctx.get("action") or ""), "content": msg, "response": msg}
            return base(
                str(ctx.get("action") or ""),
                dict(ctx.get("args") or {}),
                *(ctx.get("pargs") or ()),
                **(ctx.get("kwargs") or {}),
            )

        def _eli_exec_apply_middleware(ctx, index=0):
            if index >= len(_ELI_EXECUTOR_MIDDLEWARE_TABLE):
                if ctx.get("_pipeline_trace"):
                    _perf = __import__("time").perf_counter
                    _core_t0 = _perf()
                    out = _eli_exec_core(ctx)
                    _dt = (_perf() - _core_t0) * 1000.0
                    log.debug(
                        f"[PIPELINE][EXECUTOR] core action={ctx.get('action')} "
                        f"ok={bool(isinstance(out, dict) and out.get('ok'))} dt_ms={_dt:.2f}",
                    )
                    return out
                return _eli_exec_core(ctx)

            _name, mw = _ELI_EXECUTOR_MIDDLEWARE_TABLE[index]
            if ctx.get("_pipeline_trace"):
                _perf = __import__("time").perf_counter
                _mw_t0 = _perf()
                log.debug(
                    f"[PIPELINE][EXECUTOR] enter mw={_name} action={ctx.get('action')}",
                )
                out = mw(ctx, lambda next_ctx=None: _eli_exec_apply_middleware(ctx if next_ctx is None else next_ctx, index + 1))
                _dt = (_perf() - _mw_t0) * 1000.0
                _ok = bool(isinstance(out, dict) and out.get("ok"))
                log.debug(
                    f"[PIPELINE][EXECUTOR] exit  mw={_name} action={ctx.get('action')} "
                    f"ok={_ok} out_type={type(out).__name__} dt_ms={_dt:.2f}",
                )
                return out
            return mw(ctx, lambda next_ctx=None: _eli_exec_apply_middleware(ctx if next_ctx is None else next_ctx, index + 1))

        def execute(action, args=None, *pargs, **kwargs):  # type: ignore[no-redef]
            ctx = _eli_exec_ctx_from_call(action, args, pargs, kwargs)
            _trace = str(__import__("os").environ.get("ELI_PIPELINE_TRACE", "")).strip().lower() in {"1", "true", "yes", "on"}
            if _trace:
                ctx["_pipeline_trace"] = True
                _preview = str(ctx.get("action") or "")
                log.debug(f"[PIPELINE][EXECUTOR] begin action={_preview}")
            out = _eli_exec_apply_middleware(ctx, 0)
            if isinstance(out, dict):
                if _trace:
                    log.debug(
                        f"[PIPELINE][EXECUTOR] final action={out.get('action')} ok={out.get('ok')}",
                    )
                return out
            txt = str(out or "")
            return {
                "ok": bool(txt.strip()),
                "action": str(ctx.get("action") or ""),
                "content": txt,
                "response": txt,
            }

        execute_action = execute
        try:
            Executor.execute = lambda self, action, args=None, *a, **kw: execute(action, args, *a, **kw)  # type: ignore[name-defined]
        except Exception:
            pass

        log.debug("[EXECUTOR] canonical middleware table installed")
except Exception as _eli_executor_canonical_middleware_err:
    log.debug(f"[EXECUTOR] canonical middleware table install failed: {_eli_executor_canonical_middleware_err}")
